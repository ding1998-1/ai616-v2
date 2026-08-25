"""合同审查服务。

合同分析、问题 sidecar 和草案导出都在这里完成；路由层不直接调用 LLM 或操作文件系统。
"""

from __future__ import annotations

import io
import json
import logging
import re
from datetime import datetime
from pathlib import Path

from backend.llm_client import llm
from backend.services.document_service import (
    DOCS_DIR,
    compute_para_id,
    compute_text_hash,
    resolve_document,
)
from backend.services.knowledge_service import search_legal_provisions


logger = logging.getLogger(__name__)
CONTRACT_DATA_DIR = DOCS_DIR.parent / "contracts"
CONTRACT_DATA_DIR.mkdir(parents=True, exist_ok=True)


def issues_path(saved_name: str) -> Path:
    # 只允许存储文件名，避免 sidecar 路径穿越。
    name = Path(str(saved_name or "")).name
    if name != saved_name or not name:
        raise ValueError("文件名无效")
    return CONTRACT_DATA_DIR / f"{name}.issues.json"


def load_docx_paragraphs(saved_name: str) -> list[dict]:
    from docx import Document

    path = resolve_document(saved_name)
    if not path.exists():
        raise FileNotFoundError("文件不存在")
    if path.suffix.lower() != ".docx":
        raise ValueError("仅支持 .docx 文件")
    document = Document(path)
    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            paragraphs.append({"para_index": index, "para_id": compute_para_id(text), "text": text, "hash": compute_text_hash(text)})
    return paragraphs


def map_doc_structure(saved_name: str) -> dict:
    paragraphs = load_docx_paragraphs(saved_name)
    if not paragraphs:
        raise ValueError("文档中未找到有效段落")
    all_text = "".join(item["text"] for item in paragraphs)
    return {"success": True, "saved_name": saved_name, "doc_hash": compute_text_hash(all_text[:5000]), "paragraph_count": len(paragraphs), "paragraphs": paragraphs}


def _parse_json_array(text: str) -> list:
    match = re.search(r"\[[\s\S]*\]", str(text or ""))
    if not match:
        return []
    try:
        value = json.loads(match.group())
        return value if isinstance(value, list) else []
    except json.JSONDecodeError:
        return []


def _fallback_clauses(paragraphs: list[dict]) -> list[dict]:
    categories = (
        ("合同主体", ("甲方", "乙方", "主体", "统一社会信用代码")),
        ("权利义务", ("权利", "义务", "负责", "应当")),
        ("违约责任", ("违约", "赔偿", "违约金")),
        ("付款条款", ("付款", "支付", "价款", "发票")),
        ("解除条款", ("解除", "终止", "变更")),
        ("知识产权", ("知识产权", "著作权", "专利")),
        ("争议解决", ("争议", "仲裁", "诉讼", "法院")),
    )
    result = []
    for category, keywords in categories:
        matched = next((item for item in paragraphs if any(keyword in item["text"] for keyword in keywords)), None)
        result.append(
            {
                "category": category,
                "para_ids": [matched["para_id"]] if matched else [],
                "clause_text": matched["text"][:80] if matched else "无相关内容",
                "risk_summary": "请人工核验该类条款是否完整" if matched else "无相关内容",
            }
        )
    return result


def _classify_clauses(paragraphs: list[dict]) -> list[dict]:
    text = "\n".join(f"[{item['para_id']}] {item['text']}" for item in paragraphs)
    prompt = f"""你是资深合同审查专家。请把以下合同段落分类为合同主体、权利义务、违约责任、付款条款、解除条款、知识产权、争议解决、其他重要条款。
只输出合法 JSON 数组，每项包含 category、para_ids、clause_text、risk_summary。必须覆盖八类，没有相关内容写无相关内容。
合同段落：\n{text}"""
    if not getattr(llm, "api_key", ""):
        return _fallback_clauses(paragraphs)
    try:
        response = llm._generate([prompt], enable_thinking=False)
        data = _parse_json_array(response.generations[0].message.content)
        return data or _fallback_clauses(paragraphs)
    except Exception as exc:
        logger.warning("合同条款分类失败，使用关键词兜底：%s", exc)
        return _fallback_clauses(paragraphs)


def _best_paragraph(clause: dict, paragraphs: list[dict]) -> dict:
    ids = set(clause.get("para_ids") or [])
    matched = next((item for item in paragraphs if item["para_id"] in ids), None)
    if matched:
        return matched
    clause_text = str(clause.get("clause_text") or "")
    keywords = [char for char in clause_text if "\u4e00" <= char <= "\u9fff"]
    return max(paragraphs, key=lambda item: sum(char in item["text"] for char in keywords), default=paragraphs[0])


def _severity_and_suggestion(category: str, clause_text: str, risk_summary: str, provisions: list[dict]) -> tuple[str, dict]:
    if not getattr(llm, "api_key", ""):
        return "medium", {"issue_desc": risk_summary or "请人工核验条款", "suggested_text": "", "reason": "未配置模型，需人工复核"}
    provisions_text = "\n".join(f"- {item['content'][:300]}" for item in provisions[:3]) or "未检索到直接匹配的法条"
    severity = "medium"
    try:
        response = llm._generate([f"基于以下合同条款判断风险等级，只输出 high / medium / low：{category}\n{clause_text}\n{risk_summary}\n{provisions_text}"], enable_thinking=False)
        candidate = response.generations[0].message.content.strip().lower()
        severity = candidate if candidate in {"high", "medium", "low"} else "medium"
    except Exception:
        pass
    suggestion = {"issue_desc": risk_summary or "请人工核验条款", "suggested_text": "", "reason": "需结合合同上下文和适用制度人工复核"}
    try:
        response = llm._generate([f"请为以下合同风险给出 JSON 建议，字段 issue_desc、suggested_text、reason，不要输出其他内容。\n类别：{category}\n条款：{clause_text}\n风险：{risk_summary}\n依据：{provisions_text}"], enable_thinking=False)
        match = re.search(r"\{[\s\S]*\}", response.generations[0].message.content)
        if match:
            parsed = json.loads(match.group())
            if isinstance(parsed, dict):
                suggestion.update({key: parsed.get(key, value) for key, value in suggestion.items()})
    except Exception:
        pass
    return severity, suggestion


def analyze_contract(saved_name: str, doc_structure: list[dict] | None = None, extra_questions: list[str] | None = None) -> dict:
    paragraphs = doc_structure or load_docx_paragraphs(saved_name)
    clauses = _classify_clauses(paragraphs)
    review_points = []
    issue_id = 1
    for clause in clauses:
        category = clause.get("category", "其他重要条款")
        clause_text = str(clause.get("clause_text") or "")
        risk_summary = str(clause.get("risk_summary") or "")
        if not clause_text or clause_text == "无相关内容":
            continue
        primary = _best_paragraph(clause, paragraphs)
        provisions = search_legal_provisions(f"{category} {clause_text} {risk_summary}", top_k=3)
        severity, suggestion = _severity_and_suggestion(category, clause_text, risk_summary, provisions)
        review_points.append(
            {
                "id": issue_id,
                "para_id": primary["para_id"],
                "para_index": primary["para_index"],
                "text_hash": primary["hash"],
                "originalText": primary["text"],
                "category": category,
                "severity": severity,
                "issueDesc": suggestion.get("issue_desc", risk_summary),
                "suggestedText": suggestion.get("suggested_text", ""),
                "reason": suggestion.get("reason", ""),
                "rule": "\n".join(f"- {item['content'][:300]}" for item in provisions[:3]),
                "status": "open",
            }
        )
        issue_id += 1

    for question in extra_questions or []:
        primary = _best_paragraph({"clause_text": question}, paragraphs)
        review_points.append(
            {
                "id": issue_id,
                "para_id": primary["para_id"],
                "para_index": primary["para_index"],
                "text_hash": primary["hash"],
                "originalText": primary["text"],
                "category": "用户关注",
                "severity": "medium",
                "issueDesc": question,
                "suggestedText": "",
                "reason": "用户追加问题",
                "rule": "",
                "status": "open",
            }
        )
        issue_id += 1

    meta_path = DOCS_DIR / f"{Path(saved_name).name}.meta.json"
    bookmark_by_index = {}
    if meta_path.exists():
        try:
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
            bookmark_by_index = {item["para_index"]: item["bookmark_name"] for item in meta.get("paragraphs", [])}
        except (OSError, json.JSONDecodeError):
            pass
    for item in review_points:
        item["bookmark_name"] = bookmark_by_index.get(item["para_index"], f"audit_para_{item['para_index']}")

    payload = {"saved_name": saved_name, "doc_structure": paragraphs, "issues": review_points}
    issues_path(saved_name).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"success": True, "saved_name": saved_name, "review_points_count": len(review_points), "review_points": review_points, "message": f"审查完成，发现 {len(review_points)} 个审查点"}


def re_analyze_contract(saved_name: str, doc_structure: list[dict], extra_questions: list[str], previous_issues: list[dict] | None = None) -> dict:
    previous = previous_issues or []
    if not extra_questions:
        return {"success": True, "review_points": previous, "diff": {"new": [], "resolved": [], "modified": []}}
    current = analyze_contract(saved_name, doc_structure, extra_questions)
    new_issues = current["review_points"]
    previous_by_key = {(item.get("para_id"), item.get("text_hash")): item for item in previous}
    seen = {(item.get("para_id"), item.get("text_hash")) for item in new_issues}
    diff = {"new": [], "resolved": [], "modified": []}
    for item in new_issues:
        key = (item.get("para_id"), item.get("text_hash"))
        if key not in previous_by_key:
            diff["new"].append(item)
        elif item.get("issueDesc") != previous_by_key[key].get("issueDesc"):
            diff["modified"].append({**item, "previous_issue": previous_by_key[key]})
    for item in previous:
        key = (item.get("para_id"), item.get("text_hash"))
        if key not in seen and item.get("status") != "resolved":
            diff["resolved"].append({**item, "status": "resolved"})
    merged = []
    keys = set()
    for item in new_issues:
        key = (item.get("para_id"), item.get("text_hash"))
        if key not in keys:
            merged.append(item)
            keys.add(key)
    for item in previous:
        key = (item.get("para_id"), item.get("text_hash"))
        if key not in keys and item.get("status") != "resolved":
            merged.append({**item, "status": "resolved"})
            keys.add(key)
    payload = {"saved_name": saved_name, "doc_structure": doc_structure, "issues": merged}
    issues_path(saved_name).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"success": True, "review_points": merged, "diff": diff, "message": f"追加审查完成，新增 {len(diff['new'])} 个，修改 {len(diff['modified'])} 个，化解 {len(diff['resolved'])} 个"}


def get_contract_issues(saved_name: str) -> dict:
    path = issues_path(saved_name)
    if not path.exists():
        raise FileNotFoundError("未找到审查记录")
    return json.loads(path.read_text(encoding="utf-8"))


def delete_contract_issues(saved_name: str) -> None:
    path = issues_path(saved_name)
    if path.exists():
        path.unlink()


CONTRACT_DRAFT_PROMPT = """你是资深合同法律师。请根据用户需求起草一份完整、专业的中文合同。
合同类型：{contract_type}
用户需求：{requirements}
请按 Markdown 输出合同标题、合同双方、标的、价款与支付、双方权利义务、交付期限、验收、违约责任、变更解除、保密、争议解决、不可抗力、其他约定和签署栏。需要用户填写的内容使用[请填写]。只输出合同正文。"""


def draft_contract(contract_type: str, requirements: str) -> dict:
    prompt = CONTRACT_DRAFT_PROMPT.format(contract_type=contract_type, requirements=requirements)
    if getattr(llm, "api_key", ""):
        try:
            response = llm._generate([prompt], enable_thinking=True)
            markdown = response.generations[0].message.content.strip()
        except Exception as exc:
            logger.warning("合同草案生成失败，返回占位模板：%s", exc)
            markdown = ""
    else:
        markdown = ""
    if not markdown:
        markdown = f"# {contract_type or '合同草案'}\n\n## 一、合同需求\n{requirements}\n\n## 二、待补充事项\n- [请填写]"
    title = contract_type or "合同草案"
    for line in markdown.splitlines():
        if line.strip().startswith("# "):
            title = line.strip()[2:].strip() or title
            break
    return {"success": True, "markdown": markdown, "title": title}


def export_contract_draft(markdown: str, title: str) -> tuple[str, bytes]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    for raw_line in str(markdown or "").splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph()
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line[2:].strip() or title or "合同草案")
            run.bold = True
            run.font.size = Pt(18)
        elif line.startswith("## "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(14)
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(line[2:].strip())
        else:
            paragraph = document.add_paragraph(line)
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
            paragraph.paragraph_format.space_after = Pt(6)
    output = io.BytesIO()
    document.save(output)
    safe_title = re.sub(r"[^\w\u4e00-\u9fff-]", "_", title or "合同草案")
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"{safe_title}_{timestamp}.docx"
    return filename, output.getvalue()
