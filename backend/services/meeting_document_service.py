"""Lossless meeting document outputs for Records Pipeline v2.

The service deliberately has no FastAPI, database, or environment coupling.
It accepts the already-normalised ``generatedRecords`` payload and a list of
chronicle rows, then writes two independent files:

* the formal meeting document, containing only distilled fields; and
* the evidence draft, containing every source row (including rows classified
  as noise/excluded in an audit appendix), gap markers, chunk boundaries and
  a reverse index back to the formal entries.

This separation is the enforcement point for the quality requirement that no
raw Whisper transcript is silently copied into the formal red-head document.
The caller may use the returned manifest as the persistence boundary; this
module does not write a database or mutate a meeting record.
"""

from __future__ import annotations

import hashlib
import json
import os
import platform
import re
from difflib import SequenceMatcher
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

from backend.services.meeting_proofread_service import (
    ProofreadRequiredError,
    records_are_proofread,
)


DOCUMENT_SERVICE_VERSION = "meeting-document-v4-template-switch"
GAP_THRESHOLD_SECONDS = 120.0
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
FORMAL_MINUTES_TEMPLATE = (
    Path(__file__).resolve().parents[1]
    / "assets"
    / "templates"
    / "generic_meeting_minutes_v1.docx"
)
ENTERPRISE_MINUTES_TEMPLATE = (
    Path(__file__).resolve().parents[1]
    / "assets"
    / "templates"
    / "enterprise_redhead_minutes_v1.docx"
)
_TIME_RE = re.compile(r"^(?:(\d+):)?(\d{1,2}):(\d{2})(?:\.(\d+))?$")
_TIME_RANGE_RE = re.compile(
    r"(?P<start>\d{1,2}:\d{2}(?::\d{2})?(?:\.\d+)?)\s*[-–—]\s*"
    r"(?P<end>\d{1,2}:\d{2}(?::\d{2})?(?:\.\d+)?)"
)
_WHITESPACE_RE = re.compile(r"\s+")


@dataclass(frozen=True)
class DocumentArtifact:
    """A generated file and its audit metadata."""

    kind: str
    path: str
    filename: str
    sha256: str
    mimeType: str = DOCX_MIME

    def to_dict(self) -> dict[str, str]:
        return {
            "kind": self.kind,
            "path": self.path,
            "filename": self.filename,
            "sha256": self.sha256,
            "mimeType": self.mimeType,
        }


def _as_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _compact_text(value: Any, *, limit: int = 0) -> str:
    text = _WHITESPACE_RE.sub(" ", _as_text(value)).strip()
    return text[:limit] if limit and len(text) > limit else text


def _first(mapping: Mapping[str, Any], *keys: str) -> Any:
    for key in keys:
        value = mapping.get(key)
        if value is not None and value != "":
            return value
    return None


def _rows_from_source(source: Any) -> list[dict[str, Any]]:
    if isinstance(source, Mapping):
        source = _first(source, "segments", "chronicle", "transcripts", "rows", "items") or []
    if isinstance(source, (str, bytes)) or not isinstance(source, Iterable):
        return []
    rows: list[dict[str, Any]] = []
    for index, raw in enumerate(source):
        if isinstance(raw, Mapping):
            row = dict(raw)
        elif isinstance(raw, str):
            row = {"text": raw}
        else:
            continue
        row.setdefault("_sourceIndex", index)
        text = _as_text(_first(row, "text", "content", "transcript", "sentence"))
        row["_text"] = text
        row["_segmentId"] = _as_text(_first(row, "segmentId", "segment_id", "id", "uid")) or f"segment-{index + 1:04d}"
        row["_fileId"] = _as_text(_first(row, "fileId", "file_id", "audioFileId", "recordingId", "fileName", "filename")) or "file-1"
        row["_fileName"] = _as_text(_first(row, "fileName", "filename", "file_name")) or row["_fileId"]
        row["_speaker"] = _as_text(_first(row, "speaker", "speakerName", "speaker_name", "role")) or "说话人未识别"
        row["_time"] = _as_text(_first(row, "time", "timeRange", "timestamp", "clientTime", "serverTime"))
        row["_reason"] = _compact_text(_first(row, "excludeReason", "excludedReason", "noiseReason", "discardReason", "reason"))
        row["_excluded"] = _is_excluded(row)
        rows.append(row)
    return rows


def _is_excluded(row: Mapping[str, Any]) -> bool:
    for key in ("excluded", "isExcluded", "discarded", "filtered", "isNoise", "noise"):
        value = row.get(key)
        if isinstance(value, bool) and value:
            return True
        if isinstance(value, str) and value.strip().lower() in {"1", "true", "yes", "noise", "excluded", "剔除", "噪音"}:
            return True
    status = _as_text(row.get("status") or row.get("quality")).lower()
    return status in {"noise", "excluded", "discarded", "filtered", "剔除", "噪音", "无效"}


def _parse_seconds(value: Any) -> float | None:
    if value is None or value == "":
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = _as_text(value)
    try:
        return float(text)
    except ValueError:
        pass
    match = _TIME_RE.match(text)
    if not match:
        return None
    hours, minutes, seconds, fraction = match.groups()
    result = int(hours or 0) * 3600 + int(minutes) * 60 + int(seconds)
    return result + (float(f"0.{fraction}") if fraction else 0.0)


def _format_seconds(value: float | None) -> str:
    if value is None:
        return "未标注"
    total = max(0, int(round(value)))
    hours, remainder = divmod(total, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


def _row_bounds(row: Mapping[str, Any]) -> tuple[float | None, float | None]:
    start = _parse_seconds(_first(row, "start", "startTime", "start_time", "begin", "offsetStart"))
    end = _parse_seconds(_first(row, "end", "endTime", "end_time", "finish", "offsetEnd"))
    if start is None or end is None:
        text = _as_text(_first(row, "timeRange", "time"))
        match = _TIME_RANGE_RE.search(text)
        if match:
            start = start if start is not None else _parse_seconds(match.group("start"))
            end = end if end is not None else _parse_seconds(match.group("end"))
        elif start is None:
            start = _parse_seconds(text)
    if end is None and start is not None:
        end = start
    return start, end


def _format_gap(start: float | None, end: float | None) -> str:
    return f"{_format_seconds(start)}-{_format_seconds(end)}"


def _snapshot(records: Mapping[str, Any], rows: Sequence[Mapping[str, Any]]) -> dict[str, Any]:
    value = records.get("generationSnapshot")
    snapshot = dict(value) if isinstance(value, Mapping) else {}
    payload = [
        {
            "segmentId": row.get("_segmentId"),
            "fileId": row.get("_fileId"),
            "start": _first(row, "start", "startTime", "start_time"),
            "end": _first(row, "end", "endTime", "end_time"),
            "speaker": row.get("_speaker"),
            "text": row.get("_text"),
        }
        for row in rows
    ]
    input_hash = hashlib.sha256(
        json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
    ).hexdigest()
    snapshot.setdefault("provider", "local")
    snapshot.setdefault("model", "local-qwen")
    snapshot.setdefault("temperature", 0.1)
    snapshot.setdefault("promptVersion", "records-v2")
    snapshot.setdefault("schemaVersion", "meeting-records-v2")
    snapshot.setdefault("glossaryVersion", "1")
    snapshot.setdefault("chunkPolicy", "audio-file-boundary+time-4000-chars")
    snapshot.setdefault("pipeline", records.get("pipeline") or "records-v2")
    snapshot.setdefault("pipelineVersion", snapshot.get("pipeline") or "records-v2")
    snapshot.setdefault("inputSha256", input_hash)
    snapshot.setdefault("segmentCount", len(rows))
    snapshot.setdefault("generatedAt", datetime.now(timezone.utc).replace(microsecond=0).isoformat())
    snapshot["documentServiceVersion"] = DOCUMENT_SERVICE_VERSION
    return snapshot


def _basis(item: Mapping[str, Any]) -> tuple[str, list[str], list[str]]:
    value = item.get("basis")
    basis = value if isinstance(value, Mapping) else {}
    time_range = _compact_text(_first(basis, "timeRange", "time_range", "time"))
    quotes: list[str] = []
    segment_ids: list[str] = []
    raw_quotes = basis.get("quotes")
    if isinstance(raw_quotes, list):
        for raw_quote in raw_quotes:
            if isinstance(raw_quote, Mapping):
                text = _compact_text(_first(raw_quote, "text", "quote", "evidence"), limit=300)
                if text and text not in quotes:
                    quotes.append(text)
                sid = _as_text(_first(raw_quote, "segmentId", "segment_id", "id"))
                if sid and sid not in segment_ids:
                    segment_ids.append(sid)
            elif _as_text(raw_quote):
                text = _compact_text(raw_quote, limit=300)
                if text not in quotes:
                    quotes.append(text)
    return time_range, quotes, segment_ids


def _summary_sections(records: Mapping[str, Any]) -> tuple[list[Any], list[Any], list[Any]]:
    summary = records.get("summary")
    if isinstance(summary, Mapping):
        conclusions = list(summary.get("conclusions") or [])
        risks = list(summary.get("risks") or [])
        todos = list(summary.get("todos") or [])
    elif isinstance(summary, list):
        conclusions = list(summary)
        risks = []
        todos = []
    else:
        conclusions, risks, todos = [], [], []
    return (
        list(records.get("decisions") or conclusions),
        list(records.get("risks") or risks) + list(records.get("disclosures") or []),
        list(records.get("todos") or todos),
    )


def _item_content(item: Any, *keys: str) -> str:
    if isinstance(item, Mapping):
        return _compact_text(_first(item, *keys), limit=1000)
    return _compact_text(item, limit=1000)


def _formal_basis(item: Any) -> str:
    if not isinstance(item, Mapping):
        return ""
    basis = item.get("basis") if isinstance(item.get("basis"), Mapping) else {}
    if basis.get("evidenceValid") is False:
        return ""
    time_range, quotes, _ = _basis(item)
    details: list[str] = []
    if time_range:
        details.append(f"时间 {time_range}")
    if quotes:
        details.append("引句：" + "；".join(f"“{quote}”" for quote in quotes[:2]))
    return "；".join(details)


def render_formal_markdown(
    meeting: Mapping[str, Any],
    records: Mapping[str, Any],
) -> str:
    """Render the distilled formal content without iterating chronicle rows."""

    title = _compact_text(meeting.get("title") or meeting.get("name") or "会议纪要")
    meeting_type = _compact_text(meeting.get("type") or "会议")
    decisions, risks, todos = _summary_sections(records)
    minutes = list(records.get("minutes") or [])
    snapshot = _snapshot(records, [])
    lines = [f"# {meeting_type}会议纪要", "", "## 一、会议基本信息", ""]
    lines.extend([
        f"- 会议名称：{title}",
        f"- 会议日期：{_compact_text(meeting.get('date') or meeting.get('startTime') or '未标注')}",
        f"- 会议类型：{meeting_type}",
        f"- 所属项目：{_compact_text(meeting.get('project') or '未标注')}",
        "- 原始转写：见独立《证据底稿》附件（不纳入本正式件正文）",
        "",
        "## 二、结论与决议",
        "",
    ])
    if decisions:
        for index, item in enumerate(decisions, 1):
            if isinstance(item, Mapping):
                content = _item_content(item, "content", "decision", "title", "summary")
                kind = _compact_text(item.get("type") or "知悉")
                status = _compact_text(item.get("status") or "待确认")
                basis = _formal_basis(item)
                suffix = f"（{kind} · {status}）"
                if basis:
                    suffix += f"\n  - 依据：{basis}"
                lines.append(f"{index}. {content}{suffix}")
            else:
                lines.append(f"{index}. {_item_content(item)}")
    else:
        lines.append("（本次会议未形成有可核验原文支持的明确决议）")
    lines.extend(["", "## 三、合规风险与披露事项", ""])
    if risks:
        for index, item in enumerate(risks, 1):
            if isinstance(item, Mapping):
                content = _item_content(item, "content", "description", "title", "summary")
                severity = _compact_text(item.get("severity") or item.get("level") or "中")
                basis = _formal_basis(item)
                lines.append(f"{index}. [{severity}] {content}")
                if basis:
                    lines.append(f"   - 依据：{basis}")
            else:
                lines.append(f"{index}. {_item_content(item)}")
    else:
        lines.append("（暂无已识别风险或披露事项）")
    lines.extend(["", "## 四、会议纪要", ""])
    if minutes:
        for index, item in enumerate(minutes, 1):
            if not isinstance(item, Mapping):
                lines.append(f"{index}. {_item_content(item)}")
                continue
            agenda = _item_content(item, "agenda", "title", "topic", "content") or "议题"
            lines.append(f"### {index}. {agenda}")
            status = _compact_text(item.get("status") or "待整理")
            lines.append(f"- 状态：{status}")
            for point in item.get("keyPoints") or item.get("key_points") or []:
                if _as_text(point):
                    lines.append(f"- {_compact_text(point, limit=500)}")
            basis = _formal_basis(item)
            if basis:
                lines.append(f"- 依据：{basis}")
            lines.append("")
    else:
        lines.append("（暂无会议纪要）")
    lines.extend(["", "## 五、待办事项", ""])
    if todos:
        lines.extend(["| 序号 | 待办事项 | 责任人 | 截止时间 |", "|---|---|---|---|"])
        for index, item in enumerate(todos, 1):
            if isinstance(item, Mapping):
                task = _item_content(item, "task", "content", "title", "summary")
                owner = _compact_text(item.get("owner") or item.get("assignee") or "待确认")
                deadline = _compact_text(item.get("deadline") or "待定")
                lines.append(f"| {index} | {task} | {owner} | {deadline} |")
            else:
                lines.append(f"| {index} | {_item_content(item)} | 待确认 | 待定 |")
    else:
        lines.append("（暂无待办事项）")
    lines.extend(["", "## 六、生成参数快照", ""])
    for key in (
        "provider", "model", "temperature", "pipelineVersion", "promptVersion",
        "schemaVersion", "glossaryVersion", "chunkPolicy", "chunkCount",
        "mapCallCount", "reduceCallCount", "inputSha256", "segmentCount", "generatedAt",
    ):
        if key in snapshot and snapshot[key] not in (None, ""):
            lines.append(f"- {key}：{snapshot[key]}")
    lines.extend(["", "---", "本正式件仅包含经校对的蒸馏内容；全量原始证据请查阅独立《证据底稿》。"])
    return "\n".join(lines) + "\n"


def _gap_markers(
    rows: Sequence[Mapping[str, Any]],
    *,
    file_offsets: Mapping[str, Any] | None = None,
) -> list[dict[str, Any]]:
    offsets = {str(key): _parse_seconds(value) or 0.0 for key, value in (file_offsets or {}).items()}
    markers: list[dict[str, Any]] = []
    previous: tuple[str, float] | None = None
    for row in rows:
        start, end = _row_bounds(row)
        if start is None:
            continue
        file_id = _as_text(row.get("_fileId"))
        absolute_start = start + offsets.get(file_id, 0.0)
        if previous is not None and previous[0] == file_id:
            duration = absolute_start - previous[1]
            if duration > GAP_THRESHOLD_SECONDS:
                markers.append({
                    "fileId": file_id,
                    "fromSeconds": previous[1],
                    "toSeconds": absolute_start,
                    "durationSeconds": duration,
                    "label": f"[录音断档 {_format_gap(previous[1], absolute_start)}]",
                })
        previous = (file_id, (end if end is not None else start) + offsets.get(file_id, 0.0))
    return markers


def _map_boundaries(records: Mapping[str, Any]) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    for item in records.get("mapResults") or []:
        if not isinstance(item, Mapping):
            continue
        result.append({
            "chunkId": _as_text(item.get("chunkId") or item.get("id")),
            "fileId": _as_text(item.get("fileId")),
            "timeRange": _compact_text(item.get("timeRange")),
            "sourceSegmentIds": list(item.get("sourceSegmentIds") or [
                row.get("segmentId") for row in item.get("chunkSegments") or []
                if isinstance(row, Mapping) and row.get("segmentId")
            ]),
        })
    return result


def _record_index(records: Mapping[str, Any]) -> dict[str, list[str]]:
    index: dict[str, list[str]] = {}
    for field, label in (
        ("minutes", "会议纪要"),
        ("decisions", "决议"),
        ("risks", "风险"),
        ("disclosures", "披露"),
        ("todos", "待办"),
    ):
        for position, item in enumerate(records.get(field) or [], 1):
            if not isinstance(item, Mapping):
                continue
            _, _, ids = _basis(item)
            for segment_id in ids:
                entry = f"{label}[{position}]"
                if entry not in index.setdefault(segment_id, []):
                    index[segment_id].append(entry)
    return index


def build_evidence_manifest(
    records: Mapping[str, Any],
    chronicle: Any,
    *,
    file_offsets: Mapping[str, Any] | None = None,
    markers: Sequence[Mapping[str, Any]] | None = None,
) -> dict[str, Any]:
    """Build the lossless evidence structure used by the appendix document."""

    rows = _rows_from_source(chronicle)
    gaps = _gap_markers(rows, file_offsets=file_offsets)
    excluded = [row for row in rows if row.get("_excluded")]
    valid = [row for row in rows if not row.get("_excluded") and row.get("_text")]
    quality_flags: list[dict[str, Any]] = []
    for gap in gaps:
        quality_flags.append({"type": "audio_gap", **gap})
    for row in rows:
        if row.get("_excluded"):
            quality_flags.append({
                "type": "excluded",
                "segmentId": row.get("_segmentId"),
                "reason": row.get("_reason") or "标记为噪音/无效",
            })
        elif not row.get("_text"):
            quality_flags.append({
                "type": "empty_text",
                "segmentId": row.get("_segmentId"),
                "reason": "转写为空",
            })
    def public_row(row: Mapping[str, Any]) -> dict[str, Any]:
        return {
            "segmentId": row.get("_segmentId"),
            "fileId": row.get("_fileId"),
            "fileName": row.get("_fileName"),
            "time": row.get("_time") or _format_gap(*_row_bounds(row)),
            "start": _first(row, "start", "startTime", "start_time"),
            "end": _first(row, "end", "endTime", "end_time"),
            "speaker": row.get("_speaker"),
            "rawText": _as_text(_first(row, "rawText", "originalText", "text", "content", "transcript")),
            "correctedText": _as_text(_first(row, "correctedText", "correctedContent")) or row.get("_text"),
            "excluded": bool(row.get("_excluded")),
            "excludeReason": row.get("_reason") or ("标记为噪音/无效" if row.get("_excluded") else ""),
        }
    all_rows = [public_row(row) for row in rows]
    source_ids = [row["segmentId"] for row in all_rows if row.get("segmentId")]
    return {
        "rows": [row for row in all_rows if not row["excluded"]],
        "excludedRows": [row for row in all_rows if row["excluded"]],
        "allRows": all_rows,
        "gapMarkers": gaps,
        "qualityFlags": quality_flags,
        "mapBoundaries": _map_boundaries(records),
        "reverseIndex": _record_index(records),
        "evidenceExceptions": [
            dict(item) for item in (records.get("evidenceExceptions") or [])
            if isinstance(item, Mapping)
        ],
        "markers": [dict(item) for item in (markers or []) if isinstance(item, Mapping)],
        "coverage": {
            "sourceSegmentCount": len(source_ids),
            "evidenceSegmentCount": len(source_ids),
            "coverageRatio": 1.0 if source_ids else 0.0,
            "validSegmentCount": len(valid),
            "excludedSegmentCount": len(excluded),
        },
    }


def render_evidence_markdown(
    meeting: Mapping[str, Any],
    records: Mapping[str, Any],
    manifest: Mapping[str, Any],
) -> str:
    """Render every source row and every quality annotation for audit."""

    title = _compact_text(meeting.get("title") or meeting.get("name") or "会议")
    snapshot = _snapshot(records, list(manifest.get("allRows") or []))
    lines = [f"# {title}｜证据底稿", "", "## 一、底稿说明", "", "本附件保留会议转写的完整证据链。正式会议文件只引用蒸馏结果，本附件不作为正式正文。", ""]
    lines.extend(["## 二、质量标记", ""])
    flags = list(manifest.get("qualityFlags") or [])
    if flags:
        for flag in flags:
            kind = _compact_text(flag.get("type") or "quality")
            detail = flag.get("label") or flag.get("reason") or flag.get("segmentId") or ""
            lines.append(f"- {kind}：{_compact_text(detail)}")
    else:
        lines.append("- 未发现断档、噪音或空白标记")
    lines.extend(["", "## 三、全量有效转写", ""])
    valid_rows = list(manifest.get("rows") or [])
    if valid_rows:
        for row in valid_rows:
            lines.append(
                f"- [{_compact_text(row.get('time') or '未标注')}] "
                f"{_compact_text(row.get('speaker') or '说话人未识别')}：{_as_text(row.get('rawText'))}"
            )
            corrected = _as_text(row.get("correctedText"))
            if corrected and corrected != _as_text(row.get("rawText")):
                lines.append(f"  - 校对稿：{corrected}")
    else:
        lines.append("（暂无有效转写）")
    lines.extend(["", "## 四、已剔除行（保留审计）", ""])
    excluded_rows = list(manifest.get("excludedRows") or [])
    if excluded_rows:
        for row in excluded_rows:
            lines.append(
                f"- [{_compact_text(row.get('time') or '未标注')}] "
                f"{_compact_text(row.get('speaker') or '说话人未识别')}：{_as_text(row.get('rawText'))}"
                f"（原因：{_compact_text(row.get('excludeReason') or '标记为噪音/无效')}）"
            )
    else:
        lines.append("（无剔除行）")
    lines.extend(["", "## 五、Map 分段边界", ""])
    boundaries = list(manifest.get("mapBoundaries") or [])
    if boundaries:
        for boundary in boundaries:
            lines.append(
                f"- {_compact_text(boundary.get('chunkId') or '未命名')}｜文件 {_compact_text(boundary.get('fileId') or '未标注')}｜"
                f"{_compact_text(boundary.get('timeRange') or '未标注')}｜"
                f"{len(boundary.get('sourceSegmentIds') or [])} 个分段"
            )
    else:
        lines.append("（未提供 Map 分段信息）")
    lines.extend(["", "## 六、正式条目反向索引", ""])
    reverse = manifest.get("reverseIndex") or {}
    if reverse:
        for segment_id, entries in reverse.items():
            lines.append(f"- {segment_id}：{', '.join(_as_text(item) for item in entries)}")
    else:
        lines.append("（暂无带依据的正式条目）")
    lines.extend(["", "## 七、系统自动排除的 AI 表述", ""])
    exceptions = list(manifest.get("evidenceExceptions") or [])
    if exceptions:
        for exception in exceptions:
            item = exception.get("item") if isinstance(exception.get("item"), Mapping) else {}
            content = _item_content(item, "content", "task", "agenda", "title", "summary")
            lines.append(
                f"- {_compact_text(exception.get('field') or '未分类')}：{content or '未记录内容'}"
                f"（原因：{_compact_text(exception.get('reason') or '无可核验原文依据')}）"
            )
    else:
        lines.append("（无自动排除表述）")
    lines.extend(["", "## 八、会中标记", ""])
    custom_markers = list(manifest.get("markers") or [])
    if custom_markers:
        for marker in custom_markers:
            lines.append(f"- {_compact_text(marker.get('time') or marker.get('timestamp') or '未标注')}：{_compact_text(marker.get('note') or marker.get('text') or marker)}")
    else:
        lines.append("（无会中标记）")
    lines.extend(["", "## 九、生成参数快照", ""])
    for key in (
        "provider", "model", "temperature", "pipelineVersion", "promptVersion",
        "schemaVersion", "glossaryVersion", "chunkPolicy", "chunkCount",
        "mapCallCount", "reduceCallCount", "inputSha256", "segmentCount", "generatedAt",
    ):
        if key in snapshot and snapshot[key] not in (None, ""):
            lines.append(f"- {key}：{snapshot[key]}")
    lines.extend(["", "---", f"底稿覆盖率：{manifest.get('coverage', {}).get('coverageRatio', 0.0):.0%}（有效与剔除行均保留）"])
    return "\n".join(lines) + "\n"


def _cjk_font_name(role: str = "body") -> str:
    env_key = {
        "title": "MEETING_DOCX_TITLE_FONT",
        "heading": "MEETING_DOCX_HEADING_FONT",
        "body": "MEETING_DOCX_CJK_FONT",
    }.get(role, "MEETING_DOCX_CJK_FONT")
    configured = os.getenv(env_key, "").strip()
    if configured:
        return configured
    if platform.system() == "Darwin":
        return {"title": "FZCuSong-B09S", "heading": "Alibaba PuHuiTi", "body": "STFangsong"}.get(role, "STFangsong")
    return {"title": "Source Han Serif CN", "heading": "Source Han Sans CN", "body": "Source Han Serif CN"}.get(role, "Source Han Serif CN")


def _set_run_font(run: Any, *, size: float = 11, bold: bool = False, color: str | None = None, role: str = "body") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    font_name = _cjk_font_name(role)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _configure_document(doc: Any, *, title: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.6)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    styles = doc.styles
    normal = styles["Normal"]
    font_name = _cjk_font_name("body")
    normal.font.name = font_name
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    def set_style_fonts(style: Any) -> None:
        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{key}"), font_name)
    set_style_fonts(normal)
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = _cjk_font_name("heading")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        set_style_fonts(style)
    header = section.header.paragraphs[0]
    header.text = title
    for run in header.runs:
        _set_run_font(run, size=9, color="667085")
    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    footer.text = "AI 会议系统 · 生成文件"
    for run in footer.runs:
        _set_run_font(run, size=9, color="667085")


def _add_table_geometry(table: Any, widths_dxa: Sequence[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = 1


def _apply_table_grid_borders(table: Any) -> None:
    """Apply deterministic grid borders without depending on a template style name."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def _docx_paragraph(doc: Any, text: str, *, style: str | None = None, bold: bool = False, size: float = 11) -> Any:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    return paragraph


def _set_cell_text(cell: Any, text: Any, *, size: float = 10.5, bold: bool = False, align: int = 0, role: str = "body") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(_compact_text(text) or "待补充")
    _set_run_font(run, size=size, bold=bold, role=role)


def _meeting_value(meeting: Mapping[str, Any], *keys: str, default: str = "待补充") -> str:
    for key in keys:
        value = meeting.get(key)
        if isinstance(value, Sequence) and not isinstance(value, (str, bytes)):
            names = [
                _compact_text(item.get("name") or item.get("displayName") or item)
                if isinstance(item, Mapping) else _compact_text(item)
                for item in value
            ]
            joined = "、".join(item for item in names if item)
            if joined:
                return joined
        elif _compact_text(value):
            return _compact_text(value)
    return default


def _agenda_match_score(left: Any, right: Any) -> float:
    """Return a conservative similarity score for two Chinese agenda labels."""

    left_text = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "", _compact_text(left)).lower()
    right_text = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "", _compact_text(right)).lower()
    if not left_text or not right_text:
        return 0.0
    if left_text == right_text:
        return 1.0
    if min(len(left_text), len(right_text)) >= 4 and (
        left_text in right_text or right_text in left_text
    ):
        return 0.92
    sequence_score = SequenceMatcher(None, left_text, right_text).ratio()
    left_pairs = {left_text[index:index + 2] for index in range(len(left_text) - 1)}
    right_pairs = {right_text[index:index + 2] for index in range(len(right_text) - 1)}
    shared_pairs = left_pairs & right_pairs
    pair_score = len(shared_pairs) / max(1, min(len(left_pairs), len(right_pairs)))
    if len(shared_pairs) < 2:
        pair_score = 0.0
    return max(sequence_score, pair_score)


def _normalise_discussion_evidence(value: Any) -> str:
    text = _compact_text(value, limit=1600)
    if not text:
        return ""
    text = re.sub(r"\s*(?:\.{3,}|…+)\s*", "；", text)
    text = re.sub(r"；+", "；", text).strip("；，。 ")
    return f"{text}。" if text and text[-1] not in "。！？" else text


def _minute_discussion_points(item: Any, records: Mapping[str, Any]) -> list[str]:
    """Resolve reliable discussion text without emitting empty placeholders."""

    if not isinstance(item, Mapping):
        return []
    explicit = [
        _compact_text(point, limit=1600)
        for point in item.get("keyPoints") or item.get("key_points") or []
        if _compact_text(point)
    ]
    if explicit:
        return explicit

    agenda = _item_content(item, "agenda", "title", "topic", "content")
    basis = item.get("basis") if isinstance(item.get("basis"), Mapping) else {}
    if basis.get("evidenceValid"):
        quotes = basis.get("quotes") if isinstance(basis.get("quotes"), list) else []
        quote_text = "；".join(
            _compact_text(quote.get("text") if isinstance(quote, Mapping) else quote)
            for quote in quotes
            if _compact_text(quote.get("text") if isinstance(quote, Mapping) else quote)
        )
        normalised = _normalise_discussion_evidence(quote_text)
        if normalised and _agenda_match_score(agenda, quote_text) >= 0.28:
            return [normalised]

    best_topic: tuple[float, Mapping[str, Any]] | None = None
    for map_result in records.get("mapResults") or []:
        if not isinstance(map_result, Mapping):
            continue
        output = map_result.get("output") if isinstance(map_result.get("output"), Mapping) else {}
        for topic in output.get("topics") or []:
            if not isinstance(topic, Mapping):
                continue
            score = _agenda_match_score(
                agenda, _item_content(topic, "title", "agenda", "topic")
            )
            if score >= 0.66 and (best_topic is None or score > best_topic[0]):
                best_topic = (score, topic)
    if best_topic:
        evidence = _normalise_discussion_evidence(
            _item_content(best_topic[1], "evidence", "summary", "content")
        )
        if evidence:
            return [evidence]

    return []


def _formal_record_lines(records: Mapping[str, Any]) -> list[tuple[str, bool]]:
    decisions, risks, todos = _summary_sections(records)
    lines: list[tuple[str, bool]] = []
    minutes = [
        (item, _minute_discussion_points(item, records))
        for item in list(records.get("minutes") or [])
    ]
    minutes = [(item, points) for item, points in minutes if points]
    if minutes:
        lines.append(("一、议题及讨论记录", True))
        for index, (item, points) in enumerate(minutes, 1):
            title = _item_content(item, "agenda", "title", "topic", "content")
            lines.append((f"{index}. {title or '未命名议题'}", True))
            if isinstance(item, Mapping):
                for point in points:
                    lines.append((f"　　{_compact_text(point, limit=800)}", False))
                basis = _formal_basis(item)
                if basis:
                    lines.append((f"　　依据：{basis}", False))
    if decisions:
        lines.append(("二、结论与决议", True))
        for index, item in enumerate(decisions, 1):
            lines.append((f"{index}. {_item_content(item, 'content', 'decision', 'description', 'title', 'summary')}", False))
    if risks:
        lines.append(("三、合规风险与披露事项", True))
        for index, item in enumerate(risks, 1):
            severity = f"〔{_compact_text(item.get('severity'))}〕" if isinstance(item, Mapping) and item.get("severity") else ""
            lines.append((f"{index}. {severity}{_item_content(item, 'content', 'description', 'title', 'summary')}", False))
    if todos:
        lines.append(("四、待办事项", True))
        for index, item in enumerate(todos, 1):
            owner = _compact_text(item.get("owner")) if isinstance(item, Mapping) else "待确认"
            deadline = _compact_text(item.get("deadline")) if isinstance(item, Mapping) else "待定"
            lines.append((f"{index}. {_item_content(item, 'task', 'content', 'title', 'summary')}（责任人：{owner or '待确认'}；期限：{deadline or '待定'}）", False))
    override = records.get("latestFormalOverride") if isinstance(records.get("latestFormalOverride"), Mapping) else None
    if override:
        lines.append(("五、人工核验说明", True))
        lines.append((
            "本文件存在待人工核验内容，已由授权人员核对后放行。"
            f"操作人：{_compact_text(override.get('operator')) or '未记录'}；"
            f"时间：{_compact_text(override.get('time')) or '未记录'}；"
            f"原因：{_compact_text(override.get('reason')) or '未记录'}。",
            False,
        ))
    return lines or [("（暂无会议记录）", False)]


_FORMAL_TEMPLATE_TITLES = {
    "standard": "会 议 纪 要",
    "enterprise": "政企红头会议纪要",
    "major": "三重一大会议纪要",
    "party": "党委会议纪要",
    "board": "董事会会议纪要",
    "project": "项目推进会议纪要",
    "engineering": "工程项目会议纪要",
    "audit": "审计会议纪要",
    "concise": "会 议 纪 要",
}

_DETAILED_RECORD_TITLE = "会 议 记 录"


def _set_named_run_font(
    run: Any,
    font_name: str,
    *,
    size: float = 16,
    bold: bool = False,
    color: str | None = None,
) -> None:
    """Apply an exact font required by a retained source template."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _clear_text_runs_preserving_drawings(paragraph: Any) -> None:
    """Remove editable text while retaining anchored DrawingML/VML furniture."""

    from docx.oxml.ns import qn

    drawing_tags = {qn("w:drawing"), qn("w:pict")}
    for child in list(paragraph._p):
        if child.tag != qn("w:r"):
            continue
        if any(node.tag in drawing_tags for node in child.iter()):
            continue
        paragraph._p.remove(child)


def _remove_paragraph(paragraph: Any) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def _enterprise_issue_date(meeting: Mapping[str, Any]) -> str:
    value = _meeting_value(meeting, "date", "time", "startTime", default="")
    match = re.search(r"(20\d{2})[-/.年](\d{1,2})[-/.月](\d{1,2})", value)
    if match:
        year, month, day = match.groups()
        return f"{year}年{int(month)}月{int(day)}日"
    return value or "日期待确认"


def _enterprise_intro(meeting: Mapping[str, Any]) -> str:
    issue_date = _enterprise_issue_date(meeting)
    title = _compact_text(meeting.get("title") or meeting.get("name") or "本次会议")
    location = _meeting_value(meeting, "location", "venue", "address", default="")
    host = _meeting_value(meeting, "host", "moderator", "chairperson", default="")
    if host and location:
        opening = f"{issue_date}，{host}在{location}主持召开“{title}”会议"
    elif location:
        opening = f"{issue_date}，在{location}召开“{title}”会议"
    else:
        opening = f"{issue_date}，召开“{title}”会议"
    return f"{opening}。会议围绕既定议题进行了讨论，并形成如下意见："


def _chinese_section_number(index: int) -> str:
    numbers = "零一二三四五六七八九"
    if 1 <= index < 10:
        return numbers[index]
    if index == 10:
        return "十"
    if 10 < index < 20:
        return f"十{numbers[index % 10]}"
    if 20 <= index < 100:
        ones = numbers[index % 10] if index % 10 else ""
        return f"{numbers[index // 10]}十{ones}"
    return str(index)


def _enterprise_minutes_blocks(
    meeting: Mapping[str, Any],
    records: Mapping[str, Any],
) -> list[tuple[str, str]]:
    """Return `(role, text)` blocks for the redhead minutes body."""

    blocks: list[tuple[str, str]] = [("intro", _enterprise_intro(meeting))]
    section_index = 1
    minutes = list(records.get("minutes") or [])
    if minutes:
        for item in minutes:
            points = _minute_discussion_points(item, records)
            if not points:
                continue
            title = _item_content(item, "agenda", "title", "topic", "content") or "未命名议题"
            blocks.append(("heading", f"{_chinese_section_number(section_index)}、{title}"))
            section_index += 1
            for point in points:
                content = _compact_text(point, limit=1600)
                if content:
                    blocks.append(("body", content))
        if section_index == 1:
            summary = records.get("summary") if isinstance(records.get("summary"), Mapping) else {}
            summary_items = [
                _compact_text(item.get("content"), limit=1600)
                for item in summary.get("conclusions") or []
                if isinstance(item, Mapping)
                and isinstance(item.get("basis"), Mapping)
                and item["basis"].get("evidenceValid")
                and _compact_text(item.get("content"))
            ]
            if summary_items:
                blocks.append(("heading", "一、会议主要内容"))
                section_index = 2
                for item in summary_items:
                    blocks.append(("body", item))
    else:
        blocks.append(("heading", f"{_chinese_section_number(section_index)}、会议主要内容"))
        section_index += 1
        summary = [
            _compact_text(item, limit=1600)
            for item in list(records.get("summary") or [])
            if _compact_text(item)
        ]
        for item in summary or ["暂无已确认的会议内容。"]:
            blocks.append(("body", item))

    for title, items, fields in (
        ("会议决议", list(records.get("decisions") or []), ("content", "decision", "description", "title", "summary")),
        (
            "风险与披露事项",
            [*list(records.get("risks") or []), *list(records.get("disclosures") or [])],
            ("content", "description", "title", "summary"),
        ),
    ):
        if not items:
            continue
        blocks.append(("heading", f"{_chinese_section_number(section_index)}、{title}"))
        section_index += 1
        for item_index, item in enumerate(items, 1):
            content = _item_content(item, *fields)
            if content:
                blocks.append(("body", f"{item_index}. {content}"))

    todos = list(records.get("todos") or [])
    if todos:
        blocks.append(("heading", f"{_chinese_section_number(section_index)}、待办事项"))
        for item_index, item in enumerate(todos, 1):
            task = _item_content(item, "task", "content", "title", "summary")
            owner = _compact_text(item.get("owner")) if isinstance(item, Mapping) else ""
            deadline = _compact_text(item.get("deadline")) if isinstance(item, Mapping) else ""
            suffix = f"（责任人：{owner or '待确认'}；期限：{deadline or '待定'}）"
            blocks.append(("body", f"{item_index}. {task}{suffix}"))
    return blocks


def _set_enterprise_metadata_paragraph(paragraph: Any, label: str, value: Any) -> None:
    from docx.shared import Pt

    paragraph.clear()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(30)
    label_run = paragraph.add_run(label)
    _set_named_run_font(label_run, "黑体")
    value_run = paragraph.add_run(_compact_text(value) or "待确认")
    _set_named_run_font(value_run, "仿宋_GB2312")


def _insert_enterprise_body_paragraph(anchor: Any, role: str, text: str) -> Any:
    from docx.shared import Cm, Pt

    paragraph = anchor.insert_paragraph_before()
    paragraph.style = "Normal (Web)"
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(30)
    if role == "heading":
        paragraph.paragraph_format.space_before = Pt(7.8)
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(text)
        _set_named_run_font(run, "黑体")
    else:
        paragraph.paragraph_format.space_before = Pt(23.4 if role == "intro" else 0)
        paragraph.paragraph_format.first_line_indent = Cm(1.13)
        run = paragraph.add_run(text)
        _set_named_run_font(run, "仿宋_GB2312")
    return paragraph


def _fill_enterprise_minutes_template(
    doc: Any,
    meeting: Mapping[str, Any],
    records: Mapping[str, Any],
) -> None:
    """Fill the retained redhead template while preserving its anchored artwork."""

    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    if len(doc.paragraphs) < 20:
        raise RuntimeError("政企红头会议纪要模板结构无效")
    paragraphs = list(doc.paragraphs)
    issuer = _meeting_value(
        meeting,
        "organization",
        "company",
        "organizer",
        "issuer",
        default="单位名称待确认",
    )
    issue_date = _enterprise_issue_date(meeting)
    issuer_line = paragraphs[4]
    _clear_text_runs_preserving_drawings(issuer_line)
    issuer_line.paragraph_format.tab_stops.clear_all()
    issuer_line.paragraph_format.tab_stops.add_tab_stop(Cm(14.5), WD_TAB_ALIGNMENT.RIGHT)
    issuer_run = issuer_line.add_run(issuer)
    _set_named_run_font(issuer_run, "仿宋_GB2312")
    date_run = issuer_line.add_run(f"\t{issue_date}")
    _set_named_run_font(date_run, "仿宋_GB2312")

    title = _compact_text(meeting.get("title") or meeting.get("name") or "本次会议")
    _set_enterprise_metadata_paragraph(
        paragraphs[6], "会议时间：", _meeting_value(meeting, "time", "date", "startTime")
    )
    _set_enterprise_metadata_paragraph(
        paragraphs[7], "会议地点：", _meeting_value(meeting, "location", "venue", "address")
    )
    _set_enterprise_metadata_paragraph(paragraphs[8], "会议议题：", title)
    _set_enterprise_metadata_paragraph(
        paragraphs[9],
        "参会人员：",
        _meeting_value(meeting, "participants", "attendees", "participantNames"),
    )

    for paragraph in paragraphs[10:17]:
        _remove_paragraph(paragraph)
    closing_topic, closing_report, closing_issue = paragraphs[17:20]
    for role, text in _enterprise_minutes_blocks(meeting, records):
        _insert_enterprise_body_paragraph(closing_topic, role, text)

    closing_topic.clear()
    closing_topic.paragraph_format.space_before = Pt(0)
    closing_topic.paragraph_format.space_after = Pt(0)
    closing_topic.paragraph_format.line_spacing = Pt(30)
    topic_label = closing_topic.add_run("主题词")
    _set_named_run_font(topic_label, "仿宋_GB2312", bold=True)
    topic_value = closing_topic.add_run("：会议纪要")
    _set_named_run_font(topic_value, "仿宋_GB2312")

    _clear_text_runs_preserving_drawings(closing_report)
    report_label = closing_report.add_run("抄  报")
    _set_named_run_font(report_label, "仿宋_GB2312", bold=True)
    report_value = closing_report.add_run(
        f"：{_meeting_value(meeting, 'reportTo', 'copyTo', default='相关领导、相关部门')}"
    )
    _set_named_run_font(report_value, "仿宋_GB2312")

    print_note = _meeting_value(
        meeting, "printNote", "distributionNote", default="（印发范围待确认）"
    )
    print_note_size = "28" if len(print_note) > 8 else "32"
    for text_node in closing_issue._p.iter(qn("w:t")):
        if "一类文件" in (text_node.text or "") or "共印" in (text_node.text or ""):
            text_node.text = print_note
            run_element = text_node.getparent()
            if run_element is not None and run_element.tag == qn("w:r"):
                run_properties = run_element.find(qn("w:rPr"))
                if run_properties is not None:
                    for size_tag in ("w:sz", "w:szCs"):
                        size_node = run_properties.find(qn(size_tag))
                        if size_node is not None:
                            size_node.set(qn("w:val"), print_note_size)
    _clear_text_runs_preserving_drawings(closing_issue)
    issue_office = _meeting_value(
        meeting,
        "issuerDepartment",
        "department",
        "organization",
        default="会议管理部门",
    )
    office_run = closing_issue.add_run(issue_office)
    _set_named_run_font(office_run, "仿宋_GB2312")
    closing_issue.paragraph_format.tab_stops.clear_all()
    closing_issue.paragraph_format.tab_stops.add_tab_stop(Cm(14.5), WD_TAB_ALIGNMENT.RIGHT)
    issued_run = closing_issue.add_run(f"\t{issue_date}印发")
    _set_named_run_font(issued_run, "仿宋_GB2312")


def _set_template_run_font(run: Any, *, size: float = 10.5, bold: bool = False) -> None:
    """Apply the source template's Microsoft YaHei typography explicitly."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    font_name = "Microsoft YaHei"
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font_name)


def _set_template_cell_text(
    cell: Any,
    text: Any,
    *,
    bold: bool = False,
    align: int = 0,
    placeholder: str = "待补充",
) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(_compact_text(text) or placeholder)
    _set_template_run_font(run, bold=bold)


def _template_minutes_lines(records: Mapping[str, Any]) -> list[tuple[str, bool]]:
    """Build the concise, agenda-centred content used by the provided template."""

    lines: list[tuple[str, bool]] = []
    minutes = [
        (item, _minute_discussion_points(item, records))
        for item in list(records.get("minutes") or [])
    ]
    minutes = [(item, points) for item, points in minutes if points]
    for index, (item, points) in enumerate(minutes, 1):
        title = _item_content(item, "agenda", "title", "topic", "content")
        lines.append((f"议题{index}：{title or '未命名议题'}", True))
        for point_index, point in enumerate(points, 1):
            content = _compact_text(point, limit=1200)
            if content:
                lines.append((f"{point_index}. {content}", False))

    decisions = list(records.get("decisions") or [])
    if decisions:
        lines.append(("会议决议：", True))
        for index, item in enumerate(decisions, 1):
            content = _item_content(item, "content", "decision", "description", "title", "summary")
            if content:
                lines.append((f"{index}. {content}", False))

    risks, disclosures = list(records.get("risks") or []), list(records.get("disclosures") or [])
    if risks or disclosures:
        lines.append(("风险与披露事项：", True))
        for index, item in enumerate([*risks, *disclosures], 1):
            content = _item_content(item, "content", "description", "title", "summary")
            if content:
                lines.append((f"{index}. {content}", False))
    return lines or [("（暂无已确认的会议内容）", False)]


def _fill_minutes_template(
    doc: Any,
    meeting: Mapping[str, Any],
    records: Mapping[str, Any],
    *,
    template_id: str,
) -> None:
    from copy import deepcopy

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not doc.paragraphs or not doc.tables:
        raise RuntimeError("会议纪要模板结构无效")
    title_paragraph = doc.paragraphs[0]
    title_paragraph.text = ""
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(
        _FORMAL_TEMPLATE_TITLES.get(template_id, _FORMAL_TEMPLATE_TITLES["standard"])
    )
    _set_template_run_font(title_run, size=16, bold=True)

    table = doc.tables[0]
    if len(table.rows) < 12 or len(table.columns) < 4:
        raise RuntimeError("会议纪要模板表格结构无效")
    title = _compact_text(meeting.get("title") or meeting.get("name") or "会议")
    _set_template_cell_text(table.cell(0, 1), _meeting_value(meeting, "time", "date", "startTime"))
    _set_template_cell_text(table.cell(0, 3), _meeting_value(meeting, "location", "venue", "address"))
    _set_template_cell_text(
        table.cell(1, 1),
        _meeting_value(meeting, "type", "meetingType", "category", default="普通会议"),
    )
    _set_template_cell_text(table.cell(1, 3), _meeting_value(meeting, "recorder", "secretary"))
    _set_template_cell_text(
        table.cell(2, 1),
        _meeting_value(meeting, "participants", "attendees", "participantNames"),
    )
    _set_template_cell_text(table.cell(3, 1), title)

    content_cell = table.cell(4, 1)
    content_cell.text = ""
    for index, (text, bold) in enumerate(_template_minutes_lines(records)):
        paragraph = content_cell.paragraphs[0] if index == 0 else content_cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(text)
        _set_template_run_font(run, bold=bold)

    todos = list(records.get("todos") or [])
    required_rows = max(5, len(todos))
    while len(table.rows) < 7 + required_rows:
        table._tbl.append(deepcopy(table.rows[-1]._tr))
    for index in range(required_rows):
        row = table.rows[7 + index]
        item = todos[index] if index < len(todos) else {}
        task = _item_content(item, "task", "content", "title", "summary") if item else ""
        owner = _compact_text(item.get("owner")) if isinstance(item, Mapping) else ""
        deadline = _compact_text(item.get("deadline")) if isinstance(item, Mapping) else ""
        _set_template_cell_text(
            row.cells[0],
            str(index + 1),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            placeholder=str(index + 1),
        )
        _set_template_cell_text(row.cells[1], task, placeholder="")
        _set_template_cell_text(
            row.cells[2], owner, align=WD_ALIGN_PARAGRAPH.CENTER, placeholder=""
        )
        _set_template_cell_text(
            row.cells[3], deadline, align=WD_ALIGN_PARAGRAPH.CENTER, placeholder=""
        )


def _append_detailed_record_section(
    doc: Any,
    meeting: Mapping[str, Any],
    records: Mapping[str, Any],
) -> None:
    """Append the previous formal meeting-record layout as a new section."""

    from docx.enum.section import WD_SECTION_START
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.6)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""

    title = _compact_text(meeting.get("title") or meeting.get("name") or "会议")
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(_DETAILED_RECORD_TITLE)
    _set_run_font(run, size=18, bold=True, role="title")
    heading.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=6, cols=6)
    try:
        table.style = "Table Grid"
    except KeyError:
        table.style = None
    _apply_table_grid_borders(table)
    _add_table_geometry(table, (1260, 2100, 1260, 2100, 960, 1740))
    row = table.rows[0].cells
    _set_cell_text(row[0], "会议名称", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, role="heading")
    name_cell = row[1].merge(row[3])
    _set_cell_text(name_cell, title)
    _set_cell_text(row[4], "附页", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, role="heading")
    _set_cell_text(row[5], _meeting_value(meeting, "appendix", "page", default="无"), align=WD_ALIGN_PARAGRAPH.CENTER)

    row = table.rows[1].cells
    _set_cell_text(row[0], "会议时间", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, role="heading")
    _set_cell_text(row[1].merge(row[2]), _meeting_value(meeting, "time", "date", "startTime"))
    _set_cell_text(row[3], "会议地点", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, role="heading")
    _set_cell_text(row[4].merge(row[5]), _meeting_value(meeting, "location", "venue", "address"))

    row = table.rows[2].cells
    _set_cell_text(row[0], "主持人", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, role="heading")
    _set_cell_text(row[1].merge(row[2]), _meeting_value(meeting, "host", "moderator", "chairperson"))
    _set_cell_text(row[3], "记录人", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, role="heading")
    _set_cell_text(row[4].merge(row[5]), _meeting_value(meeting, "recorder", "secretary"))

    row = table.rows[3].cells
    participants = row[0].merge(row[5])
    _set_cell_text(participants, f"参加单位及人员：{_meeting_value(meeting, 'participants', 'attendees', 'participantNames')}")

    body = table.rows[4].cells[0].merge(table.rows[4].cells[5])
    body.text = ""
    body.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for index, (line, bold) in enumerate(_formal_record_lines(records)):
        paragraph = body.paragraphs[0] if index == 0 else body.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.15
        line_run = paragraph.add_run(line)
        _set_run_font(line_run, size=10.5, bold=bold, role="heading" if bold else "body")
    body_row = table.rows[4]
    body_row.height = Cm(14.2)
    body_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    row = table.rows[5].cells
    _set_cell_text(row[0], "编制人", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, role="heading")
    _set_cell_text(row[1].merge(row[2]), _meeting_value(meeting, "compiler", "createdBy", "creator"), align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(row[3], "日期", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, role="heading")
    _set_cell_text(row[4].merge(row[5]), _meeting_value(meeting, "compiledDate", "date", default=datetime.now().strftime("%Y.%m.%d")), align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_index, table_row in enumerate(table.rows):
        tr_pr = table_row._tr.get_or_add_trPr()
        if row_index != 4:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in table_row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            margin = tc_pr.first_child_found_in("w:tcMar")
            if margin is None:
                margin = OxmlElement("w:tcMar")
                tc_pr.append(margin)
            for edge in ("top", "left", "bottom", "right"):
                node = margin.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    margin.append(node)
                node.set(qn("w:w"), "100")
                node.set(qn("w:type"), "dxa")


def _write_formal_docx(
    path: Path,
    meeting: Mapping[str, Any],
    records: Mapping[str, Any],
    *,
    template_id: str = "standard",
) -> None:
    """Write one formal DOCX containing concise minutes and the detailed record."""

    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - depends on deployment package set
        raise RuntimeError("python-docx is required to generate meeting documents") from exc
    template_path = (
        ENTERPRISE_MINUTES_TEMPLATE if template_id == "enterprise" else FORMAL_MINUTES_TEMPLATE
    )
    if not template_path.is_file():
        raise RuntimeError(f"会议纪要模板不存在：{template_path}")
    doc = Document(template_path)
    if template_id == "enterprise":
        _fill_enterprise_minutes_template(doc, meeting, records)
    else:
        _fill_minutes_template(doc, meeting, records, template_id=template_id)
    _append_detailed_record_section(doc, meeting, records)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def _write_evidence_docx(path: Path, meeting: Mapping[str, Any], records: Mapping[str, Any], manifest: Mapping[str, Any]) -> None:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required to generate meeting documents") from exc
    doc = Document()
    title = _compact_text(meeting.get("title") or meeting.get("name") or "会议")
    _configure_document(doc, title=f"{title}｜证据底稿")
    heading = doc.add_paragraph()
    run = heading.add_run(f"{title}｜证据底稿")
    _set_run_font(run, size=20, bold=True, color="0B2545")
    _docx_paragraph(doc, "本附件保留全量有效转写、剔除行和质量标记；正式文件只引用蒸馏结果。", size=10)
    doc.add_heading("一、质量标记", level=1)
    flags = list(manifest.get("qualityFlags") or [])
    if flags:
        for flag in flags:
            _docx_paragraph(doc, f"{_compact_text(flag.get('type') or 'quality')}：{_compact_text(flag.get('label') or flag.get('reason') or flag.get('segmentId') or '')}", style="List Bullet", size=9)
    else:
        _docx_paragraph(doc, "未发现断档、噪音或空白标记", size=9)
    doc.add_heading("二、全量有效转写", level=1)
    rows = list(manifest.get("rows") or [])
    for row in rows:
        _docx_paragraph(doc, f"[{_compact_text(row.get('time') or '未标注')}] {_compact_text(row.get('speaker') or '说话人未识别')}：{_as_text(row.get('rawText'))}", size=9)
        corrected = _as_text(row.get("correctedText"))
        if corrected and corrected != _as_text(row.get("rawText")):
            _docx_paragraph(doc, f"校对稿：{corrected}", size=9)
    if not rows:
        _docx_paragraph(doc, "（暂无有效转写）", size=9)
    doc.add_heading("三、已剔除行（保留审计）", level=1)
    excluded = list(manifest.get("excludedRows") or [])
    for row in excluded:
        _docx_paragraph(doc, f"[{_compact_text(row.get('time') or '未标注')}] {_compact_text(row.get('speaker') or '说话人未识别')}：{_as_text(row.get('rawText'))}（原因：{_compact_text(row.get('excludeReason') or '标记为噪音/无效')}）", size=9)
    if not excluded:
        _docx_paragraph(doc, "（无剔除行）", size=9)
    doc.add_heading("四、Map 分段边界", level=1)
    for boundary in manifest.get("mapBoundaries") or []:
        _docx_paragraph(doc, f"{_compact_text(boundary.get('chunkId') or '未命名')}｜文件 {_compact_text(boundary.get('fileId') or '未标注')}｜{_compact_text(boundary.get('timeRange') or '未标注')}｜{len(boundary.get('sourceSegmentIds') or [])} 个分段", size=9)
    if not manifest.get("mapBoundaries"):
        _docx_paragraph(doc, "（未提供 Map 分段信息）", size=9)
    doc.add_heading("五、正式条目反向索引", level=1)
    reverse = manifest.get("reverseIndex") or {}
    for segment_id, entries in reverse.items():
        _docx_paragraph(doc, f"{segment_id}：{', '.join(_as_text(item) for item in entries)}", size=9)
    if not reverse:
        _docx_paragraph(doc, "（暂无带依据的正式条目）", size=9)
    doc.add_heading("六、系统自动排除的 AI 表述", level=1)
    exceptions = list(manifest.get("evidenceExceptions") or [])
    for exception in exceptions:
        item = exception.get("item") if isinstance(exception.get("item"), Mapping) else {}
        content = _item_content(item, "content", "task", "agenda", "title", "summary")
        _docx_paragraph(
            doc,
            f"{_compact_text(exception.get('field') or '未分类')}：{content or '未记录内容'}"
            f"（原因：{_compact_text(exception.get('reason') or '无可核验原文依据')}）",
            size=9,
        )
    if not exceptions:
        _docx_paragraph(doc, "（无自动排除表述）", size=9)
    doc.add_heading("七、生成参数快照", level=1)
    snapshot = _snapshot(records, list(manifest.get("allRows") or []))
    for key in ("provider", "model", "pipelineVersion", "promptVersion", "schemaVersion", "glossaryVersion", "chunkPolicy", "inputSha256", "segmentCount", "generatedAt"):
        if key in snapshot and snapshot[key] not in (None, ""):
            _docx_paragraph(doc, f"{key}：{snapshot[key]}", size=9)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def generate_document_bundle(
    meeting_id: str,
    meeting: Mapping[str, Any],
    records: Mapping[str, Any],
    chronicle: Any,
    output_dir: str | Path,
    *,
    file_offsets: Mapping[str, Any] | None = None,
    markers: Sequence[Mapping[str, Any]] | None = None,
    require_proofread: bool = True,
    timestamp: str | None = None,
    template_id: str = "standard",
) -> dict[str, Any]:
    """Generate the formal document and independent evidence attachment."""

    if require_proofread and not records_are_proofread(records):
        raise ProofreadRequiredError("正式文件必须在校对通过后生成")
    rows = _rows_from_source(chronicle)
    snapshot = _snapshot(records, rows)
    manifest = build_evidence_manifest(records, chronicle, file_offsets=file_offsets, markers=markers)
    stamp = timestamp or datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    safe_id = re.sub(r"[^A-Za-z0-9_.-]+", "_", _as_text(meeting_id) or "meeting")
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    formal_path = output_path / f"{safe_id}_会议纪要及记录_{stamp}.docx"
    evidence_path = output_path / f"{safe_id}_证据底稿_{stamp}.docx"
    template_id = template_id if template_id in _FORMAL_TEMPLATE_TITLES else "standard"
    _write_formal_docx(
        formal_path,
        meeting,
        {**dict(records), "generationSnapshot": snapshot},
        template_id=template_id,
    )
    _write_evidence_docx(evidence_path, meeting, {**dict(records), "generationSnapshot": snapshot}, manifest)
    return {
        "serviceVersion": DOCUMENT_SERVICE_VERSION,
        "meetingId": str(meeting_id),
        "generationSnapshot": snapshot,
        "coverage": manifest["coverage"],
        "templateId": template_id,
        "templateTitle": _FORMAL_TEMPLATE_TITLES[template_id],
        "evidenceManifest": manifest,
        "formal": DocumentArtifact(
            kind="formal",
            path=str(formal_path),
            filename=formal_path.name,
            sha256=hashlib.sha256(formal_path.read_bytes()).hexdigest(),
        ).to_dict(),
        "evidence": DocumentArtifact(
            kind="evidence",
            path=str(evidence_path),
            filename=evidence_path.name,
            sha256=hashlib.sha256(evidence_path.read_bytes()).hexdigest(),
        ).to_dict(),
    }


__all__ = [
    "DOCUMENT_SERVICE_VERSION",
    "DOCX_MIME",
    "DocumentArtifact",
    "GAP_THRESHOLD_SECONDS",
    "build_evidence_manifest",
    "generate_document_bundle",
    "render_evidence_markdown",
    "render_formal_markdown",
]
