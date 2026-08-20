import os as _os
_os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
_os.environ.setdefault("HF_HUB_OFFLINE", "1")
_os.environ.setdefault("SENTENCE_TRANSFORMERS_HOME", _os.path.expanduser("~/.cache/huggingface/hub"))
del _os

from fastapi import FastAPI, HTTPException, Request, UploadFile, File, Form, Response, WebSocket, WebSocketDisconnect
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List, Any, AsyncIterator, Dict
import json
import logging
import os
import asyncio
import httpx
import io
import re
import asyncio
import concurrent.futures
import time
from contextlib import asynccontextmanager
import base64
import contextvars
import hashlib
import jwt
import uuid
import sqlite3
from pathlib import Path
import socket
import threading
from threading import Lock, RLock
from urllib.parse import quote, urlparse
from dotenv import load_dotenv
from legal_case_db import LegalCaseDatabase, LegalCase
from case_similarity import CaseSimilarityMatcher
from ongoing_case_tracker import OngoingCaseTracker
from audit_persistence import persistence
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from demo_content import build_archive_history, get_demo_assets, get_seed_knowledge_files
from datetime import datetime, timedelta


load_dotenv()

# ═══ 模块化导入 — 替代原有内联定义 ═════════════════════════════════════════════
# 以下符号已迁移到 backend/ 子模块，通过 import 引入保持兼容。
# 如果同名符号在本文后续重新定义，Python 的"后定义覆盖前导入"规则会使本地定义生效。
import sys as _sys
from pathlib import Path as _Path
_PROJECT_ROOT = _Path(__file__).resolve().parent
if str(_PROJECT_ROOT) not in _sys.path:
    _sys.path.insert(0, str(_PROJECT_ROOT))
del _sys, _Path, _PROJECT_ROOT

from backend.config import (  # noqa: E402
    # 环境变量
    LLM_CONCURRENCY, AUTH_SECRET, DASHSCOPE_FUN_ASR_WS_URL,
    DASHSCOPE_API_KEY, DASHSCOPE_WORKSPACE, PERSIST_DIR,
    # 路径常量
    CUSTOM_RULES_DIR, CUSTOM_RULES_DB, MEETING_DATA_DIR,
    APP_DB, MEETING_FILES_DIR, RULES_IMAGES_DIR,
    AUTH_DATA_DIR, USERS_DB, DEPARTMENTS_DB, MEETINGS_DB, MEETING_TRANSCRIPTS_DB,
    ASR_CONFIG_DIR, ASR_HOTWORDS_DB, ASR_CORRECTIONS_DB,
    # 锁
    APP_DB_LOCK, MEETINGS_LOCK, MEETING_TRANSCRIPTS_LOCK,
    # 全局对象
    llm_semaphore,
    # 上传限制
    MAX_UPLOAD_BYTES, MAX_AUDIO_BYTES, MAX_EXCEL_BYTES,
    # ASR 重连
    ASR_RECONNECT_BASE_DELAY, ASR_RECONNECT_MAX_DELAY, ASR_RECONNECT_MAX_RETRIES,
    # Qwen3-ASR
    QWEN_ASR_URL, ASR_BACKEND,
    # 网络工具
    get_public_host, get_browser_backend_base, PUBLIC_HOST,
    now_text as _now_text, today_text as _today_text,
)
from backend.models import (  # noqa: E402
    ChatResponse, LegalCompareRequest, KBQueryRequest,
    LoginRequest, MeetingRegisterRequest, UserUpsertRequest,
    MeetingTranscriptChunkRequest, MeetingRecorderSessionRequest,
    MeetingRecorderAudioMetaRequest, MeetingUpsertRequest, MeetingPatchRequest,
    MeetingIssueRequest, MeetingStageRequest, MeetingAgendaRealtimeCheckRequest,
    MeetingTranscriptCorrectionRequest, MeetingMarkerRequest,
    MeetingRecordsUpdateRequest, ChatRequest,
    AgendaCreateRequest, AgendaPatchRequest,
    AgendaRecordRequest, AgendaDecisionRequest, AgendaDecisionPatchRequest,
    MeetingSignatureRequest,
)
from backend.db import (  # noqa: E402
    _db_connect, _init_app_db, _db_fetch_meetings, _db_save_meetings,
    _db_upsert_meeting, _db_delete_meeting_rows, _db_insert_meeting_rows,
    _db_insert_issue_sources, _db_insert_agenda_drafts,
    _db_insert_materials, _db_insert_events,
    _db_load_transcripts, _db_save_transcripts, _db_insert_transcript_row,
    _db_upsert_transcript, _db_load_transcripts_for_meeting,
    _db_delete_meeting_by_id, _metadata_get, _metadata_set,
    _meeting_from_row, _migrate_legacy_meeting_json_once,
    _wal_checkpoint, _maybe_checkpoint_wal,
    # 缓存层
    _load_meetings, _save_meetings,
    _load_meeting_transcripts, _save_meeting_transcripts,
    _invalidate_transcripts_cache, _invalidate_meetings_cache,
    # 用户管理
    _default_users, _load_users, _save_users,
    # 工具函数
    _default_meetings, _default_agenda_drafts, _default_issue_sources,
    _safe_meeting_id, _check_meeting_access, _creator_from_user,
    _phase_color, _normalize_meeting, _public_meeting,
    _json_loads, _json_dumps, _safe_storage_filename,
    _derive_agenda_drafts, _clean_agenda_check_transcript,
)
from backend.llm_client import (  # noqa: E402
    DeepSeekThinkingLLM, llm,
    _get_httpx_async, _get_httpx_sync,
)
from backend.deps import (  # noqa: E402
    _now_text, _today_text,
    _get_public_host, _get_browser_backend_base,
    _issue_auth_token, _get_user_from_auth_token,
    _get_request_user, _require_admin, _public_user,
    _resolve_meeting_role, _build_meeting_from_request,
    _append_meeting_activity_light,
    _save_departments, _load_departments,
)
from backend.services.agenda_service import (  # noqa: E402
    list_meeting_agendas, get_meeting_agenda, get_meeting_active_agenda,
    create_meeting_agenda, update_meeting_agenda, delete_meeting_agenda,
    activate_meeting_agenda,
    list_agenda_records, create_agenda_record,
    list_agenda_decisions, create_agenda_decision,
    update_agenda_decision, delete_agenda_decision,
    generate_decisions_for_agenda,
)
from backend.services.signature_service import (  # noqa: E402
    list_signatures, sign_target, invalidate_target_signatures,
    is_fully_signed, signed_signer_count, required_signer_count,
    compute_content_hash,
)
from backend.services.permission_service import (  # noqa: E402
    get_user_roles, add_user_role, remove_user_role,
    list_agenda_acl, grant_agenda_acl, revoke_agenda_acl,
    filter_agendas_for_user, can_view_agenda,
)
from backend.services.knowledge_service import search_agenda_knowledge  # noqa: E402
# ═══ 模块化导入结束 ═════════════════════════════════════════════════════════════


# Context variable to hold the current FastAPI Request for cancellation checks
current_request = contextvars.ContextVar("current_request", default=None)

# ── LangChain imports ────────────────────────────────────────────────────────
from langchain.agents import create_react_agent, AgentExecutor
from langchain_core.prompts import PromptTemplate
from langchain_core.tools import tool
from langchain_core.language_models.chat_models import BaseChatModel
from langchain_core.messages import AIMessage, BaseMessage, HumanMessage, SystemMessage, AIMessageChunk
from langchain_core.outputs import ChatGeneration, ChatResult, ChatGenerationChunk
from langchain_chroma import Chroma
# 使用直接包装 sentence_transformers 绕开 langchain_huggingface 版本冲突
try:
    from sentence_transformers import SentenceTransformer as _SentTransformer
    class _STLangChainEmbeddings:
        def __init__(self, model_name: str):
            self._model = _SentTransformer(model_name)
        def embed_documents(self, texts: List[str]) -> List[List[float]]:
            return self._model.encode(texts, normalize_embeddings=True).tolist()
        def embed_query(self, text: str) -> List[float]:
            return self._model.encode([text], normalize_embeddings=True).tolist()[0]
    _embedding_fn_cls = _STLangChainEmbeddings
except Exception as _emb_import_err:
    _embedding_fn_cls = None
    _emb_import_err_msg = str(_emb_import_err)


# Configure logging — rotating file handler, 10MB × 10 backups
from logging.handlers import RotatingFileHandler as _RotatingFileHandler
_log_dir = Path(os.path.dirname(os.path.abspath(__file__))) / "logs"
_log_dir.mkdir(parents=True, exist_ok=True)
_file_handler = _RotatingFileHandler(
    _log_dir / "backend.log", maxBytes=10 * 1024 * 1024, backupCount=10, encoding="utf-8",
)
_file_handler.setFormatter(logging.Formatter(
    "[%(asctime)s] %(levelname)s %(name)s: %(message)s", datefmt="%Y-%m-%d %H:%M:%S",
))
logging.basicConfig(level=logging.INFO, handlers=[_file_handler, logging.StreamHandler()])
logger = logging.getLogger(__name__)




class ReportGeneratedException(Exception):
    def __init__(self, report_md: str):
        self.report_md = report_md

# ── FastAPI lifespan: pre-warm heavy resources ─────────────────────────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    """Pre-warm ChromaDB, SentenceTransformer, and db migration on startup."""
    if not AUTH_SECRET:
        logger.critical("APP_AUTH_SECRET 未设置！服务器拒绝启动。请在 .env 中配置。")
        raise RuntimeError("APP_AUTH_SECRET is required")
    import concurrent.futures as _cf
    _warm_executor = _cf.ThreadPoolExecutor(max_workers=2, thread_name_prefix="warmup-")
    _loop = asyncio.get_event_loop()
    # 声纹引擎初始化（基于 resemblyzer，无需 HF_TOKEN）— 与其他预热并行
    from backend.voiceprint import init_voiceprint_engine
    _warm_tasks = [
        _loop.run_in_executor(_warm_executor, _migrate_legacy_meeting_json_once),
        _loop.run_in_executor(_warm_executor, lambda: _get_vectorstore(False)),
        _loop.run_in_executor(_warm_executor, _get_case_db),
        _loop.run_in_executor(_warm_executor, init_voiceprint_engine),
    ]

    results = await asyncio.gather(*_warm_tasks, return_exceptions=True)
    _warm_executor.shutdown(wait=False)

    # 检查声纹引擎初始化结果
    _vp_result = results[-1]
    if isinstance(_vp_result, Exception):
        logger.warning("【启动】声纹引擎初始化失败（不影响其他功能）: %s", _vp_result)
    else:
        logger.info("【启动】声纹引擎已就绪")

    # v5: 启动后台 ASR pending store 清理任务
    _asr_cleanup_task = asyncio.create_task(_cleanup_asr_pending_store())
    # 恢复上次重启未持久化的 Whisper 结果
    _recover_orphaned_whisper_results()
    # 清理孤儿录音 chunk 文件（超过 1 小时未合并的 chunk）
    _cleanup_orphaned_audio_chunks()
    logger.info("【启动】预热完成 — ChromaDB、模型、数据迁移已就绪")
    yield
    # v5: 取消后台任务
    _asr_cleanup_task.cancel()
    try:
        await _asr_cleanup_task
    except asyncio.CancelledError:
        pass
    # ── Shutdown: clean up resources ──
    global _httpx_async_client, _httpx_sync_client
    if _httpx_async_client and not _httpx_async_client.is_closed:
        await _httpx_async_client.aclose()
    if _httpx_sync_client and not _httpx_sync_client.is_closed:
        _httpx_sync_client.close()
    _llm_executor.shutdown(wait=True)
    logger.info("【关闭】HTTP 连接池和线程池已清理")


# ── Upload limits 已从 backend.config 导入 ────────────────────────────────────

async def _read_upload_safe(file: UploadFile, max_bytes: int = MAX_UPLOAD_BYTES) -> bytes:
    """流式读取上传文件，硬限制内存占用。超出上限立即拒绝，不缓冲全量。"""
    # 先读 max_bytes+1 字节检查是否超限
    content = await file.read(max_bytes + 1)
    if len(content) > max_bytes:
        raise HTTPException(
            status_code=413,
            detail=f"文件过大（>{max_bytes / 1024 / 1024:.0f}MB），上限 {max_bytes / 1024 / 1024:.0f}MB",
        )
    if not content:
        raise HTTPException(status_code=400, detail="文件为空")
    # 读入剩余部分（Starlite 内部已缓冲，通常无需额外读取）
    remaining = await file.read()
    return content + remaining if remaining else content

# Initialize FastAPI
app = FastAPI(title="三重一大合规审核 API", lifespan=lifespan)


# ── Health check ──────────────────────────────────────────────────────────────
@app.get("/health")
async def health_check():
    status = {"status": "ok", "checks": {}}
    # DB check
    try:
        with _db_connect() as conn:
            conn.execute("SELECT 1").fetchone()
        status["checks"]["db"] = "ok"
    except Exception as e:
        status["checks"]["db"] = f"fail: {e}"
        status["status"] = "degraded"
    # ChromaDB check
    try:
        vs = _get_vectorstore(False)
        status["checks"]["chromadb"] = "ok" if vs is not None else "not_loaded"
    except Exception as e:
        status["checks"]["chromadb"] = f"fail: {e}"
        status["status"] = "degraded"
    # LLM key check
    status["checks"]["llm"] = "ok" if os.environ.get("DEEPSEEK_API_KEY", "") else "no_api_key"
    # ASR key check
    status["checks"]["asr_dashscope"] = "ok" if DASHSCOPE_API_KEY else "no_api_key"

    # Qwen3-ASR local check
    try:
        qwen_health = await (await _get_qwen_client()).health()
        status["checks"]["asr_qwen"] = qwen_health.get("status", "unknown")
    except Exception:
        status["checks"]["asr_qwen"] = "unavailable"
    # Data counts
    try:
        with _db_connect() as conn:
            status["checks"]["meetings"] = conn.execute("SELECT COUNT(*) FROM meetings").fetchone()[0]
            status["checks"]["transcripts"] = conn.execute("SELECT COUNT(*) FROM meeting_transcripts").fetchone()[0]
    except Exception:
        pass
    status["uptime_seconds"] = round(time.monotonic(), 1)
    return JSONResponse(status)


# Global semaphore to limit concurrent LLM requests
llm_semaphore = asyncio.Semaphore(int(os.environ.get("LLM_CONCURRENCY", "5")))
# Whisper 终审全局并发锁 —— 确保同一时刻只有一个 Whisper 任务在跑，防止 GPU OOM
_whisper_semaphore = asyncio.Semaphore(1)
# Dedicated thread pool for LLM calls (avoids saturating asyncio default executor)
_llm_executor = concurrent.futures.ThreadPoolExecutor(max_workers=8, thread_name_prefix="llm-")
# Shared HTTP clients — connection pooling avoids TCP+TLS handshake per request
_httpx_async_client: Optional[httpx.AsyncClient] = None
_httpx_sync_client: Optional[httpx.Client] = None

def _get_httpx_async() -> httpx.AsyncClient:
    global _httpx_async_client
    if _httpx_async_client is None or _httpx_async_client.is_closed:
        _httpx_async_client = httpx.AsyncClient(
            timeout=httpx.Timeout(180.0, connect=10.0),
            limits=httpx.Limits(max_keepalive_connections=20, max_connections=50),
            http2=False,
            trust_env=False,
        )
    return _httpx_async_client

def _get_httpx_sync() -> httpx.Client:
    global _httpx_sync_client
    if _httpx_sync_client is None or _httpx_sync_client.is_closed:
        _httpx_sync_client = httpx.Client(
            timeout=httpx.Timeout(180.0, connect=10.0),
            limits=httpx.Limits(max_keepalive_connections=10, max_connections=30),
            http2=False,
            trust_env=False,
        )
    return _httpx_sync_client
AUTH_SECRET = os.environ.get("APP_AUTH_SECRET", "")
DASHSCOPE_FUN_ASR_WS_URL = os.environ.get("DASHSCOPE_FUN_ASR_WS_URL", "wss://dashscope.aliyuncs.com/api-ws/v1/inference")
DASHSCOPE_API_KEY = (
    os.environ.get("DASHSCOPE_API_KEY")
    or os.environ.get("DASHSCOPE_BAILIAN_API_KEY")
    or os.environ.get("BAILIAN_API_KEY")
    or ""
)
DASHSCOPE_WORKSPACE = os.environ.get("DASHSCOPE_WORKSPACE") or os.environ.get("DASHSCOPE_WORKSPACE_ID") or ""

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── 路由注册 ──
from backend.routes.auth import router as auth_router
app.include_router(auth_router)

from backend.routes.voiceprint import router as voiceprint_router
app.include_router(voiceprint_router)

# ── 请求审计中间件 — 记录所有 API 请求的 method/path/status/duration/user/IP ──
@app.middleware("http")
async def audit_request_middleware(request: Request, call_next):
    start = time.monotonic()
    try:
        response = await call_next(request)
    except Exception:
        elapsed_ms = (time.monotonic() - start) * 1000
        logger.exception("【审计】%s %s → 500 UNHANDLED  %dms", request.method, request.url.path, round(elapsed_ms))
        return JSONResponse({"detail": "Internal server error"}, status_code=500)
    elapsed_ms = (time.monotonic() - start) * 1000
    # 提取用户（从 Authorization header）
    user_id = "-"
    auth = request.headers.get("authorization", "")
    if auth.startswith("Bearer "):
        try:
            payload = jwt.decode(auth[7:], AUTH_SECRET, algorithms=["HS256"])
            user_id = payload.get("sub", "-")
        except Exception:
            user_id = "invalid"
    # 提取真实 IP
    ip = request.headers.get("x-forwarded-for", "").split(",")[0].strip() or request.client.host if request.client else "-"
    logger.info(
        "【审计】%s %s → %s %s  %dms  user=%s  ip=%s",
        request.method, request.url.path,
        response.status_code,
        "stream" if response.headers.get("content-type") == "text/event-stream" else "",
        round(elapsed_ms),
        user_id, ip,
    )
    return response

# ── Request models ────────────────────────────────────────────────────────────
class ChatResponse(BaseModel):
    success: bool
    message: str
    report: Optional[str] = None
    legal_analysis: Optional[dict] = None

class LegalCompareRequest(BaseModel):
    case_type: str
    case_description: str
    case_amount: float = 0.0

class KBQueryRequest(BaseModel):
    query: str








# MeetingTranscriptChunkRequest、MeetingRecorderSessionRequest 已从 backend.models 导入
# 此处只保留 backend_full.py 独用的 MeetingRecorderAudioMetaRequest


class MeetingRecorderAudioMetaRequest(BaseModel):
    meeting_id: str
    meeting_title: str = ""
    agenda: str = ""
    duration_seconds: Optional[int] = None
    device_type: Optional[str] = None
    device_id: Optional[str] = None
    device_label: Optional[str] = None
    channel: Optional[str] = None
    transport: Optional[str] = None
    firmware_version: Optional[str] = None


class MeetingUpsertRequest(BaseModel):
    id: Optional[str] = None
    title: str = ""
    project: str = ""
    projectCode: str = ""
    project_code: str = ""
    agenda: str = ""
    date: str = ""
    type: str = "普通企业会议"
    meetingNo: str = ""
    meetingMode: str = "normal"
    meeting_mode: str = ""
    phase: str = "问题收集中"
    creator: str = ""
    requireFullSignature: bool = False
    issueSources: Optional[List[dict]] = None
    agendaDrafts: Optional[List[dict]] = None
    materials: Optional[List[dict]] = None


class MeetingPatchRequest(BaseModel):
    title: Optional[str] = None
    project: Optional[str] = None
    projectCode: Optional[str] = None
    project_code: Optional[str] = None
    agenda: Optional[str] = None
    date: Optional[str] = None
    type: Optional[str] = None
    meetingNo: Optional[str] = None
    meetingMode: Optional[str] = None
    meeting_mode: Optional[str] = None
    phase: Optional[str] = None
    archived: Optional[bool] = None
    projectBound: Optional[bool] = None
    agendaFrozen: Optional[bool] = None
    reviewDone: Optional[bool] = None
    archiveDone: Optional[bool] = None
    requireFullSignature: Optional[bool] = None
    issueSources: Optional[List[dict]] = None
    agendaDrafts: Optional[List[dict]] = None
    materials: Optional[List[dict]] = None


class MeetingIssueRequest(BaseModel):
    name: str = "当前用户"
    content: str
    type: str = "text"
    meta: str = ""
    source: str = "manual"


class MeetingStageRequest(BaseModel):
    stage: str
    phase: str = ""


class MeetingAgendaRealtimeCheckRequest(BaseModel):
    agendaDrafts: List[dict] = []
    latestTranscripts: List[dict] = []
    meetingMode: str = "normal"


class MeetingTranscriptCorrectionRequest(BaseModel):
    corrected_transcript: str
    signature_data: str
    client_time: Optional[str] = None


class MeetingMarkerRequest(BaseModel):
    marker_type: str  # decision / todo / dispute / material
    agenda_id: str = ""
    agenda_title: str = ""
    transcript_id: str = ""
    transcript_text: str = ""
    transcript_time: str = ""
    transcript_speaker: str = ""
    note: str = ""


class MeetingRecordsUpdateRequest(BaseModel):
    summary: Optional[List[str]] = None
    minutes: Optional[List[dict]] = None
    decisions: Optional[List[dict]] = None
    todos: Optional[List[dict]] = None


class ChatRequest(BaseModel):
    matter_type: str
    material_text: str
    custom_rule_ids: Optional[List[str]] = None

# ── Custom DeepSeek LLM (supports reasoning_content) ─────────────────────────
# 注释掉本地 Qwen 模型，改为线上 DeepSeek
# class QwenThinkingLLM(BaseChatModel):
#     """兼容 reasoning_content 字段的 Qwen 客户端"""
#     api_base: str = "http://192.168.66.44:8088/v1"
#     model_name: str = "Qwen3.5-35B-A3B"
#     temperature: float = 0.1
#     max_tokens: int = 80000
#     timeout: float = 180.0
# ...
# llm = QwenThinkingLLM()

class DeepSeekThinkingLLM(BaseChatModel):
    """兼容易 reasoning_content 字段的 DeepSeek V3 (支持思考模式和 agent) 客户端"""
    # 按照文档说明，这里调用官方 api，如果需要在环境变量注入密钥则取消注释或在这里明文修改
    api_key: str = os.environ.get("DEEPSEEK_API_KEY", "")
    api_base: str = "https://api.deepseek.com/chat/completions"
    model_name: str = "deepseek-chat"
    temperature: float = 0.1
    max_tokens: int = 8000
    timeout: float = 180.0

    @property
    def _llm_type(self) -> str:
        return "deepseek-thinking"

    def _convert_messages(self, messages: List[Any]) -> List[dict]:
        result = []
        for m in messages:
            if isinstance(m, str):
                result.append({"role": "user", "content": m})
            elif isinstance(m, SystemMessage):
                result.append({"role": "system", "content": m.content})
            elif isinstance(m, HumanMessage):
                result.append({"role": "user", "content": m.content})
            elif isinstance(m, AIMessage):
                msg_dict = {"role": "assistant", "content": m.content or ""}
                
                # Retrieve reasoning_content if saved in additional_kwargs (multi-turn tool invoking)
                if m.additional_kwargs and "reasoning_content" in m.additional_kwargs:
                    msg_dict["reasoning_content"] = m.additional_kwargs["reasoning_content"]
                    
                result.append(msg_dict)
            else:
                result.append({"role": "user", "content": str(getattr(m, 'content', m))})
        return result

    def _generate(self, messages: List[BaseMessage], stop=None, run_manager=None, enable_thinking: bool = True, **kwargs) -> ChatResult:
        payload = {
            "model": self.model_name,
            "messages": self._convert_messages(messages),
            "temperature": self.temperature,
            "max_tokens": self.max_tokens,
        }
        if enable_thinking:
            payload["thinking"] = {"type": "enabled"}
            
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }

        client = _get_httpx_sync()
        resp = client.post(self.api_base, json=payload, headers=headers)
        if resp.status_code >= 400:
            logger.error(f"DeepSeek 错误响应: {resp.text} payload: {json.dumps(payload, ensure_ascii=False)}")
        resp.raise_for_status()
        data = resp.json()

        msg = data["choices"][0]["message"]
        # 重要：Agent 调用必须使用 content 字段。
        text = msg.get("content") or ""
        if not text:
            # 如果 content 为空，说明模型只输出了思考链而没有实质回答，记录警告便于调试
            rc = msg.get("reasoning_content", "")
            logger.warning(
                f"【LLM】content 为空，模型可能仅输出了思考链（reasoning_content 长度={len(rc)}）。"
                f"建议检查模型是否支持当前请求格式，当前返回空字符串。"
            )
        return ChatResult(generations=[ChatGeneration(message=AIMessage(content=text))])

    async def _agenerate(self, messages, stop=None, run_manager=None, **kwargs):
        return await asyncio.get_event_loop().run_in_executor(
            _llm_executor, lambda: self._generate(messages, stop, run_manager, **kwargs)
        )

    async def _astream(
        self, messages: List[BaseMessage], stop=None, run_manager=None, enable_thinking: bool = True, **kwargs
    ) -> AsyncIterator[ChatGenerationChunk]:
        """Stream chunks from DeepSeek, tagging reasoning vs content separately."""
        payload = {
            "model": self.model_name,
            "messages": self._convert_messages(messages),
            "temperature": self.temperature,
            "max_tokens": self.max_tokens,
            "stream": True,
        }
        if enable_thinking:
            payload["thinking"] = {"type": "enabled"}

        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }

        client = _get_httpx_async()
        async with client.stream("POST", self.api_base, json=payload, headers=headers) as response:
            if response.status_code >= 400:
                text_err = await response.aread()
                logger.error(f"DeepSeek 流式响应错误 {response.status_code}: {text_err.decode('utf-8')} Payload: {json.dumps(payload, ensure_ascii=False)}")
            response.raise_for_status()
            async for line in response.aiter_lines():
                # Check for client disconnection context
                req = current_request.get()
                if req and getattr(req, "is_disconnected", None):
                    if await req.is_disconnected():
                        logger.info("【LLM】Client disconnected, aborting LLM stream.")
                        raise asyncio.CancelledError()

                line = line.strip()
                if not line or not line.startswith("data: "):
                    continue
                data_str = line[6:]
                if data_str == "[DONE]":
                    break
                try:
                    data = json.loads(data_str)
                    delta = data["choices"][0].get("delta", {})
                    r_text = delta.get("reasoning_content") or ""
                    c_text = delta.get("content") or ""

                    # Yield thinking chunk (reasoning_content) tagged as 'thinking'
                    if r_text:
                        chunk = ChatGenerationChunk(
                            message=AIMessageChunk(
                                content=r_text,
                                additional_kwargs={"chunk_type": "thinking"},
                            )
                        )
                        if run_manager:
                            await run_manager.on_llm_new_token(r_text, chunk=chunk)
                        yield chunk

                    # Yield answer chunk (content) tagged as 'content'
                    if c_text:
                        chunk = ChatGenerationChunk(
                            message=AIMessageChunk(
                                content=c_text,
                                additional_kwargs={"chunk_type": "content"},
                            )
                        )
                        if run_manager:
                            await run_manager.on_llm_new_token(c_text, chunk=chunk)
                        yield chunk
                except Exception:
                    pass

llm = DeepSeekThinkingLLM()

# Initialize legal modules
case_db = None
similarity_matcher = None
case_tracker = None


def _get_case_db() -> LegalCaseDatabase:
    global case_db
    if case_db is None:
        case_db = LegalCaseDatabase()
    return case_db


def _get_similarity_matcher() -> CaseSimilarityMatcher:
    global similarity_matcher
    if similarity_matcher is None:
        similarity_matcher = CaseSimilarityMatcher()
    return similarity_matcher


def _get_case_tracker() -> OngoingCaseTracker:
    global case_tracker
    if case_tracker is None:
        case_tracker = OngoingCaseTracker()
    return case_tracker


# ── Compliance rules DB ───────────────────────────────────────────────────────
RULES_DB = {
    "重大决策": {
        "强制要求": "必须经党委前置研究讨论；集体决策；法律审查；会议纪要存档。",
        "禁止事项": "禁止个人或少数人决定；无会议纪要；未经法律审查。",
        "决策程序": ["提出书面建议书", "党支部审查列入", "承办部门拟方案", "征求意见", "院办公室报告", "院务会议集体讨论表决", "实施与监督"],
        "责任主体": "党委书记/董事长主持；主管领导论证；法律合规部审查；纪检监督部监督。"
    },
    "重大项目安排": {
        "强制要求": "必须可行性报告、风险评估、法律审查；党委前置；【城投特色】必须包含关联项目编码（Project ID）以穿透项目全生命周期库资金链；【城投特色】必须经公共资源交易中心招投标系统自动比对，决策时间必须合规且早于实际招标时间。",
        "禁止事项": "禁止超预算、无审批；【城投特色】严禁重复立项或超概算支付；严禁先斩后奏的违规招标。",
        "决策程序": ["项目审查", "专家论证", "征求意见", "会议决策", "公示", "正式审批"],
        "责任主体": "战略规划部论证；法律合规部审查。"
    },
    "大额度资金运作": {
        "强制要求": "必须资金使用计划；双人签字或集体审批；【城投特色】必须与财务系统联动，自动校验提取资金数据，核查是否触碰隐性债务红线及融资成本上限。",
        "禁止事项": "禁止私下转账、无审计记录；【城投特色】严禁违规新增地方政府隐性债务。",
        "决策程序": ["安排预算", "党组集体研究", "公开公示", "资金拨付"],
        "责任主体": "财务部门执行；审计部监督。"
    },
    "重要人事任免": {
        "强制要求": "坚持党管干部；事先征求纪检意见；集体决定；任前公示；试用期考核。",
        "禁止事项": "禁止个人决定。",
        "决策程序": ["民主推荐", "组织考察", "会议决定", "任前公示", "试用1年", "正式任免"],
        "责任主体": "人力资源部考察；纪检监督部意见。"
    }
}

# ── Agent Tools ───────────────────────────────────────────────────────────────
@tool
def extract_rules(matter_type: str, custom_rules_text: str = "") -> str:
    """从制度提取当前事项的强制要求、禁止事项、决策程序、责任主体"""
    data = dict(RULES_DB.get(matter_type.strip(), {"error": "未匹配事项类型"}))
    if "error" in data or not custom_rules_text.strip():
        return json.dumps(data, ensure_ascii=False, indent=2)

    custom_summary = _compose_custom_rule_summary(custom_rules_text, matter_type)
    if custom_summary["mandatory"]:
        data["强制要求"] = f"{data.get('强制要求', '')}；自定义制度补充：{'；'.join(custom_summary['mandatory'])}"
    if custom_summary["forbidden"]:
        data["禁止事项"] = f"{data.get('禁止事项', '')}；自定义制度补充：{'；'.join(custom_summary['forbidden'])}"
    if custom_summary["procedures"]:
        merged_steps = list(dict.fromkeys((data.get("决策程序", []) + custom_summary["procedures"])))
        data["决策程序"] = merged_steps
    data["自定义制度摘要"] = custom_summary["summary_lines"]
    return json.dumps(data, ensure_ascii=False, indent=2)

@tool
def validate_material(material_text: str, rules_text: str) -> str:
    """与规则交叉校验，返回每条规则状态 + 证据（关键词匹配）"""
    try:
        rules = json.loads(rules_text)
        if "error" in rules:
            return json.dumps({"error": rules["error"]}, ensure_ascii=False)
        report = []
        sentences = [s.strip() for s in material_text.split('。') if s.strip()]
        for key in ["强制要求", "禁止事项"]:
            if key not in rules:
                continue
            rule_text = rules[key]
            keywords = [w for w in rule_text.replace('；', '、').replace('，', '、').split('、') if len(w) >= 2]
            best_evidence, best_hits = "无明显证据", 0
            for sent in sentences:
                hits = sum(1 for kw in keywords if kw in sent)
                if hits > best_hits:
                    best_hits, best_evidence = hits, sent
            hit_rate = best_hits / max(len(keywords), 1)
            report.append({"规则": rule_text, "状态": "合规" if hit_rate > 0.3 else "⚠️ 不合规",
                           "关键词命中率": round(hit_rate, 2), "证据": best_evidence})
        return json.dumps(report, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)

@tool
def check_procedure_completeness(material_text: str, rules_text: str) -> str:
    """检查决策程序是否完整覆盖"""
    try:
        rules = json.loads(rules_text)
        if "error" in rules:
            return json.dumps({"error": rules["error"]}, ensure_ascii=False)
        steps = rules.get("决策程序", [])
        report = [{"环节": step, "状态": "已覆盖" if step in material_text else "缺失"} for step in steps]
        return json.dumps(report, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)

@tool
def identify_responsibility(material_text: str, rules_text: str) -> str:
    """检查责任主体和监督是否落实"""
    try:
        rules = json.loads(rules_text)
        if "error" in rules:
            return json.dumps({"error": rules["error"]}, ensure_ascii=False)
        subject = rules.get("责任主体", "")
        if subject and subject in material_text:
            return "责任主体明确，监督机制提及。"
        return "责任主体或监督缺失，请补充。"
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)

@tool
def generate_compliance_report(results: str) -> str:
    """综合所有工具结果，生成国企公文规范的 Markdown 合规审核报告"""
    try:
        prompt = f"""你是资深三重一大合规审核专家，请严格按照国企公文规范，基于以下工具执行结果，输出一份排版精美的 Markdown 合规审核报告。

【工具执行结果原始数据（仅供参考，不得在报告中出现）】
{results}

=========
【输出格式要求】
1. 报告全程使用中文。
2. 输出格式为纯 Markdown 文本，章节用 ## 标题区隔，数据区域用表格呈现。
3. 在报告最开头，直接输出以下风险雷达 XML 数据块（根据实际审查结果，将 status 改为 green/yellow/red）：
   <risk_radar>
   <item status="green">党委前置审查(通过)</item>
   <item status="yellow">程序完整性(缺少某环节)</item>
   <item status="red">大额资金审批(违规)</item>
   </risk_radar>
4. 紧接 XML 块之后，输出以下九个章节：
   - ## 一、审核基本信息（表格形式：审核类型、审核日期、审核结论）
   - ## 二、风险等级评定（使用 ⚠️ 高风险 / 🔶 中风险 / 🟢 低风险，说明评定理由）
   - ## 三、违规事项与证据清单（务必在叙述违规/合规项时，用特殊的 Markdown 链接语法标出证据原文。例如：[此项目缺乏财务审计报告](evidence:"未见财务部门签字的审计文本")）
   - ## 四、程序完整性核查（Markdown 表格：程序环节 | 状态 | 备注）
   - ## 五、责任主体认定（说明责任人/监督部门落实情况）
   - ## 六、整改建议（分条陈述，每条格式：**建议N**：具体要求。关键要求：在每条具体建议的末尾，必须加上一个特殊的生成按钮链接。例如：[🔨 一键生成《整改通知单》及补正依据](remediate:"为这个项目起草缺少前置审查环节的情况说明及补发文的模版")）
   - ## 七、决策溯源档案（电子档案）（说明：系统已具备"一键生成迎检报告"功能，支持按年度、按事项类型一键导出决策记录、参会人员名单、表决结果及原始文件，形成完整的电子档案库，以应对国资委巡视和审计。请以此口吻进行简要功能描述提示。）
   - ## 八、整改闭环管理（说明：除审核外，系统具备"整改通知下发 -> 整改报告上报 -> 复核销号"的闭环管理模块，保障违规问题闭环跟踪。请在此简要描述本次审核所涉及的闭环流程或系统功能赋能。）
   - ## 九、统计分析与驾驶舱（说明：系统具备决策效能分析看板，实时展示本年度党委会召开次数、重大项目研究数量、资金总额、"紧急上会/临时动议"等异常事项占比，支撑向国资委汇报的关键 KPI。请简要描述本系统驾驶舱模块是如何实时监控统计及展示此类指标的。）
5. 报告末尾附：> 📋 本报告由 AI 合规审核系统自动生成，仅供参考，最终结论以人工复核为准。

现在从 <risk_radar> 开始输出："""
        response = llm.invoke(prompt)
        # Force a stop of the ReAct agent loop by raising our custom exception
        raise ReportGeneratedException(response.content)
    except Exception as e:
        logger.error(f"Error in generate_compliance_report: {e}")
        # If it's our own stop exception, re-raise it so the upstream try-except catches it
        if isinstance(e, ReportGeneratedException):
            raise e
        return f"报告生成失败: {e}"

tools = [extract_rules, validate_material, check_procedure_completeness, identify_responsibility, generate_compliance_report]

# ── Removed ReAct Agent due to infinite loop bugs ──
# We will use a deterministic 5-step procedural pipeline instead in audit_stream

# ── ChromaDB setup ────────────────────────────────────────────────────────────
PERSIST_DIR = str(Path(os.path.dirname(os.path.abspath(__file__))) / "chroma_db")
vectorstore = None
_vectorstore_error = ""
_vectorstore_lock = Lock()


def _get_vectorstore(create_if_missing: bool = False):
    """Lazy-load Chroma so auth and basic pages are available immediately."""
    global vectorstore, _vectorstore_error
    if vectorstore is not None:
        return vectorstore

    with _vectorstore_lock:
        if vectorstore is not None:
            return vectorstore

        persist_path = Path(PERSIST_DIR)
        if not persist_path.exists():
            if not create_if_missing:
                return None
            persist_path.mkdir(parents=True, exist_ok=True)

        if _embedding_fn_cls is None:
            _vectorstore_error = f"Embedding 模块加载失败：{_emb_import_err_msg}"
            logger.warning(_vectorstore_error)
            return None

        try:
            logger.info(f"Loading ChromaDB from {persist_path}")
            _embed_fn = _embedding_fn_cls("shibing624/text2vec-base-chinese")
            import chromadb
            _client = chromadb.PersistentClient(path=str(persist_path))
            vectorstore = Chroma(
                client=_client,
                collection_name="langchain",
                embedding_function=_embed_fn,
            )
            _vectorstore_error = ""
            logger.info("ChromaDB loaded successfully.")
            return vectorstore
        except Exception as e:
            _vectorstore_error = str(e)
            logger.error(f"Failed to load ChromaDB: {e}")
            return None

# Legal Comparison Skills
@tool
def find_similar_legal_cases(case_description: str, case_type: str) -> str:
    """在历史案例库中查找相似案例并计算胜诉率"""
    try:
        result = _get_similarity_matcher().get_case_comparison_details(case_description, case_type, top_k=3)
        return json.dumps(result, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Error in find_similar_legal_cases: {str(e)}")
        return json.dumps({"error": str(e)}, ensure_ascii=False)

@tool
def calculate_win_rate(case_description: str, case_type: str) -> str:
    """基于相似案例计算胜诉率和风险分析"""
    try:
        matcher = _get_similarity_matcher()
        win_rate = matcher.calculate_win_rate(case_description, case_type, top_k=5)
        risk_analysis = matcher.analyze_case_risk(case_description, case_type)
        return json.dumps({
            "win_rate_analysis": win_rate,
            "risk_analysis": risk_analysis
        }, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Error in calculate_win_rate: {str(e)}")
        return json.dumps({"error": str(e)}, ensure_ascii=False)

@tool
def get_ongoing_cases_summary() -> str:
    """获取当前正在进行的案件概览和开庭提醒"""
    try:
        tracker = _get_case_tracker()
        summary = tracker.get_case_summary()
        upcoming = tracker.get_upcoming_hearings(30)
        return json.dumps({
            "summary": summary,
            "upcoming_hearings": [
                {
                    "case_id": case.case_id,
                    "case_title": case.case_title,
                    "hearing_date": case.next_hearing_date,
                    "court_name": case.court_name,
                    "case_type": case.case_type
                }
                for case in upcoming
            ]
        }, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Error in get_ongoing_cases_summary: {str(e)}")
        return json.dumps({"error": str(e)}, ensure_ascii=False)

@tool
def generate_legal_report(legal_analysis: str) -> str:
    """综合法务分析结果生成最终报告"""
    try:
        prompt = f"""
        你是法务分析专家。
        基于以下法律分析结果：
        {legal_analysis}

        生成结构化的法务分析报告（Markdown 格式）：
        1. 案件类型与基本信息
        2. 相似案例分析（引用历史案例）
        3. 胜诉率预测与置信度
        4. 风险评估（高/中/低）
        5. 关键风险因素
        6. 应诉建议与策略
        7. 当前相关案件提醒
        严格基于数据，客观分析。
        """
        response = llm.invoke(prompt)
        return response.content
    except Exception as e:
        logger.error(f"Error in generate_legal_report: {str(e)}")
        return f"报告生成失败: {str(e)}"

# Compliance tools
compliance_tools = [
    extract_rules,
    validate_material,
    check_procedure_completeness,
    identify_responsibility,
    generate_compliance_report
]

# Legal comparison tools
legal_tools = [
    find_similar_legal_cases,
    calculate_win_rate,
    get_ongoing_cases_summary,
    generate_legal_report
]

# Compliance Agent prompt
compliance_prompt = ChatPromptTemplate.from_messages([
    ("system", """你是严格的三重一大合规审核 Agent。
必须按顺序执行：
1. extract_rules
2. validate_material
3. check_procedure_completeness
4. identify_responsibility
5. generate_compliance_report
不要添加任何外部司法/案件内容。

可以使用以下工具：
{tools}
工具名称列表：{tool_names}"""),
    ("human", "{input}"),
    MessagesPlaceholder("agent_scratchpad"),
])

# Legal Comparison Agent prompt
legal_prompt = ChatPromptTemplate.from_messages([
    ("system", """你是专业的法务对比分析 Agent。
必须按顺序执行：
1. find_similar_legal_cases
2. calculate_win_rate
3. get_ongoing_cases_summary
4. generate_legal_report
基于历史案例库和胜诉率预测，提供客观的法务分析建议。

可以使用以下工具：
{tools}
工具名称列表：{tool_names}"""),
    ("human", "{input}"),
    MessagesPlaceholder("agent_scratchpad"),
])

# Initialize compliance agent
compliance_agent = create_react_agent(llm, compliance_tools, compliance_prompt)
compliance_executor = AgentExecutor(
    agent=compliance_agent,
    tools=compliance_tools,
    verbose=True,
    max_iterations=10,
    handle_parsing_errors=True
)

# Initialize legal comparison agent
legal_agent = create_react_agent(llm, legal_tools, legal_prompt)
legal_executor = AgentExecutor(
    agent=legal_agent,
    tools=legal_tools,
    verbose=True,
    max_iterations=10,
    handle_parsing_errors=True
)


# ── API Routes ────────────────────────────────────────────────────────────────
@app.get("/legal-case-types")
async def get_legal_case_types(request: Request):
    """获取可用的案件类型"""
    _get_request_user(request, required=True)
    cases = _get_case_db().get_all_cases()
    types = list(set(case.case_type for case in cases))
    return {"legal_case_types": types}

@app.post("/legal-compare", response_model=ChatResponse)
async def legal_compare(request: LegalCompareRequest):
    """执行法务对比分析"""
    try:
        logger.info(f"收到法务对比请求 - 案件类型: {request.case_type}")

        query = f"""
        案件类型：{request.case_type}
        案件描述：{request.case_description}
        案件金额：{request.case_amount} 元
        """

        logger.info("开始执行法务对比分析...")
        result = legal_executor.invoke({"input": query})

        logger.info("法务对比分析完成")
        return ChatResponse(
            success=True,
            message="分析完成",
            report=result["output"]
        )

    except Exception as e:
        logger.error(f"法务分析失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/ongoing-cases")
async def get_ongoing_cases(request: Request):
    """获取所有正在进行的案件"""
    _get_request_user(request, required=True)
    try:
        cases = _get_case_tracker().get_all_ongoing_cases()
        return {
            "success": True,
            "total": len(cases),
            "cases": [case.to_dict() for case in cases]
        }
    except Exception as e:
        logger.error(f"获取案件失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/upcoming-hearings")
async def get_upcoming_hearings(request: Request):
    """获取即将开庭的案件"""
    _get_request_user(request, required=True)
    try:
        hearings = _get_case_tracker().get_upcoming_hearings(30)
        return {
            "success": True,
            "total": len(hearings),
            "hearings": [case.to_dict() for case in hearings]
        }
    except Exception as e:
        logger.error(f"获取开庭信息失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/")
async def root():
    return {"message": "三重一大合规审核 API 服务已启动"}

@app.get("/matter-types")
async def get_matter_types(request: Request):
    _get_request_user(request, required=True)
    return {"matter_types": list(RULES_DB.keys())}


@app.get("/api/demo_assets")
async def get_demo_assets_api(request: Request):
    _get_request_user(request, required=True)
    return JSONResponse({"success": True, **get_demo_assets()})


@app.get("/api/audit_history")
async def get_audit_history_api(request: Request):
    _get_request_user(request, required=True)
    try:
        history = persistence.get_history()
        return JSONResponse({
            "success": True,
            "history": build_archive_history(history),
        })
    except Exception as e:
        logger.exception(f"归档记录加载失败: {e}")
        return JSONResponse({
            "success": True,
            "history": build_archive_history([]),
            "warning": "归档历史存在异常，已自动回退为内置档案样例。",
        })


AUTH_DATA_DIR = Path(os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "auth"))
AUTH_DATA_DIR.mkdir(parents=True, exist_ok=True)
USERS_DB = AUTH_DATA_DIR / "users.json"

CUSTOM_RULES_DIR = Path(os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "custom_rules"))
CUSTOM_RULES_DIR.mkdir(parents=True, exist_ok=True)
CUSTOM_RULES_DB = CUSTOM_RULES_DIR / "files.json"

# 以下常量均从 backend.config 导入，此处不再重复定义，
# 避免锁对象重复导致 backend/db.py 与 backend_full.py 同步失效。

# ── WAL checkpoint — prevent unbounded WAL file growth ───────────────────────
_WAL_CHECKPOINT_INTERVAL = 300  # seconds between auto-checkpoints






















def _meeting_role_for_user(user: dict) -> dict:
    """从用户真实数据解析会议角色——不再硬编码任何人。"""
    return {
        "displayName": user.get("name") or user.get("username") or "参会人",
        "meetingRole": user.get("meetingRole") or "参会代表",
        "seat": user.get("meetingSeat") or (user.get("dept") or "参会部门"),
        "voiceprint": 90,
    }




def _append_meeting_event(meeting_id: str, event: dict) -> dict:
    safe_id = re.sub(r"[^a-zA-Z0-9_-]", "_", meeting_id or "default-meeting")
    with MEETING_TRANSCRIPTS_LOCK:
        data = _load_meeting_transcripts()
        meeting = data.setdefault(safe_id, {"meetingId": safe_id, "events": [], "updatedAt": ""})
        meeting["events"].append(event)
        meeting["updatedAt"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        meeting["events"] = meeting["events"][-200:]
        _save_meeting_transcripts(data)
        return meeting













































































def _extract_json_object(text: str, required_keys: list = None) -> Optional[dict]:
    """从文本中提取 JSON 对象。

    Args:
        text: 包含 JSON 的文本
        required_keys: 如果指定，只返回包含这些 key 的 dict
    """
    content = (text or "").strip()
    if not content:
        return None
    # 移除 markdown 代码块
    content = re.sub(r"^```(?:json)?\s*", "", content, flags=re.IGNORECASE)
    content = re.sub(r"\s*```$", "", content)
    # 尝试直接解析整个内容
    try:
        data = json.loads(content)
        if isinstance(data, dict):
            if not required_keys or any(k in data for k in required_keys):
                return data
    except Exception:
        pass
    # 查找所有 JSON 对象
    decoder = json.JSONDecoder()
    candidates = []
    for match in re.finditer(r"\{", content):
        try:
            data, _ = decoder.raw_decode(content[match.start():])
            if isinstance(data, dict):
                candidates.append(data)
        except Exception:
            continue
    # 优先返回包含 required_keys 的候选
    if required_keys:
        for d in candidates:
            if all(k in d for k in required_keys):
                return d
        # 如果没有完全匹配的，返回包含部分 required_keys 的最大候选
        partial_matches = [d for d in candidates if any(k in d for k in required_keys)]
        if partial_matches:
            return max(partial_matches, key=lambda d: len(json.dumps(d, ensure_ascii=False)))
    # 否则返回最大的候选（通常是最完整的）
    if candidates:
        return max(candidates, key=lambda d: len(json.dumps(d, ensure_ascii=False)))
    return None




def _is_low_value_meeting_utterance(text: str) -> bool:
    compact = re.sub(r"[\s，。,.!?！？、；;：:（）()【】\[\]\"'""‘’]", "", str(text or "")).lower()
    if not compact or len(compact) <= 2:
        return True
    patterns = [
        r"^我是[\u4e00-\u9fa5a-z0-9]{1,12}$",
        r"^我叫[\u4e00-\u9fa5a-z0-9]{1,12}$",
        r"^喂+$",
        r"^嗯+$",
        r"^啊+$",
        r"^测试$",
        r"^test$",
        r"^yeah$",
        r"^seven$",
        r"^hello$",
    ]
    return any(re.match(pattern, compact, flags=re.IGNORECASE) for pattern in patterns)


def _is_garbled_asr(text: str) -> bool:
    """检测 ASR 乱码噪音，避免送入 DeepSeek 干扰分析。"""
    t = str(text or "").strip()
    if len(t) <= 6:
        return True
    compact = re.sub(r"[\s，。,.!?！？、；;：:（）()【】\[\]\"'""''""]", "", t).lower()
    if len(compact) < 6:
        return True
    # 单字重复率 > 50%
    from collections import Counter
    char_counts = Counter(compact)
    if char_counts and char_counts.most_common(1)[0][1] / len(compact) > 0.5 and len(compact) > 8:
        return True
    # 连续相同字符 >= 6 次
    if re.search(r"(.)\1{5,}", compact):
        return True
    # 碎片化严重：平均词长 < 2.5 且碎片 > 6
    fragments = re.findall(r"[一-龥a-z0-9]+", t)
    if fragments and sum(len(f) for f in fragments) / len(fragments) < 2.5 and len(fragments) > 6:
        return True
    return False


def _extract_agenda_keywords(title: str) -> List[str]:
    raw = re.split(r"[\s，,、；;：:。.!?！？（）()【】\[\]《》\"'""‘’/\\-]+", str(title or ""))
    stop_words = {"讨论", "审议", "研究", "关于", "事项", "安排", "工作", "会议", "问题", "进行", "后续", "发布"}
    keywords = []
    for word in raw:
        word = word.strip()
        if len(word) < 2 or word in stop_words:
            continue
        keywords.append(word)
    for word in re.findall(r"[A-Za-z][A-Za-z0-9._-]{2,}", str(title or "")):
        if word not in keywords:
            keywords.append(word)
    return keywords[:12]


def _local_realtime_agenda_check(agenda_drafts: List[dict], latest_transcripts: List[dict]) -> List[dict]:
    transcript_text = " ".join(_clean_agenda_check_transcript(item.get("transcript") or item.get("text")) for item in latest_transcripts)
    useful_texts = [
        _clean_agenda_check_transcript(item.get("transcript") or item.get("text"))
        for item in latest_transcripts
        if not _is_low_value_meeting_utterance(item.get("transcript") or item.get("text"))
    ]
    useful_joined = " ".join(useful_texts)
    results = []
    for index, item in enumerate(agenda_drafts[:8]):
        agenda_id = str(item.get("id") or f"agenda-{index}")
        title = str(item.get("title") or item.get("agenda") or "待确认议题")
        if not transcript_text:
            relation, status, confidence, reason = "waiting", "等待讨论", 0.0, "尚未收到手机端实时转写。"
        elif not useful_joined:
            relation, status, confidence, reason = "irrelevant", "未命中", 0.08, "当前发言主要是身份确认或测试语句，不能判定为正在讨论该议题。"
        else:
            keywords = _extract_agenda_keywords(title)
            hit_words = [word for word in keywords if word and word.lower() in useful_joined.lower()]
            # 即使关键词没命中，只要有实质发言且不是测试语句，也判定为正在讨论
            if hit_words:
                relation = "matched"
                status = "正在比对"
                confidence = min(0.92, 0.42 + len(hit_words) * 0.14)
                reason = f"最近发言命中议题关键词：{'、'.join(hit_words[:4])}。"
            elif len(useful_joined) > 30:
                # 有实质内容但没命中关键词 —— 判定为正在讨论（语义相关但关键词不同）
                relation = "matched"
                status = "正在比对"
                confidence = 0.55
                reason = "检测到实质发言内容，可能与议题相关。"
            else:
                relation, status, confidence, reason = "irrelevant", "未命中", 0.18, "最近发言未命中当前议题关键词，暂不判定为议题讨论。"
        results.append({
            "agendaId": agenda_id,
            "relation": relation,
            "status": status,
            "confidence": round(float(confidence), 2),
            "reason": reason,
            "suggestion": "继续监听" if relation in {"waiting", "irrelevant"} else "继续比对",
            "evidence": useful_texts[-2:],
        })
    return results


def _rag_search_meeting_context(meeting: dict, query_extra: str = "", top_k: int = 5) -> str:
    """从 ChromaDB 知识库检索与会议议题/项目相关的制度条款。
    不可用时静默降级，返回空字符串。"""
    try:
        vs = _get_vectorstore(False)
        if not vs:
            return ""
        query_parts = []
        for key in ("agenda", "project", "title"):
            val = str(meeting.get(key, "")).strip()
            if val and val not in query_parts:
                query_parts.append(val)
        if query_extra and query_extra not in query_parts:
            query_parts.append(query_extra)
        query = " ".join(query_parts)
        if not query:
            return ""
        results = vs.similarity_search(query, k=top_k)
        if not results:
            return ""
        contexts = []
        for i, doc in enumerate(results, 1):
            source = doc.metadata.get("source", "未知来源")
            page = doc.metadata.get("page")
            page_info = f" 第{page}页" if page else ""
            contexts.append(f"【参考 {i}】来源: {source}{page_info}\n{doc.page_content[:800]}")
        return "\n\n" + "\n\n---\n".join(contexts)
    except Exception:
        return ""


async def _deepseek_realtime_agenda_check(meeting: dict, agenda_drafts: List[dict], latest_transcripts: List[dict], meeting_mode: str) -> List[dict]:
    compact_agendas = [
        {"id": str(item.get("id") or f"agenda-{index}"), "title": str(item.get("title") or item.get("agenda") or "待确认议题")[:120]}
        for index, item in enumerate(agenda_drafts[:8])
    ]
    compact_transcripts = [
        {
            "speaker": item.get("speakerName") or item.get("speaker") or item.get("username") or "",
            "time": item.get("clientTime") or item.get("serverTime") or item.get("time") or "",
            "text": _clean_agenda_check_transcript(item.get("transcript") or item.get("text"))[:240],
        }
        for item in latest_transcripts[-12:]
    ]
    if not compact_agendas or not any(item.get("text") for item in compact_transcripts):
        return []
    prompt = f"""
你是政府会议实时纪要助手。请判断"最近发言"是否正在讨论各个会议议题。

会议信息：
- 会议名称：{meeting.get("title") or "AI 会议"}
- 会议性质：{"三重一大会议" if meeting_mode == "major" else "普通会议"}


议题 JSON：
{json.dumps(compact_agendas, ensure_ascii=False)}

最近转写 JSON：
{json.dumps(compact_transcripts, ensure_ascii=False)}

规则：
1. "我是某某""喂""测试""Yeah""Seven"等身份确认或测试语句，判定为 irrelevant。
2. 发言内容与议题标题存在语义相关即可判定为 matched，不需要逐字匹配。同一领域的话题（如讨论足球俱乐部和球员，与"足球媒体"议题相关）应判定为正在讨论。
3. 如果发言已经形成结论、责任人、动作或时间，relation 可为 resolved。
4. 普通会议不要输出三重一大、材料拦截、金额阈值等内容。
5. 每个议题都要返回结果，禁止编造发言之外的信息。
6. 宽松判定: 发言内容只要与议题有合理关联就应该判定为 matched，不要过度严格。

只返回 JSON，不要 Markdown：
{{
  "results": [
    {{
      "agendaId": "议题id",
      "relation": "waiting|irrelevant|matched|resolved|disputed",
      "status": "等待讨论|未命中|正在比对|已回应|有争议",
      "confidence": 0.0,
      "reason": "一句话说明判断依据",
      "suggestion": "继续监听|切换议题|形成待办|会后复核",
      "evidence": ["引用最相关的短发言，最多2条"]
    }}
  ]
}}
"""
    try:
        async with llm_semaphore:
            result = await llm._agenerate(
                [SystemMessage(content="你只输出可解析 JSON，字段必须稳定，适合前端直接渲染。"), HumanMessage(content=prompt)],
                enable_thinking=False,
            )
        text = result.generations[0].message.content if result.generations else ""
        payload = _extract_json_object(text)
        rows = payload.get("results") if isinstance(payload, dict) else None
        if not isinstance(rows, list):
            return []
        valid_ids = {item["id"] for item in compact_agendas}
        normalized = []
        for row in rows:
            if not isinstance(row, dict):
                continue
            agenda_id = str(row.get("agendaId") or "")
            if agenda_id not in valid_ids:
                continue
            relation = str(row.get("relation") or "irrelevant")
            if relation not in {"waiting", "irrelevant", "matched", "resolved", "disputed"}:
                relation = "irrelevant"
            normalized.append({
                "agendaId": agenda_id,
                "relation": relation,
                "status": str(row.get("status") or ("正在比对" if relation == "matched" else "未命中"))[:20],
                "confidence": max(0, min(1, float(row.get("confidence") or 0))),
                "reason": str(row.get("reason") or "")[:120],
                "suggestion": str(row.get("suggestion") or "继续监听")[:40],
                "evidence": [str(item)[:120] for item in row.get("evidence", [])[:2]] if isinstance(row.get("evidence"), list) else [],
            })
        return normalized
    except Exception as exc:
        logger.warning("DeepSeek 实时议题比对失败：%s", exc)
        return []


def _local_generate_meeting_records(meeting: dict, transcripts: List[dict], events: List[dict]) -> dict:
    clean_transcripts = [
        {
            "time": item.get("clientTime") or str(item.get("serverTime") or "")[11:19],
            "speaker": item.get("speakerName") or item.get("username") or "参会人",
            "role": item.get("speakerRole") or item.get("seat") or "参会代表",
            "text": _clean_agenda_check_transcript(item.get("transcript")),
            "signed": bool(item.get("correctionSigned")),
        }
        for item in transcripts
        if _clean_agenda_check_transcript(item.get("transcript"))
    ]
    audio_events = [item for item in events if item.get("type") == "audio" and item.get("playbackUrl")]
    if not clean_transcripts:
        return {
            "generated": False,
            "aiProvider": "none",
            "source": "暂无真实转写",
            "message": "当前会议还没有手机端实时转写，不能生成真实会议记录。",
            "transcriptCount": 0,
            "audioCount": len(audio_events),
            "chronicle": [],
            "summary": [],
            "minutes": [],
            "decisions": [],
            "todos": [],
        }

    agenda_titles = [
        str(item.get("title") or "").strip()
        for item in (meeting.get("agendaDrafts") if isinstance(meeting.get("agendaDrafts"), list) else [])
        if str(item.get("title") or "").strip()
    ] or [str(meeting.get("agenda") or "本次会议议题").strip()]
    all_text = " ".join(item["text"] for item in clean_transcripts)
    summary = [
        f'本次会议围绕“{agenda_titles[0]}”形成了 {len(clean_transcripts)} 条真实转写底稿。',
        f"系统已留存 {len(audio_events)} 段手机端录音原件，可与转写底稿交叉核验。",
    ]
    # 检测结论性表达 — 排除否定形式（不同意、未通过等）
    _negation_re = re.compile(r'(不|未|没|反对|否决)\s*(同意|通过|决定|确认)')
    _conclusion_words = ["同意", "通过", "决定", "确认", "原则同意", "暂缓", "再议"]
    _has_conclusion = False
    for word in _conclusion_words:
        if word in all_text:
            # 检查该词前面是否有否定前缀
            pos = all_text.find(word)
            prefix = all_text[max(0, pos-3):pos]
            if not re.search(r'(不|未|没|反对|否决)\s*$', prefix):
                _has_conclusion = True
                break
    if _has_conclusion:
        summary.append("发言中出现明确结论性表达，建议秘书复核后固化为会议决议。")
    else:
        summary.append('发言中尚未出现稳定的"同意/通过/决定"表述，决议需由秘书会后确认。')

    chronicle = [
        {
            "time": item["time"] or "--:--",
            "speaker": item["speaker"],
            "role": item["role"],
            "content": item["text"],
            "signed": item["signed"],
        }
        for item in clean_transcripts
    ]
    minutes = []
    for index, title in enumerate(agenda_titles[:6], start=1):
        keywords = _extract_agenda_keywords(title)
        related = [
            item for item in clean_transcripts
            if any(word and word.lower() in item["text"].lower() for word in keywords)
        ]
        minutes.append({
            "agenda": title,
            "status": "已讨论" if related else "待人工确认",
            "basis": f"命中 {len(related)} 条相关转写" if related else "真实转写中暂未发现明显相关发言",
            "keyPoints": [item["text"][:120] for item in related[:3]],
        })

    # ── 决议提取（增加上下文过滤 + 议题关联）──
    decision_words = ("同意", "通过", "决定", "确认", "原则同意", "暂缓", "再议")
    negate_re = re.compile(r'(不|未|没|反对|否决)\s*(同意|通过|决定|确认)')
    _context_negate = re.compile(r'(但是|不过|可是|然而|如果|条件是|前提|还是|还要|还需要|先)')
    # 预构建议题关键词索引，用于决议→议题关联
    _agenda_kw_index = []
    agenda_drafts = meeting.get("agendaDrafts") or []
    for _idx, _item in enumerate(agenda_drafts[:8]):
        _title = str(_item.get("title") or _item.get("agenda") or "")
        _aid = str(_item.get("id") or f"agenda-{_idx}")
        _kws = _extract_agenda_keywords(_title)
        _agenda_kw_index.append({"id": _aid, "title": _title, "keywords": _kws})

    def _match_agenda_for_decision(text):
        text_lower = text.lower()
        best_id, best_title, best_hits = "", "", 0
        for ag in _agenda_kw_index:
            hits = sum(1 for kw in ag["keywords"] if kw and kw.lower() in text_lower)
            if hits > best_hits:
                best_hits = hits
                best_id = ag["id"]
                best_title = ag["title"]
        return (best_id, best_title) if best_hits > 0 else ("", "")

    decisions = []
    for item in clean_transcripts:
        text = item["text"]
        if not any(word in text for word in decision_words):
            continue
        if negate_re.search(text):
            continue
        if len(text) < 10 or len(text) > 200:
            continue
        for word in decision_words:
            if word in text:
                pos = text.find(word)
                prefix = text[max(0, pos - 6):pos]
                if _context_negate.search(prefix):
                    break
        else:
            _mid, _mtitle = _match_agenda_for_decision(text)
            decisions.append({
                "time": item["time"] or "--:--",
                "speaker": item["speaker"],
                "content": text[:150],
                "status": "待秘书确认",
                "agendaId": _mid,
                "agenda": _mtitle,
            })
        if len(decisions) >= 8:
            break
        else:
            decisions.append({
                "time": item["time"] or "--:--",
                "speaker": item["speaker"],
                "content": text[:150],
                "status": "待秘书确认",
            })
        if len(decisions) >= 8:
            break

    # ── 待办提取（更宽泛的关键词 + 正确提取责任人）──
    todo_patterns = [
        r'(?:需要|必须|应该|要|得)\s*(?:安排|负责|提交|跟进|落实|补|做|完成|准备|整理|确认)',
        r'(?:安排|由|让|叫|请)\s*\S+\s*(?:负责|跟进|落实|处理|做|盯)',
        r'(?:下周|明天|会后|尽快|抓紧|马上|月底|一周内|三天内)',
        r'(?:你来|来负责|来跟进|来处理)',
        r'(?:记得|别忘了|注意)',
    ]
    _todo_re = re.compile('|'.join(todo_patterns))
    # 责任人提取 — 从发言中找被指派人，而非发言人
    _owner_patterns = [
        re.compile(r'(?:由|让|请|叫)\s*([一-龥]{2,4})(?:\s*(?:负责|跟进|落实|处理|做|来|盯))'),
        re.compile(r'([一-龥]{2,4})\s*(?:你来|来负责|来跟进|来处理|来做|来盯)'),
        re.compile(r'(?:安排)\s*([一-龥]{2,4})'),
    ]
    # 截止时间提取（扩充）
    _deadline_patterns = {
        r'下周[一二三四五六日]': None,  # 动态处理
        r'下周': '下周', r'明天': '明天', r'后天': '后天',
        r'月底': '月底前', r'会后': '会后',
        r'尽快': '尽快', r'马上': '立即', r'立即': '立即', r'抓紧': '尽快',
        r'三天内': '3天内', r'一周内': '一周内', r'两周内': '两周内',
        r'月底前': '月底前', r'月底': '月底前',
        r'下个月': '下个月', r'本月底': '本月底前',
    }
    _deadline_re = re.compile('|'.join(k for k in _deadline_patterns.keys()))
    # 优先级判断（扩充）
    _high_priority = re.compile(r'立即|马上|抓紧|紧急|火速|尽快.*不能拖')
    _medium_priority = re.compile(r'尽快|尽量|尽早|不能拖')
    todos = []
    _todo_seen = set()  # 去重
    for item in clean_transcripts:
        if not _todo_re.search(item["text"]):
            continue
        text = item["text"][:150]
        # 去重：同一发言人相似内容只保留一条
        dedup_key = f"{item['speaker']}:{text[:30]}"
        if dedup_key in _todo_seen:
            continue
        _todo_seen.add(dedup_key)
        # 提取责任人（从发言中找被指派人）
        owner = item["speaker"]  # 默认为发言人
        for pat in _owner_patterns:
            m = pat.search(item["text"])
            if m:
                candidate = m.group(1)
                # 排除常见非人名词汇
                if candidate not in ("我们", "大家", "你们", "他们", "这个", "那个", "所有"):
                    owner = candidate
                    break
        # 提取截止时间
        deadline_match = _deadline_re.search(item["text"])
        deadline = '待定'
        if deadline_match:
            matched = deadline_match.group()
            deadline = _deadline_patterns.get(matched) or matched
        # 判断优先级
        if _high_priority.search(item["text"]):
            priority = '高'
        elif _medium_priority.search(item["text"]):
            priority = '中'
        else:
            priority = '低'
        todos.append({
            "task": text,
            "owner": owner,
            "deadline": deadline,
            "priority": priority,
            "reference": text[:30],
            "timestamp": item["time"] or "",
            "status": "待确认",
        })
        if len(todos) >= 8:
            break
    return {
        "generated": True,
        "aiProvider": "local-rule",
        "source": "真实转写本地整理",
        "message": "DeepSeek 不可用时，系统按真实转写做本地规则整理。",
        "transcriptCount": len(clean_transcripts),
        "audioCount": len(audio_events),
        "chronicle": chronicle,
        "summary": summary,
        "minutes": minutes,
        "decisions": decisions,
        "todos": todos,
    }


def _clean_ai_payload(payload: dict) -> dict:
    """清洗 DeepSeek 输出中看起来像原始 ASR 噪音的条目。"""
    cleaned = dict(payload)
    filler_pattern = re.compile(
        r"(对对+|是是是+|嗯嗯+|啊啊+|就就+|卖卖+|返返+|将将+|去去+|做做做+|已经全部做进去了|"
        r"你好像你得你你你|那可能创业人数项目之啊|他他这个结债过来)",
    )
    for field in ("decisions", "todos"):
        items = cleaned.get(field)
        if not isinstance(items, list):
            continue
        kept = []
        for item in items:
            content = str(item.get("content") or item.get("task") or "")
            # 丢弃明显是 ASR 碎片的条目
            if filler_pattern.search(content):
                continue
            if len(content) > 300:  # 太长通常是从转写原文复制过来的
                continue
            kept.append(item)
        cleaned[field] = kept
    return cleaned


def _clean_raw_decisions(decisions: list) -> list:
    """过滤看起来像转写原文而非正式决议的条目。"""
    return [
        d for d in decisions
        if len(str(d.get("content", ""))) < 300
        and not re.search(r"(对对+|是是是+|嗯嗯+|卖卖卖+|返返+|将将+|你好像你得你你)", str(d.get("content", "")))
    ]


def _clean_raw_todos(todos: list) -> list:
    """过滤看起来像转写原文的待办条目。"""
    return [
        t for t in todos
        if len(str(t.get("task", ""))) < 300
        and not re.search(r"(对对+|是是是+|嗯嗯+|卖卖卖+|返返+|将将+|你好像你得你你|他他这个)", str(t.get("task", "")))
    ]


async def _deepseek_generate_meeting_records(meeting: dict, transcripts: List[dict], events: List[dict]) -> dict:
    all_clean = [
        {
            "time": item.get("clientTime") or str(item.get("serverTime") or "")[11:19],
            "speaker": item.get("speakerName") or item.get("username") or "参会人",
            "role": item.get("speakerRole") or item.get("seat") or "参会代表",
            "text": _clean_agenda_check_transcript(item.get("transcript"))[:300],
            "signed": bool(item.get("correctionSigned")),
        }
        for item in transcripts
        if _clean_agenda_check_transcript(item.get("transcript"))
    ]
    # ── 加入 Whisper 终审转写 ──
    whisper_texts = []
    for item in events:
        if item.get("type") == "transcript" and item.get("action") == "whisper-review":
            t = item.get("text", "")
            if t:
                whisper_texts.append(t)
    if whisper_texts:
        whisper_combined = "\n".join(whisper_texts)
        all_clean.append({
            "time": "终审",
            "speaker": "Whisper-large-v3",
            "role": "AI兜底转写（高精度）",
            "text": whisper_combined[:15000],
            "signed": True,
        })
    if not all_clean:
        return {}

    # 智能采样：≤100 条全送，>100 条取头部+均匀采样+尾部，保证时间全覆盖
    total = len(all_clean)
    if total <= 100:
        clean_transcripts = all_clean
    else:
        head = all_clean[:20]
        tail = all_clean[-30:]
        mid = all_clean[20:-30]
        step = max(1, len(mid) // 60)  # 中间取 ~60 条
        sampled_mid = mid[::step]
        clean_transcripts = head + sampled_mid + tail
        logger.info("DeepSeek 转写采样：%d → %d 条（头部20+中部%d+尾部30）", total, len(clean_transcripts), len(sampled_mid))
    agenda_titles = [
        str(item.get("title") or "").strip()
        for item in (meeting.get("agendaDrafts") if isinstance(meeting.get("agendaDrafts"), list) else [])
        if str(item.get("title") or "").strip()
    ] or [str(meeting.get("agenda") or "本次会议议题").strip()]
    audio_count = len([item for item in events if item.get("type") == "audio" and item.get("playbackUrl")])
    prompt = f"""
你是专业会议纪要秘书。以下是一段会议的真实语音转写，请**提炼**为结构化会议记录。

⚠️ 核心要求：
- summary/decisions/todos 是你**归纳提炼**的结果，每条 1-3 句，**严禁**复制粘转写原文长段落
- **严禁编造**：所有内容必须来源于转写原文，不得添加转写中未出现的人名、产品名、工具名、数字等。不确定的宁可不写
- decisions 每条 ≤ 80 字，格式：「会议确定…」或「会议决定…」
- todos 每条包含完整信息：任务、责任人、截止时间、优先级、原文引用
- 跳过纯试音（喂喂喂、大家好、测试），但涉及业务的对话都要提炼
- 如果多段转写讨论同一事项，合并为一条

会议信息：
{json.dumps({
    "title": meeting.get("title"),
    "date": meeting.get("date"),
    "type": meeting.get("type"),
    "agendaTitles": agenda_titles,
    "totalTranscripts": total,
    "note": f"共{total}条转写，{'全部已提供' if total <= 100 else f'已采样{len(clean_transcripts)}条覆盖全时段'}",
}, ensure_ascii=False)}

转写内容（按时间排列）：
{json.dumps(clean_transcripts, ensure_ascii=False)}

严格按以下 JSON 格式输出（不要输出 chronicle，纪实由系统单独生成）：
{{
  "summary": ["会议主要讨论了什么，3-5条概括，每条≤50字"],
  "minutes": [{{"agenda":"议题","status":"已讨论","keyPoints":["要点1","要点2"]}}],
  "decisions": [{{"content":"会议确定/决定…（≤80字）","status":"待秘书确认","agenda":"关联的议题名称（如有）"}}],
  "todos": [{{
    "task": "具体任务描述（≤60字）",
    "owner": "责任人姓名",
    "deadline": "截止时间（未提及填'待定'）",
    "priority": "高/中/低",
    "reference": "相关发言原文片段（≤30字）",
    "timestamp": "HH:MM:SS"
  }}]
}}
"""
    try:
        async with llm_semaphore:
            result = await llm._agenerate(
                [SystemMessage(content="你是会议纪要专家。你的输出是提炼后的正式文本，不是转写原文。每条 decision/todo 一句话。直接输出最终结果，不要输出思考过程。"), HumanMessage(content=prompt)],
                enable_thinking=False,
            )
        text = result.generations[0].message.content if result.generations else ""
        # 提取 JSON
        payload = _extract_json_object(text, required_keys=["summary", "decisions", "todos"])
        if not isinstance(payload, dict):
            logger.warning("【AI纪要】未提取到有效 JSON，返回空")
            return {}
        payload = _clean_ai_payload(payload)
        ai_decisions = _clean_raw_decisions(payload.get("decisions") or [])
        ai_todos = _clean_raw_todos(payload.get("todos") or [])
        logger.info(f"【AI纪要】清洗后 decisions={len(ai_decisions)}, todos={len(ai_todos)}")
        local = _local_generate_meeting_records(meeting, transcripts, events)
        # 优先用 DeepSeek 的提炼结果，空的才用本地兜底
        ai_summary = payload.get("summary")
        use_ai = isinstance(ai_summary, list) and len(ai_summary) > 0 and not any(
            "形成了" in str(s) and "真实转写底稿" in str(s) for s in ai_summary
        )
        return {
            **local,
            "aiProvider": "deepseek",
            "source": f"DeepSeek 根据{total}条转写生成",
            "message": "会议记录由后端读取真实转写后生成，未出现的内容不会写入结论。",
            "summary": ai_summary if use_ai else local["summary"],
            "chronicle": local["chronicle"],
            "minutes": payload.get("minutes") if isinstance(payload.get("minutes"), list) and payload.get("minutes") else local["minutes"],
            "decisions": ai_decisions if ai_decisions else local["decisions"],
            "todos": ai_todos if ai_todos else local["todos"],
        }
    except Exception as exc:
        logger.warning("DeepSeek 会议记录生成失败：%s", exc)
        return {}


def _normalize_ai_agenda_drafts(
    payload: dict,
    imported_count: int,
    fallback_project: str,
    meeting_mode: str = "major",
    source_label: str = "DeepSeek 根据收集素材真实提炼",
) -> List[dict]:
    if any(key in payload for key in ("title", "议题名称", "议题标题", "议题")):
        raw_items = [payload]
    else:
        raw_items = (
            payload.get("agendaDrafts")
            or payload.get("issues")
            or payload.get("items")
            or payload.get("议题池")
            or payload.get("议题列表")
            or payload.get("待办议题")
            or payload.get("上会议题")
            or []
        )
    if not isinstance(raw_items, list):
        return []
    drafts = []
    for index, item in enumerate(raw_items[: max(1, min(imported_count, 6))], start=1):
        if not isinstance(item, dict):
            continue
        title = re.sub(r"\s+", " ", str(item.get("title") or item.get("meetingTodo") or item.get("开会待办") or item.get("议题名称") or item.get("议题标题") or item.get("议题") or "").strip())
        title = re.sub(r"等\d+个项目", "", title).strip(" ，、；;")
        if not title:
            continue
        project = re.sub(r"\s+", " ", str(item.get("project") or item.get("本地项目") or item.get("关联项目") or fallback_project or "本地项目").strip())
        issue_type = re.sub(r"\s+", " ", str(item.get("type") or item.get("businessType") or item.get("业务类型") or item.get("三重一大类型") or "重要事项决策").strip())
        risk = re.sub(r"\s+", " ", str(item.get("risk") or item.get("riskLevel") or item.get("风险等级") or "中风险").strip())
        status = re.sub(r"\s+", " ", str(item.get("status") or item.get("状态") or item.get("处理状态") or "待确认").strip())
        source = re.sub(r"\s+", " ", str(item.get("source") or item.get("依据") or item.get("素材来源") or item.get("来源说明") or f"{imported_count} 条素材").strip())
        todos = item.get("meetingTodo") or item.get("开会待办") or item.get("todos") or item.get("materials") or item.get("待办") or item.get("待办事项") or item.get("材料缺口") or []
        if isinstance(todos, str):
            todo_text = todos
        elif isinstance(todos, list):
            todo_text = "、".join(str(todo).strip() for todo in todos if str(todo).strip())
        else:
            todo_text = ""
        changes = item.get("changes") or item.get("变更明细") or item.get("提炼依据") or item.get("依据") or []
        if isinstance(changes, str):
            changes = [changes]
        if not isinstance(changes, list):
            changes = []
        clean_changes = [re.sub(r"\s+", " ", str(change).strip()) for change in changes if str(change).strip()]
        if todo_text:
            clean_changes.append(f"生成待办：{todo_text}")
        clean_changes.append(source_label)
        is_normal_meeting = meeting_mode == "normal"
        if is_normal_meeting:
            project = project if project and project not in {"本地项目", "高新区二期厂房消防改造", "高新区二期厂房改造"} else "本次会议"
            issue_type = "普通会议议题"
            risk = "普通"
            status = "待确认"
            if not todo_text or any(word in todo_text for word in ["材料", "三重一大", "合规", "表决前", "金额", "资金测算"]):
                todo_text = "确认讨论范围，安排会议讨论"
        drafts.append({
            "id": f"issue-{index:03d}",
            "title": title[:80],
            "source": source[:80],
            "project": project[:60],
            "type": issue_type[:60],
            "risk": risk if is_normal_meeting else (risk if risk in ("高风险", "中风险", "低风险") else "中风险"),
            "status": status[:20],
            "todoText": todo_text[:100] if todo_text else "补材料、定责任人、安排会议审议",
            "changes": clean_changes[:5],
        })
    return drafts


async def _deepseek_extract_agenda_drafts(meeting: dict, imported: List[dict]) -> List[dict]:
    if not imported:
        return []
    compact_rows = []
    for index, item in enumerate(imported[:120], start=1):
        compact_rows.append({
            "序号": index,
            "提交人": item.get("name", ""),
            "时间": item.get("time", ""),
            "内容": item.get("content", ""),
            "来源": item.get("meta", ""),
        })
    meeting_mode = meeting.get("meetingMode") if meeting.get("meetingMode") in {"normal", "major"} else "normal"
    if meeting_mode == "major":
        role_instruction = '你是政府单位"三重一大"会前议题秘书。'
        mode_rules = """
6. 判断"三重一大"业务类型：重大事项决策、重要人事任免、重大项目安排、大额度资金运作，可组合。
7. 提炼本地项目名称；如果无法确认，写"待绑定本地项目"。项目名称只放在 project 字段，不要用项目名堆叠标题。
8. 给出风险等级：高风险/中风险/低风险。
9. 给出会议前必须完成的开会待办，如补材料、定责任人、安排上会、表决前补齐。每个字段要短，避免长篇解释。
"""
        output_type = "三重一大业务类型"
        output_risk = "高风险"
        output_todo = "会议前要完成的动作"
    else:
        role_instruction = "你是政府单位普通会议的会前议题秘书。"
        mode_rules = """
6. 这是普通会议，不是"三重一大"会议。禁止输出"三重一大"、合规审查、材料缺口、金额阈值、表决前补齐等内容。
7. type 固定输出"普通会议议题"，risk 固定输出"普通"，project 如果无法从素材确认就写"本次会议"。
8. meetingTodo 只写会议讨论动作，例如"确认问题现状与责任分工""讨论上线范围与验收口径""安排会议讨论"。
9. 输出要让秘书一眼知道下一步要拿到会上讨论什么，不要变成项目清单或材料清单。
"""
        output_type = "普通会议议题"
        output_risk = "普通"
        output_todo = "确认讨论范围，安排会议讨论"
    prompt = f"""
{role_instruction}请基于收集到的问题素材做真实提炼，不要编造素材外事实。

当前会议：
- 会议名称：{meeting.get("title") or "待创建 AI 会议"}
- 当前项目：{meeting.get("project") or "本地项目"}
- 当前议题：{meeting.get("agenda") or "待确认议题"}
- 会议性质：{"三重一大会议" if meeting_mode == "major" else "普通会议"}


收集素材 JSON：
{json.dumps(compact_rows, ensure_ascii=False)}

请完成：
1. 输出目标是"秘书下一步要创建的开会待办事项"，不是项目清单，也不是单纯识别多个项目。
2. 同一事项的多条消息要合并；不同问题必须拆开。不要把不相关问题合并成一个议题。最多输出与输入素材数量相同的待办，最少1个，上限6个。如果用户只提交了一条素材，就只输出一个议题，不要拆分。
3. 标题必须像待办事项，例如"讨论 AI 视频剪辑平台开发进度与测试安排"，不要写成"某某项目等3个项目问题"。
4. 多个项目即使问题类型相似，也不要合并成"等 N 个项目"；要按真正要开会处理的事项拆成独立待办。
5. 如果一句或一行中同时包含项目资金、合同采购、干部任免等不同事项，需要拆成独立待办。
{mode_rules}

只返回 JSON，不要 Markdown，不要解释。格式：
{{
  "agendaDrafts": [
    {{
      "title": "开会待办事项标题，不要出现"等N个项目"",
      "source": "多少条素材 + 关键来源",
      "project": "本地项目名称或待绑定本地项目",
      "type": "{output_type}",
      "risk": "{output_risk}",
      "status": "待确认",
      "meetingTodo": "{output_todo}",
      "durationMinutes": 15,
      "todos": ["待办动作1", "待办动作2"],
      "changes": ["首次识别依据", "聚类/拆分依据"]
    }}
  ]
}}
"""
    try:
        async with llm_semaphore:
            result = await llm._agenerate(
                [
                    SystemMessage(content="你只输出可解析 JSON，字段必须稳定，适合前端直接渲染。"),
                    HumanMessage(content=prompt),
                ],
                enable_thinking=False,
            )
        text = result.generations[0].message.content if result.generations else ""
        payload = _extract_json_object(text)
        if not payload:
            logger.warning("DeepSeek 议题提炼未返回可解析 JSON")
            return []
        drafts = _normalize_ai_agenda_drafts(
            payload,
            len(imported),
            meeting.get("project") or "本地项目",
            meeting_mode=meeting_mode,
            source_label="DeepSeek 根据收集素材真实提炼",
        )
        if not drafts:
            logger.warning("DeepSeek 议题提炼 JSON 字段未匹配，keys=%s", list(payload.keys()))
        return drafts
    except Exception as exc:
        logger.warning("DeepSeek 议题提炼失败：%s", exc)
        return []


def _append_meeting_activity(meeting_id: str, event: dict):
    safe_id = _safe_meeting_id(meeting_id)
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id)
        if not meeting:
            speaker = event.get("speaker") or {}
            title = event.get("meetingTitle") or "手机接入的 AI 会议"
            agenda = event.get("agenda") or "待确认议题"
            now = _now_text()
            meeting = {
                "id": safe_id,
                "title": title,
                "project": "本地项目",
                "projectCode": f"LOCAL-{datetime.now().strftime('%Y%m%d')}-001",
                "agenda": agenda,
                "date": _today_text(),
                "type": "普通企业会议",
                "creator": f"{speaker.get('dept') or '参会部门'} {speaker.get('displayName') or speaker.get('username') or '参会人'}",
                "createdAt": now,
                "updatedAt": now,
                "phase": "会中记录" if event.get("action") in ("join", "start", "stop") else "问题收集中",
                "issueSources": [],
                "agendaDrafts": _default_agenda_drafts("本地项目", agenda)[:1],
                "materials": [],
                "events": [],
                "archived": False,
                "projectBound": False,
                "agendaFrozen": False,
                "reviewDone": False,
                "archiveDone": False,
            }
        events = meeting.setdefault("events", [])
        events.append({**event, "serverTime": event.get("serverTime") or _now_text()})
        meeting["events"] = events[-200:]
        meeting["updatedAt"] = _now_text()
        meetings[safe_id] = meeting
        _save_meetings(meetings)




def _load_custom_rules() -> List[dict]:
    if CUSTOM_RULES_DB.exists():
        try:
            with open(CUSTOM_RULES_DB, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []
    return []


def _save_custom_rules(files: List[dict]):
    with open(CUSTOM_RULES_DB, "w", encoding="utf-8") as f:
        json.dump(files, f, ensure_ascii=False, indent=2)


def _extract_text_from_raw(filename: str, raw: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    text = ""
    if ext == "docx":
        import docx
        doc = docx.Document(io.BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
    elif ext == "pdf":
        import pdfplumber
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            pages = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
        text = "\n\n".join(pages)
    elif ext in ("txt", "md"):
        text = raw.decode("utf-8", errors="replace")
    else:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型：.{ext}")
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _compose_custom_rule_summary(raw_text: str, matter_type: str) -> dict:
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    focused = []
    for line in lines:
        if matter_type in line or any(word in line for word in ["必须", "禁止", "程序", "审议", "审批", "公示", "表决", "集体研究"]):
            focused.append(line)
        if len(focused) >= 8:
            break
    if not focused:
        focused = lines[:8]

    mandatory = [line for line in focused if "必须" in line or "应当" in line][:4]
    forbidden = [line for line in focused if "禁止" in line or "不得" in line][:4]
    procedures = []
    for line in focused:
        parts = re.split(r"[、，；。]", line)
        procedures.extend([part.strip() for part in parts if len(part.strip()) >= 2 and any(word in part for word in ["审查", "论证", "公示", "表决", "审批", "研究"])])
        if len(procedures) >= 6:
            break

    return {
        "summary_lines": focused,
        "mandatory": mandatory,
        "forbidden": forbidden,
        "procedures": procedures[:6],
    }


def _resolve_custom_rules_text(custom_rule_ids: Optional[List[str]], matter_type: str) -> str:
    if not custom_rule_ids:
        return ""
    files = _load_custom_rules()
    selected = [item for item in files if item["id"] in custom_rule_ids and (item.get("matterType") in (matter_type, "通用", None))]
    if not selected:
        return ""
    return "\n\n".join(
        f"【{item['name']}】\n{item.get('parsedText', '')[:4000]}"
        for item in selected
    )






def _default_departments() -> List[dict]:
    """默认部门列表（仅首次初始化时使用）。"""
    return [
        {"id": "dept_admin", "name": "信息管理中心"},
        {"id": "dept_general", "name": "总经理办公室"},
        {"id": "dept_finance", "name": "财务部"},
        {"id": "dept_legal", "name": "合规法务部"},
        {"id": "dept_project", "name": "项目管理部"},
        {"id": "dept_audit", "name": "审计监察部"},
        {"id": "dept_hr", "name": "人力资源部"},
    ]


















@app.put("/api/users/{user_id}")
async def update_user(request: Request, user_id: str, body: UserUpsertRequest):
    _require_admin(request)
    users = _load_users()
    for index, user in enumerate(users):
        if user["id"] != user_id:
            continue
        if any(item["username"] == body.username and item["id"] != user_id for item in users):
            raise HTTPException(status_code=409, detail="用户名已存在")
        users[index] = {
            **user,
            "username": body.username,
            "name": body.name,
            "role": body.role,
            "dept": body.dept,
            "status": body.status,
            "password": body.password or user["password"],
            "meetingRole": body.meetingRole if body.meetingRole is not None else user.get("meetingRole", "参会代表"),
            "meetingSeat": body.meetingSeat if body.meetingSeat is not None else user.get("meetingSeat", ""),
        }
        _save_users(users)
        return JSONResponse({"success": True, "user": _public_user(users[index])})
    raise HTTPException(status_code=404, detail="用户不存在")


@app.delete("/api/users/{user_id}")
async def delete_user(request: Request, user_id: str):
    current_user = _require_admin(request)
    if current_user["id"] == user_id:
        raise HTTPException(status_code=400, detail="不能删除当前登录管理员")
    users = _load_users()
    filtered = [user for user in users if user["id"] != user_id]
    if len(filtered) == len(users):
        raise HTTPException(status_code=404, detail="用户不存在")
    _save_users(filtered)
    return JSONResponse({"success": True})


@app.get("/api/meetings")
async def list_meetings(request: Request, include_archived: bool = False, limit: int = 50, offset: int = 0):
    _get_request_user(request, required=True)
    meetings = _load_meetings()
    rows = [_public_meeting(item, include_detail=False) for item in meetings.values()]
    if not include_archived:
        rows = [item for item in rows if not item.get("archived")]
    rows.sort(key=lambda item: item.get("updatedAt") or item.get("createdAt") or "", reverse=True)
    total = len(rows)
    page = rows[offset:offset + limit] if limit > 0 else rows
    return JSONResponse({"success": True, "meetings": page, "total": total, "limit": limit, "offset": offset})


@app.post("/api/meetings")
async def create_or_update_meeting(request: Request, body: MeetingUpsertRequest):
    user = _get_request_user(request, required=True)
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting_id = _safe_meeting_id(body.id)
        existing = meetings.get(meeting_id)
        explicit_fields = {**body.dict(exclude_unset=True), "id": meeting_id}
        meeting = _build_meeting_from_request(MeetingUpsertRequest(**{**body.dict(), "id": meeting_id}), user, existing, explicit_fields)
        meeting["agendaDrafts"] = _derive_agenda_drafts(meeting)
        meetings[meeting["id"]] = meeting
        _save_meetings(meetings)
    # 遗留待办自动带入：查历史未完成待办，写入新会议
    if not existing:  # 仅新建会议时带入
        try:
            with _db_connect() as conn:
                pending_todos = conn.execute(
                    """SELECT * FROM meeting_todos
                       WHERE status IN ('待处理', '进行中') AND meeting_id != ?
                       ORDER BY CASE priority WHEN '高' THEN 0 WHEN '中' THEN 1 ELSE 2 END, created_at DESC
                       LIMIT 20""",
                    (meeting["id"],)
                ).fetchall()
            if pending_todos:
                carryover = []
                for r in pending_todos:
                    carryover.append({
                        "task": r["task"], "owner": r["owner"], "deadline": r["deadline"],
                        "priority": r["priority"], "status": "待处理", "source": "carryover",
                        "reference": f"来自：{r['meeting_title'] or '历史会议'}",
                    })
                _sync_todos_to_table(meeting["id"], meeting.get("title", ""), carryover, source="carryover")
        except Exception as e:
            logger.warning("遗留待办带入失败: %s", e)
    return JSONResponse({"success": True, "meeting": _public_meeting(meeting, include_detail=True)})


@app.get("/api/meetings/{meeting_id}/carryover-todos")
async def get_carryover_todos(request: Request, meeting_id: str):
    """获取该会议的遗留待办列表。"""
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    try:
        with _db_connect() as conn:
            rows = conn.execute(
                "SELECT * FROM meeting_todos WHERE meeting_id = ? AND source = 'carryover' ORDER BY created_at DESC",
                (safe_id,)
            ).fetchall()
        todos = [{"id": r["id"], "task": r["task"], "owner": r["owner"], "deadline": r["deadline"],
                  "priority": r["priority"], "status": r["status"], "reference": r["reference"]}
                 for r in rows]
        return JSONResponse({"success": True, "todos": todos, "total": len(todos)})
    except Exception as e:
        return JSONResponse({"success": True, "todos": [], "total": 0})


@app.get("/api/meetings/{meeting_id}")
async def get_meeting(request: Request, meeting_id: str):
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    meetings = _load_meetings()
    meeting = meetings.get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    return JSONResponse({"success": True, "meeting": _public_meeting(meeting, include_detail=True)})


@app.patch("/api/meetings/{meeting_id}")
async def patch_meeting(request: Request, meeting_id: str, body: MeetingPatchRequest):
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        if safe_id not in meetings:
            raise HTTPException(status_code=404, detail="会议不存在")
        meeting = meetings[safe_id]
        _check_meeting_access(user, meeting)
        patch = body.dict(exclude_unset=True)
        # ── 议题冻结校验：冻结后禁止修改议题相关字段 ──
        if meeting.get("agendaFrozen"):
            _agenda_fields = {"agendaDrafts", "agenda", "issueSources", "agendaTitle", "agendaFrozen"}
            _blocked = [k for k in patch if k in _agenda_fields]
            if _blocked:
                # admin 可以强制修改（特殊权限）
                if user.get("role") != "admin":
                    raise HTTPException(
                        status_code=403,
                        detail=f"议题已冻结，无法修改：{', '.join(_blocked)}。如需修改请联系管理员。"
                    )
        if "meeting_mode" in patch and patch.get("meeting_mode"):
            patch["meetingMode"] = patch.pop("meeting_mode")
        if "project_code" in patch and patch.get("project_code"):
            patch["projectCode"] = patch.pop("project_code")
        if "projectCode" in patch and patch["projectCode"] is None:
            patch.pop("projectCode", None)
        for key, value in patch.items():
            if value is not None:
                meeting[key] = value
        meeting["updatedAt"] = _now_text()
        if "issueSources" in patch and "agendaDrafts" not in patch:
            meeting["agendaDrafts"] = _derive_agenda_drafts(meeting)
        meetings[safe_id] = meeting
        _save_meetings(meetings)
    return JSONResponse({"success": True, "meeting": _public_meeting(meeting, include_detail=True)})


# ═══ 正式议题（meeting_agendas）—— 议题为最小业务单元 ═══

def _can_manage_agenda(user: dict, meeting: dict) -> bool:
    """agenda:activate / agenda:create_temporary —— admin、主持人、会议秘书。"""
    if not user:
        return False
    if user.get("role") == "admin":
        return True
    creator = meeting.get("creator") or ""
    if creator and (user.get("name") in creator or creator in (user.get("name") or "")):
        return True
    mr = (user.get("meetingRole") or user.get("role") or "").strip()
    return mr in {"主持人", "会议秘书", "秘书", "host", "secretary"}


@app.get("/api/meetings/{meeting_id}/agendas")
async def list_agendas(request: Request, meeting_id: str):
    """列出会议全部正式议题（含从 agendaDrafts 的兼容物化）。

    保密议题在 API 层过滤（§57）：无权限用户只拿到脱敏占位，内容不下发。
    """
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    meeting = _load_meetings().get(safe_id) or {}
    _check_meeting_access(user, meeting)
    agendas = list_meeting_agendas(safe_id)
    agendas = filter_agendas_for_user(user, meeting, agendas)
    active = get_meeting_active_agenda(safe_id)
    if active and not can_view_agenda(user, meeting, active):
        active = None
    return JSONResponse({
        "success": True,
        "agendas": agendas,
        "activeAgendaId": (active or {}).get("id", ""),
        "activeAgenda": active,
    })


@app.get("/api/meetings/{meeting_id}/agendas/{agenda_id}")
async def get_agenda_detail(request: Request, meeting_id: str, agenda_id: str):
    """获取单个议题详情（保密议题无权限时返回脱敏占位）。"""
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    meeting = _load_meetings().get(safe_id) or {}
    _check_meeting_access(user, meeting)
    agenda = get_meeting_agenda(safe_id, agenda_id)
    if not agenda:
        raise HTTPException(status_code=404, detail="议题不存在")
    if not can_view_agenda(user, meeting, agenda):
        agenda["title"] = "（保密议题）"
        agenda["description"] = ""
        agenda["restricted"] = True
    return JSONResponse({"success": True, "agenda": agenda})


@app.post("/api/meetings/{meeting_id}/agendas")
async def create_agenda(request: Request, meeting_id: str, body: AgendaCreateRequest):
    """新增议题；会中临时议题 agendaType=temporary 立即持久化。"""
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    meeting = _load_meetings().get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    _check_meeting_access(user, meeting)
    if body.agendaType == "temporary" and not _can_manage_agenda(user, meeting):
        raise HTTPException(status_code=403, detail="仅主持人或会议秘书可以创建临时议题")
    try:
        agenda = create_meeting_agenda(
            safe_id, body.title, body.description,
            agenda_type=body.agendaType or "standard",
            source=body.source or ("in_meeting" if body.agendaType == "temporary" else "prepared"),
            confidentiality_level=body.confidentialityLevel or "normal",
            permission_level=body.permissionLevel,
            proposer_user_id=body.proposerUserId,
            owner_user_id=body.ownerUserId,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    _append_meeting_activity_light(safe_id, {"type": "agenda.created", "payload": {"agendaId": agenda["id"], "title": agenda["title"]}})
    return JSONResponse({"success": True, "agenda": agenda})


@app.patch("/api/meetings/{meeting_id}/agendas/{agenda_id}")
async def patch_agenda(request: Request, meeting_id: str, agenda_id: str, body: AgendaPatchRequest):
    """字段级更新议题（主持人/秘书/管理员）。"""
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    meeting = _load_meetings().get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    _check_meeting_access(user, meeting)
    if not _can_manage_agenda(user, meeting):
        raise HTTPException(status_code=403, detail="仅主持人或会议秘书可以修改议题")
    try:
        agenda = update_meeting_agenda(safe_id, agenda_id, body.dict(exclude_unset=True))
    except KeyError:
        raise HTTPException(status_code=404, detail="议题不存在")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return JSONResponse({"success": True, "agenda": agenda})


@app.delete("/api/meetings/{meeting_id}/agendas/{agenda_id}")
async def remove_agenda(request: Request, meeting_id: str, agenda_id: str):
    """删除议题（主持人/秘书/管理员）。"""
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    meeting = _load_meetings().get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    _check_meeting_access(user, meeting)
    if not _can_manage_agenda(user, meeting):
        raise HTTPException(status_code=403, detail="仅主持人或会议秘书可以删除议题")
    delete_meeting_agenda(safe_id, agenda_id)
    _append_meeting_activity_light(safe_id, {"type": "agenda.deleted", "payload": {"agendaId": agenda_id}})
    return JSONResponse({"success": True})


@app.post("/api/meetings/{meeting_id}/agendas/{agenda_id}/activate")
async def activate_agenda(request: Request, meeting_id: str, agenda_id: str):
    """切换当前议题：结束旧议题 → 新议题 started_at → active_agenda_id 持久化。"""
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    meeting = _load_meetings().get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    _check_meeting_access(user, meeting)
    if not _can_manage_agenda(user, meeting):
        raise HTTPException(status_code=403, detail="仅主持人或会议秘书可以切换议题")
    try:
        agenda = activate_meeting_agenda(safe_id, agenda_id)
    except KeyError:
        raise HTTPException(status_code=404, detail="议题不存在")
    _append_meeting_activity_light(safe_id, {"type": "agenda.activated", "payload": {"agendaId": agenda_id, "previous": agenda.get("previousAgendaId")}})
    return JSONResponse({"success": True, "agenda": agenda})


# ═══ 议题级会议记录与决议（§37-41：记录=讨论过程，决议=最终结果）═══

@app.get("/api/meetings/{meeting_id}/agendas/{agenda_id}/records")
async def list_agenda_records_route(request: Request, meeting_id: str, agenda_id: str):
    """列出议题讨论记录；agenda_id=all 时返回整场。"""
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    records = list_agenda_records(safe_id, "" if agenda_id == "all" else agenda_id)
    return JSONResponse({"success": True, "records": records, "total": len(records)})


@app.post("/api/meetings/{meeting_id}/agendas/{agenda_id}/records")
async def create_agenda_record_route(request: Request, meeting_id: str, agenda_id: str, body: AgendaRecordRequest):
    """新增议题讨论记录。"""
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    try:
        records = create_agenda_record(
            safe_id, agenda_id, body.content,
            speaker_name=body.speakerName or user.get("name") or user.get("username") or "",
            speaker_user_id=user.get("id") or "",
            record_type=body.recordType,
            transcript_id=body.transcriptId,
            source=body.source,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return JSONResponse({"success": True, "records": records})


@app.get("/api/meetings/{meeting_id}/agendas/{agenda_id}/decisions")
async def list_agenda_decisions_route(request: Request, meeting_id: str, agenda_id: str):
    """列出议题决议；agenda_id=all 时返回整场。"""
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    decisions = list_agenda_decisions(safe_id, "" if agenda_id == "all" else agenda_id)
    return JSONResponse({"success": True, "decisions": decisions, "total": len(decisions)})


@app.post("/api/meetings/{meeting_id}/agendas/{agenda_id}/decisions/generate")
async def generate_agenda_decisions_route(request: Request, meeting_id: str, agenda_id: str):
    """按议题生成决议草稿（逐议题，不整场猜测归属）。"""
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    result = generate_decisions_for_agenda(safe_id, agenda_id, created_by=user.get("name") or user.get("username") or "")
    return JSONResponse({"success": True, **result})


@app.post("/api/meetings/{meeting_id}/agendas/{agenda_id}/decisions")
async def create_agenda_decision_route(request: Request, meeting_id: str, agenda_id: str, body: AgendaDecisionRequest):
    """新增议题决议（正式绑定 agenda_id）。"""
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    try:
        decisions = create_agenda_decision(
            safe_id, agenda_id, body.title, body.content,
            created_by=user.get("name") or user.get("username") or "",
            source=body.source, status=body.status,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return JSONResponse({"success": True, "decisions": decisions})


@app.patch("/api/meetings/{meeting_id}/agendas/{agenda_id}/decisions/{decision_id}")
async def patch_agenda_decision_route(request: Request, meeting_id: str, agenda_id: str, decision_id: str, body: AgendaDecisionPatchRequest):
    """更新决议；title/content 变更自动递增 version。"""
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    try:
        decision = update_agenda_decision(safe_id, agenda_id, decision_id, body.dict(exclude_unset=True))
    except KeyError:
        raise HTTPException(status_code=404, detail="决议不存在")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return JSONResponse({"success": True, "decision": decision})


@app.delete("/api/meetings/{meeting_id}/agendas/{agenda_id}/decisions/{decision_id}")
async def remove_agenda_decision_route(request: Request, meeting_id: str, agenda_id: str, decision_id: str):
    """删除决议。"""
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    delete_agenda_decision(safe_id, agenda_id, decision_id)
    return JSONResponse({"success": True})


@app.post("/api/meetings/{meeting_id}/issues")
async def append_meeting_issue(request: Request, meeting_id: str, body: MeetingIssueRequest):
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    now = _now_text()
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id)
        if not meeting:
            meeting = _build_meeting_from_request(
                MeetingUpsertRequest(id=safe_id, title="待创建 AI 会议", project="本地项目", agenda="待确认议题", phase="问题收集中"),
                user,
                None,
            )
        issue_sources = meeting.setdefault("issueSources", [])
        issue = {
            "id": f"issue_source_{uuid.uuid4().hex[:10]}",
            "name": body.name or _creator_from_user(user),
            "time": datetime.now().strftime("%H:%M"),
            "type": body.type or "text",
            "content": re.sub(r"\s+", " ", body.content or "").strip(),
            "meta": body.meta,
            "source": body.source or "manual",
            "serverTime": now,
            "userId": user.get("id"),
        }
        if not issue["content"]:
            raise HTTPException(status_code=400, detail="问题内容不能为空")
        issue_sources.append(issue)
        meeting["issueSources"] = issue_sources[-300:]
        meeting["agendaDrafts"] = _derive_agenda_drafts(meeting)
        meeting["phase"] = meeting.get("phase") or "问题收集中"
        meeting["updatedAt"] = now
        meetings[safe_id] = meeting
        _save_meetings(meetings)
    return JSONResponse({"success": True, "issue": issue, "meeting": _public_meeting(meeting, include_detail=True)})


@app.get("/api/meetings/issues/template")
async def download_meeting_issue_template(request: Request):
    _get_request_user(request, required=True)
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    from openpyxl.worksheet.datavalidation import DataValidation

    wb = Workbook()
    ws = wb.active
    ws.title = "问题台账"
    headers = ["问题描述", "来源部门", "提交人", "发生时间", "关联项目", "涉及金额(万元)", "材料缺口", "备注"]
    examples = [
        ["消防改造现场发现新增隐患，预计需要追加预算。", "项目管理部", "王明", "2026-06-09 09:30", "高新区二期厂房消防改造", 860, "资金来源测算表；可研修订说明", "建议纳入董事会或总经理办公会审议"],
        ["合同变更签证资料不完整，需要法务复核。", "法务合规部", "周法务", "2026-06-09 10:10", "高新区二期厂房消防改造", "", "合同变更草案；法务审查意见", "表决前补齐"],
    ]
    ws.append(headers)
    for row in examples:
        ws.append(row)

    header_fill = PatternFill("solid", fgColor="1D5FD7")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(style="thin", color="D9E2EF")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)
    for row in ws.iter_rows(min_row=2, max_row=80, min_col=1, max_col=len(headers)):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)
    widths = [34, 14, 12, 18, 24, 14, 30, 28]
    for index, width in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + index)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:H80"
    dv = DataValidation(type="decimal", operator="greaterThanOrEqual", formula1="0", allow_blank=True)
    ws.add_data_validation(dv)
    dv.add("F2:F80")

    guide = wb.create_sheet("填写说明")
    guide_rows = [
        ["字段", "填写说明"],
        ["问题描述", "必填。把浙政钉、企业微信、现场发现的问题写成一句完整描述。"],
        ["来源部门", "建议填写业务部门，如项目管理部、财务管理部、法务合规部。"],
        ["提交人", "建议填写真实姓名，便于后续催办。"],
        ["发生时间", "建议格式：2026-06-09 09:30。"],
        ["关联项目", "没有外部项目库时，填写本系统内的本地项目名称。"],
        ["涉及金额(万元)", "可选。涉及预算、采购、合同变更时填写数字。"],
        ["材料缺口", "可选。多个材料用分号隔开，如：资金来源测算表；法务审查意见。"],
        ["备注", "可选。填写上会建议、紧急程度或其他说明。"],
    ]
    for row in guide_rows:
        guide.append(row)
    guide.column_dimensions["A"].width = 18
    guide.column_dimensions["B"].width = 70
    for cell in guide[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    filename = "AI会议问题收集模板.xlsx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )


@app.post("/api/meetings/{meeting_id}/issues/import-excel")
async def import_meeting_issues_excel(request: Request, meeting_id: str, file: UploadFile = File(...)):
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    filename = file.filename or "问题台账.xlsx"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ("xlsx", "csv"):
        raise HTTPException(status_code=400, detail="仅支持 .xlsx 或 .csv 问题台账")
    raw = await _read_upload_safe(file, MAX_EXCEL_BYTES)
    if not raw:
        raise HTTPException(status_code=400, detail="上传文件不能为空")

    rows = []
    if ext == "xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        ws = wb["问题台账"] if "问题台账" in wb.sheetnames else wb.active
        header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), None)
        if not header_row:
            raise HTTPException(status_code=400, detail="Excel 模板缺少表头")
        headers = [str(value or "").strip() for value in header_row]
        for values in ws.iter_rows(min_row=2, max_row=301, values_only=True):
            row = {headers[index]: values[index] if index < len(values) else "" for index in range(len(headers))}
            rows.append(row)
    else:
        import csv
        text = raw.decode("utf-8-sig", errors="replace")
        reader = csv.DictReader(io.StringIO(text))
        rows.extend(reader)

    now = _now_text()
    imported = []
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id)
        if not meeting:
            meeting = _build_meeting_from_request(
                MeetingUpsertRequest(
                    id=safe_id,
                    title="待创建 AI 会议",
                    project="本地项目",
                    agenda="待确认议题",
                    phase="问题收集中",
                ),
                user,
                None,
            )
        issue_sources = meeting.setdefault("issueSources", [])
        for row in rows:
            desc = re.sub(r"\s+", " ", str(row.get("问题描述") or row.get("问题") or row.get("描述") or "").strip())
            if not desc:
                continue
            dept = str(row.get("来源部门") or "").strip()
            submitter = str(row.get("提交人") or "").strip()
            occurred_at = str(row.get("发生时间") or "").strip()
            project = str(row.get("关联项目") or "").strip()
            amount = str(row.get("涉及金额(万元)") or "").strip()
            gap = str(row.get("材料缺口") or "").strip()
            note = str(row.get("备注") or "").strip()
            parts = [desc]
            if project:
                parts.append(f"关联项目：{project}")
            if amount:
                parts.append(f"涉及金额：{amount} 万元")
            if gap:
                parts.append(f"材料缺口：{gap}")
            if note:
                parts.append(f"备注：{note}")
            issue = {
                "id": f"issue_source_{uuid.uuid4().hex[:10]}",
                "name": f"{dept} {submitter}".strip() or _creator_from_user(user),
                "time": occurred_at[:16] if occurred_at else datetime.now().strftime("%H:%M"),
                "type": "text",
                "content": "；".join(parts),
                "meta": f"Excel：{filename}",
                "source": "excel",
                "serverTime": now,
                "userId": user.get("id"),
            }
            issue_sources.append(issue)
            imported.append(issue)
        if not imported:
            raise HTTPException(status_code=400, detail='未读取到有效问题，请检查"问题描述"列')
        meeting["issueSources"] = issue_sources[-300:]
        meeting["agendaDrafts"] = _derive_agenda_drafts(meeting)
        if meeting["agendaDrafts"]:
            first_draft = meeting["agendaDrafts"][0]
            placeholder_agendas = {"", "待确认议题", "高新区二期厂房改造追加预算审议"}
            if str(meeting.get("agenda") or "").strip() in placeholder_agendas:
                meeting["agenda"] = first_draft.get("title") or meeting.get("agenda") or "待确认议题"
            placeholder_projects = {"", "本地项目", "高新区二期厂房消防改造", "高新区二期厂房改造"}
            if str(meeting.get("project") or "").strip() in placeholder_projects and first_draft.get("project"):
                meeting["project"] = first_draft.get("project")
            if str(meeting.get("title") or "").strip() in {"", "待创建 AI 会议", "高新区二期厂房消防改造专题会"}:
                meeting["title"] = f"{meeting.get('project') or '本地项目'}专题会"
        meeting["updatedAt"] = now
        meetings[safe_id] = meeting
        _save_meetings(meetings)

    ai_drafts = await _deepseek_extract_agenda_drafts(meeting, imported)
    ai_provider = "local-rule"
    if ai_drafts:
        ai_provider = "deepseek"
        with MEETINGS_LOCK:
            meetings = _load_meetings()
            meeting = meetings.get(safe_id, meeting)
            meeting["agendaDrafts"] = ai_drafts
            first_draft = ai_drafts[0]
            if first_draft.get("title"):
                meeting["agenda"] = first_draft["title"]
            if first_draft.get("project") and first_draft.get("project") != "待绑定本地项目":
                meeting["project"] = first_draft["project"]
            if str(meeting.get("title") or "").strip() in {"", "待创建 AI 会议", "高新区二期厂房消防改造专题会"}:
                meeting["title"] = f"{meeting.get('project') or '本地项目'}专题会"
            meeting["updatedAt"] = _now_text()
            meetings[safe_id] = meeting
            _save_meetings(meetings)

    return JSONResponse({
        "success": True,
        "importedCount": len(imported),
        "aiProvider": ai_provider,
        "issues": imported,
        "meeting": _public_meeting(meeting, include_detail=True),
    })


@app.post("/api/meetings/{meeting_id}/agenda/generate")
async def generate_meeting_agenda(request: Request, meeting_id: str):
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        if safe_id not in meetings:
            raise HTTPException(status_code=404, detail="会议不存在")
        meeting = meetings[safe_id]
        issue_sources = meeting.get("issueSources") if isinstance(meeting.get("issueSources"), list) else []
    if not issue_sources:
        raise HTTPException(status_code=400, detail="请先收集至少一条问题或素材")

    ai_drafts = await _deepseek_extract_agenda_drafts(meeting, issue_sources)
    ai_provider = "deepseek" if ai_drafts else "local-rule"
    drafts = ai_drafts or _derive_agenda_drafts(meeting)
    if not drafts:
        raise HTTPException(status_code=400, detail="当前素材不足，无法生成开会待办")

    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id, meeting)
        meeting["agendaDrafts"] = drafts
        first_draft = drafts[0]
        next_agenda = "；".join(str(item.get("title") or "").strip() for item in drafts if item.get("title"))
        if next_agenda:
            meeting["agenda"] = next_agenda[:180]
        if meeting.get("meetingMode") == "normal":
            meeting["project"] = meeting.get("project") if meeting.get("project") not in {"", "本地项目", "高新区二期厂房消防改造", "高新区二期厂房改造"} else "本次会议"
        elif first_draft.get("project") and first_draft.get("project") != "待绑定本地项目":
            meeting["project"] = first_draft["project"]
        if str(meeting.get("title") or "").strip() in {"", "待创建 AI 会议", "AI 会议问题收集", "高新区二期厂房消防改造专题会"}:
            if meeting.get("meetingMode") == "normal":
                meeting["title"] = "AI 会议问题讨论会"
            else:
                meeting["title"] = f"{meeting.get('project') or '本地项目'}专题会"
        meeting["phase"] = meeting.get("phase") or "问题收集中"
        meeting["updatedAt"] = _now_text()
        meetings[safe_id] = meeting
        _save_meetings(meetings)

    return JSONResponse({
        "success": True,
        "aiProvider": ai_provider,
        "sourceCount": len(issue_sources),
        "agendaDrafts": drafts,
        "meeting": _public_meeting(meeting, include_detail=True),
    })


@app.post("/api/meetings/{meeting_id}/agenda/realtime-check")
async def realtime_check_meeting_agenda(request: Request, meeting_id: str, body: MeetingAgendaRealtimeCheckRequest):
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    agenda_drafts = body.agendaDrafts[:8] if isinstance(body.agendaDrafts, list) else []
    latest_transcripts = body.latestTranscripts[-12:] if isinstance(body.latestTranscripts, list) else []
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id, {})
        if not agenda_drafts:
            agenda_drafts = meeting.get("agendaDrafts") if isinstance(meeting.get("agendaDrafts"), list) else []
        if not agenda_drafts and meeting.get("agenda"):
            agenda_drafts = [{"id": "agenda-current", "title": meeting.get("agenda")}]
    if not agenda_drafts:
        raise HTTPException(status_code=400, detail="缺少会议议题，无法实时比对")

    meeting_mode = body.meetingMode if body.meetingMode in {"normal", "major"} else (meeting.get("meetingMode") if meeting.get("meetingMode") in {"normal", "major"} else "normal")
    ai_results = await _deepseek_realtime_agenda_check(meeting, agenda_drafts, latest_transcripts, meeting_mode)
    provider = "deepseek" if ai_results else "local-rule"
    results = ai_results or _local_realtime_agenda_check(agenda_drafts, latest_transcripts)
    return JSONResponse({
        "success": True,
        "aiProvider": provider,
        "meetingId": safe_id,
        "results": results,
    })


@app.post("/api/meetings/{meeting_id}/stage")
async def update_meeting_stage(request: Request, meeting_id: str, body: MeetingStageRequest):
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    phase_by_stage = {
        "collect": "会前确认",
        "meeting": "会中记录",
        "audit": "会后终审",
        "archive": "已归档",
    }
    stage = body.stage if body.stage in phase_by_stage else "collect"
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        if safe_id not in meetings:
            raise HTTPException(status_code=404, detail="会议不存在")
        meeting = meetings[safe_id]
        _check_meeting_access(user, meeting)
        meeting["phase"] = body.phase or phase_by_stage[stage]
        if stage in ("meeting", "audit", "archive"):
            meeting["projectBound"] = True
            meeting["agendaFrozen"] = True
        if stage in ("audit",):
            # 会议结束进入终审时自动触发 Whisper（用户审核时就能看到高精度纪实）
            has_whisper = any(
                e.get("type") == "transcript" and e.get("action") == "whisper-review"
                for e in meeting.get("events", [])
            )
            if not has_whisper:
                audio_dir = MEETING_FILES_DIR / "recordings" / safe_id
                if audio_dir.exists() and any(audio_dir.iterdir()):
                    _safe_create_task(_run_whisper_review_for_meeting(safe_id), name=f"whisper-review-{safe_id}")
                    logger.info("Whisper终审: meeting=%s, audit 阶段自动触发", safe_id)
        if stage in ("archive",):
            # ── 签字完成率检查 ──
            _t_data = _load_meeting_transcripts().get(safe_id, {"transcripts": []})
            _t_list = _t_data.get("transcripts", [])
            _speakers = set()
            _signed_speakers = set()
            for _t in _t_list:
                _name = _t.get("speakerName") or _t.get("username") or ""
                if _name:
                    _speakers.add(_name)
                    if _t.get("correctionSigned"):
                        _signed_speakers.add(_name)
            _total = len(_speakers)
            _signed = len(_signed_speakers)
            meeting["signedCount"] = _signed
            meeting["participantsCount"] = _total
            if _total > 0 and _signed < _total:
                _unsigned = _speakers - _signed_speakers
                logger.warning("归档签字拦截: meeting=%s, %d/%d 人已签, 未签: %s",
                               safe_id, _signed, _total, ", ".join(_unsigned))
                # 允许归档但标记警告，不硬性阻断（改为软拦截 + 警告信息）
                meeting["signWarning"] = f"{_signed}/{_total} 人已签字确认，未签字：{'、'.join(_unsigned)}"
            else:
                meeting.pop("signWarning", None)
            meeting["reviewDone"] = True
            meeting["archiveDone"] = True
        event = {
            "id": f"stage_{uuid.uuid4().hex[:10]}",
            "type": "stage",
            "stage": stage,
            "phase": meeting["phase"],
            "serverTime": _now_text(),
        }
        meeting.setdefault("events", []).append(event)
        meeting["events"] = meeting["events"][-200:]
        meeting["updatedAt"] = event["serverTime"]
        meetings[safe_id] = meeting
        _save_meetings(meetings)
    return JSONResponse({"success": True, "meeting": _public_meeting(meeting, include_detail=True)})


@app.delete("/api/meetings/{meeting_id}")
async def archive_meeting(request: Request, meeting_id: str):
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        if safe_id not in meetings:
            raise HTTPException(status_code=404, detail="会议不存在")
        meeting = meetings[safe_id]
        _check_meeting_access(user, meeting)
        # ── 归档闭环（§53）：要求全员签字时，未签齐禁止正式归档 ──
        if meeting.get("requireFullSignature"):
            from backend.services.signature_service import is_fully_signed, signed_signer_count, required_signer_count
            if not is_fully_signed(safe_id):
                raise HTTPException(
                    status_code=409,
                    detail=f"尚未全员签字（已签 {signed_signer_count(safe_id)} / 应签 {required_signer_count(safe_id)}），无法正式归档",
                )
        meeting["archived"] = True
        meeting["updatedAt"] = _now_text()
        meetings[safe_id] = meeting
        _save_meetings(meetings)
    return JSONResponse({"success": True, "meeting": _public_meeting(meeting, include_detail=False)})


# ═══ 会议成果签字（§50-54：绑定版本与 content_hash，未签齐禁止归档）═══

@app.get("/api/meetings/{meeting_id}/signatures")
async def list_meeting_signatures_route(request: Request, meeting_id: str):
    """列出会议签字记录；?agenda_id=&target_type=&target_id= 可选过滤。"""
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    params = request.query_params
    signatures = list_signatures(
        safe_id,
        agenda_id=params.get("agenda_id", ""),
        target_type=params.get("target_type", ""),
        target_id=params.get("target_id", ""),
    )
    return JSONResponse({
        "success": True,
        "signatures": signatures,
        "signedCount": signed_signer_count(safe_id),
        "requiredCount": required_signer_count(safe_id),
        "fullySigned": is_fully_signed(safe_id),
    })


@app.post("/api/meetings/{meeting_id}/signatures")
async def sign_meeting_target_route(request: Request, meeting_id: str, body: MeetingSignatureRequest):
    """签署会议成果：校验内容哈希与版本，签名绑定 version + content_hash。"""
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    try:
        signature = sign_target(
            safe_id, body.agendaId, body.targetType, body.targetId,
            body.version, body.content,
            signer_user_id=user.get("id") or "",
            signer_name=body.signerName or user.get("name") or user.get("username") or "",
            signer_role=body.signerRole or user.get("meetingRole") or "",
            signature_data=body.signatureData,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return JSONResponse({"success": True, "signature": signature})


@app.post("/api/meetings/{meeting_id}/signatures/hash")
async def signature_hash_route(request: Request, meeting_id: str, body: MeetingSignatureRequest):
    """返回目标内容在指定版本下的期望 content_hash（前端签字前比对）。"""
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    return JSONResponse({
        "success": True,
        "contentHash": compute_content_hash(safe_id, body.agendaId, body.targetId, body.version, body.content),
    })


# ═══ 权限模型（§25-26 多角色；§57-59 保密议题 ACL）═══

@app.get("/api/users/{user_id}/roles")
async def list_user_roles_route(request: Request, user_id: str):
    """用户全局角色列表。"""
    _get_request_user(request, required=True)
    return JSONResponse({"success": True, "roles": get_user_roles(user_id)})


@app.post("/api/users/{user_id}/roles")
async def add_user_role_route(request: Request, user_id: str):
    """添加用户全局角色（admin 专属）。"""
    _require_admin(request)
    body = await request.json()
    role = str(body.get("role", "")).strip()
    try:
        roles = add_user_role(user_id, role, granted_by=request.state.user.get("username") if hasattr(request.state, "user") else "")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return JSONResponse({"success": True, "roles": roles})


@app.delete("/api/users/{user_id}/roles/{role}")
async def remove_user_role_route(request: Request, user_id: str, role: str):
    """移除用户全局角色（admin 专属）。"""
    _require_admin(request)
    remove_user_role(user_id, role)
    return JSONResponse({"success": True, "roles": get_user_roles(user_id)})


@app.get("/api/agendas/{agenda_id}/acl")
async def list_agenda_acl_route(request: Request, agenda_id: str):
    """议题访问控制列表。"""
    _get_request_user(request, required=True)
    return JSONResponse({"success": True, "acl": list_agenda_acl(agenda_id)})


@app.post("/api/agendas/{agenda_id}/acl")
async def grant_agenda_acl_route(request: Request, agenda_id: str):
    """授予议题权限（view/edit/sign/admin），admin/主持人/秘书专属。"""
    user = _get_request_user(request, required=True)
    body = await request.json()
    meeting_id = str(body.get("meetingId", ""))
    target_user_id = str(body.get("userId", ""))
    permission = str(body.get("permission", "view"))
    if not target_user_id:
        raise HTTPException(status_code=400, detail="缺少 userId")
    meeting = _load_meetings().get(_safe_meeting_id(meeting_id)) or {}
    if not _can_manage_agenda(user, meeting):
        raise HTTPException(status_code=403, detail="仅管理员、主持人或会议秘书可以授予议题权限")
    try:
        acl = grant_agenda_acl(agenda_id, _safe_meeting_id(meeting_id), target_user_id, permission, granted_by=user.get("name") or user.get("username") or "")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return JSONResponse({"success": True, "acl": acl})


@app.delete("/api/agendas/{agenda_id}/acl/{user_id}/{permission}")
async def revoke_agenda_acl_route(request: Request, agenda_id: str, user_id: str, permission: str):
    """撤销议题权限。"""
    user = _get_request_user(request, required=True)
    meeting = _load_meetings().get("") or {}
    # 权限校验：通过查询该议题所属会议判断治理者身份
    from backend.services.agenda_service import get_meeting_agenda
    target = None
    for mid, m in _load_meetings().items():
        ag = get_meeting_agenda(mid, agenda_id)
        if ag:
            target = m
            break
    if not _can_manage_agenda(user, target or {}):
        raise HTTPException(status_code=403, detail="仅管理员、主持人或会议秘书可以撤销议题权限")
    revoke_agenda_acl(agenda_id, user_id, permission)
    return JSONResponse({"success": True})


@app.get("/api/meetings/{meeting_id}/materials")
async def list_meeting_materials(request: Request, meeting_id: str):
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    meetings = _load_meetings()
    meeting = meetings.get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    return JSONResponse({"success": True, "materials": _normalize_meeting(meeting).get("materials", [])})


@app.post("/api/meetings/{meeting_id}/materials/upload")
async def upload_meeting_material(
    request: Request,
    meeting_id: str,
    file: UploadFile = File(...),
    material_name: str = Form("支撑材料"),
):
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    original_name = _safe_storage_filename(file.filename or "meeting-material")
    raw = await _read_upload_safe(file, MAX_UPLOAD_BYTES)
    if not raw:
        raise HTTPException(status_code=400, detail="上传文件不能为空")
    if len(raw) > 80 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="单个材料不能超过 80MB")

    material_id = f"mat_{uuid.uuid4().hex[:12]}"
    storage_dir = MEETING_FILES_DIR / safe_id
    storage_dir.mkdir(parents=True, exist_ok=True)
    stored_name = f"{material_id}_{original_name}"
    storage_path = storage_dir / stored_name
    storage_path.write_bytes(raw)

    now = _now_text()
    record = {
        "id": material_id,
        "name": material_name or original_name,
        "status": "已上传",
        "tone": "green",
        "fileName": original_name,
        "storedName": stored_name,
        "size": len(raw),
        "uploadedAt": now,
        "uploader": _creator_from_user(user),
        "downloadUrl": f"/api/meetings/{safe_id}/materials/{material_id}/download",
    }

    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id)
        if not meeting:
            storage_path.unlink(missing_ok=True)
            raise HTTPException(status_code=404, detail="会议不存在")
        materials = meeting.setdefault("materials", [])
        materials = [item for item in materials if item.get("name") != record["name"]]
        materials.append(record)
        meeting["materials"] = materials
        meeting.setdefault("events", []).append({
            "id": f"material_{uuid.uuid4().hex[:10]}",
            "type": "material",
            "serverTime": now,
            "materialId": material_id,
            "materialName": record["name"],
            "fileName": original_name,
            "uploader": record["uploader"],
        })
        meeting["updatedAt"] = now
        meetings[safe_id] = meeting
        _save_meetings(meetings)

    return JSONResponse({"success": True, "material": record, "meeting": _public_meeting(meeting, include_detail=True)})


@app.get("/api/meetings/{meeting_id}/materials/{material_id}/download")
async def download_meeting_material(request: Request, meeting_id: str, material_id: str):
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    meetings = _load_meetings()
    meeting = meetings.get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    material = next((item for item in _normalize_meeting(meeting).get("materials", []) if item.get("id") == material_id), None)
    if not material:
        raise HTTPException(status_code=404, detail="材料不存在")
    stored_name = _safe_storage_filename(material.get("storedName") or "")
    file_path = MEETING_FILES_DIR / safe_id / stored_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="材料文件不存在")
    return FileResponse(file_path, filename=material.get("fileName") or stored_name)


@app.get("/api/meetings/{meeting_id}/whisper-review")
async def get_whisper_review(request: Request, meeting_id: str):
    """获取 Whisper 终审转写结果"""
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    transcripts_data = _load_meeting_transcripts().get(safe_id, {"transcripts": [], "events": []})
    events = transcripts_data.get("events", [])
    whisper_results = []
    for item in events:
        if item.get("type") == "transcript" and item.get("action") == "whisper-review":
            whisper_results.append({
                "id": item.get("id", ""),
                "text": item.get("text", ""),
                "model": item.get("model", "Whisper-large-v3"),
                "serverTime": item.get("serverTime", ""),
                "sourceFiles": item.get("sourceFiles", 0),
            })
    return JSONResponse({"meetingId": safe_id, "whisperReview": whisper_results})


async def _generate_whisper_meeting_docx(meeting_id: str) -> Optional[str]:
    safe_id = _safe_meeting_id(meeting_id)
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id)
        if meeting is not None:
            whisper_docx = dict(meeting.get("whisperDocx") or {})
            whisper_docx.update({"status": "generating", "updatedAt": _now_text()})
            meeting["whisperDocx"] = whisper_docx
            meetings[safe_id] = meeting
            _save_meetings(meetings)
    if not meeting:
        return None
    try:
        with MEETING_TRANSCRIPTS_LOCK:
            transcripts_data = _load_meeting_transcripts().get(safe_id, {"events": [], "transcripts": []})
        transcripts = transcripts_data.get("transcripts", [])
        events = transcripts_data.get("events", [])

        # ── 提取 Whisper 终审文本（取最新一条） ──
        whisper_segments = []
        whisper_full_text = ""
        whisper_event_data = None
        for item in events:
            if item.get("type") == "transcript" and item.get("action") == "whisper-review":
                segs = item.get("segments", [])
                txt = item.get("text", "")
                if segs:
                    whisper_segments = segs
                    whisper_event_data = item
                if txt and len(txt) > len(whisper_full_text):
                    whisper_full_text = txt

        # 如果有 segments 但 text 很短，从 segments 拼接
        if whisper_segments and len(whisper_full_text) < 100:
            whisper_full_text = "".join(s.get("text", "") for s in whisper_segments)

        local_records = _local_generate_meeting_records(meeting, transcripts, events)
        ai_records = await _deepseek_generate_meeting_records(meeting, transcripts, events)
        records = ai_records if ai_records else local_records
        if not records.get("generated"):
            return None

        # 如果有 Whisper 文本，替换纪实部分（时间戳对齐会议时间轴）
        # 校验：Whisper 段数必须 >= 转写条数的 30%，否则视为垃圾数据跳过
        if whisper_full_text and whisper_segments:
            transcript_count = len(transcripts)
            whisper_count = len(whisper_segments)
            if transcript_count > 0 and whisper_count < transcript_count * 0.3:
                logger.warning("Whisper 段数(%d) < 转写条数(%d) 的 30%%，跳过替换", whisper_count, transcript_count)
            else:
                whisper_chronicle = _build_whisper_chronicle(whisper_segments, whisper_event_data or {})
                if whisper_chronicle:
                    records["chronicle"] = whisper_chronicle
                    records["whisperText"] = whisper_full_text
                    records["audioCount"] = 1

        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        def _cjk(run, font="仿宋", size=Pt(12), bold=False, color=None):
            run.font.name = font
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement("w:rFonts")
            for k in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rFonts.set(qn(k), font)
            rPr.insert(0, rFonts)
            run.font.size = size
            run.bold = bold
            if color:
                run.font.color.rgb = color

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(3.2)
        section.bottom_margin = Cm(2.8)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.6)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"{meeting.get('type') or '会议'}会议纪实、决议与待办事项")
        _cjk(run, "方正小标宋简体", Pt(22), bold=True, color=RGBColor(180, 35, 24))

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _cjk(meta.add_run(f"会议名称：{meeting.get('title', '')}    日期：{meeting.get('date', '')}"), "仿宋", Pt(10))

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for i, h in enumerate(["项目", "内容"]):
            cell = table.rows[0].cells[i]
            cell.text = ""
            _cjk(cell.paragraphs[0].add_run(h), "黑体", Pt(11), bold=True)
        rows = [
            ("会议名称", meeting.get("title", "")),
            ("所属项目", meeting.get("project", "")),
            ("审议议题", meeting.get("agenda", "")),
            ("Whisper 终审", "已接收完整音频并生成终审转写"),
        ]
        for k, v in rows:
            row = table.add_row().cells
            row[0].text = ""; _cjk(row[0].paragraphs[0].add_run(k), "黑体", Pt(11), bold=True)
            row[1].text = ""; _cjk(row[1].paragraphs[0].add_run(str(v)), "仿宋", Pt(11))
        doc.add_paragraph("")

        def heading(text):
            p = doc.add_paragraph()
            _cjk(p.add_run(text), "黑体", Pt(14), bold=True)

        def para(text):
            p = doc.add_paragraph()
            _cjk(p.add_run(text), "仿宋", Pt(12))
            return p

        heading("一、会议纪实")
        for item in records.get("chronicle", []):
            para(f"{item.get('time', '')} {item.get('speaker', '')}（{item.get('role', '')}）：{item.get('content', '')}")

        heading("二、会议纪要")
        for item in records.get("minutes", [])[:20]:
            para(f"{item.get('agenda', '')}｜{item.get('status', '')}｜{item.get('basis', '')}")
            for point in (item.get("keyPoints") or [])[:5]:
                para(f"• {point}")

        heading("三、决议事项")
        for item in records.get("decisions", [])[:20]:
            para(f"会议决定：{item.get('content', '')}")

        heading("四、待办事项")
        for item in records.get("todos", [])[:20]:
            para(f"{item.get('owner', '待确认')}：{item.get('task', '')}")

        heading("五、Whisper 终审摘要")
        audio_count = records.get("audioCount", len([e for e in events if e.get("type") == "audio"]))
        if whisper_full_text:
            para(f"本纪实基于 Whisper-large-v3 高精度转写生成（{len(whisper_segments)} 段，{len(whisper_full_text)} 字）。实时转写 {len(transcripts)} 条作为参考。")
        else:
            para(f"本纪实基于 {len(transcripts)} 条实时转写与 {audio_count} 段录音生成。")

        output_dir = MEETING_FILES_DIR / safe_id
        output_dir.mkdir(parents=True, exist_ok=True)
        filename = f"{safe_id}_whisper会议纪实_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        output_path = output_dir / filename
        doc.save(output_path)
        # 保存 Whisper 增强的 records 回 generatedRecords（覆盖旧缓存）
        # 这样前端 /api/meetings/{id}/records 会拿到 Whisper 高质量版本
        with MEETINGS_LOCK:
            meetings = _load_meetings()
            meeting2 = meetings.get(safe_id, {})
            gr = meeting2.get("generatedRecords") or {}
            # 保留 whisperDocx 元数据
            gr["whisperDocx"] = {
                "fileName": filename,
                "path": str(output_path),
                "generatedAt": _now_text(),
                "status": "done",
                "updatedAt": _now_text(),
            }
            # 用 Whisper 增强的 records 覆盖主缓存
            records["cachedAt"] = _now_text()
            records["whisperEnhanced"] = True
            gr.update(records)
            meeting2["generatedRecords"] = gr
            meetings[safe_id] = meeting2
            _save_meetings(meetings)
            _invalidate_meetings_cache()
        logger.info("Whisper终审: meeting=%s, records 已更新（纪实用 Whisper 高精度转写）", safe_id)
        return str(output_path)
    except Exception as exc:
        logger.warning("Whisper docx 生成失败 meeting=%s: %s", meeting_id, exc)
        with MEETINGS_LOCK:
            meetings = _load_meetings()
            meeting2 = meetings.get(safe_id, {})
            gr = meeting2.get("generatedRecords") or {}
            whisper_docx = dict(gr.get("whisperDocx") or {})
            whisper_docx.update({"status": "failed", "error": str(exc), "updatedAt": _now_text()})
            gr["whisperDocx"] = whisper_docx
            meeting2["generatedRecords"] = gr
            meetings[safe_id] = meeting2
            _save_meetings(meetings)
        return None


@app.get("/api/meetings/{meeting_id}/whisper-docx")
async def download_whisper_meeting_docx(request: Request, meeting_id: str):
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    meetings = _load_meetings()
    meeting = meetings.get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    _check_meeting_access(user, meeting)
    gr = meeting.get("generatedRecords") or {}
    whisper_docx = (gr.get("whisperDocx") or meeting.get("whisperDocx") or {})
    file_path = Path(whisper_docx.get("path") or "")
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Whisper 纪实文件尚未生成")
    filename = whisper_docx.get("fileName") or file_path.name
    return FileResponse(file_path, filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.get("/api/meetings/{meeting_id}/archive/docx")
async def download_meeting_archive_docx(request: Request, meeting_id: str):
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    meetings = _load_meetings()
    meeting = meetings.get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    _check_meeting_access(user, meeting)
    meeting = _public_meeting(meeting, include_detail=True)
    transcripts_data = _load_meeting_transcripts().get(safe_id, {"transcripts": [], "events": []})
    transcripts = transcripts_data.get("transcripts", [])
    events = transcripts_data.get("events", [])

    # 优先用缓存的会议记录（进入归档阶段时已生成），避免下载时重复调 AI
    cached = meeting.get("generatedRecords")
    if cached and cached.get("generated") and cached.get("cachedAt"):
        records = cached
    else:
        local_records = _local_generate_meeting_records(meeting, transcripts, events)
        ai_records = await _deepseek_generate_meeting_records(meeting, transcripts, events)
        records = ai_records if ai_records else local_records
        # 缓存到 meeting 数据中
        records["cachedAt"] = _now_text()
        with MEETINGS_LOCK:
            ms = _load_meetings()
            m2 = ms.get(safe_id, {})
            m2["generatedRecords"] = records
            ms[safe_id] = m2
            _save_meetings(ms)
            _invalidate_meetings_cache()

    # ── 先生成 Markdown ──
    md_lines = []
    meeting_type = (meeting.get("type") or "会议").strip()
    md_lines.append(f"# {meeting_type}会议纪要")
    md_lines.append("")
    md_lines.append(f"**密级**：内部 | **编号**：{safe_id} | **日期**：{meeting.get('date', '')}")
    md_lines.append("")
    md_lines.append("## 一、基本信息")
    md_lines.append("")
    md_lines.append(f"| 项目 | 内容 |")
    md_lines.append(f"|------|------|")
    md_lines.append(f"| 会议名称 | {meeting.get('title', '')} |")
    md_lines.append(f"| 会议日期 | {meeting.get('date', '')} |")
    md_lines.append(f"| 会议类型 | {meeting_type} |")
    md_lines.append(f"| 所属项目 | {meeting.get('project', '')} |")
    md_lines.append(f"| 审议议题 | {meeting.get('agenda', '')} |")
    md_lines.append(f"| 转写底稿 | {len(transcripts)} 条 |")
    md_lines.append(f"| 录音片段 | {len([e for e in events if e.get('type') == 'audio'])} 段 |")
    md_lines.append("")

    md_lines.append("## 二、会议议题")
    md_lines.append("")
    for i, item in enumerate(meeting.get("agendaDrafts", []) or [], 1):
        md_lines.append(f"{i}. **{item.get('title', '')}**（{item.get('type', '')} · {item.get('risk', '')}）")
    if not meeting.get("agendaDrafts"):
        md_lines.append(f"- {meeting.get('agenda', '待确认')}")
    md_lines.append("")

    md_lines.append("## 三、会议摘要")
    md_lines.append("")
    for item in records.get("summary", []) or []:
        md_lines.append(f"- {item}")
    if not records.get("summary"):
        md_lines.append("（暂无 AI 生成摘要）")
    md_lines.append("")

    md_lines.append("## 四、会议纪要")
    md_lines.append("")
    for item in records.get("minutes", []) or []:
        md_lines.append(f"### {item.get('agenda', '议题')}")
        md_lines.append(f"- 状态：{item.get('status', '')}")
        md_lines.append(f"- 依据：{item.get('basis', '')}")
        for point in (item.get("keyPoints") or [])[:5]:
            md_lines.append(f"  - {point}")
        md_lines.append("")
    if not records.get("minutes"):
        md_lines.append("（暂无 AI 生成纪要）")
        md_lines.append("")

    md_lines.append("## 五、决议草案")
    md_lines.append("")
    decisions = records.get("decisions", []) or []
    if decisions:
        for item in decisions:
            md_lines.append(f"- {item.get('content', '')}（{item.get('status', '待秘书确认')}）")
    else:
        md_lines.append("（未识别到明确决议，需秘书人工确认）")
    md_lines.append("")

    md_lines.append("## 六、待办事项")
    md_lines.append("")
    todos = records.get("todos", []) or []
    if todos:
        md_lines.append("| 序号 | 待办事项 | 责任人 | 截止时间 | 优先级 |")
        md_lines.append("|------|----------|--------|----------|--------|")
        for i, item in enumerate(todos, 1):
            md_lines.append(f"| {i} | {item.get('task', '')} | {item.get('owner', '待确认')} | {item.get('deadline', '待定')} | {item.get('priority', '低')} |")
    else:
        md_lines.append("（暂无待办）")
    md_lines.append("")

    md_lines.append("## 七、会议纪实")
    md_lines.append("")
    chronicle = records.get("chronicle", []) or []
    if chronicle:
        for item in chronicle:
            md_lines.append(f"- **{item.get('time', '')}** {item.get('speaker', '')}（{item.get('role', '')}）：{item.get('content', '')}")
    else:
        for item in transcripts[-30:]:
            speaker = item.get("speakerName", "参会人")
            time = item.get("clientTime") or item.get("serverTime", "")
            text = item.get("transcript", "")
            md_lines.append(f"- **{time}** {speaker}：{text}")
    md_lines.append("")

    signed = [t for t in transcripts if t.get("correctionSigned")]
    md_lines.append("## 八、签署确认")
    md_lines.append("")
    md_lines.append(f"共 {len(signed)} 位参会人已通过手机手写签名确认发言内容。")
    md_lines.append("")

    md_lines.append("---")
    md_lines.append(f'*本文件由 AI 会议与"三重一大"合规管理系统自动生成，基于 {len(transcripts)} 条真实转写和 {len([e for e in events if e.get("type") == "audio"])} 段录音。可在 Word 中继续编辑完善。*')

    md_content = "\n".join(md_lines)

    # ── Markdown → 精美 DOCX ──
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    CJK_FONT = "仿宋"
    HEADING_FONT = "黑体"
    TITLE_FONT = "方正小标宋简体"

    def _cjk(run, font=CJK_FONT, size=Pt(12), bold=False, color=None):
        run.font.name = font
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        for k in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(k), font)
        rPr.insert(0, rFonts)
        run.font.size = size
        run.bold = bold
        if color:
            run.font.color.rgb = color

    def _heading(doc, text, level=1):
        h = doc.add_heading(level=level)
        run = h.add_run(text)
        _cjk(run, HEADING_FONT, Pt(16 if level == 1 else 14), bold=True)

    def _para(doc, text, indent=False):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(28)
        run = p.add_run(text)
        _cjk(run)
        return p

    def _table(doc, headers, rows):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(headers):
            cell = t.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            _cjk(run, HEADING_FONT, Pt(11), bold=True)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = t.rows[i + 1].cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(val))
                _cjk(run, CJK_FONT, Pt(11))
        doc.add_paragraph("")
        return t

    doc = Document()
    section = doc.sections[0]
    # A4 纸张
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # 页脚（页码）
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_dash1 = fp.add_run("— ")
    _cjk(run_dash1, CJK_FONT, Pt(9))
    # 插入 PAGE 字段
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run_fld1 = fp.add_run()
    run_fld1._element.append(fldChar_begin)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run_fld2 = fp.add_run()
    run_fld2._element.append(instrText)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run_fld3 = fp.add_run()
    run_fld3._element.append(fldChar_end)
    run_dash2 = fp.add_run(" —")
    _cjk(run_dash2, CJK_FONT, Pt(9))

    # 发文机关（红头上方）
    with _db_connect() as _conn:
        _org_name = _metadata_get(_conn, "org_name") or ""
    if _org_name:
        org_p = doc.add_paragraph()
        org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        org_p.paragraph_format.space_after = Pt(4)
        run_org = org_p.add_run(_org_name)
        _cjk(run_org, HEADING_FONT, Pt(16), color=RGBColor(180, 35, 24))

    # 红头标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run(f"{meeting_type}会议纪要")
    _cjk(run, TITLE_FONT, Pt(22), bold=True, color=RGBColor(180, 35, 24))

    # 发文字号（标题和红线之间）
    doc_no_p = doc.add_paragraph()
    doc_no_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_no_p.paragraph_format.space_before = Pt(4)
    doc_no_p.paragraph_format.space_after = Pt(2)
    _doc_no_text = f"〔{datetime.now().strftime('%Y')}〕第{meeting.get('docNumber', 'XX')}号"
    run_doc_no = doc_no_p.add_run(_doc_no_text)
    _cjk(run_doc_no, CJK_FONT, Pt(12))

    # 红线
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_before = Pt(2)
    line.paragraph_format.space_after = Pt(12)
    run = line.add_run("━" * 35)
    _cjk(run, CJK_FONT, Pt(10), color=RGBColor(180, 35, 24))

    # 基本信息表
    _table(doc,
        ["项目", "内容"],
        [
            ["会议名称", meeting.get("title", "")],
            ["会议日期", meeting.get("date", "")],
            ["会议类型", meeting_type],
            ["所属项目", meeting.get("project", "")],
            ["审议议题", meeting.get("agenda", "")],
            ["转写底稿", f"{len(transcripts)} 条"],
            ["录音片段", f"{len([e for e in events if e.get('type') == 'audio'])} 段"],
        ]
    )

    # 会议议题
    _heading(doc, "一、会议议题")
    for i, item in enumerate(meeting.get("agendaDrafts", []) or [], 1):
        _para(doc, f"{i}. {item.get('title', '')}（{item.get('type', '')}，{item.get('risk', '')}）", indent=True)
    if not meeting.get("agendaDrafts"):
        _para(doc, meeting.get("agenda", "待确认"), indent=True)

    # 会议摘要
    _heading(doc, "二、会议摘要")
    for item in records.get("summary", []) or []:
        _para(doc, f"　　{item}", indent=True)
    if not records.get("summary"):
        _para(doc, "（暂无 AI 生成摘要）")

    # 会议纪要（AI 整理，放前面）
    _heading(doc, "三、会议纪要")
    for item in records.get("minutes", []) or []:
        h3 = doc.add_heading(level=3)
        run = h3.add_run(item.get("agenda", "议题"))
        _cjk(run, HEADING_FONT, Pt(13), bold=True)
        _para(doc, f"状态：{item.get('status', '')}　依据：{item.get('basis', '')}", indent=True)
        for point in (item.get("keyPoints") or [])[:5]:
            _para(doc, f"　• {point}", indent=True)
    if not records.get("minutes"):
        _para(doc, "（暂无 AI 生成纪要）")

    # 决议草案（AI 整理）
    _heading(doc, "四、决议草案")
    decisions = records.get("decisions", []) or []
    if decisions:
        for item in decisions:
            _para(doc, f"{item.get('content', '')}　[{item.get('status', '待秘书确认')}]", indent=True)
    else:
        _para(doc, "（未识别到明确决议，需秘书人工确认）")

    # 待办事项（AI 整理）
    _heading(doc, "五、待办事项")
    todos = records.get("todos", []) or []
    if todos:
        _table(doc,
            ["序号", "待办事项", "责任人", "截止时间", "优先级"],
            [[str(i + 1), t.get("task", ""), t.get("owner", ""), t.get("deadline", "待定"), t.get("priority", "低")] for i, t in enumerate(todos)]
        )
    else:
        _para(doc, "（暂无待办事项）")

    # 会议纪实（原文，放最后）
    _heading(doc, "六、会议纪实")
    chronicle = records.get("chronicle", []) or []
    if chronicle:
        for item in chronicle:
            _para(doc, f"{item.get('time', '')} {item.get('speaker', '')}（{item.get('role', '')}）：{item.get('content', '')}", indent=True)
    else:
        for item in transcripts[-30:]:
            speaker = item.get("speakerName", "参会人")
            time = item.get("clientTime") or item.get("serverTime", "")
            _para(doc, f"{time} {speaker}：{item.get('transcript', '')}", indent=True)

    # 签署确认（含手写签名图片）
    signed = [t for t in transcripts if t.get("correctionSigned")]
    if signed:
        _heading(doc, "七、签署确认")
        for t in signed:
            p = doc.add_paragraph()
            run = p.add_run(f"{t.get('speakerName', '参会人')}（{t.get('speakerRole', '参会代表')}）")
            _cjk(run, HEADING_FONT, Pt(11), bold=True)
            sig = t.get("signatureData", "")
            if sig and sig.startswith("data:image/"):
                try:
                    import base64
                    from io import BytesIO as _BytesIO
                    _header, _b64 = sig.split(",", 1)
                    _img_bytes = base64.b64decode(_b64)
                    p2 = doc.add_paragraph()
                    p2.add_run().add_picture(_BytesIO(_img_bytes), width=Cm(4))
                except Exception as _sig_err:
                    logger.warning("签名图片渲染失败: %s", _sig_err)
            sign_time = t.get("correctionSignedAt") or t.get("serverTime", "")
            p3 = doc.add_paragraph()
            run3 = p3.add_run(f"签署时间：{sign_time}")
            _cjk(run3, CJK_FONT, Pt(10), color=RGBColor(100, 100, 100))

    # ── 防伪码 ──
    # 发文字号自增
    with _db_connect() as _conn:
        _doc_seq = _metadata_get(_conn, "doc_number_seq") or "0"
        _new_seq = int(_doc_seq) + 1
        _metadata_set(_conn, "doc_number_seq", str(_new_seq))
        _doc_number = f"{_new_seq:04d}"
    # 更新发文字号
    for p in doc.paragraphs:
        for run in p.runs:
            if "第XX号" in (run.text or ""):
                run.text = run.text.replace("第XX号", f"第{_doc_number}号")

    # 计算文档防伪哈希（SHA256）
    import hashlib as _hashlib
    _doc_content_for_hash = json.dumps({
        "meetingId": safe_id,
        "title": meeting.get("title", ""),
        "date": meeting.get("date", ""),
        "transcriptCount": len(transcripts),
        "recordsSummary": [s[:50] for s in (records.get("summary") or [])[:3]],
        "docNumber": _doc_number,
        "generatedAt": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
    }, ensure_ascii=False)
    _verify_code = _hashlib.sha256(_doc_content_for_hash.encode()).hexdigest()[:16].upper()
    _verify_code_formatted = "-".join([_verify_code[i:i+4] for i in range(0, 16, 4)])

    # 防伪信息区域
    doc.add_paragraph("")
    verify_p = doc.add_paragraph()
    verify_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    verify_p.paragraph_format.space_before = Pt(12)
    run_v1 = verify_p.add_run(f"防伪编码：{_verify_code_formatted}")
    _cjk(run_v1, CJK_FONT, Pt(10), color=RGBColor(100, 100, 100))

    # 生成二维码（含验证链接）
    try:
        import qrcode
        from io import BytesIO as _BytesIO2
        _verify_url = f"VERIFY:{safe_id}|{_doc_number}|{_verify_code}"
        qr = qrcode.QRCode(version=1, box_size=4, border=1)
        qr.add_data(_verify_url)
        qr.make(fit=True)
        qr_img = qr.make_image(fill_color="black", back_color="white")
        qr_buf = _BytesIO2()
        qr_img.save(qr_buf, format="PNG")
        qr_buf.seek(0)
        qr_p = doc.add_paragraph()
        qr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qr_p.add_run().add_picture(qr_buf, width=Cm(3))
        qr_note = doc.add_paragraph()
        qr_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_qr = qr_note.add_run("扫码验证文件真伪")
        _cjk(run_qr, CJK_FONT, Pt(8), color=RGBColor(150, 150, 150))
    except Exception as _qr_err:
        logger.warning("二维码生成失败: %s", _qr_err)

    # 生成说明
    gen_note = doc.add_paragraph()
    gen_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_note.paragraph_format.space_before = Pt(8)
    run_gen = gen_note.add_run(f"— 本文件由 AI 会议系统自动生成 · {datetime.now().strftime('%Y-%m-%d %H:%M')} —")
    _cjk(run_gen, CJK_FONT, Pt(9), color=RGBColor(150, 150, 150))

    # 保存
    output_dir = MEETING_FILES_DIR / safe_id
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{safe_id}_会议纪要_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    output_path = output_dir / filename
    doc.save(output_path)

    # 同时保存 MD 版本
    md_path = output_dir / f"{safe_id}_会议纪要_{datetime.now().strftime('%Y%m%d%H%M%S')}.md"
    md_path.write_text(md_content, encoding="utf-8")

    return FileResponse(output_path, filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ── Transcripts in-memory cache (1s TTL, max 50 meetings, 5000 rows total) ──
_transcripts_cache: Optional[dict] = None
_transcripts_cache_time: float = 0.0
_transcripts_cache_ttl: float = 1.0
_transcripts_cache_max_keys: int = 50
_transcripts_cache_max_rows: int = 5000












def _cleanup_orphaned_audio_chunks():
    """启动时清理超过 1 小时未合并的孤儿录音 chunk 文件。"""
    import time as _time
    recordings_base = MEETING_FILES_DIR / "recordings"
    if not recordings_base.exists():
        return
    max_age_sec = 3600  # 1 小时
    now = _time.time()
    total_cleaned = 0
    for meeting_dir in recordings_base.iterdir():
        if not meeting_dir.is_dir():
            continue
        chunk_files = list(meeting_dir.glob("chunk_*.webm"))
        if not chunk_files:
            continue
        # 检查最新的 chunk 文件是否超过 max_age
        newest_mtime = max(f.stat().st_mtime for f in chunk_files)
        if now - newest_mtime < max_age_sec:
            continue  # 还在录音中，跳过
        # 已过期的 chunk，直接清理
        for f in chunk_files:
            f.unlink(missing_ok=True)
        total_cleaned += len(chunk_files)
        logger.info("清理孤儿 chunk: %s, %d 个文件", meeting_dir.name, len(chunk_files))
    if total_cleaned > 0:
        logger.info("【启动】共清理 %d 个孤儿 chunk 文件", total_cleaned)


def _recover_orphaned_whisper_results():
    """启动时检查是否有上次重启未持久化的 Whisper 结果，恢复写入 DB。"""
    import json as _json
    recordings_base = MEETING_FILES_DIR / "recordings"
    if not recordings_base.exists():
        return
    recovered = 0
    for whisper_cache in recordings_base.rglob("_whisper_result.json"):
        try:
            data = _json.loads(whisper_cache.read_text(encoding="utf-8"))
            meeting_id = data.get("meetingId", "")
            if not meeting_id:
                whisper_cache.unlink(missing_ok=True)
                continue
            # 检查是否已持久化（避免重复）
            with MEETINGS_LOCK:
                meetings = _load_meetings()
                meeting = meetings.get(meeting_id, {})
            existing_ids = {
                e.get("id") for e in meeting.get("events", [])
                if e.get("action") == "whisper-review"
            }
            if data.get("id") not in existing_ids:
                _append_meeting_activity_light(meeting_id, data)
                recovered += 1
            whisper_cache.unlink(missing_ok=True)
        except Exception as e:
            logger.warning("恢复 Whisper 结果失败 %s: %s", whisper_cache, e)
    if recovered:
        logger.info("启动恢复: %d 个 Whisper 结果已持久化", recovered)


async def _run_whisper_review_for_meeting(meeting_id: str):
    """Whisper 终审转写 —— 仅在归档阶段触发。
    全局并发锁确保同一时刻只有一个 Whisper 任务在跑，防止 GPU OOM。
    """
    async with _whisper_semaphore:
        await _do_whisper_review(meeting_id)


async def _do_whisper_review(meeting_id: str):
    """实际的 Whisper 终审逻辑（受 _whisper_semaphore 保护）。"""
    try:
        # 检查是否已有完整的 whisper-review 结果，避免重复运行
        safe_id_check = _safe_meeting_id(meeting_id)
        with MEETINGS_LOCK:
            meetings_check = _load_meetings()
            meeting_check = meetings_check.get(safe_id_check, {})
        has_full_whisper = any(
            e.get("type") == "transcript" and e.get("action") == "whisper-review"
            and e.get("duration", 0) > 30  # 忽略极短的早期 review
            for e in meeting_check.get("events", [])
        )
        if has_full_whisper:
            logger.info("Whisper终审: meeting=%s 已有完整结果，跳过", safe_id_check)
            return
        import subprocess as _sp
        from backend.whisper_transcribe import transcribe_file
        safe_id = _safe_meeting_id(meeting_id)
        audio_dir = MEETING_FILES_DIR / "recordings" / safe_id
        if not audio_dir.exists():
            return
        # 排除 chunk 分片文件（chunk_*.webm 是流式上传的中间文件）
        files = (
            sorted(f for f in audio_dir.glob("*.webm") if not f.name.startswith("chunk_")) +
            sorted(audio_dir.glob("*.mp4")) +
            sorted(audio_dir.glob("*.m4a"))
        )
        if not files:
            return

        # 过滤出有效的音频文件（跳过损坏文件）
        valid_files = []
        for f in files:
            try:
                r = _sp.run(["ffprobe", "-v", "quiet", "-show_entries",
                             "format=duration", "-of", "csv=p=0", str(f)],
                            capture_output=True, timeout=10)
                if r.returncode == 0 and r.stdout.strip():
                    valid_files.append(f)
            except Exception:
                pass

        if not valid_files:
            logger.warning("Whisper终审: 无有效音频文件 meeting=%s", safe_id)
            return

        logger.info("Whisper终审: meeting=%s, %d/%d 个有效文件", safe_id, len(valid_files), len(files))

        # ── 计算每个录音文件的时间偏移量（用于 Whisper 时间戳对齐会议时间轴）──
        events: list = []
        try:
            with MEETING_TRANSCRIPTS_LOCK:
                transcripts_data = _load_meeting_transcripts().get(safe_id, {"events": [], "transcripts": []})
            events = transcripts_data.get("events", [])
        except Exception as _evt_err:
            logger.warning("Whisper终审加载 events 失败 meeting=%s: %s", safe_id, _evt_err)
        audio_events = {
            e.get("fileName"): e for e in events
            if e.get("type") == "audio" and e.get("action") == "audio-uploaded"
        }
        # 为每个有效文件找到对应的事件，获取 recordingStartTime
        file_starts = []
        for f in valid_files:
            evt = audio_events.get(f.name, {})
            start_str = evt.get("recordingStartTime")
            if start_str:
                try:
                    start_dt = datetime.fromisoformat(start_str.replace("Z", "+00:00"))
                    file_starts.append((f, start_dt.timestamp()))
                except Exception:
                    file_starts.append((f, None))
            else:
                file_starts.append((f, None))
        # 按 startTime 排序（有时间的优先，无时间的按文件名排）
        file_starts.sort(key=lambda x: (x[1] is None, x[1] or 0, x[0].name))
        valid_files = [fs[0] for fs in file_starts]
        # 计算每个文件在合并音频中的起始偏移（秒）和在会议时间轴中的偏移
        file_durations = []  # 每个文件的时长
        for f, _ in file_starts:
            try:
                r = _sp.run(["ffprobe", "-v", "quiet", "-show_entries",
                             "format=duration", "-of", "csv=p=0", str(f)],
                            capture_output=True, timeout=10)
                dur = float(r.stdout.strip()) if r.returncode == 0 and r.stdout.strip() else 0
            except Exception:
                dur = 0
            file_durations.append(dur)
        # 记录整个录音的起始时间（最早文件的 startTime）
        meeting_start_ts = None
        for _, ts in file_starts:
            if ts is not None:
                if meeting_start_ts is None or ts < meeting_start_ts:
                    meeting_start_ts = ts
        # 计算每个文件相对于最早录音的时间偏移
        file_offsets = {}  # filename → offset in seconds from meeting start
        if meeting_start_ts is not None:
            for (f, ts), dur in zip(file_starts, file_durations):
                if ts is not None:
                    file_offsets[f.name] = ts - meeting_start_ts
                else:
                    file_offsets[f.name] = 0
        # file_durations 用于后续 Whisper segment → 原始文件的映射

        # 合并有效录音文件为临时 WAV
        merged_wav = audio_dir / "_merged_whisper.wav"
        try:
            if len(valid_files) == 1:
                r = _sp.run(
                    ["ffmpeg", "-y", "-i", str(valid_files[0]),
                     "-af", "loudnorm=I=-16:TP=-1.5:LRA=11",
                     "-acodec", "pcm_s16le", "-ac", "1", "-ar", "16000",
                     str(merged_wav)],
                    capture_output=True, timeout=120,
                )
            else:
                input_args = []
                for f in valid_files:
                    input_args.extend(["-i", str(f)])
                r = _sp.run(
                    ["ffmpeg", "-y"] + input_args +
                    ["-filter_complex", f"concat=n={len(valid_files)}:v=0:a=1[out];[out]loudnorm=I=-16:TP=-1.5:LRA=11[norm]",
                     "-map", "[norm]", "-acodec", "pcm_s16le", "-ac", "1", "-ar", "16000",
                     str(merged_wav)],
                    capture_output=True, timeout=120,
                )
            if r.returncode != 0:
                logger.warning("Whisper终审 ffmpeg 合并失败: %s", r.stderr.decode()[:200])
                return
        except Exception as e:
            logger.warning("Whisper终审 ffmpeg 异常: %s", e)
            return

        if not merged_wav.exists() or merged_wav.stat().st_size < 32000:
            merged_wav.unlink(missing_ok=True)
            return

        logger.info("Whisper终审: meeting=%s, %d文件, 本地 openai-whisper 转写", safe_id, len(valid_files))

        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            None,
            lambda: transcribe_file(str(merged_wav), model_size="large-v3", language="zh"),
        )

        # ── 声纹 Speaker Diarization（会后校准）──
        voiceprint_diarization = None
        try:
            from backend.voiceprint import get_voiceprint_engine as _get_vp_eng, deserialize_embedding as _deser_emb
            _vp_eng = _get_vp_eng()
            if _vp_eng and _vp_eng.is_ready:
                from backend.db import _db_load_voiceprint_profiles
                _vp_profiles = _db_load_voiceprint_profiles()
                _vp_enrolled = {}
                for _p in _vp_profiles:
                    _vp_enrolled[_p["user_id"]] = _deser_emb(_p["embedding"])
                if _vp_enrolled:
                    logger.info("会后声纹校准: 开始 diarization meeting=%s", safe_id)
                    diar_segments = await loop.run_in_executor(
                        None,
                        lambda: _vp_eng.diarize(str(merged_wav)),
                    )
                    if diar_segments:
                        label_map = await loop.run_in_executor(
                            None,
                            lambda: _vp_eng.match_diarization_to_enrolled(
                                diar_segments, str(merged_wav), _vp_enrolled,
                            ),
                        )
                        voiceprint_diarization = {
                            "segments": diar_segments,
                            "label_map": label_map,
                        }
                        # 回填到 meeting_transcripts：按时间戳对齐
                        _now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        with APP_DB_LOCK:
                            with _db_connect() as _vp_conn:
                                _vp_rows = _vp_conn.execute(
                                    "SELECT id, server_time, speaker_name, identified_by FROM meeting_transcripts "
                                    "WHERE meeting_id = ? ORDER BY server_time",
                                    (safe_id,),
                                ).fetchall()
                        # 将 diarization segment 按时间映射到 transcript
                        # 需要 meeting_start_ts 来对齐
                        if meeting_start_ts and _vp_rows:
                            for _vp_row in _vp_rows:
                                _row_id = _vp_row[0]
                                _row_time_str = _vp_row[1]
                                _row_identified_by = _vp_row[3] or "manual"
                                # 会后 diarization 覆盖所有记录（包括 voiceprint-realtime），
                                # 因为实时识别可能首次匹配错误，会后全局校准更准确
                                try:
                                    _row_dt = datetime.strptime(_row_time_str, "%Y-%m-%d %H:%M:%S")
                                    _row_offset = _row_dt.timestamp() - meeting_start_ts
                                except Exception:
                                    continue
                                # 找到对应的 diarization segment
                                _matched_label = None
                                for seg in diar_segments:
                                    if seg["start"] <= _row_offset <= seg["end"]:
                                        _matched_label = seg["speaker"]
                                        break
                                if _matched_label and _matched_label in label_map:
                                    _matched_user_id = label_map[_matched_label]
                                    if _matched_user_id:
                                        # 查 display_name
                                        _matched_profile = next(
                                            (p for p in _vp_profiles if p["user_id"] == _matched_user_id), None
                                        )
                                        _matched_name = _matched_profile["display_name"] if _matched_profile else _matched_user_id
                                        with APP_DB_LOCK:
                                            with _db_connect() as _upd_conn:
                                                _upd_conn.execute(
                                                    "UPDATE meeting_transcripts SET speaker_name = ?, identified_by = ? WHERE id = ?",
                                                    (_matched_name, "voiceprint-diarization", _row_id),
                                                )
                        logger.info("会后声纹校准: 完成 meeting=%s, %d segments, %d labels mapped",
                                    safe_id, len(diar_segments), len(label_map))
        except Exception as vp_diar_err:
            logger.warning("会后声纹校准异常 meeting=%s: %s", safe_id, vp_diar_err)

        merged_wav.unlink(missing_ok=True)

        text = result.get("text", "")
        if text:
            logger.info("Whisper终审结果: %r", text[:200])
            whisper_event = {
                "id": f"whisper_review_{uuid.uuid4().hex[:8]}",
                "type": "transcript",
                "action": "whisper-review",
                "meetingId": safe_id,
                "serverTime": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "text": text,
                "model": "faster-whisper-large-v3",
                "sourceFiles": len(files),
                "segments": result.get("segments", []),
                "duration": result.get("duration", 0),
                "fileOffsets": {k: v for k, v in file_offsets.items()},
                "fileDurations": file_durations,
                "meetingStartTs": meeting_start_ts,
                "voiceprintDiarization": voiceprint_diarization,
            }
            # 先写入磁盘临时文件，防止服务器重启丢失结果
            import json as _json
            whisper_cache = audio_dir / "_whisper_result.json"
            try:
                whisper_cache.write_text(_json.dumps(whisper_event, ensure_ascii=False), encoding="utf-8")
            except Exception:
                pass
            _append_meeting_activity_light(safe_id, whisper_event)
            # 持久化成功后删除临时文件
            whisper_cache.unlink(missing_ok=True)
            _safe_create_task(_generate_whisper_meeting_docx(safe_id), name=f"whisper-docx-{safe_id}")
    except Exception as e:
        logger.warning("Whisper终审异常 meeting=%s: %s", meeting_id, e)


@app.post("/api/meeting/recorder/session")
async def meeting_recorder_session(request: Request, body: MeetingRecorderSessionRequest):
    user = _get_request_user(request, required=True)
    role = _resolve_meeting_role(user)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    # ── 录音客户端注册（user → meeting_participant → audio_client 链）──
    client_id = (body.device_id or "").strip()
    if client_id and body.action in ("start", "join", "resume"):
        try:
            from backend.db import _db_upsert_audio_client
            _db_upsert_audio_client(_safe_meeting_id(body.meeting_id), client_id, user, {
                "device_type": body.device_type or "mobile",
                "device_label": body.device_label or "手机麦克风",
                "firmware_version": body.firmware_version or "",
                "transport": body.transport or "web-mobile",
            })
        except Exception:
            pass
    event = {
        "id": f"session_{uuid.uuid4().hex[:12]}",
        "type": "session",
        "action": body.action,
        "meetingId": body.meeting_id,
        "meetingTitle": body.meeting_title,
        "agenda": body.agenda,
        "audioSize": body.audio_size,
        "durationSeconds": body.duration_seconds,
        "deviceType": body.device_type or "mobile",
        "deviceId": body.device_id or "",
        "deviceLabel": body.device_label or "手机麦克风",
        "channel": body.channel or "primary",
        "transport": body.transport or "web-mobile",
        "firmwareVersion": body.firmware_version or "",
        "serverTime": now,
        "speaker": role,
    }
    _append_meeting_activity_light(body.meeting_id, event)

    # ── Whisper 终审仅在归档阶段触发（不在每次 session stop 时重复运行）──

    # SSE 推送
    from backend.config import sse_manager as _sse_mgr
    _safe_create_task(_sse_mgr.publish(body.meeting_id, "session", {
        "meetingId": body.meeting_id,
        "action": body.action,
        "speaker": role,
        "serverTime": now,
    }), name=f"sse-session-{body.meeting_id}")
    return JSONResponse({"success": True, "event": event, "speaker": role})


@app.post("/api/meeting/recorder/audio/chunk")
async def upload_meeting_recorder_audio_chunk(
    request: Request,
    meeting_id: str = Form(...),
    chunk_index: int = Form(0),
    client_id: str = Form(""),
    file: UploadFile = File(...),
):
    """流式上传录音片段 —— 前端每 3 秒上传一次，避免浏览器内存溢出。

    P0-6 幂等: 同一 meeting_id + client_id + chunk_index 重复上传不产生重复文件。
    P0-4 ACK: 返回 ack=chunk_index 确认已落盘。
    P0-18 MIME: 根据上传文件的 MIME 类型保存正确扩展名。
    """
    user = _get_request_user(request, required=True)
    username = (user.get("username") or user.get("name") or "unknown").strip()
    safe_id = _safe_meeting_id(meeting_id)
    audio_dir = MEETING_FILES_DIR / "recordings" / safe_id
    audio_dir.mkdir(parents=True, exist_ok=True)

    # P0-18: 根据 MIME 选择扩展名
    mime = (file.content_type or "").lower()
    ext = ".webm"
    if "mp4" in mime:
        ext = ".mp4"
    elif "ogg" in mime:
        ext = ".ogg"

    # P0-3: 文件名 = chunk_{client_id}_{index} 或 chunk_{username}_{index}
    cid = re.sub(r'[^a-zA-Z0-9_-]', '', client_id)[:32] if client_id else ""
    if cid:
        chunk_name = f"chunk_{cid}_{chunk_index:06d}{ext}"
    else:
        chunk_name = f"chunk_{username}_{chunk_index:06d}{ext}"
    chunk_path = audio_dir / chunk_name

    # P0-6 幂等: 文件已存在且大小>0 → 不重复写入，直接返回 ACK
    if chunk_path.exists() and chunk_path.stat().st_size > 0:
        logger.info("录音 chunk 幂等跳过: meeting=%s, chunk=%s, size=%d", safe_id, chunk_name, chunk_path.stat().st_size)
        return JSONResponse({"success": True, "ack": chunk_index, "chunkIndex": chunk_index, "duplicate": True})

    content = await _read_upload_safe(file, 50 * 1024 * 1024)  # 单 chunk 最大 50MB
    # 原子写入：先写临时文件再 rename，防止崩溃留下半截文件
    tmp_chunk = chunk_path.with_suffix(".tmp")
    with open(tmp_chunk, "wb") as f:
        f.write(content)
        f.flush()
        os.fsync(f.fileno())
    tmp_chunk.rename(chunk_path)
    logger.info("[AUDIO] meeting=%s, client=%s, chunk=%d, mime=%s, bytes=%d, saved=%s",
                safe_id, cid or username, chunk_index, mime, len(content), chunk_name)
    # P0-4: ACK 确认已落盘
    return JSONResponse({"success": True, "ack": chunk_index, "chunkIndex": chunk_index, "size": len(content), "user": username})


@app.post("/api/meeting/recorder/audio/complete")
async def complete_meeting_recorder_audio(
    request: Request,
    meeting_id: str = Form(...),
    meeting_title: str = Form(""),
    agenda: str = Form(""),
    duration_seconds: Optional[int] = Form(None),
    total_chunks: int = Form(0),
    recording_start_time: Optional[str] = Form(None),
):
    """录音完成 —— 合并所有 chunk 为完整文件。

    修复记录 (2026-07-13):
    - 用 ffmpeg remux 替代二进制拼接，修复损坏的 MP4/WebM 容器
    - 校验 chunk 数量，缺失时记录警告
    - 原子写入（temp + rename），防止崩溃留下半截文件
    """
    user = _get_request_user(request, required=True)
    role = _resolve_meeting_role(user)
    username = (user.get("username") or user.get("name") or "unknown").strip()
    safe_id = _safe_meeting_id(meeting_id)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    audio_dir = MEETING_FILES_DIR / "recordings" / safe_id
    if not audio_dir.exists():
        return JSONResponse({"success": False, "error": "录音目录不存在"}, status_code=404)

    # 按用户收集 chunk（新格式 chunk_{user}_{index}.ext），兼容旧格式 chunk_{index}.webm
    # P0-3: 支持 client_id 格式和多种扩展名
    user_chunks = sorted(audio_dir.glob(f"chunk_{username}_*.*"))
    if not user_chunks:
        # 降级：兼容旧格式（无用户名前缀）
        user_chunks = sorted(audio_dir.glob("chunk_*.*"))
    # 只收集音频文件，排除临时文件和合并后的 mp4
    chunks = [c for c in user_chunks if c.suffix in ('.webm', '.mp4', '.ogg') and not c.name.startswith('audio_')]
    if not chunks:
        return JSONResponse({"success": False, "error": "无录音片段"}, status_code=404)

    # chunk 数量校验
    if total_chunks > 0 and len(chunks) != total_chunks:
        logger.warning(
            "录音 chunk 数量不匹配: meeting=%s, user=%s, expected=%d, found=%d",
            safe_id, username, total_chunks, len(chunks),
        )

    # 合并为完整文件（ffmpeg remux 修复容器头）
    audio_id = f"audio_{uuid.uuid4().hex[:12]}"
    stored_name = f"{audio_id}.mp4"
    file_path = audio_dir / stored_name
    tmp_path = audio_dir / f".tmp_{audio_id}.mp4"

    try:
        # 1) 创建 ffmpeg concat 列表
        concat_list = audio_dir / f".concat_{audio_id}.txt"
        with open(concat_list, "w") as f:
            for chunk in chunks:
                f.write(f"file '{chunk}'\n")

        # 2) ffmpeg remux：合并所有 chunk 并重新封装为完整容器
        proc = await asyncio.create_subprocess_exec(
            "ffmpeg", "-y",
            "-f", "concat", "-safe", "0",
            "-i", str(concat_list),
            "-c", "copy",        # 不重新编码，只重新封装
            "-movflags", "+faststart",
            str(tmp_path),
            stdout=asyncio.subprocess.DEVNULL,
            stderr=asyncio.subprocess.PIPE,
        )
        _, stderr = await proc.communicate()

        if proc.returncode != 0 or not tmp_path.exists() or tmp_path.stat().st_size == 0:
            # ffmpeg 失败，降级为二进制拼接（兼容无 ffmpeg 环境）
            logger.warning("ffmpeg remux 失败，降级为二进制拼接: %s", stderr.decode()[-200:])
            with open(tmp_path, "wb") as out:
                for chunk in chunks:
                    out.write(chunk.read_bytes())

        # 3) 原子写入：rename tmp → final
        tmp_path.rename(file_path)
        logger.info("录音合并完成: meeting=%s, chunks=%d, size=%d", safe_id, len(chunks), file_path.stat().st_size)

    except Exception as e:
        logger.error("录音合并异常: meeting=%s, error=%s", safe_id, e)
        tmp_path.unlink(missing_ok=True)
        # 降级：直接拼接
        with open(file_path, "wb") as out:
            for chunk in chunks:
                out.write(chunk.read_bytes())
    finally:
        # 清理临时文件
        (audio_dir / f".concat_{audio_id}.txt").unlink(missing_ok=True)

    # 清理 chunk 文件
    for chunk in chunks:
        chunk.unlink(missing_ok=True)

    file_size = file_path.stat().st_size

    event = {
        "id": audio_id,
        "type": "audio",
        "action": "audio-uploaded",
        "meetingId": safe_id,
        "meetingTitle": meeting_title,
        "agenda": agenda,
        "serverTime": now,
        "speaker": role,
        "fileName": stored_name,
        "storedName": stored_name,
        "audioSize": file_size,
        "durationSeconds": duration_seconds,
        "recordingStartTime": recording_start_time,
        "playbackUrl": f"/api/meeting/recorder/audio/{safe_id}/{audio_id}",
    }
    try:
        _append_meeting_activity_light(safe_id, event)
    except Exception as e:
        logger.warning("录音事件记录失败（不影响音频合并）: meeting=%s, error=%s", safe_id, e)
    return JSONResponse({"success": True, "event": event, "audioSize": file_size})


@app.post("/api/meeting/recorder/audio")
async def upload_meeting_recorder_audio(
    request: Request,
    meeting_id: str = Form(...),
    meeting_title: str = Form(""),
    agenda: str = Form(""),
    duration_seconds: Optional[int] = Form(None),
    file: UploadFile = File(...),
):
    user = _get_request_user(request, required=True)
    role = _resolve_meeting_role(user)
    safe_id = _safe_meeting_id(meeting_id)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    ext = Path(file.filename or "recording.webm").suffix.lower() or ".webm"
    if ext not in {".webm", ".mp3", ".m4a", ".wav", ".ogg", ".mp4"}:
        ext = ".webm"
    audio_id = f"audio_{uuid.uuid4().hex[:12]}"
    audio_dir = MEETING_FILES_DIR / "recordings" / safe_id
    audio_dir.mkdir(parents=True, exist_ok=True)
    stored_name = f"{audio_id}{ext}"
    file_path = audio_dir / stored_name
    content = await _read_upload_safe(file, MAX_AUDIO_BYTES)
    file_path.write_bytes(content)
    event = {
        "id": audio_id,
        "type": "audio",
        "action": "audio-uploaded",
        "meetingId": safe_id,
        "meetingTitle": meeting_title,
        "agenda": agenda,
        "serverTime": now,
        "speaker": role,
        "fileName": file.filename or stored_name,
        "storedName": stored_name,
        "audioSize": len(content),
        "durationSeconds": duration_seconds,
        "playbackUrl": f"/api/meeting/recorder/audio/{safe_id}/{audio_id}",
    }
    _append_meeting_activity_light(safe_id, event)
    return JSONResponse({"success": True, "event": event})


@app.get("/api/meeting/recorder/audio/{meeting_id}/{audio_id}")
async def download_meeting_recorder_audio(request: Request, meeting_id: str, audio_id: str):
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    # Direct DB lookup — no need to load all transcripts
    with _db_connect() as conn:
        row = conn.execute(
            "SELECT payload_json FROM meeting_events WHERE meeting_id = ? AND id = ? AND type = 'audio'",
            (safe_id, audio_id),
        ).fetchone()
    target = _json_loads(row["payload_json"], {}) if row else None
    if not target:
        raise HTTPException(status_code=404, detail="录音文件不存在")
    stored_name = target.get("storedName") or ""
    file_path = MEETING_FILES_DIR / "recordings" / safe_id / stored_name
    if not stored_name or not file_path.exists():
        raise HTTPException(status_code=404, detail="录音文件不存在")
    _ext_media = {".webm": "audio/webm", ".mp4": "audio/mp4", ".m4a": "audio/mp4", ".mp3": "audio/mpeg", ".wav": "audio/wav", ".ogg": "audio/ogg"}
    _suffix = Path(stored_name).suffix.lower()
    return FileResponse(file_path, filename=target.get("fileName") or stored_name, media_type=_ext_media.get(_suffix, "audio/webm"))


# ═══ ASR 文本清洗管线 ═══════════════════════════════════════════════════

def _clean_asr_text(text: str) -> str:
    """纪实原则：保留发言原貌。只做最小清洗——连续语气词去重、连续标点归一。"""
    if not text or not text.strip():
        return ""
    # 连续重复语气词去重（嗯嗯嗯→嗯）
    text = re.sub(r'([嗯啊哎哦呢吧啦哈嘿哟哇]){2,}', r'\1', text)
    # 连续重复标点归一化
    text = re.sub(r'([。.？?！!，,、]){2,}', r'\1', text)
    # 去除多余空白
    text = re.sub(r'\s+', ' ', text)
    # 标点统一
    text = text.replace('..', '.').replace('。。', '。').replace('？？', '？').replace('！！', '！')
    return text.strip()


async def _llm_refine_text(raw_text: str) -> str:
    """LLM 语义规整：去语气词、同音字修正、公文风分段。
    后台异步运行，不阻塞实时转写入库。
    """
    if len(raw_text) < 6:
        return raw_text
    try:
        from langchain_core.messages import HumanMessage
        prompt = f"""政企会议文本规整：去语气词和口吃重复，修正同音错字（如"司上班"→"公司上班"），补全标点。
仅输出规整后文本，不要任何解释。
输入：{raw_text}
输出："""
        result = await llm._agenerate([HumanMessage(content=prompt)], enable_thinking=False)
        refined = (result.generations[0].message.content or "").strip()
        return refined if refined and len(refined) >= len(raw_text) * 0.5 else raw_text
    except Exception as e:
        logger.warning("_llm_refine_text 失败: %s", e)
        return raw_text


@app.post("/api/meeting/transcripts/chunk")
async def post_meeting_transcript_chunk(request: Request, body: MeetingTranscriptChunkRequest):
    user = _get_request_user(request, required=True)
    # 自动创建会议（手机录音的 meeting-local-* 会议可能不在 meetings 表中）
    safe_id = _safe_meeting_id(body.meeting_id)
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        if safe_id not in meetings:
            now_str = _now_text()
            new_meeting = {
                "id": safe_id,
                "title": body.meeting_title or "手机录音会议",
                "project": "",
                "agenda": body.agenda or "",
                "date": now_str[:10],
                "type": "普通企业会议",
                "meetingMode": "normal",
                "creator": user.get("displayName") or user.get("username", ""),
                "createdAt": now_str,
                "updatedAt": now_str,
                "phase": "会中记录",
                "agendaFrozen": True,
                "meetingCreated": True,
                "reviewDone": False,
                "archiveDone": False,
                "projectBound": False,
                "issueSources": [],
                "agendaDrafts": [],
                "materials": [],
                "events": [],
            }
            _db_upsert_meeting(new_meeting)
            _invalidate_meetings_cache()
            logger.info("自动创建会议: %s (%s)", safe_id, body.meeting_title)
    raw_transcript = re.sub(r"\s+", " ", body.transcript or "").strip()
    if not raw_transcript:
        raise HTTPException(status_code=400, detail="转写内容不能为空")
    # 规则清洗
    transcript = _clean_asr_text(raw_transcript)
    if not transcript:
        raise HTTPException(status_code=400, detail="转写内容无效（纯语气词或标点）")
    role = _resolve_meeting_role(user)
    # 手机端可主动声明当前发言人（覆盖登录账号角色）
    speaker_name = getattr(body, "speaker_name", None) or role["displayName"]
    speaker_role = getattr(body, "speaker_role", None) or role["meetingRole"]
    speaker_dept = getattr(body, "speaker_dept", None) or role["dept"]
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    # ── 议题绑定：优先客户端显式 agenda_id，否则取后端持久化的 active_agenda_id（§32 手机不选议题）──
    active_agenda_id = ""
    try:
        _m = _load_meetings().get(safe_id) or {}
        active_agenda_id = _m.get("activeAgendaId") or ""
    except Exception:
        pass
    agenda_id = getattr(body, "agenda_id", None) or active_agenda_id or ""
    participant_id = getattr(body, "participant_id", None) or ""
    audio_client_id = getattr(body, "audio_client_id", None) or ""
    if not participant_id:
        try:
            from backend.db import _db_find_participant_row
            participant_id = _db_find_participant_row(safe_id, role.get("userId") or "")
        except Exception:
            pass
    record = {
        "id": f"tr_{uuid.uuid4().hex[:12]}",
        "meetingId": body.meeting_id,
        "meetingTitle": body.meeting_title,
        "agenda": body.agenda,
        "speakerName": speaker_name,
        "speakerRole": speaker_role,
        "speakerDept": speaker_dept,
        "seat": role["seat"],
        "username": role["username"],
        "transcript": transcript,
        "isFinal": body.is_final,
        "clientTime": body.client_time,
        "serverTime": now,
        "confidence": body.confidence if body.confidence is not None else 0.92,
        "source": "mobile-recorder",
        "speakerConfidence": body.speaker_confidence or 0,
        "identifiedBy": body.identified_by or "manual",
        "agendaId": agenda_id,
        "speakerUserId": role.get("userId") or "",
        "participantId": participant_id,
        "audioClientId": audio_client_id,
    }
    # 客户端心跳：更新 meeting_audio_clients 的 last_seen_at
    if audio_client_id:
        try:
            from backend.db import _db_upsert_audio_client
            _db_upsert_audio_client(safe_id, audio_client_id, user, {
                "device_type": "mobile", "transport": "transcript-push",
            })
        except Exception:
            pass
    # 去重：仅查询近 30 秒内的转写（避免大会议全量加载）
    cutoff = (datetime.now() - timedelta(seconds=30)).strftime("%Y-%m-%d %H:%M:%S")
    try:
        with APP_DB_LOCK:
            with _db_connect() as _dedup_conn:
                _dedup_rows = _dedup_conn.execute(
                    "SELECT transcript, server_time, username FROM meeting_transcripts "
                    "WHERE meeting_id = ? AND server_time >= ? "
                    "ORDER BY server_time DESC LIMIT 50",
                    (body.meeting_id, cutoff),
                ).fetchall()
        existing_recent = [
            {"transcript": r[0], "serverTime": r[1], "username": r[2]}
            for r in _dedup_rows
        ]
    except Exception:
        existing_recent = []
    is_duplicate = False
    for prev in existing_recent:
        prev_time = prev.get("serverTime", "")
        if prev_time and now < prev_time:
            continue  # 跳过未来时间（时钟偏差）
        if prev_time:
            try:
                delta = (datetime.strptime(now, "%Y-%m-%d %H:%M:%S") - datetime.strptime(prev_time, "%Y-%m-%d %H:%M:%S")).total_seconds()
            except Exception:
                delta = 999
            if abs(delta) > 5:
                continue
        prev_text = re.sub(r"\s+", "", str(prev.get("transcript", "")))
        curr_text = re.sub(r"\s+", "", transcript)
        if not prev_text or not curr_text:
            continue
        # 仅跳过完全相同的文本（精确去重）
        # 不再做包含关系判断，避免误杀合法的短句跟进发言
        if prev_text == curr_text:
            is_duplicate = True
            logger.info("跳过完全重复转写: %s", transcript[:60])
            break
    if is_duplicate:
        return JSONResponse({"success": True, "duplicate": True})

    # 连续发言合并：同一说话人 10 秒内的 newText 拼接到同一条记录，
    # 或 Jaccard 相似度 > 0.7 时视为 ASR 修订，替换上一条而非新增。
    # 上限 200 字，超长自动分段避免一条记录吃进整个发言。
    MAX_MERGE_CHARS = 200
    merged = False
    prev_id = None
    prev_text = ""
    try:
        cutoff_merge = (datetime.now() - timedelta(seconds=30)).strftime("%Y-%m-%d %H:%M:%S")
        with APP_DB_LOCK:
            with _db_connect() as _mrg_conn:
                _mrg_row = _mrg_conn.execute(
                    "SELECT id, transcript, server_time FROM meeting_transcripts "
                    "WHERE meeting_id = ? AND username = ? AND server_time >= ? "
                    "ORDER BY server_time DESC LIMIT 1",
                    (body.meeting_id, role["username"], cutoff_merge),
                ).fetchone()
        if _mrg_row:
            prev_id, prev_text, prev_stime = _mrg_row[0], _mrg_row[1], _mrg_row[2]
            delta = (datetime.strptime(now, "%Y-%m-%d %H:%M:%S")
                     - datetime.strptime(prev_stime, "%Y-%m-%d %H:%M:%S")).total_seconds()
            # 10 秒内同说话人 且 合并后不超限 → 合并
            if 0 < delta <= 10 and len(prev_text) + len(transcript) <= MAX_MERGE_CHARS:
                merged = True
            else:
                # 超过 10 秒但文本高度相似 → ASR 修订 → 替换上一条
                import difflib
                sim = difflib.SequenceMatcher(None, prev_text, transcript).ratio()
                if sim > 0.85:
                    merged = True
                    logger.info("ASR修订合并(说话人=%s, delta=%.0fs, 文本相似度=%.0f%%)",
                                speaker_name, delta, sim * 100)
    except Exception:
        pass

    if merged:
        record["id"] = prev_id
        record["transcript"] = prev_text + transcript
        _db_upsert_transcript(record)
        _invalidate_transcripts_cache()
        logger.info("连续发言合并(说话人=%s): +%d字 → 共%d字",
                    speaker_name, len(transcript), len(record["transcript"]))
    else:
        _db_upsert_transcript(record)
        _invalidate_transcripts_cache()
        logger.info("转写保存: meeting=%s, speaker=%s, %d字, final=%s, text=%s",
                    body.meeting_id, speaker_name, len(transcript),
                    body.is_final, transcript[:80])
        # 事件中同时保存转写文本 — 防止 meeting_transcripts 表意外清空后数据全丢
        _append_meeting_activity_light(body.meeting_id, {
            "id": f"transcript_event_{uuid.uuid4().hex[:10]}",
            "type": "transcript",
            "transcriptId": record["id"],
            "speakerName": speaker_name,
            "speakerRole": speaker_role,
            "transcript": transcript,
            "serverTime": now,
        })

    # SSE 推送 — 合并时发送完整文本，新增时发送当前文本
    from backend.config import sse_manager
    sse_text = record["transcript"]  # 合并后是完整文本，新增时即 transcript
    _safe_create_task(sse_manager.publish(body.meeting_id, "transcript", {
        "id": record["id"],
        "meeting_id": body.meeting_id,
        "speakerName": speaker_name,
        "speakerRole": speaker_role,
        "transcript": sse_text,
        "time": record["clientTime"] or now,
        "isFinal": body.is_final,
    }), name=f"sse-transcript-{body.meeting_id}")
    return JSONResponse({"success": True, "record": record})


@app.get("/api/meetings/{meeting_id}/transcripts/sse")
async def sse_meeting_transcripts(request: Request, meeting_id: str):
    """SSE 实时推送 — 桌面端替换轮询，手机端写入后秒级推送到桌面。
    EventSource 不支持自定义 header，token 通过 ?token= 查询参数传递。"""
    # 先尝试标准 Authorization header，失败则回退到查询参数 token
    token = request.query_params.get("token")
    if token:
        user = _get_user_from_auth_token(token, required=True)
    else:
        _get_request_user(request, required=True)
    from backend.config import sse_manager
    import json, asyncio as _asyncio

    q = sse_manager.subscribe(_safe_meeting_id(meeting_id))

    async def event_generator():
        try:
            yield f"data: {json.dumps({'type': 'connected', 'meetingId': meeting_id}, ensure_ascii=False)}\n\n"
            while True:
                if await request.is_disconnected():
                    break
                try:
                    payload = await _asyncio.wait_for(q.get(), timeout=30.0)
                    yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"
                except _asyncio.TimeoutError:
                    yield ": heartbeat\n\n"
        except _asyncio.CancelledError:
            pass
        finally:
            sse_manager.unsubscribe(_safe_meeting_id(meeting_id), q)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


@app.get("/api/meeting/transcripts/{meeting_id}")
async def get_meeting_transcripts(request: Request, meeting_id: str, limit: int = 200, offset: int = 0):
    _get_request_user(request, required=True)
    meeting = _db_load_transcripts_for_meeting(meeting_id)
    all_events = meeting.get("events", [])
    all_transcripts = meeting.get("transcripts", [])
    if offset:
        paginated_events = all_events[-(limit + offset):-offset]
        paginated_transcripts = all_transcripts[-(limit + offset):-offset]
    else:
        paginated_events = all_events[-limit:]
        paginated_transcripts = all_transcripts[-limit:]
    return JSONResponse({
        "success": True,
        "meetingId": meeting_id,
        "meetingTitle": meeting.get("meetingTitle", ""),
        "agenda": meeting.get("agenda", ""),
        "meetingPhase": meeting.get("phase", ""),
        "updatedAt": meeting.get("updatedAt"),
        "events": paginated_events,
        "transcripts": paginated_transcripts,
        "totalEvents": len(all_events),
        "totalTranscripts": len(all_transcripts),
    })


@app.post("/api/meetings/{meeting_id}/realtime-todos")
async def extract_realtime_todos(request: Request, meeting_id: str):
    """实时待办提取 —— 会议进行中调用，从最近转写中提取待办事项。"""
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    body = await request.json()
    recent_transcripts = body.get("transcripts", [])
    if not recent_transcripts:
        return JSONResponse({"success": True, "todos": []})
    # 取最近 20 条转写
    compact = [
        {
            "time": t.get("clientTime") or t.get("serverTime", ""),
            "speaker": t.get("speakerName") or t.get("speaker") or "",
            "text": (t.get("transcript") or t.get("text") or "")[:200],
        }
        for t in recent_transcripts[-20:]
    ]
    prompt = f"""从以下会议转写中提取待办事项（最多5条）。只提取明确的任务分配，忽略闲聊和试音。

转写内容：
{json.dumps(compact, ensure_ascii=False)}

严格按 JSON 格式输出：
{{"todos": [{{"task": "任务描述（≤40字）", "owner": "责任人", "priority": "高/中/低"}}]}}"""
    try:
        async with llm_semaphore:
            result = await llm._agenerate(
                [SystemMessage(content="你是待办提取助手。只输出JSON，不要解释。"), HumanMessage(content=prompt)],
                enable_thinking=False,
            )
        text = result.generations[0].message.content if result.generations else ""
        payload = _extract_json_object(text)
        todos = payload.get("todos", []) if payload else []
        return JSONResponse({"success": True, "todos": todos[:5]})
    except Exception as e:
        logger.warning("实时待办提取失败: %s", e)
        return JSONResponse({"success": True, "todos": []})


# ══════════════════════════════════════════════════════════════════════════════════
# 待办管理 API
# ══════════════════════════════════════════════════════════════════════════════════

def _sync_todos_to_table(meeting_id: str, meeting_title: str, todos: list, source: str = "ai"):
    """将 generatedRecords.todos 同步写入 meeting_todos 表。
    幂等：已有相同 task+owner 的待办保留用户修改的状态，只插入新增项。
    """
    safe_id = _safe_meeting_id(meeting_id)
    now = _now_text()
    try:
        with _db_connect() as conn:
            # 查出现有待办（task+owner 作为去重键）
            existing = {}
            for r in conn.execute(
                "SELECT id, task, owner FROM meeting_todos WHERE meeting_id = ? AND source = ?",
                (safe_id, source)
            ).fetchall():
                key = f"{(r['task'] or '').strip()}|{(r['owner'] or '').strip()}"
                existing[key] = r["id"]
            for t in todos:
                key = f"{(t.get('task') or '').strip()}|{(t.get('owner') or '').strip()}"
                if key in existing:
                    # 已存在，跳过（保留用户可能修改的状态）
                    continue
                todo_id = f"todo_{uuid.uuid4().hex[:10]}"
                conn.execute(
                    """INSERT INTO meeting_todos
                       (id, meeting_id, meeting_title, task, owner, deadline, priority,
                        status, source, reference, created_at, updated_at, payload_json)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                    (
                        todo_id, safe_id, meeting_title,
                        t.get("task", ""), t.get("owner", ""), t.get("deadline", ""),
                        t.get("priority", "中"), t.get("status", "待处理"),
                        source, t.get("reference", ""),
                        now, now, _json_dumps(t),
                    )
                )
    except Exception as e:
        logger.warning("待办同步失败 meeting=%s: %s", meeting_id, e)


@app.get("/api/todos")
async def list_todos(request: Request, status: str = "", owner: str = "",
                     priority: str = "", limit: int = 100, offset: int = 0):
    """跨会议待办列表。支持按 status/owner/priority 筛选。"""
    _get_request_user(request, required=True)
    conditions, params = [], []
    if status:
        conditions.append("status = ?")
        params.append(status)
    if owner:
        conditions.append("owner LIKE ?")
        params.append(f"%{owner}%")
    if priority:
        conditions.append("priority = ?")
        params.append(priority)
    where = "WHERE " + " AND ".join(conditions) if conditions else ""
    try:
        with _db_connect() as conn:
            rows = conn.execute(
                f"SELECT * FROM meeting_todos {where} ORDER BY "
                f"CASE priority WHEN '高' THEN 0 WHEN '中' THEN 1 ELSE 2 END, "
                f"created_at DESC LIMIT ? OFFSET ?",
                params + [limit, offset]
            ).fetchall()
            total = conn.execute(
                f"SELECT COUNT(*) as cnt FROM meeting_todos {where}", params
            ).fetchone()["cnt"]
        todos = []
        for r in rows:
            todos.append({
                "id": r["id"], "meetingId": r["meeting_id"], "meetingTitle": r["meeting_title"],
                "task": r["task"], "owner": r["owner"], "deadline": r["deadline"],
                "priority": r["priority"], "status": r["status"], "source": r["source"],
                "reference": r["reference"], "createdAt": r["created_at"],
                "updatedAt": r["updated_at"], "completedAt": r["completed_at"],
            })
        return JSONResponse({"success": True, "todos": todos, "total": total})
    except Exception as e:
        logger.error("待办列表查询失败: %s", e)
        return JSONResponse({"success": True, "todos": [], "total": 0})


@app.post("/api/meetings/{meeting_id}/todos")
async def create_meeting_todo(request: Request, meeting_id: str):
    """手动新增会议待办。"""
    _get_request_user(request, required=True)
    body = await request.json()
    task = (body.get("task") or "").strip()
    if not task:
        raise HTTPException(status_code=400, detail="待办内容不能为空")
    safe_id = _safe_meeting_id(meeting_id)
    now = _now_text()
    todo_id = f"todo_{uuid.uuid4().hex[:10]}"
    # 获取会议标题
    meeting_title = ""
    try:
        with _db_connect() as conn:
            row = conn.execute("SELECT title FROM meetings WHERE id = ?", (safe_id,)).fetchone()
            if row:
                meeting_title = row["title"] or ""
    except Exception:
        pass
    try:
        with _db_connect() as conn:
            conn.execute(
                """INSERT INTO meeting_todos
                   (id, meeting_id, meeting_title, task, owner, deadline, priority,
                    status, source, reference, created_at, updated_at, payload_json)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (
                    todo_id, safe_id, meeting_title, task,
                    body.get("owner", ""), body.get("deadline", ""),
                    body.get("priority", "中"), "待处理", "manual",
                    body.get("reference", ""), now, now, _json_dumps(body),
                )
            )
        return JSONResponse({"success": True, "id": todo_id})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/todos/{todo_id}")
async def update_todo(request: Request, todo_id: str):
    """更新待办（状态、截止时间等）。"""
    _get_request_user(request, required=True)
    body = await request.json()
    now = _now_text()
    updates, params = ["updated_at = ?"], [now]
    for field, col in [("task", "task"), ("owner", "owner"), ("deadline", "deadline"),
                       ("priority", "priority"), ("status", "status")]:
        if field in body:
            updates.append(f"{col} = ?")
            params.append(body[field])
    if body.get("status") in ("已完成", "已取消"):
        updates.append("completed_at = ?")
        params.append(now)
    params.append(todo_id)
    try:
        with _db_connect() as conn:
            conn.execute(f"UPDATE meeting_todos SET {', '.join(updates)} WHERE id = ?", params)
        return JSONResponse({"success": True})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/todos/{todo_id}")
async def delete_todo(request: Request, todo_id: str):
    """删除待办。"""
    _get_request_user(request, required=True)
    try:
        with _db_connect() as conn:
            conn.execute("DELETE FROM meeting_todos WHERE id = ?", (todo_id,))
        return JSONResponse({"success": True})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════════════════════════════
# 系统设置 API
# ══════════════════════════════════════════════════════════════════════════════════

@app.get("/api/settings")
async def get_settings(request: Request):
    """读取系统设置。"""
    _get_request_user(request, required=True)
    try:
        with _db_connect() as conn:
            org_name = _metadata_get(conn, "org_name") or ""
        return JSONResponse({"success": True, "orgName": org_name})
    except Exception as e:
        return JSONResponse({"success": True, "orgName": ""})


@app.put("/api/settings")
async def update_settings(request: Request):
    """更新系统设置（仅 admin）。"""
    user = _get_request_user(request, required=True)
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可修改系统设置")
    body = await request.json()
    try:
        with _db_connect() as conn:
            if "orgName" in body:
                _metadata_set(conn, "org_name", str(body["orgName"]).strip())
        return JSONResponse({"success": True})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def _build_whisper_chronicle(segments: list, whisper_event: dict) -> list:
    """从 Whisper segments 构建纪实列表，时间戳对齐会议时间轴。

    Whisper segments 的 start 是相对于合并音频的偏移。
    如果有 fileOffsets/fileDurations，说明合并了多个录音文件，
    需要将 segment 的 start 映射回原始文件，再加上该文件在会议时间轴中的偏移。
    """
    f_offsets = whisper_event.get("fileOffsets", {})
    f_durations = whisper_event.get("fileDurations", [])
    f_names = list(f_offsets.keys()) if f_offsets else []

    # 构建每个文件在合并音频中的起始位置
    cumulative = []
    cum = 0.0
    for dur in f_durations:
        cumulative.append(cum)
        cum += dur

    chronicle = []
    for seg in segments:
        t = seg.get("text", "").strip()
        if not t:
            continue
        seg_start = seg.get("start", 0)

        if f_names and cumulative:
            # 找到该 segment 属于哪个原始文件
            file_idx = 0
            for i in range(len(cumulative)):
                if i + 1 < len(cumulative):
                    if cumulative[i] <= seg_start < cumulative[i + 1]:
                        file_idx = i
                        break
                else:
                    file_idx = i
            # 会议时间轴时间 = (segment 在合并音频中的位置 - 文件在合并音频中的起始) + 文件在会议中的偏移
            fname = f_names[file_idx] if file_idx < len(f_names) else ""
            file_offset = f_offsets.get(fname, 0)
            file_start_in_merged = cumulative[file_idx] if file_idx < len(cumulative) else 0
            aligned_secs = max(0, int(seg_start - file_start_in_merged + file_offset))
        else:
            # 单文件或无偏移信息，直接用 segment 的 start
            aligned_secs = max(0, int(seg_start))

        mins, secs = divmod(aligned_secs, 60)
        chronicle.append({
            "time": f"{mins:02d}:{secs:02d}",
            "speaker": "Whisper 终审",
            "role": "AI 高精度转写",
            "content": t,
        })
    return chronicle


@app.get("/api/meetings/{meeting_id}/records")
async def generate_meeting_records(request: Request, meeting_id: str, force: bool = False):
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    # 优先用缓存（force=true 时跳过缓存强制重新生成）
    cached = meeting.get("generatedRecords")
    if not force and cached and cached.get("generated") and cached.get("cachedAt"):
        records = cached
    else:
        if force:
            _invalidate_transcripts_cache()
        with MEETING_TRANSCRIPTS_LOCK:
            transcripts_data = _load_meeting_transcripts().get(safe_id, {"events": [], "transcripts": []})
        transcripts = transcripts_data.get("transcripts", [])
        events = transcripts_data.get("events", [])
        local_records = _local_generate_meeting_records(meeting, transcripts, events)
        if not local_records.get("generated"):
            return JSONResponse({"success": True, "meetingId": safe_id, "records": local_records})
        ai_records = await _deepseek_generate_meeting_records(meeting, transcripts, events)
        records = ai_records or local_records
        records["cachedAt"] = _now_text()
        # 标记 Whisper 终审是否已注入，并用 Whisper segments 替换纪实
        whisper_event = None
        for e in events:
            if e.get("type") == "transcript" and e.get("action") == "whisper-review":
                whisper_event = e
                break
        if whisper_event:
            records["whisperEnhanced"] = True
            # 用 Whisper segments 构建高精度纪实（时间戳对齐会议时间轴）
            segs = whisper_event.get("segments", [])
            if segs:
                whisper_chronicle = _build_whisper_chronicle(segs, whisper_event)
                if whisper_chronicle:
                    records["chronicle"] = whisper_chronicle
        with MEETINGS_LOCK:
            ms = _load_meetings()
            m2 = ms.get(safe_id, {})
            old_gr = m2.get("generatedRecords") or {}
            # 保留 whisperDocx 元数据（force 重新生成时不应丢失）
            if old_gr.get("whisperDocx"):
                records["whisperDocx"] = old_gr["whisperDocx"]
            m2["generatedRecords"] = records
            ms[safe_id] = m2
            _save_meetings(ms)
            _invalidate_meetings_cache()
            # 同步待办到 meeting_todos 表
            _todos = records.get("todos") or []
            if _todos:
                _sync_todos_to_table(safe_id, m2.get("title", ""), _todos, source="ai")
    # merge user overrides from events
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting_latest = meetings.get(safe_id, meeting)
    overrides = None
    for event in reversed(meeting_latest.get("events", [])):
        if event.get("type") == "records_override":
            overrides = event.get("payload") if isinstance(event.get("payload"), dict) else {}
            break
    if overrides:
        records = {**records, **overrides}
    return JSONResponse({
        "success": True,
        "meetingId": safe_id,
        "records": records,
    })


@app.post("/api/meetings/{meeting_id}/records/update")
async def update_meeting_records(request: Request, meeting_id: str, body: MeetingRecordsUpdateRequest):
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    _check_meeting_access(user, meeting)
    override = {}
    if body.summary is not None:
        override["summary"] = body.summary
    if body.minutes is not None:
        override["minutes"] = body.minutes
    if body.decisions is not None:
        override["decisions"] = body.decisions
    if body.todos is not None:
        override["todos"] = body.todos
    event = {
        "id": f"records_override_{_now_text()}",
        "type": "records_override",
        "serverTime": _now_text(),
        "payload": override,
    }
    _append_meeting_activity(safe_id, event)
    # 保存版本记录
    try:
        now = _now_text()
        editor_name = user.get("name") or user.get("username", "")
        with _db_connect() as conn:
            max_ver = conn.execute(
                "SELECT COALESCE(MAX(version), 0) as v FROM meeting_record_versions WHERE meeting_id = ?", (safe_id,)
            ).fetchone()
            new_ver = (max_ver["v"] if max_ver else 0) + 1
            conn.execute(
                """INSERT INTO meeting_record_versions
                   (id, meeting_id, version, editor, edit_summary, records_json, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (f"ver_{uuid.uuid4().hex[:10]}", safe_id, new_ver, editor_name,
                 f"编辑了{', '.join(k for k in override.keys())}", json.dumps(override, ensure_ascii=False), now)
            )
    except Exception as e:
        logger.warning("版本保存失败: %s", e)
    return JSONResponse({"success": True, "meetingId": safe_id, "message": "记录已保存"})


@app.get("/api/meetings/{meeting_id}/versions")
async def get_meeting_versions(request: Request, meeting_id: str):
    """获取会议纪要版本列表。"""
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    try:
        with _db_connect() as conn:
            rows = conn.execute(
                "SELECT id, version, editor, edit_summary, created_at FROM meeting_record_versions WHERE meeting_id = ? ORDER BY version DESC",
                (safe_id,)
            ).fetchall()
        versions = [dict(r) for r in rows]
        return JSONResponse({"success": True, "versions": versions})
    except Exception as e:
        return JSONResponse({"success": True, "versions": []})


@app.get("/api/meetings/{meeting_id}/versions/{version}")
async def get_meeting_version(request: Request, meeting_id: str, version: int):
    """获取指定版本的纪要内容。"""
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    try:
        with _db_connect() as conn:
            row = conn.execute(
                "SELECT * FROM meeting_record_versions WHERE meeting_id = ? AND version = ?",
                (safe_id, version)
            ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="版本不存在")
        records = json.loads(row["records_json"]) if row["records_json"] else {}
        return JSONResponse({"success": True, "version": row["version"], "editor": row["editor"],
                             "editSummary": row["edit_summary"], "createdAt": row["created_at"], "records": records})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/meetings/{meeting_id}/markers")
async def add_meeting_marker(request: Request, meeting_id: str, body: MeetingMarkerRequest):
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    _check_meeting_access(user, meeting)
    marker_event = {
        "id": f"marker_{uuid.uuid4().hex[:12]}",
        "type": "marker",
        "serverTime": _now_text(),
        "payload": {
            "markerType": body.marker_type,
            "agendaId": body.agenda_id,
            "agendaTitle": body.agenda_title,
            "transcriptId": body.transcript_id,
            "transcriptText": body.transcript_text,
            "transcriptTime": body.transcript_time,
            "transcriptSpeaker": body.transcript_speaker,
            "note": body.note,
            "createdBy": _resolve_meeting_role(user).get("displayName", "记录员"),
        },
    }
    _append_meeting_activity(safe_id, marker_event)
    return JSONResponse({"success": True, "meetingId": safe_id, "markerId": marker_event["id"], "message": "标记已保存"})


@app.get("/api/meetings/{meeting_id}/markers")
async def get_meeting_markers(request: Request, meeting_id: str):
    _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    markers = [
        {"id": event.get("id"), **event.get("payload", {})}
        for event in meeting.get("events", [])
        if event.get("type") == "marker"
    ]
    return JSONResponse({"success": True, "meetingId": safe_id, "markers": markers})


@app.delete("/api/meetings/{meeting_id}/markers/{marker_id}")
async def delete_meeting_marker(request: Request, meeting_id: str, marker_id: str):
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    safe_marker_id = re.sub(r"[^a-zA-Z0-9_-]", "", marker_id)
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id)
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    _check_meeting_access(user, meeting)
    events = meeting.get("events", [])
    new_events = [e for e in events if e.get("id") != safe_marker_id]
    if len(new_events) == len(events):
        raise HTTPException(status_code=404, detail="标记不存在")
    meeting["events"] = new_events
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meetings[safe_id] = meeting
        _save_meetings(meetings)
    return JSONResponse({"success": True, "meetingId": safe_id, "markerId": safe_marker_id, "message": "标记已删除"})


@app.post("/api/meeting/transcripts/{meeting_id}/{transcript_id}/correction")
async def correct_meeting_transcript(request: Request, meeting_id: str, transcript_id: str, body: MeetingTranscriptCorrectionRequest):
    user = _get_request_user(request, required=True)
    corrected = re.sub(r"\s+", " ", body.corrected_transcript or "").strip()
    signature_data = (body.signature_data or "").strip()
    if not corrected:
        raise HTTPException(status_code=400, detail="修正后的发言不能为空")
    if not signature_data.startswith("data:image/"):
        raise HTTPException(status_code=400, detail="请先完成手机手写签名")

    now = _now_text()
    role = _resolve_meeting_role(user)
    with MEETING_TRANSCRIPTS_LOCK:
        data = _load_meeting_transcripts()
        meeting = data.get(meeting_id)
        if not meeting:
            raise HTTPException(status_code=404, detail="会议转写不存在")
        target = None
        for record in meeting.get("transcripts", []):
            if record.get("id") == transcript_id:
                target = record
                break
        if not target:
            raise HTTPException(status_code=404, detail="转写记录不存在")
        owner_username = (target.get("username") or "").lower()
        current_username = (user.get("username") or "").lower()
        if user.get("role") != "admin" and owner_username and owner_username != current_username:
            raise HTTPException(status_code=403, detail="只能修正并签署本人发言")

        original_text = target.get("originalTranscript") or target.get("transcript") or ""
        target["originalTranscript"] = original_text
        target["transcript"] = corrected
        target["correctedTranscript"] = corrected
        target["correctionSigned"] = True
        target["correctionSignedAt"] = now
        target["correctionClientTime"] = body.client_time
        target["correctionAuthor"] = role["displayName"]
        target["correctionUsername"] = role["username"]
        target["signatureData"] = signature_data[:800000]
        # 同步写入专用列（主存储），payload_json 由 _save_meeting_transcripts 处理
        try:
            with _db_connect() as conn:
                conn.execute(
                    """UPDATE meeting_transcripts SET
                       correction_signed=1, correction_signed_at=?,
                       signature_data=?, corrected_transcript=?
                       WHERE id=? AND meeting_id=?""",
                    (now, signature_data[:800000], corrected, transcript_id, _safe_meeting_id(meeting_id))
                )
        except Exception as e:
            logger.warning("签名专用列写入失败（payload_json 已保存）: %s", e)
        event = {
            "id": f"correction_{uuid.uuid4().hex[:10]}",
            "type": "transcript-correction",
            "meetingId": meeting_id,
            "transcriptId": transcript_id,
            "speakerName": target.get("speakerName"),
            "originalTranscript": original_text,
            "correctedTranscript": corrected,
            "signedAt": now,
            "signer": {
                "displayName": role["displayName"],
                "username": role["username"],
                "meetingRole": role["meetingRole"],
                "seat": role["seat"],
            },
            "serverTime": now,
        }
        meeting.setdefault("events", []).append(event)
        meeting["updatedAt"] = now
        _save_meeting_transcripts(data)
    _append_meeting_activity(meeting_id, event)
    return JSONResponse({"success": True, "record": target, "event": event})


@app.post("/api/meeting/transcripts/{meeting_id}/{transcript_id}/speaker")
async def update_transcript_speaker(request: Request, meeting_id: str, transcript_id: str, body: dict):
    """Update speaker identity on a transcript — used for post-hoc diarization correction."""
    user = _get_request_user(request, required=True)
    speaker_name = str(body.get("speakerName") or "").strip()
    speaker_role = str(body.get("speakerRole") or "").strip()
    speaker_dept = str(body.get("speakerDept") or "").strip()
    if not speaker_name:
        raise HTTPException(status_code=400, detail="发言人姓名不能为空")

    now = _now_text()
    with _db_connect() as conn:
        row = conn.execute(
            "SELECT id, speaker_name, speaker_role, speaker_dept FROM meeting_transcripts WHERE id = ? AND meeting_id = ?",
            (transcript_id, _safe_meeting_id(meeting_id)),
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="转写记录不存在")

    old_name, old_role, old_dept = row["speaker_name"], row["speaker_role"], row["speaker_dept"]
    with _db_connect() as conn:
        conn.execute(
            "UPDATE meeting_transcripts SET speaker_name = ?, speaker_role = ?, speaker_dept = ?, payload_json = json_set(payload_json, '$.speakerCorrected', ?) WHERE id = ?",
            (speaker_name, speaker_role, speaker_dept, "true", transcript_id),
        )
        conn.execute("COMMIT")
    _invalidate_transcripts_cache()

    role = _resolve_meeting_role(user)
    event = {
        "id": f"speaker_correction_{uuid.uuid4().hex[:10]}",
        "type": "speaker-correction",
        "meetingId": meeting_id,
        "transcriptId": transcript_id,
        "oldSpeaker": {"name": old_name, "role": old_role, "dept": old_dept},
        "newSpeaker": {"name": speaker_name, "role": speaker_role, "dept": speaker_dept},
        "correctedBy": role["displayName"],
        "serverTime": now,
    }
    _append_meeting_activity_light(meeting_id, event)
    return JSONResponse({"success": True, "transcriptId": transcript_id, "speakerName": speaker_name, "speakerRole": speaker_role, "speakerDept": speaker_dept, "event": event})


_ASR_RECONNECT_BASE_DELAY = 1.0   # seconds
_ASR_RECONNECT_MAX_DELAY = 16.0   # cap exponential backoff
_ASR_RECONNECT_MAX_RETRIES = 5    # per connection attempt window


async def _dashscope_connect(url: str, headers: Dict[str, str], retries: int = 3):
    """Connect to DashScope with retry on transient failures."""
    import websockets
    last_err = None
    for attempt in range(retries):
        try:
            return await websockets.connect(
                url,
                additional_headers=headers,
                max_size=8 * 1024 * 1024,
                ping_interval=20,
                ping_timeout=20,
            )
        except TypeError:
            try:
                return await websockets.connect(
                    url,
                    extra_headers=headers,
                    max_size=8 * 1024 * 1024,
                    ping_interval=20,
                    ping_timeout=20,
                )
            except Exception as e:
                last_err = e
        except Exception as e:
            last_err = e
        if attempt < retries - 1:
            delay = min(_ASR_RECONNECT_BASE_DELAY * (2 ** attempt), _ASR_RECONNECT_MAX_DELAY)
            await asyncio.sleep(delay)
    raise last_err or RuntimeError("Failed to connect to DashScope")


def _dashscope_text_result(message: dict) -> dict:
    output = ((message.get("payload") or {}).get("output") or {})
    sentence = output.get("sentence") or {}
    text = sentence.get("text") or output.get("text") or ""
    return {
        "text": text,
        "isFinal": bool(sentence.get("sentence_end") or sentence.get("sentenceEnd") or output.get("is_final")),
        "beginTime": sentence.get("begin_time") or sentence.get("beginTime"),
        "endTime": sentence.get("end_time") or sentence.get("endTime"),
    }


@app.websocket("/api/meeting/asr/ws")
async def meeting_asr_websocket(websocket: WebSocket):
    await websocket.accept()
    token = websocket.query_params.get("token", "")
    try:
        user = _get_user_from_auth_token(token, required=True)
    except HTTPException as exc:
        await websocket.send_json({"type": "error", "message": exc.detail})
        await websocket.close(code=4401)
        return

    api_key = DASHSCOPE_API_KEY
    if not api_key:
        await websocket.send_json({
            "type": "error",
            "message": "未配置 DASHSCOPE_API_KEY，已无法连接 Fun-ASR。",
        })
        await websocket.close(code=1011)
        return

    meeting_id = websocket.query_params.get("meetingId") or "meeting-gxq-fc-2026-02"
    meeting_title = websocket.query_params.get("meetingTitle") or ""
    agenda = websocket.query_params.get("agenda") or ""
    user_role = _resolve_meeting_role(user)
    task_id = f"asr_{uuid.uuid4().hex}"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "user-agent": "ai-compliance-demo/fun-asr-mobile-recorder",
    }
    if DASHSCOPE_WORKSPACE:
        headers["X-DashScope-WorkSpace"] = DASHSCOPE_WORKSPACE

    dash_ws = None
    dash_to_client_task = None
    task_started = asyncio.Event()
    reconnect_attempt = 0

    async def _run_dashscope_session():
        """Connect to DashScope and relay results to client. Returns True on clean finish."""
        nonlocal dash_ws, task_id
        dash_ws = await _dashscope_connect(DASHSCOPE_FUN_ASR_WS_URL, headers)
        task_id = f"asr_{uuid.uuid4().hex}"  # new task per session
        run_task = {
            "header": {"action": "run-task", "task_id": task_id, "streaming": "duplex"},
            "payload": {
                "task_group": "audio", "task": "asr", "function": "recognition",
                "model": os.environ.get("DASHSCOPE_FUN_ASR_MODEL", "paraformer-realtime-v2"),
                "parameters": {"format": "pcm", "sample_rate": 16000, "language_hints": ["zh"]},
                "input": {},
            },
        }
        await dash_ws.send(json.dumps(run_task, ensure_ascii=False))

        async def dashscope_to_client():
            async for raw_message in dash_ws:
                if isinstance(raw_message, bytes):
                    continue
                try:
                    message = json.loads(raw_message)
                except Exception:
                    await websocket.send_json({"type": "raw", "message": raw_message})
                    continue
                event = (message.get("header") or {}).get("event") or ""
                if event == "task-started":
                    task_started.set()
                    await websocket.send_json({
                        "type": "ready", "taskId": task_id,
                        "meetingId": meeting_id, "speaker": user_role,
                    })
                elif event == "result-generated":
                    result = _dashscope_text_result(message)
                    await websocket.send_json({
                        "type": "result", "taskId": task_id,
                        "meetingId": meeting_id, "meetingTitle": meeting_title,
                        "agenda": agenda, **result,
                    })
                elif event == "task-finished":
                    await websocket.send_json({"type": "finished", "taskId": task_id})
                    return True
                elif event in ("task-failed", "error"):
                    err = (message.get("header") or {}).get("error_message") or "Fun-ASR 错误"
                    await websocket.send_json({"type": "error", "message": err})
                    return False

        nonlocal dash_to_client_task
        dash_to_client_task = asyncio.create_task(dashscope_to_client())
        await asyncio.wait_for(task_started.wait(), timeout=15)
        return True

    async def _cleanup_session():
        nonlocal dash_to_client_task, dash_ws
        if dash_to_client_task and not dash_to_client_task.done():
            dash_to_client_task.cancel()
        if dash_ws:
            try:
                await dash_ws.close()
            except Exception:
                pass

    try:
        await _run_dashscope_session()

        # ── Audio forwarding loop with automatic DashScope reconnection ──────
        audio_buffer = []  # queue audio during reconnection
        while True:
            try:
                client_message = await websocket.receive()
            except WebSocketDisconnect:
                break
            if client_message.get("type") == "websocket.disconnect":
                break
            audio_bytes = client_message.get("bytes")
            if audio_bytes:
                # 缓冲音频（不管连接状态都缓存，重连后回放）
                audio_buffer.append(audio_bytes)
                if len(audio_buffer) > 100:
                    audio_buffer = audio_buffer[-50:]  # keep last ~2.5s at 20ms/frame
                if not dash_ws:
                    # 连接已断，等待重连完成
                    continue
                try:
                    await dash_ws.send(audio_bytes)
                except Exception:
                    # DashScope connection lost — reconnect transparently
                    reconnect_attempt += 1
                    if reconnect_attempt > _ASR_RECONNECT_MAX_RETRIES:
                        await websocket.send_json({"type": "error", "message": "Fun-ASR 重连次数超限"})
                        break
                    delay = min(_ASR_RECONNECT_BASE_DELAY * (2 ** (reconnect_attempt - 1)), _ASR_RECONNECT_MAX_DELAY)
                    logger.warning("Fun-ASR 连接断开，%ds 后重连 (第 %d 次)", delay, reconnect_attempt)
                    await websocket.send_json({"type": "reconnecting", "attempt": reconnect_attempt, "delay": delay})
                    await _cleanup_session()
                    task_started.clear()
                    await asyncio.sleep(delay)
                    try:
                        await _run_dashscope_session()
                        reconnect_attempt = 0  # reset on success
                        await websocket.send_json({"type": "reconnected", "taskId": task_id})
                        # Re-send buffered audio (last 2s worth)
                        for buf_bytes in audio_buffer[-40:]:
                            try:
                                await dash_ws.send(buf_bytes)
                            except Exception:
                                break
                    except Exception as e:
                        logger.exception("Fun-ASR 重连失败")
                        await websocket.send_json({"type": "error", "message": f"重连失败: {e}"})
                        continue
                continue
            text = client_message.get("text")
            if not text:
                continue
            try:
                command = json.loads(text)
            except Exception:
                continue
            if command.get("type") == "finish":
                finish_task = {
                    "header": {"action": "finish-task", "task_id": task_id, "streaming": "duplex"},
                    "payload": {"input": {}},
                }
                try:
                    await dash_ws.send(json.dumps(finish_task, ensure_ascii=False))
                except Exception:
                    pass
                break
        if dash_to_client_task:
            try:
                await asyncio.wait_for(dash_to_client_task, timeout=8)
            except asyncio.TimeoutError:
                pass
    except Exception as exc:
        logger.exception("Fun-ASR websocket proxy failed")
        try:
            await websocket.send_json({"type": "error", "message": f"Fun-ASR 连接失败：{exc}"})
        except Exception:
            pass
    finally:
        await _cleanup_session()
        try:
            await websocket.close()
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════════════════════════
# Qwen3-ASR WebSocket 端点 — 本地实时语音识别
# ═══════════════════════════════════════════════════════════════════════════════

# SenseVoice 标签清洗：只杀模型内部尖括号标签，不碰任何一个汉字
_TOK_RE = re.compile(r'<\|.*?\|>')

# ── 工业级流水线：Energy Gate + VAD-lite ──

_GOVERNMENT_ASR_HOTWORDS = [
    "党委会", "党组会", "三重一大", "党委前置", "前置研究", "集体决策", "会议纪要", "议题",
    "审议", "研究", "讨论", "表决", "通过", "同意", "原则同意", "暂缓", "再议",
    "法务审查", "合规审查", "纪检监督", "风险评估", "可研报告", "资金测算", "预算控制",
    "重大项目安排", "大额度资金运作", "重要人事任免", "重大事项决策", "项目立项",
    "招投标", "合同签订", "工程变更", "付款审批", "资产处置", "安全生产", "消防改造",
    "党委办公室", "总经理办公室", "财务部", "法务部", "合规部", "审计部", "项目管理部",
    "人力资源部", "纪检监察部", "党群工作部", "综合管理部",
]

_UNIT_ASR_HOTWORD_TEMPLATE = [
    "党委会议", "党委会审议", "党组会议", "前置研究讨论", "书记办公会", "总经理办公会", "班子会",
    "纪委", "纪检组", "监督检查", "整改落实", "责任追究", "议题申报", "会前沟通",
    "可研", "立项", "招采", "采购审批", "合同审批", "工程签证", "预算调整", "资金计划",
    "投标文件", "中标通知书", "法务意见", "审计意见", "风险提示", "合规意见",
    "人事任免", "干部考察", "民主推荐", "任前公示", "试用期", "绩效考核",
    "消防验收", "安全生产责任", "资产盘点", "资产处置", "信息化建设", "系统上线", "数据治理",
]

_SYNONYM_CORRECTIONS = {
    "三种一大": "三重一大",
    "党外前置": "党委前置",
    "党委钱置": "党委前置",
    "法务审查": "法务审查",
    "发务审查": "法务审查",
    "合贵审查": "合规审查",
    "合规市查": "合规审查",
    "纪检监督": "纪检监督",
    "纪检监都": "纪检监督",
    "项目立响": "项目立项",
    "项目立项": "项目立项",
    "招投标": "招投标",
    "招头标": "招投标",
    "合同签定": "合同签订",
    "合同签订": "合同签订",
    "资金测算": "资金测算",
    "资金额算": "资金测算",
    "预算控置": "预算控制",
    "预算控制": "预算控制",
    "可行性研究报告": "可行性研究报告",
    "可研报告": "可研报告",
    "党委办": "党委办公室",
    "总办": "总经理办公室",
    "法务办": "法务部",
    "合规办": "合规部",
}


def _build_asr_hotwords(meeting_title: str = "", agenda: str = "", project: str = "", extra: Optional[list[str]] = None) -> list[str]:
    words = list(dict.fromkeys(_GOVERNMENT_ASR_HOTWORDS + _UNIT_ASR_HOTWORD_TEMPLATE + _load_asr_custom_hotwords()))
    for text in (meeting_title, agenda, project):
        for part in re.split(r"[\s，,、；;：:。.!?！？（）()【】\[\]《》\"'‘’/\\-]+", str(text or "")):
            part = part.strip()
            if len(part) >= 2 and part not in words:
                words.append(part)
    if extra:
        for word in extra:
            word = str(word or "").strip()
            if len(word) >= 2 and word not in words:
                words.append(word)
    return words[:160]


def _apply_asr_homophone_corrections(text: str) -> str:
    t = str(text or "")
    corrections = dict(_SYNONYM_CORRECTIONS)
    corrections.update(_load_asr_corrections())
    for wrong, right in corrections.items():
        t = t.replace(wrong, right)
    t = re.sub(r"\b(党委|党组|法务|合规|纪检|预算|资金|项目|合同|招投标|人事)\s*([办部处]|审查|前置|立项|控制|测算)\b", lambda m: m.group(0).replace(" ", ""), t)
    return t


def _load_json_list(path: Path) -> list:
    try:
        if path.exists():
            return json.loads(path.read_text(encoding="utf-8")) or []
    except Exception:
        pass
    return []


def _save_json_list(path: Path, data: list) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _load_asr_custom_hotwords() -> list[str]:
    items = _load_json_list(ASR_HOTWORDS_DB)
    words: list[str] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        if item.get("enabled", True) is False:
            continue
        word = str(item.get("word") or "").strip()
        if word and word not in words:
            words.append(word)
    return words


def _load_asr_corrections() -> dict[str, str]:
    items = _load_json_list(ASR_CORRECTIONS_DB)
    mapping: dict[str, str] = {}
    for item in items:
        if not isinstance(item, dict):
            continue
        if item.get("enabled", True) is False:
            continue
        wrong = str(item.get("wrong") or "").strip()
        right = str(item.get("right") or "").strip()
        if wrong and right:
            mapping[wrong] = right
    return mapping


def _energy_gate(pcm_bytes: bytes, threshold: int = 150) -> bool:
    """动态能量门 + ZCR 联合拦截：
    能量 AND ZCR 双重判定 — 噪声即使有能量，ZCR 不过也拦；语音即使轻，ZCR 过就能放。
    Returns False = 拦截，True = 放行。"""
    import struct as _st
    n = min(len(pcm_bytes) // 2, 200)
    if n < 10:
        return False
    samples = _st.unpack(f'<{n}h', pcm_bytes[:n*2])
    peak = max(abs(s) for s in samples)
    rms = int((sum(s*s for s in samples) / n) ** 0.5)
    # ZCR：噪声 ZCR 极端（<0.02 直流偏移 或 >0.6 高频噪声），人声居中
    zcr = sum(1 for i in range(1, n) if (samples[i] >= 0) != (samples[i-1] >= 0))
    zcr_norm = zcr / n
    # 联合判定：必须同时满足能量 AND 语音 ZCR
    has_energy = rms >= threshold or peak >= 400
    has_voice_zcr = 0.02 < zcr_norm < 0.65
    return has_energy and has_voice_zcr

def _vad_lite(pcm_bytes: bytes) -> bool:
    """已合并到 _energy_gate 联合判定，保留别名。"""
    return True

# 重复 Token 熔断正则：连续 6 个相同汉字
_REPEAT_RE = re.compile(r'(.)\1{5,}')
# v5: 断连时暂存未提交文本，值格式 {"text": str, "timestamp": float, "version": int}
_asr_pending_store: dict = {}
_ASR_PENDING_TTL_SEC = 5 * 60  # 5 分钟过期

# 僵尸 Session 清洗：以 meeting_id+username 为 Key，防止重连产生僵尸缓存
_ACTIVE_ASR_SESSIONS: dict = {}  # key → session_id

# v5: QwenASRClient 应用级单例（避免每个 WS 连接实例化独立 httpx client）
_qwen_asr_client = None  # type: Optional[QwenASRClient]
_qwen_client_lock = asyncio.Lock()

async def _get_qwen_client() -> "QwenASRClient":
    """获取 QwenASRClient 单例（double-check locking）。"""
    from backend.qwen_asr_client import QwenASRClient
    global _qwen_asr_client
    if _qwen_asr_client is not None:
        return _qwen_asr_client
    async with _qwen_client_lock:
        if _qwen_asr_client is not None:
            return _qwen_asr_client
        _qwen_asr_client = QwenASRClient(base_url=QWEN_ASR_URL, chunk_timeout=5.0)
        return _qwen_asr_client

async def _cleanup_asr_pending_store():
    """后台任务：定期清理 _asr_pending_store 中过期条目（每 60s）。"""
    while True:
        await asyncio.sleep(60)
        now = time.time()
        expired = [
            k for k, v in _asr_pending_store.items()
            if isinstance(v, dict) and now - v.get("timestamp", 0) > _ASR_PENDING_TTL_SEC
        ]
        for k in expired:
            _asr_pending_store.pop(k, None)
        if expired:
            logger.info("ASR pending store: cleaned %d expired entries", len(expired))

@app.websocket("/api/meeting/asr/qwen/ws")
async def meeting_asr_qwen_websocket(websocket: WebSocket):
    """本地 ASR WebSocket（FunASR SenseVoiceSmall）。

    协议: type: "final" + newText/fullText（与 DashScope 兼容，前端无需改动）
    先精准清洗标签 → _fuzzy_lcp 增量提取 → 句尾或 ≥6 字 → 推送 final
    """
    from backend.qwen_asr_client import (
        QwenASRClient,
        ASRError, ASRUnavailableError, ASRSessionExpiredError, ASRChunkTimeoutError,
    )

    await websocket.accept()

    token = websocket.query_params.get("token", "")
    try:
        user = _get_user_from_auth_token(token, required=True)
    except HTTPException as exc:
        await websocket.send_json({"type": "error", "message": exc.detail})
        await websocket.close(code=4401)
        return

    meeting_id = websocket.query_params.get("meetingId") or "meeting-gxq-fc-2026-02"
    meeting_title = websocket.query_params.get("meetingTitle") or ""
    agenda = websocket.query_params.get("agenda") or ""
    user_role = _resolve_meeting_role(user)
    task_id = f"asr_{uuid.uuid4().hex}"

    qwen_client = await _get_qwen_client()
    if not await qwen_client.is_available():
        await websocket.send_json({"type": "error", "message": f"本地 ASR 不可用 ({QWEN_ASR_URL})"})
        await websocket.close(code=1011)
        return

    def _clean(text: str) -> str:
        t = _TOK_RE.sub('', text).strip()
        t = re.sub(r'([。，？！,?!])\1+', r'\1', t)
        t = _apply_asr_homophone_corrections(t)
        return t

    def _fuzzy_lcp(a: str, b: str, tolerance: int = 2) -> int:
        n, mismatches = min(len(a), len(b)), 0
        i = 0
        while i < n:
            if a[i] != b[i]:
                mismatches += 1
                if mismatches > tolerance:
                    return i - mismatches + 1 if i >= mismatches else 0
            i += 1
        return i

    session_id = None
    chunk_count = 0
    stale_count = 0       # 流式失活计数
    prev_full_text = ""   # 上一帧全量文本，熔断规则 B
    repeat_burst = 0
    committed_text = ""   # LCP 参考锚点 — 与 session_id 同生共死
    pending_buffer = ""
    last_full_text = ""
    last_change_time = time.monotonic()
    sent_tail = ""
    spk_id = ""
    ui_bubble_start_idx = 0  # 当前气泡在全文中的起始位置
    recv_loop_alive = asyncio.Event()
    # 滑动音频缓冲区：攒够 600ms（19200 bytes）再发包
    audio_float_buffer = bytearray()
    ASR_CHUNK_BYTES = 16000  # 8000 samples * 2 = 500ms @ 16kHz
    # Silero VAD（替代 Energy Gate）
    from backend.silero_vad import SileroVAD
    silero_vad = SileroVAD(threshold=0.5)

    # 声纹识别（可选）
    from backend.voiceprint import get_voiceprint_engine as _get_vp_engine
    _vp_engine = _get_vp_engine()
    _vp_enrolled: dict = {}  # {user_id: embedding} 缓存
    _vp_audio_buffer = bytearray()  # 声纹音频缓冲区
    _vp_identified_user: Optional[str] = None  # 当前识别结果
    _vp_identified_name: str = ""
    _vp_identified_confidence: float = 0.0
    _vp_identified_by: str = "manual"
    VP_BUFFER_BYTES = 64000  # ~2秒 @ 16kHz 16-bit

    if _vp_engine and _vp_engine.is_ready:
        try:
            from backend.db import _db_load_voiceprint_profiles
            from backend.voiceprint import deserialize_embedding
            _profiles = _db_load_voiceprint_profiles()
            for _p in _profiles:
                _vp_enrolled[_p["user_id"]] = deserialize_embedding(_p["embedding"])
            if _vp_enrolled:
                logger.info("声纹引擎就绪，已加载 %d 个声纹", len(_vp_enrolled))
        except Exception as vp_load_err:
            logger.warning("声纹加载失败: %s", vp_load_err)
            _vp_engine = None  # 降级：不做声纹识别
    else:
        _vp_engine = None

    # Session 门禁：清理同一用户+会议下的僵尸 session
    session_key = f"{user.get('username', '')}_{meeting_id}"
    old_sid = _ACTIVE_ASR_SESSIONS.pop(session_key, None)
    if old_sid:
        logger.warning("Session门禁: 发现僵尸 %s，finish 旧会话", old_sid)
        try:
            await qwen_client.finish(old_sid)
        except Exception:
            pass

    asr_task = None
    try:
        # Lazy start：延后到检测到真正语音时才创建 ASR session
        # 避免环境噪声被 Paraformer 幻觉成 "对对对""没有没有"
        await websocket.send_json({
            "type": "ready", "taskId": task_id,
            "meetingId": meeting_id, "speaker": user_role, "backend": "paraformer",
        })

        # 断连恢复
        if websocket.query_params.get("resume") and session_key in _asr_pending_store:
            entry = _asr_pending_store.pop(session_key)
            pending_resume = entry["text"] if isinstance(entry, dict) else entry
            if pending_resume:
                await websocket.send_json(_vp_enrich({
                    "type": "final", "taskId": task_id,
                    "meetingId": meeting_id, "meetingTitle": meeting_title,
                    "agenda": agenda, "newText": pending_resume, "fullText": pending_resume,
                    "isFinal": True, "backend": "paraformer",
                }))
                sent_tail = (sent_tail + pending_resume)[-50:]

        # ── 800ms 静默超时：动作 A — 同生共死硬切断 ──
        async def _silence_flush_monitor():
            nonlocal pending_buffer, sent_tail, last_change_time
            nonlocal committed_text, last_full_text, ui_bubble_start_idx
            nonlocal session_id
            while recv_loop_alive.is_set():
                await asyncio.sleep(0.2)
                if not recv_loop_alive.is_set():
                    return
                if time.monotonic() - last_change_time > 0.8 and pending_buffer:
                    if not (pending_buffer == sent_tail[-len(pending_buffer):] and len(pending_buffer) >= 3):
                        # 发完最后的气泡
                        bubble_text = last_full_text[ui_bubble_start_idx:]
                        if bubble_text.strip():
                            await websocket.send_json(_vp_enrich({
                                "type": "final", "taskId": task_id,
                                "meetingId": meeting_id, "meetingTitle": meeting_title,
                                "agenda": agenda, "newText": bubble_text, "fullText": last_full_text,
                                "isFinal": True, "backend": "paraformer", "spk": spk_id,
                            }))
                    # 💥 同生共死：GPU 重置，Python 同步归零
                    try:
                        await qwen_client.finish(session_id)
                    except Exception:
                        pass
                    session_id = await qwen_client.start(hotwords=_build_asr_hotwords(meeting_title=meeting_title, agenda=agenda, project="", extra=[user_role.get("displayName", ""), user_role.get("meetingRole", "")]))
                    _ACTIVE_ASR_SESSIONS[session_key] = session_id
                    committed_text = ""
                    last_full_text = ""
                    pending_buffer = ""
                    sent_tail = ""
                    ui_bubble_start_idx = 0
                    last_change_time = time.monotonic()
                    silero_vad.reset()

        monitor_task = asyncio.create_task(_silence_flush_monitor())
        recv_loop_alive.set()

        # ── P0-15: Receiver / ASR Worker 双协程架构 ──
        # receiver() 只负责 receive → 入队（不阻塞）
        # asr_worker() 从队列消费 → VAD → ASR → 文本处理
        # 录音持久化由前端 HTTP POST /audio/chunk 独立完成（ACK 已有）

        # P0-16: 有界队列，满了丢 ASR 不丢录音
        asr_queue: asyncio.Queue[Optional[bytes]] = asyncio.Queue(maxsize=120)

        # 声纹信息注入辅助函数
        _vp_identified_user_ref = [None]  # 用 list 包装以便 worker 内修改
        _vp_identified_name_ref = [""]
        _vp_identified_confidence_ref = [0.0]
        _vp_identified_by_ref = ["manual"]
        _pending_buffer_ref = [""]  # worker 退出时保存待恢复文本

        def _vp_enrich(msg: dict) -> dict:
            """向消息中注入声纹识别结果（如有）。"""
            if _vp_identified_user_ref[0] and msg.get("type") in ("final", "preview"):
                msg["speaker_name"] = _vp_identified_name_ref[0]
                msg["speaker_confidence"] = round(_vp_identified_confidence_ref[0], 4)
                msg["identified_by"] = _vp_identified_by_ref[0]
            return msg

        async def receiver():
            """P0-15: 只负责接收音频并入队，不阻塞于 ASR。"""
            while True:
                try:
                    msg = await websocket.receive()
                except WebSocketDisconnect:
                    break
                if msg.get("type") == "websocket.disconnect":
                    break
                if msg.get("text"):
                    try:
                        cmd = json.loads(msg["text"])
                        if cmd.get("type") == "finish":
                            break
                        # P0-8: 心跳 pong 响应
                        if cmd.get("type") == "ping":
                            await websocket.send_json({"type": "pong", "timestamp": cmd.get("timestamp")})
                    except Exception:
                        pass
                    continue
                audio_bytes = msg.get("bytes")
                if not audio_bytes:
                    continue
                # P0-16: 队列满了丢 ASR chunk，录音已由 HTTP 保存
                try:
                    asr_queue.put_nowait(audio_bytes)
                except asyncio.QueueFull:
                    logger.warning("[AUDIO] ASR queue full, dropping chunk (录音已保存)")

        async def asr_worker():
            """P0-15: 从队列消费音频 → VAD → ASR → 文本处理。"""
            nonlocal session_id

            session_id_local = session_id
            chunk_count = 0
            stale_count = 0
            # P0-14: ASR 健康状态跟踪
            consecutive_failures = 0
            asr_degraded = False
            last_recovery_check = 0.0
            prev_full_text = ""
            repeat_burst = 0
            committed_text = ""
            pending_buffer = ""
            last_full_text = ""
            last_change_time = time.monotonic()
            sent_tail = ""
            spk_id = ""
            ui_bubble_start_idx = 0
            audio_float_buffer = bytearray()

            _vp_audio_buffer = bytearray()

            while recv_loop_alive.is_set():
                try:
                    audio_bytes = await asyncio.wait_for(asr_queue.get(), timeout=1.0)
                except asyncio.TimeoutError:
                    continue
                if audio_bytes is None:  # 哨兵值，退出信号
                    break

                # Stage 1: Silero VAD
                is_speech, vad_prob = silero_vad.process(audio_bytes)
                if not is_speech:
                    continue

                # P0-14: degraded 状态 — 跳过 ASR，定期检查恢复
                if asr_degraded:
                    now = time.monotonic()
                    if now - last_recovery_check > 10.0:  # 每 10 秒检查一次
                        last_recovery_check = now
                        if await qwen_client.is_available():
                            logger.info("ASR 恢复可用，重建 session")
                            asr_degraded = False
                            consecutive_failures = 0
                            session_id_local = await qwen_client.start(hotwords=_build_asr_hotwords(meeting_title=meeting_title, agenda=agenda, project="", extra=[user_role.get("displayName", ""), user_role.get("meetingRole", "")]))
                            session_id = session_id_local
                            _ACTIVE_ASR_SESSIONS[session_key] = session_id_local
                            audio_float_buffer = bytearray()
                        else:
                            logger.debug("ASR 仍不可用，继续录音")
                    continue  # degraded 期间跳过 ASR 处理

                # 声纹识别
                if _vp_engine and _vp_enrolled and _vp_identified_user_ref[0] is None:
                    _vp_audio_buffer.extend(audio_bytes)
                    if len(_vp_audio_buffer) >= VP_BUFFER_BYTES:
                        vp_pcm = bytes(_vp_audio_buffer[:VP_BUFFER_BYTES])
                        _vp_audio_buffer.clear()
                        try:
                            loop_now = asyncio.get_event_loop()
                            _uid, _conf = await loop_now.run_in_executor(
                                None,
                                lambda: _vp_engine.identify_speaker_from_bytes(vp_pcm, _vp_enrolled),
                            )
                            if _uid:
                                _vp_identified_user_ref[0] = _uid
                                _vp_identified_confidence_ref[0] = _conf
                                _vp_identified_by_ref[0] = "voiceprint-realtime"
                                from backend.db import _db_get_voiceprint_by_user
                                _vp_profile = _db_get_voiceprint_by_user(_uid)
                                if _vp_profile:
                                    _vp_identified_name_ref[0] = _vp_profile.get("display_name", _uid)
                                logger.info("声纹识别: %s (conf=%.3f)", _vp_identified_name_ref[0] or _uid, _conf)
                        except Exception as vp_err:
                            logger.debug("声纹识别异常（忽略）: %s", vp_err)

                # 第一次检测到语音 → 延迟创建 ASR session
                if session_id_local is None:
                    session_id_local = await qwen_client.start(hotwords=_build_asr_hotwords(meeting_title=meeting_title, agenda=agenda, project="", extra=[user_role.get("displayName", ""), user_role.get("meetingRole", "")]))
                    session_id = session_id_local
                    _ACTIVE_ASR_SESSIONS[session_key] = session_id_local
                    audio_float_buffer = bytearray()
                    logger.info("ASR会话(延迟启动): %s (meeting=%s)", session_id_local, meeting_id)

                # Stage 3: 滑动缓冲区 → 对齐
                audio_float_buffer.extend(audio_bytes)
                if len(audio_float_buffer) < ASR_CHUNK_BYTES:
                    continue
                send_bytes = bytes(audio_float_buffer[:ASR_CHUNK_BYTES])
                audio_float_buffer = audio_float_buffer[ASR_CHUNK_BYTES:]

                # Stage 4: ASR 推理（P0-14: 分类异常 + 自动重建 + degraded）
                chunk_count += 1
                try:
                    result = await qwen_client.send_chunk(session_id_local, send_bytes)
                    raw_text = result.get("text", "").strip()
                    spk_id = result.get("spk", "")
                    consecutive_failures = 0  # 成功 → 重置计数
                except ASRSessionExpiredError as e:
                    consecutive_failures += 1
                    logger.warning("Session失效，自动重建 (chunk %d, 失败%d次): %s", chunk_count, consecutive_failures, e)
                    try:
                        session_id_local = await qwen_client.start()
                        session_id = session_id_local
                        _ACTIVE_ASR_SESSIONS[session_key] = session_id_local
                        result = await qwen_client.send_chunk(session_id_local, send_bytes)
                        raw_text = result.get("text", "").strip()
                        spk_id = result.get("spk", "")
                        consecutive_failures = 0  # 重建成功
                    except Exception as e2:
                        logger.warning("重建失败: %s", e2)
                        if consecutive_failures >= 3:
                            asr_degraded = True
                            last_recovery_check = time.monotonic()
                            logger.warning("ASR 进入 degraded 状态（连续 %d 次失败），录音继续", consecutive_failures)
                        continue
                except ASRChunkTimeoutError as e:
                    consecutive_failures += 1
                    logger.warning("ASR chunk %d 超时 (失败%d次): %s", chunk_count, consecutive_failures, e)
                    if consecutive_failures >= 3:
                        asr_degraded = True
                        last_recovery_check = time.monotonic()
                        logger.warning("ASR 进入 degraded 状态（连续 %d 次超时），录音继续", consecutive_failures)
                    continue
                except ASRUnavailableError as e:
                    consecutive_failures += 1
                    logger.warning("ASR 不可用 (chunk %d, 失败%d次): %s", chunk_count, consecutive_failures, e)
                    if consecutive_failures >= 3:
                        asr_degraded = True
                        last_recovery_check = time.monotonic()
                        logger.warning("ASR 进入 degraded 状态（连续 %d 次不可用），录音继续", consecutive_failures)
                    continue
                except ASRError as e:
                    consecutive_failures += 1
                    logger.warning("ASR chunk %d 错误 (失败%d次): %s", chunk_count, consecutive_failures, e)
                    continue

                # 物理清洗
                clean_text = _clean(raw_text)

                # Stage 5: LCP 裁剪
                if not clean_text or clean_text == last_full_text:
                    stale_count += 1
                else:
                    prefix_len = _fuzzy_lcp(committed_text, clean_text, tolerance=2)
                    new_content = clean_text[prefix_len:]
                    if not new_content:
                        stale_count += 1
                    else:
                        stale_count = 0
                        last_full_text = clean_text
                        last_change_time = time.monotonic()

                        # Stage 6: 重复 Token 熔断
                        if _REPEAT_RE.search(clean_text):
                            logger.warning("熔断A: 重复Token '%s'，重启", clean_text[-20:])
                            try: await qwen_client.finish(session_id_local)
                            except Exception: pass
                            session_id_local = await qwen_client.start()
                            session_id = session_id_local
                            _ACTIVE_ASR_SESSIONS[session_key] = session_id_local
                            committed_text = ""; last_full_text = ""; pending_buffer = ""
                            sent_tail = ""; ui_bubble_start_idx = 0
                            silero_vad.reset()
                            continue

                        if clean_text == prev_full_text:
                            repeat_burst += 1
                            if repeat_burst >= 3:
                                logger.warning("熔断B: 连续3帧相同，重启")
                                try: await qwen_client.finish(session_id_local)
                                except Exception: pass
                                session_id_local = await qwen_client.start()
                                session_id = session_id_local
                                _ACTIVE_ASR_SESSIONS[session_key] = session_id_local
                                committed_text = ""; last_full_text = ""; prev_full_text = ""
                                pending_buffer = ""; sent_tail = ""
                                repeat_burst = 0; ui_bubble_start_idx = 0
                                silero_vad.reset()
                                continue
                        else:
                            prev_full_text = clean_text
                            repeat_burst = 0

                        # Preview
                        bubble_text = clean_text[ui_bubble_start_idx:]
                        await websocket.send_json(_vp_enrich({
                            "type": "preview", "taskId": task_id,
                            "meetingId": meeting_id, "meetingTitle": meeting_title,
                            "agenda": agenda, "text": bubble_text,
                            "isFinal": False, "backend": "paraformer", "spk": spk_id,
                        }))

                        # 标点 → 换行
                        END_PUNCTS = ('。', '？', '！', '…')
                        if clean_text.endswith(END_PUNCTS):
                            bubble = clean_text[ui_bubble_start_idx:]
                            if bubble.strip():
                                await websocket.send_json(_vp_enrich({
                                    "type": "final", "taskId": task_id,
                                    "meetingId": meeting_id, "meetingTitle": meeting_title,
                                    "agenda": agenda, "newText": bubble, "fullText": clean_text,
                                    "isFinal": True, "backend": "paraformer", "spk": spk_id,
                                }))
                            ui_bubble_start_idx = len(clean_text)
                            committed_text = clean_text
                            pending_buffer = ""

                        # 无标点但 ≥8 字 → 累积提交
                        else:
                            pending_buffer = pending_buffer + new_content
                            if len(pending_buffer) >= 8:
                                if not (pending_buffer == sent_tail[-len(pending_buffer):] and len(pending_buffer) >= 3):
                                    await websocket.send_json(_vp_enrich({
                                        "type": "final", "taskId": task_id,
                                        "meetingId": meeting_id, "meetingTitle": meeting_title,
                                        "agenda": agenda, "newText": pending_buffer, "fullText": clean_text,
                                        "isFinal": True, "backend": "paraformer", "spk": spk_id,
                                    }))
                                    sent_tail = (sent_tail + pending_buffer)[-50:]
                                committed_text = clean_text
                                ui_bubble_start_idx = len(clean_text)
                                pending_buffer = ""

                # Stage 7: 流式失活检测
                if stale_count >= 30:
                    logger.warning("失活: %d chunks无新字，重启", stale_count)
                    try: await qwen_client.finish(session_id_local)
                    except Exception: pass
                    session_id_local = await qwen_client.start()
                    session_id = session_id_local
                    _ACTIVE_ASR_SESSIONS[session_key] = session_id_local
                    committed_text = ""; last_full_text = ""; prev_full_text = ""
                    pending_buffer = ""; sent_tail = ""
                    ui_bubble_start_idx = 0; stale_count = 0; repeat_burst = 0

                # 每30帧诊断
                if chunk_count % 30 == 0:
                    import struct as _st
                    n = min(len(send_bytes)//2, 100)
                    s = _st.unpack(f'<{n}h', send_bytes[:n*2])
                    p = max(abs(x) for x in s) if s else 0
                    r = int((sum(x*x for x in s)/len(s))**0.5) if s else 0
                    logger.info("ASR diag: chunk=%d bytes=%d peak=%d rms=%d stale=%d raw='%s'",
                                chunk_count, len(send_bytes), p, r, stale_count, raw_text[:50])

            # worker 退出前保存待恢复文本
            _pending_buffer_ref[0] = pending_buffer

        # 启动双协程
        asr_task = asyncio.create_task(asr_worker())
        await receiver()
        # receiver 退出 → 发送哨兵值通知 worker 退出
        try:
            asr_queue.put_nowait(None)
        except asyncio.QueueFull:
            pass

    except Exception as exc:
        logger.exception("本地 ASR WS 异常")
        try:
            await websocket.send_json({"type": "error", "message": f"ASR 错误: {exc}"})
        except Exception:
            pass
    finally:
        recv_loop_alive.clear()
        # 取消 monitor 和 asr_worker
        monitor_task.cancel()
        try:
            await monitor_task
        except asyncio.CancelledError:
            pass
        if asr_task is not None:
            asr_task.cancel()
            try:
                await asyncio.wait_for(asr_task, timeout=2.0)
            except (asyncio.CancelledError, asyncio.TimeoutError):
                pass
        _ACTIVE_ASR_SESSIONS.pop(session_key, None)
        if session_id:
            try:
                await qwen_client.finish(session_id)
            except Exception:
                pass
        if _pending_buffer_ref[0]:
            _asr_pending_store[session_key] = {
                "text": _pending_buffer_ref[0], "timestamp": time.time(), "version": 1,
            }
        try:
            await websocket.send_json({"type": "finished", "taskId": task_id})
        except Exception:
            pass
        try:
            await websocket.close()
        except Exception:
            pass


@app.get("/api/custom_rules")
async def list_custom_rules(request: Request):
    _get_request_user(request, required=True)
    return JSONResponse({"success": True, "files": _load_custom_rules()})


@app.post("/api/custom_rules/upload")
async def upload_custom_rule(request: Request, file: UploadFile = File(...), matter_type: str = "通用"):
    _get_request_user(request, required=True)
    filename = file.filename or "未命名制度文件"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext != "pdf":
        raise HTTPException(status_code=400, detail="制度文件仅支持 PDF 上传")
    raw = await file.read()
    parsed_text = _extract_text_from_raw(filename, raw)
    summary = _compose_custom_rule_summary(parsed_text, matter_type)
    record = {
        "id": f"rule_{uuid.uuid4().hex[:12]}",
        "name": filename,
        "matterType": matter_type,
        "uploadedAt": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "charCount": len(parsed_text),
        "parsedText": parsed_text,
        "summaryLines": summary["summary_lines"],
    }
    files = _load_custom_rules()
    files.insert(0, record)
    _save_custom_rules(files)
    return JSONResponse({"success": True, "file": record})


@app.delete("/api/custom_rules/{rule_id}")
async def delete_custom_rule(request: Request, rule_id: str):
    _get_request_user(request, required=True)
    files = _load_custom_rules()
    filtered = [item for item in files if item["id"] != rule_id]
    if len(filtered) == len(files):
        raise HTTPException(status_code=404, detail="制度文件不存在")
    _save_custom_rules(filtered)
    return JSONResponse({"success": True})


@app.get("/api/rules_gallery")
async def get_rules_gallery(request: Request):
    _get_request_user(request, required=True)
    gallery = get_demo_assets().get("rulesGallery", [])
    for item in gallery:
        item["imageUrl"] = f"/api/rules_images/{item['filename']}"
    return JSONResponse({"success": True, "items": gallery})


@app.get("/api/rules_images/{filename}")
async def get_rules_image(filename: str):
    # 防止路径遍历攻击
    safe_name = re.sub(r"[^a-zA-Z0-9_.-]", "", filename)
    if safe_name != filename:
        raise HTTPException(status_code=400, detail="文件名包含非法字符")
    image_path = (RULES_IMAGES_DIR / safe_name).resolve()
    rules_dir = RULES_IMAGES_DIR.resolve()
    if not str(image_path).startswith(str(rules_dir) + os.sep) and image_path != rules_dir:
        raise HTTPException(status_code=403, detail="禁止访问目录外文件")
    if not image_path.exists():
        raise HTTPException(status_code=404, detail="规则图片不存在")
    return FileResponse(path=image_path)

def sanitize_report(text: str) -> str:
    """Strip any ReAct/JSON artifacts that may leak into the final report."""
    # Remove fenced code blocks (```...```)
    text = re.sub(r'```[\s\S]*?```', '', text)
    # Remove Action / Observation / Thought lines (ReAct scaffolding)
    text = re.sub(r'^(Action|Observation|Thought|Action Input)\s*:.*$', '', text, flags=re.MULTILINE)
    # Remove lines that are pure JSON objects/arrays
    text = re.sub(r'^\s*[\{\[][\s\S]*?[\}\]]\s*$', '', text, flags=re.MULTILINE)
    # Remove "Final Answer:" prefix if the model echoed it
    text = re.sub(r'^Final Answer\s*:\s*', '', text, flags=re.MULTILINE)
    # Collapse excess blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


@app.post("/audit_stream")
@app.post("/api/audit_stream")
async def audit_stream(http_request: Request, body: ChatRequest):
    """流式执行合规审核 (SSE)"""
    _get_request_user(http_request, required=True)
    logger.info(f"收到流式审核请求 - 事项类型: {body.matter_type}")
    material = body.material_text  # Allow full context length
    query = f"审核事项：{body.matter_type}\n材料内容：{material}"

    # Set the current request in context so the LLM can check for disconnection
    current_request.set(http_request)

    async def event_generator():
        try:
            if llm_semaphore.locked():
                yield f"data: {json.dumps({'type': 'queue_warning', 'content': '模型当前繁忙，您的请求已进入排队，请稍候...'}, ensure_ascii=False)}\n\n"
            
            async with llm_semaphore:
                custom_rules_text = _resolve_custom_rules_text(body.custom_rule_ids, body.matter_type)
                # ── Deterministic Procedural Pipeline ──
                # Step 1: extract_rules
                yield f"data: {json.dumps({'type': 'tool_start', 'tool': 'extract_rules'}, ensure_ascii=False)}\n\n"
                rules_res = extract_rules.invoke({"matter_type": body.matter_type, "custom_rules_text": custom_rules_text})
                yield f"data: {json.dumps({'type': 'tool_end', 'tool': 'extract_rules', 'result': rules_res}, ensure_ascii=False)}\n\n"

                # Step 2-4 can run in parallel once rules are ready
                yield f"data: {json.dumps({'type': 'tool_start', 'tool': 'validate_material'}, ensure_ascii=False)}\n\n"
                yield f"data: {json.dumps({'type': 'tool_start', 'tool': 'check_procedure_completeness'}, ensure_ascii=False)}\n\n"
                yield f"data: {json.dumps({'type': 'tool_start', 'tool': 'identify_responsibility'}, ensure_ascii=False)}\n\n"

                val_task = asyncio.to_thread(validate_material.invoke, {"material_text": material, "rules_text": rules_res})
                proc_task = asyncio.to_thread(check_procedure_completeness.invoke, {"material_text": material, "rules_text": rules_res})
                resp_task = asyncio.to_thread(identify_responsibility.invoke, {"material_text": material, "rules_text": rules_res})
                val_res, proc_res, resp_res = await asyncio.gather(val_task, proc_task, resp_task)

                yield f"data: {json.dumps({'type': 'tool_end', 'tool': 'validate_material', 'result': val_res}, ensure_ascii=False)}\n\n"
                yield f"data: {json.dumps({'type': 'tool_end', 'tool': 'check_procedure_completeness', 'result': proc_res}, ensure_ascii=False)}\n\n"
                yield f"data: {json.dumps({'type': 'tool_end', 'tool': 'identify_responsibility', 'result': resp_res}, ensure_ascii=False)}\n\n"

                # Step 5: generate_compliance_report (streamed directly)
                yield f"data: {json.dumps({'type': 'tool_start', 'tool': 'generate_compliance_report'}, ensure_ascii=False)}\n\n"
                
                combined_results = f"""
1. 规则提取结果：\n{rules_res}
2. 材料合规校验结果：\n{val_res}
3. 程序完整性核查：\n{proc_res}
4. 责任主体落实：\n{resp_res}
                """
                
                prompt_text = f"""你是资深三重一大合规审核专家，请严格按照国企公文规范，基于以下前置工具执行结果，输出一份排版精美的 Markdown 合规审核报告。材料内容详见下文。

【材料内容】
{material}

【工具执行结果原始数据（仅供参考，不得在报告中出现）】
{combined_results}

=========
【输出格式要求】
1. 报告使用专业严谨的中文表述。
2. 严禁输出任何 JSON、代码块（```）、Action/Observation 标签、工具调用原始内容。
3. 全文顶部必须先输出一个隐藏的雷达数据块（用于前端红绿灯看板渲染），格式如下（严格单行或多行 XML，根据检查结果输出对应的红黄绿灯）：
   <risk_radar>
   <item status="green">党委前置审查(通过)</item>
   <item status="yellow">程序完整性(缺少某环节)</item>
   <item status="red">大额资金审批(违规)</item>
   </risk_radar>
4. 紧接 XML 块之后，输出以下九个章节：
   - ## 一、审核基本信息（表格形式：审核类型、审核日期、审核结论）
   - ## 二、风险等级评定（使用 ⚠️ 高风险 / 🔶 中风险 / 🟢 低风险，说明评定理由）
   - ## 三、违规事项与证据清单（务必在叙述违规/合规项时，用特殊的 Markdown 链接语法标出证据原文。例如：[此项目缺乏财务审计报告](evidence:"未见财务部门签字的审计文本")）
   - ## 四、程序完整性核查（Markdown 表格：程序环节 | 状态 | 备注）
   - ## 五、责任主体认定（说明责任人/监督部门落实情况）
   - ## 六、整改建议（分条陈述，每条格式：**建议N**：具体要求。关键要求：在每条具体建议的末尾，必须加上一个特殊的生成按钮链接。例如：[🔨 一键生成《整改通知单》及补正依据](remediate:"为这个项目起草缺少前置审查环节的情况说明及补发文的模版")）
   - ## 七、决策溯源档案（电子档案）（说明：系统已具备"一键生成迎检报告"功能，支持按年度、按事项类型一键导出决策记录、参会人员名单、表决结果及原始文件，形成完整的电子档案库，以应对国资委巡视和审计。请以此口吻进行简要功能描述提示。）
   - ## 八、整改闭环管理（说明：除审核外，系统具备"整改通知下发 -> 整改报告上报 -> 复核销号"的闭环管理模块，保障违规问题闭环跟踪。请在此简要描述本次审核所涉及的闭环流程或系统功能赋能。）
   - ## 九、统计分析与驾驶舱（说明：系统具备决策效能分析看板，实时展示本年度党委会召开次数、重大项目研究数量、资金总额、"紧急上会/临时动议"等异常事项占比，支撑向国资委汇报的关键 KPI。请简要描述本系统驾驶舱模块是如何实时监控统计及展示此类指标的。）
5. 报告末尾附：> 📋 本报告由 AI 合规审核系统自动生成，仅供参考，最终结论以人工复核为准。

现在从 <risk_radar> 开始输出："""
                
                # 强制阻断思考链，直接让大模型吐出前端需要的 XML 标签打头
                messages = [
                    SystemMessage(content="你是资深三重一大合规审核专家。直接输出最终的风险雷达数据块及审计报告全文。请务必以 <risk_radar> 开头。"),
                    HumanMessage(content=prompt_text)
                ]
                
                full_report = "" 
                async for chunk in llm._astream(messages, enable_thinking=False):
                    if await http_request.is_disconnected():
                        logger.info("【Audit】前台断开，中断生成。")
                        raise asyncio.CancelledError()
                    
                    chunk_text = chunk.message.content
                    chunk_type = chunk.message.additional_kwargs.get("chunk_type", "content")
                    
                    if chunk_text:
                        if chunk_type == "thinking":
                            yield f"data: {json.dumps({'type': 'thinking_chunk', 'content': chunk_text}, ensure_ascii=False)}\n\n"
                        else:
                            full_report += chunk_text
                            yield f"data: {json.dumps({'type': 'llm_chunk', 'content': chunk_text}, ensure_ascii=False)}\n\n"
                
                # Report finished
                yield f"data: {json.dumps({'type': 'tool_end', 'tool': 'generate_compliance_report'}, ensure_ascii=False)}\n\n"
                
                clean = sanitize_report(full_report)
                try:
                    persistence.save_audit(
                        matter_type=body.matter_type,
                        material=material,
                        report=clean,
                        results={
                            "rules": rules_res,
                            "validation": val_res,
                            "procedure": proc_res,
                            "responsibility": resp_res,
                        },
                    )
                except Exception as persist_err:
                    logger.warning(f"审核记录归档失败: {persist_err}")
                yield f"data: {json.dumps({'type': 'report', 'content': clean}, ensure_ascii=False)}\n\n"
                yield 'data: {"type": "done"}\n\n'

        except Exception as e:
            logger.error(f"Audit stream error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'detail': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")

@app.get("/api/knowledge/agenda-search")
async def agenda_knowledge_search(request: Request, q: str = "", limit: int = 20):
    """议题级历史知识检索（§60-64）：返回 具体会议 + 具体议题 + 最终决议。

    保密议题按当前用户权限过滤（内容不下发）。
    """
    user = _get_request_user(request, required=True)
    try:
        limit = max(1, min(int(limit), 50))
    except Exception:
        limit = 20
    result = search_agenda_knowledge(q, limit=limit, user=user)
    return JSONResponse(result)


@app.post("/kb_stream")
@app.post("/api/kb_stream")
async def kb_stream(http_request: Request, body: KBQueryRequest):
    """流式查询企业知识库 (RAG)"""
    _get_request_user(http_request, required=True)
    logger.info(f"收到知识库查询: {body.query}")

    # Set the current request in context so the LLM can check for disconnection
    current_request.set(http_request)

    async def kb_generator():
        try:
            store = _get_vectorstore()
            if store is None:
                detail = _vectorstore_error or "知识库未初始化，请联系管理员建立索引。"
                raise ValueError(detail)

            # 1. Retrieve relevant docs
            yield f"data: {json.dumps({'type': 'tool_start', 'tool': '检索本地文档库(MMR)'}, ensure_ascii=False)}\n\n"
            docs = store.max_marginal_relevance_search(body.query, k=4, fetch_k=12)
            context = "\n\n".join([f"【参考资料 {i+1}】\n{d.page_content}" for i, d in enumerate(docs)])

            # 提取来源元数据（文件名 + 页码 + 片段预览）
            sources = []
            for i, d in enumerate(docs):
                meta = d.metadata if hasattr(d, 'metadata') and d.metadata else {}
                page = meta.get("page")
                total = meta.get("total_pages")
                loc = f"第{page}页" if page else (f"片段{meta.get('chunk', 0)}")
                if page and total:
                    loc = f"第{page}/{total}页"
                sources.append({
                    "index": i + 1,
                    "source": meta.get("source", "未知来源"),
                    "location": loc,
                    "page": page,
                    "total_pages": total,
                    "chunk": meta.get("chunk", 0),
                    "doc_id": meta.get("doc_id", ""),
                    "snippet": d.page_content[:200] + ("..." if len(d.page_content) > 200 else ""),
                })
            yield f"data: {json.dumps({'type': 'sources', 'sources': sources}, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'type': 'tool_end', 'tool': f'检索完毕，找到 {len(docs)} 条相关资料'}, ensure_ascii=False)}\n\n"

            prompt_text = f"""请根据以下内部资料回答用户的提问。

【内部资料开始】
{context}
【内部资料结束】

要求：
1. 请详细、专业地回答。
2. 答案必须严格基于上述资料，不得捏造不存在的政策；如果资料中没有相关直接信息，请说明当前本地知识库并未涉及此部分细节。
3. 在回答中标出参考了哪些资料（使用 [参考N] 标记，其中 N 是参考资料编号）。
4. 必须全部使用中文（简体）进行思考和回答，禁止出现英文。

用户提问：{body.query}"""
            messages = [
                SystemMessage(content="你是一个专业的城投企业合规法务专家。请始终严格使用中文（简体）回答问题，包括思考过程和最终结论，绝对禁止使用英文。"),
                HumanMessage(content=prompt_text)
            ]

            # 3. Stream LLM response
            full_response = ""
            
            if llm_semaphore.locked():
                yield f"data: {json.dumps({'type': 'queue_warning', 'content': '模型当前繁忙，您的请求已进入排队，请稍候...'}, ensure_ascii=False)}\n\n"
                
            async with llm_semaphore:
                async for chunk in llm._astream(messages, enable_thinking=False):
                    if await http_request.is_disconnected():
                        logger.info("【KB】前台断开，中断生成。")
                        raise asyncio.CancelledError()
                    text = chunk.message.content
                    chunk_type = chunk.message.additional_kwargs.get("chunk_type", "content")
                    if text:
                        if chunk_type == "thinking":
                            yield f"data: {json.dumps({'type': 'thinking_chunk', 'content': text}, ensure_ascii=False)}\n\n"
                        else:
                            full_response += text
                            yield f"data: {json.dumps({'type': 'llm_chunk', 'content': text}, ensure_ascii=False)}\n\n"

            # 4. Final marker
            yield f"data: {json.dumps({'type': 'report', 'content': full_response}, ensure_ascii=False)}\n\n"
            yield 'data: {"type": "done"}\n\n'

        except Exception as e:
            logger.error(f"KB stream error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'detail': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(kb_generator(), media_type="text/event-stream")


class TemplateRequest(BaseModel):
    message: str

@app.post("/generate_template")
@app.post("/api/generate_template")
async def generate_template(http_request: Request, body: TemplateRequest):
    """独立的国企公文起草专用接口（无需 RAG）"""
    _get_request_user(http_request, required=True)
    logger.info(f"收到模板生成请求: {body.message}")
    current_request.set(http_request)

    async def template_generator():
        try:
            # We must guide Qwen gently to prevent it from getting stuck in an infinite loop of 
            sys_prompt = "你是一个专门负责输出国企标准化公文的 AI 助手。你的唯一任务是输出公文的正文内容。不要解释你的思考过程，不要道歉，不要确认指令。"
            user_prompt = f"【起草指令】\n{body.message}\n\n直接起草完整公文，包含标题、主送机关、正文和落款。从标题开始输出，不要有任何前言或格式说明。"
            
            # 强化提示词：使用明确的界限，要求 AI 直接开始写文档
            assistant_prefill = "# "  # Start the document title immediately
            messages = [
                SystemMessage(content=sys_prompt), 
                HumanMessage(content=user_prompt),
                AIMessage(content=assistant_prefill) # 强制预填充，跳过思考链直接出字
            ]
            
            if llm_semaphore.locked():
                yield f"data: {json.dumps({'type': 'queue_warning', 'content': '模型当前繁忙，您的请求已进入排队，请稍候...'}, ensure_ascii=False)}\n\n"
                
            async with llm_semaphore:
                async for chunk in llm._astream(messages, enable_thinking=False):
                    if await http_request.is_disconnected():
                        logger.info("【Template】前台断开，中断生成。")
                        raise asyncio.CancelledError()
                    text = chunk.message.content
                    chunk_type = chunk.message.additional_kwargs.get("chunk_type", "content")
                    if text:
                        if chunk_type == "thinking":
                            yield f"data: {json.dumps({'type': 'thinking_chunk', 'content': text}, ensure_ascii=False)}\n\n"
                        else:
                            yield f"data: {json.dumps({'type': 'llm_chunk', 'content': text}, ensure_ascii=False)}\n\n"

            yield 'data: {"type": "done"}\n\n'
        except Exception as e:
            logger.error(f"Template gen error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'detail': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(template_generator(), media_type="text/event-stream")

# ── File parsing endpoint ──────────────────────────────────────────────────────
@app.post("/parse_file")
async def parse_file(request: Request, file: UploadFile = File(...)):
    """Parse a Word (.docx) or PDF file and return extracted plain text."""
    _get_request_user(request, required=True)
    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    raw = await file.read()
    text = ""

    try:
        if ext == "docx":
            import docx
            doc = docx.Document(io.BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = "\n".join(paragraphs)

        elif ext == "pdf":
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                pages = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
            text = "\n\n".join(pages)

        elif ext in ("txt", "md"):
            text = raw.decode("utf-8", errors="replace")

        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件类型：.{ext}，仅支持 .docx / .pdf / .txt")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文件解析失败: {e}")
        raise HTTPException(status_code=500, detail=f"文件解析失败：{str(e)}")

    # Clean up whitespace
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    char_count = len(text)
    logger.info(f"文件解析完成: {filename}, {char_count} 字符")
    return JSONResponse({"text": text, "filename": filename, "char_count": char_count})


@app.post("/api/ocr/image")
async def ocr_image(request: Request, file: UploadFile = File(...)):
    _get_request_user(request, required=True)
    api_key = DASHSCOPE_API_KEY
    if not api_key:
        raise HTTPException(status_code=500, detail="未配置 DASHSCOPE_API_KEY，无法调用百炼 OCR")

    filename = file.filename or "image"
    content_type = file.content_type or "image/jpeg"
    if not content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="OCR 仅支持图片文件")

    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="图片内容为空")
    if len(raw) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="图片超过 10MB，请压缩后上传")

    data_url = f"data:{content_type};base64,{base64.b64encode(raw).decode('ascii')}"
    model = os.environ.get("DASHSCOPE_OCR_MODEL", "qwen-vl-ocr")
    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "请对图片做 OCR，提取所有可见中文和数字。只返回纯文本，不要解释，不要 Markdown。",
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": data_url},
                    },
                ],
            }
        ],
        "temperature": 0,
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    if DASHSCOPE_WORKSPACE:
        headers["X-DashScope-WorkSpace"] = DASHSCOPE_WORKSPACE

    try:
        async with httpx.AsyncClient(timeout=60.0, trust_env=False) as client:
            response = await client.post(
                "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions",
                headers=headers,
                json=payload,
            )
        if response.status_code >= 400:
            logger.error("百炼 OCR 调用失败: %s %s", response.status_code, response.text[:500])
            raise HTTPException(status_code=502, detail=f"百炼 OCR 调用失败：HTTP {response.status_code}")
        data = response.json()
        text = ""
        choices = data.get("choices") if isinstance(data, dict) else None
        if choices:
            message = choices[0].get("message") or {}
            content = message.get("content")
            if isinstance(content, str):
                text = content
            elif isinstance(content, list):
                text = "\n".join(str(part.get("text", "")) for part in content if isinstance(part, dict))
        text = re.sub(r"\n{3,}", "\n\n", text or "").strip()
        return JSONResponse({
            "success": True,
            "filename": filename,
            "model": model,
            "text": text,
            "char_count": len(text),
        })
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("百炼 OCR 处理异常")
        raise HTTPException(status_code=500, detail=f"OCR 处理失败：{exc}")



# ── Shared chunker: preserve page numbers ────────────────────────────────────
def _chunk_text_with_pages(text: str, pages: Optional[List[str]] = None,
                           chunk_size: int = 500, overlap: int = 50,
                           source: str = "", doc_id_prefix: str = "") -> tuple:
    """按页切分文本，保留页码元数据。

    PDF 模式（pages 不为空）：每页独立切 chunk，标注页码（从 1 开始）。
    纯文本模式：按字符切分，page 字段为 None。

    Returns:
        (chunks: List[str], metadatas: List[dict], ids: List[str])
    """
    chunks = []
    metadatas = []
    ids = []

    if pages:
        total_pages = len(pages)
        for page_num, page_text in enumerate(pages, start=1):
            page_text = page_text.strip()
            if not page_text:
                continue
            start = 0
            while start < len(page_text):
                end = min(start + chunk_size, len(page_text))
                chunk_text = page_text[start:end]
                if chunk_text.strip():
                    cid = f"{doc_id_prefix}_p{page_num}_c{len(chunks)}"
                    chunks.append(chunk_text)
                    metadatas.append({
                        "source": source, "doc_id": doc_id_prefix,
                        "page": page_num, "chunk": len(chunks) - 1,
                        "total_pages": total_pages,
                    })
                    ids.append(cid)
                start += chunk_size - overlap
    else:
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk_text = text[start:end]
            if chunk_text.strip():
                cid = f"{doc_id_prefix}_c{len(chunks)}"
                chunks.append(chunk_text)
                metadatas.append({
                    "source": source, "doc_id": doc_id_prefix,
                    "page": None, "chunk": len(chunks) - 1,
                })
                ids.append(cid)
            start += chunk_size - overlap

    return chunks, metadatas, ids


# ── File ingest (parse + vectorize into ChromaDB) ─────────────────────────────
class IngestResponse(BaseModel):
    success: bool
    filename: str
    chunks: int
    char_count: int
    message: str

@app.post("/ingest_file", response_model=IngestResponse)
async def ingest_file(request: Request, file: UploadFile = File(...)):
    """Parse a file and add it to the ChromaDB vector knowledge base.
    The document will then be searchable via /kb_stream Q&A."""
    _get_request_user(request, required=True)
    global vectorstore

    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    raw = await file.read()
    text = ""
    pages = None  # PDF 页面列表

    try:
        if ext == "docx":
            import docx
            doc = docx.Document(io.BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = "\n".join(paragraphs)
        elif ext == "pdf":
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
            text = "\n\n".join(p for p in pages if p.strip())
        elif ext in ("txt", "md"):
            text = raw.decode("utf-8", errors="replace")
        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件类型：.{ext}")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"ingest_file parse error: {e}")
        raise HTTPException(status_code=500, detail=f"文件解析失败：{e}")

    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        raise HTTPException(status_code=400, detail="文件内容为空，无法入库。")

    # 保留页码的智能切分（PDF 按页切，DOCX 按字符切）
    doc_id_prefix = f"ingest_{uuid.uuid4().hex[:10]}"
    chunks, metadatas, ids = _chunk_text_with_pages(
        text, pages=pages, source=filename, doc_id_prefix=doc_id_prefix,
    )

    # Add to ChromaDB
    store = _get_vectorstore(create_if_missing=True)
    if store is None:
        raise HTTPException(status_code=503, detail=f"向量数据库初始化失败：{_vectorstore_error or '未知错误'}")

    try:
        store.add_texts(texts=chunks, metadatas=metadatas, ids=ids)
        logger.info(f"ingest_file: {filename} → {len(chunks)} chunks added to ChromaDB")
    except Exception as e:
        logger.error(f"vectorstore.add_texts failed: {e}")
        raise HTTPException(status_code=500, detail=f"向量入库失败：{e}")

    return IngestResponse(
        success=True,
        filename=filename,
        chunks=len(chunks),
        char_count=len(text),
        message=f"文件已成功入库，生成 {len(chunks)} 个语义片段，可在合规问答中检索。",
    )


@app.get("/kb_stats")
@app.get("/api/kb_stats")
async def kb_stats(request: Request):
    """Return the current ChromaDB knowledge base statistics."""
    _get_request_user(request, required=True)
    store = _get_vectorstore()
    if store is None:
        return {"available": False, "count": 0, "message": _vectorstore_error or "向量数据库未就绪"}
    try:
        count = store._collection.count()
        return {"available": True, "count": count, "message": f"知识库已就绪，共 {count} 个语义片段"}
    except Exception as e:
        return {"available": False, "count": 0, "message": str(e)}


# ── Knowledge Library File Management (Persistent) ──────────────────────────
class KnowledgeFile(BaseModel):
    id: str
    name: str
    type: str  # "pdf" | "docx" | "txt"
    size: str  # e.g. "128 KB"
    date: str  # ISO date string
    tags: List[str] = []
    linked: bool = False
    vectorized: bool = False
    uploader: str
    uploaderRole: str
    dept: str
    libraryCategory: Optional[str] = None  # "cases" | "knowledge" | "shared"
    parsedText: Optional[str] = None
    savedName: Optional[str] = None  # for editable docs


class KnowledgeFileUpdate(BaseModel):
    """Partial update model — all fields optional."""
    id: Optional[str] = None
    name: Optional[str] = None
    type: Optional[str] = None
    size: Optional[str] = None
    date: Optional[str] = None
    tags: Optional[List[str]] = None
    linked: Optional[bool] = None
    vectorized: Optional[bool] = None
    uploader: Optional[str] = None
    uploaderRole: Optional[str] = None
    dept: Optional[str] = None
    libraryCategory: Optional[str] = None
    parsedText: Optional[str] = None
    savedName: Optional[str] = None

KNOWLEDGE_FILES_DIR = Path(os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "knowledge_files"))
KNOWLEDGE_FILES_DIR.mkdir(parents=True, exist_ok=True)
KNOWLEDGE_FILES_DB = KNOWLEDGE_FILES_DIR / "files.json"

def _load_knowledge_files() -> List[dict]:
    if KNOWLEDGE_FILES_DB.exists():
        try:
            with open(KNOWLEDGE_FILES_DB, "r", encoding="utf-8") as f:
                data = json.load(f)
                if data:
                    return data
        except Exception:
            pass
    return get_seed_knowledge_files()

def _save_knowledge_files(files: List[dict]):
    with open(KNOWLEDGE_FILES_DB, "w", encoding="utf-8") as f:
        json.dump(files, f, ensure_ascii=False, indent=2)


def _resolve_editable_saved_name(record: dict) -> Optional[str]:
    if not isinstance(record, dict):
        return None

    current_saved_name = record.get("savedName")
    if current_saved_name and (DOCS_DIR / current_saved_name).exists():
        return current_saved_name

    file_type = str(record.get("type") or "").lower()
    source_name = str(record.get("name") or "").strip()
    if file_type not in {"docx", "doc"} or not source_name:
        return None

    suffix = f"_{source_name}"
    candidates: List[Path] = []
    for path in DOCS_DIR.iterdir():
        if not path.is_file():
            continue
        if path.suffix.lower() not in {".docx", ".doc"}:
            continue
        if path.name == source_name or path.name.endswith(suffix):
            candidates.append(path)

    if not candidates:
        return None

    def sort_key(path: Path):
        has_meta = (DOCS_DIR / f"{path.name}.meta.json").exists()
        return (1 if has_meta else 0, path.stat().st_mtime)

    return sorted(candidates, key=sort_key, reverse=True)[0].name


def _repair_knowledge_file_records(files: List[dict], persist: bool = False) -> List[dict]:
    changed = False
    for record in files:
        resolved_saved_name = _resolve_editable_saved_name(record)
        if not resolved_saved_name:
            continue

        if record.get("savedName") != resolved_saved_name:
            record["savedName"] = resolved_saved_name
            changed = True

        tags = list(record.get("tags") or [])
        if "可编辑" not in tags:
            record["tags"] = [*tags, "可编辑"]
            changed = True

    if changed and persist:
        _save_knowledge_files(files)
    return files


def _discover_orphaned_docs(files: List[dict], persist: bool = False) -> List[dict]:
    """Find .docx files in the docs directory that have no knowledge_files record yet."""
    if not DOCS_DIR.exists():
        return files

    existing_saved_names = {f.get("savedName") for f in files if f.get("savedName")}
    existing_names = {f.get("name") for f in files}
    discovered = []

    for path in sorted(DOCS_DIR.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True):
        if not path.is_file():
            continue
        if path.suffix.lower() not in {".docx", ".doc"}:
            continue
        if path.name in existing_saved_names:
            continue
        # Skip files that look like meeting materials (contain meeting- prefix)
        if path.name.startswith("meeting-"):
            continue
        # Check if the original filename is already tracked
        safe_name = path.name
        # Extract original filename from saved_name format: "uuid_originalname.docx"
        parts = safe_name.split("_", 1)
        original_name = parts[1] if len(parts) > 1 else safe_name
        if original_name in existing_names:
            continue
        # Skip exported/reviewed derivative files
        if "审查版" in original_name or "留痕审查版" in original_name:
            continue

        file_stat = path.stat()
        record = {
            "id": f"discovered_{safe_name.replace('.', '_')[:40]}",
            "name": original_name,
            "type": "docx",
            "size": f"{round(file_stat.st_size / 1024)} KB",
            "date": datetime.fromtimestamp(file_stat.st_mtime).strftime("%Y-%m-%d"),
            "tags": ["合同审查", "可编辑"],
            "linked": False,
            "vectorized": False,
            "uploader": "系统",
            "uploaderRole": "admin",
            "dept": "信息管理中心",
            "libraryCategory": "cases",
            "parsedText": None,
            "savedName": safe_name,
        }
        discovered.append(record)
        existing_saved_names.add(safe_name)

    if discovered:
        files = discovered + files
        if persist:
            _save_knowledge_files(files)

    return files

@app.get("/api/knowledge_files")
async def list_knowledge_files(request: Request):
    """List all knowledge library files (persistent)."""
    _get_request_user(request, required=True)
    files = _repair_knowledge_file_records(_load_knowledge_files(), persist=True)
    files = _discover_orphaned_docs(files, persist=True)
    return JSONResponse({"files": files})

@app.post("/api/knowledge_files")
async def add_knowledge_file(request: Request, file: KnowledgeFile):
    """Add a new knowledge library file record."""
    _get_request_user(request, required=True)
    files = _load_knowledge_files()
    # Check if already exists (by id)
    if any(f["id"] == file.id for f in files):
        return JSONResponse({"success": False, "message": "文件已存在"}, status_code=409)
    files.insert(0, file.model_dump(exclude_none=True))
    _save_knowledge_files(files)
    return JSONResponse({"success": True, "file": file.model_dump(exclude_none=True)})

@app.put("/api/knowledge_files/{file_id}")
async def update_knowledge_file(request: Request, file_id: str, update: KnowledgeFileUpdate):
    """Update a knowledge library file record (e.g., after parsing/vectorizing)."""
    _get_request_user(request, required=True)
    files = _load_knowledge_files()
    for i, f in enumerate(files):
        if f["id"] == file_id:
            # Merge update into existing record (only non-None fields)
            update_data = update.model_dump(exclude_none=True)
            files[i] = {**f, **update_data, "id": file_id}
            _save_knowledge_files(files)
            return JSONResponse({"success": True, "file": files[i]})
    raise HTTPException(status_code=404, detail="文件不存在")


@app.delete("/api/knowledge_files/{file_id}")
async def delete_knowledge_file(request: Request, file_id: str):
    """Delete a knowledge library file record."""
    _get_request_user(request, required=True)
    files = _load_knowledge_files()
    original_len = len(files)
    files = [f for f in files if f["id"] != file_id]
    if len(files) == original_len:
        raise HTTPException(status_code=404, detail="文件不存在")
    _save_knowledge_files(files)
    return JSONResponse({"success": True})

@app.post("/api/knowledge_files/{file_id}/vectorize")
async def toggle_vectorize(request: Request, file_id: str):
    """Toggle vectorized flag for a file. If enabling, re-ingest into ChromaDB."""
    _get_request_user(request, required=True)
    files = _load_knowledge_files()
    target = None
    for f in files:
        if f["id"] == file_id:
            target = f
            break
    if not target:
        raise HTTPException(status_code=404, detail="文件不存在")

    if target["vectorized"]:
        # Just toggle off (don't remove from ChromaDB, just mark)
        target["vectorized"] = False
        target["linked"] = False
        _save_knowledge_files(files)
        return JSONResponse({"success": True, "vectorized": False, "linked": False})

    # Need to vectorize - re-parse and ingest
    if not target.get("parsedText"):
        raise HTTPException(status_code=400, detail="文件未解析，无法向量化")

    store = _get_vectorstore(create_if_missing=True)
    if store is None:
        raise HTTPException(status_code=503, detail=f"向量数据库初始化失败：{_vectorstore_error or '未知错误'}")

    try:
        store.add_texts(
            texts=[target["parsedText"]],
            metadatas=[{"source": target["name"], "file_id": target["id"]}],
            ids=[target["id"]]
        )
        target["vectorized"] = True
        target["linked"] = True
        target["tags"] = [t for t in target.get("tags", []) if t != "待入库"] + ["已入库"]
        _save_knowledge_files(files)
        return JSONResponse({"success": True, "vectorized": True, "linked": True})
    except Exception as e:
        logger.error(f"Vectorize failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/knowledge_files/{file_id}/link")
async def toggle_link(request: Request, file_id: str):
    """Toggle linked flag (does not affect ChromaDB, only UI state)."""
    _get_request_user(request, required=True)
    files = _load_knowledge_files()
    for f in files:
        if f["id"] == file_id:
            f["linked"] = not f["linked"]
            _save_knowledge_files(files)
            return JSONResponse({"success": True, "linked": f["linked"]})
    raise HTTPException(status_code=404, detail="文件不存在")


# ── OnlyOffice Document Server Integration ───────────────────────────────────
# OnlyOffice DS address (container name resolves via Docker network, or use localhost for dev)
OO_HOST = os.environ.get("ONLYOFFICE_HOST", "localhost")
OO_PORT = os.environ.get("ONLYOFFICE_PORT", "8081")
ONLYOFFICE_URL = f"http://{OO_HOST}:{OO_PORT}"
# Internal backend address for OnlyOffice container to reach this backend service
BACKEND_INTERNAL_HOST = os.environ.get("BACKEND_INTERNAL_HOST", "ia-audit-backend")
INTERNAL_BACKEND_URL = f"http://{BACKEND_INTERNAL_HOST}:8000"
ONLYOFFICE_API_PATH = "/web-apps/apps/api/documents/api.js"

# JWT secret for OnlyOffice (must match ONLYOFFICE_JWT_SECRET env var)
OO_JWT_SECRET = os.environ.get("ONLYOFFICE_JWT_SECRET", "your-secret-key-change-in-production")

# Directory for storing editable documents
DOCS_DIR = Path(os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "docs"))
DOCS_DIR.mkdir(parents=True, exist_ok=True)

# Directory for storing plugin files
PLUGINS_DIR = Path(os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "docs_plugins"))
PLUGINS_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {"docx", "xlsx", "pptx", "doc", "xls", "ppt", "pdf", "txt"}


def _get_file_ext(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _get_document_type(ext: str) -> str:
    if ext in {"docx", "doc", "odt", "txt"}:
        return "word"
    if ext in {"xlsx", "xls", "ods", "csv"}:
        return "cell"
    if ext in {"pptx", "ppt", "odp"}:
        return "slide"
    if ext == "pdf":
        return "pdf"
    return "word"


# ── Bookmark insertion for .docx files ────────────────────────────────────────

def _add_bookmarks_to_docx(docx_bytes: bytes, doc_uuid: str, risk_map: dict = None) -> tuple[bytes, list]:
    """
    Insert ascending bookmarks at each non-empty paragraph in a .docx.
    Returns (modified_docx_bytes, paragraph_index_map).
    paragraph_index_map: list of {para_index, bookmark_name, text_preview}

    risk_map: optional dict {para_index: risk_id} — if provided, bookmarks use
    'risk_{risk_id}' naming so they can be referenced by risk ID directly.
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        return docx_bytes, []

    doc = Document(io.BytesIO(docx_bytes))
    paragraph_map = []

    # We use w:bookmarkStart and w:bookmarkEnd OOXML elements
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Bookmark name: use risk_id if provided, otherwise fall back to para_index
        if risk_map and idx in risk_map:
            bookmark_name = f"risk_{risk_map[idx]}"
        else:
            bookmark_name = f"audit_para_{idx}"
        text_preview = text[:80] + ("..." if len(text) > 80 else "")

        paragraph_map.append({
            "para_index": idx,
            "bookmark_name": bookmark_name,
            "text_preview": text_preview,
        })

        # Create bookmarkStart element
        bm_start = OxmlElement('w:bookmarkStart')
        bm_start.set(qn('w:id'), str(idx))
        bm_start.set(qn('w:name'), bookmark_name)

        # Create bookmarkEnd element with same id
        bm_end = OxmlElement('w:bookmarkEnd')
        bm_end.set(qn('w:id'), str(idx))

        # Insert at the beginning of the paragraph (before all runs)
        # Using insert(0, ...) ensures the bookmark is at the very start
        para._p.insert(0, bm_start)
        para._p.insert(1, bm_end)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue(), paragraph_map


# ── Plugin generation for a specific document

def _generate_audit_plugin_js(saved_name: str, issues: list) -> str:
    """Generate a OnlyOffice plugin JS that handles navigation and replace commands.
    Uses Asc.plugin.executeMethod (the correct OnlyOffice plugin API) to:
    - GoToBookmark: navigate to a bookmarked paragraph
    - SearchAndReplace: find and replace text
    Receives messages from React app via Asc.plugin.onExternalPluginMessage."""
    tmpl = """
(function(window, undefined) {
    var _lastBookmark = "";

    // Extract paraIndex from bookmark name like "audit_para_7" or raw "7"
    function bmNameToParaIndex(bmName) {
        if (!bmName) return null;
        if (/^\\d+$/.test(bmName)) return parseInt(bmName, 10);
        var m = bmName.match(/^audit_para_(\\d+)$/);
        if (m) return parseInt(m[1], 10);
        return null;
    }

    // Navigate to a bookmarked paragraph
    function goToBookmark(bmName) {
        if (!bmName) return;
        _lastBookmark = bmName;
        var idx = bmNameToParaIndex(bmName);
        var target = (idx !== null) ? ('audit_para_' + idx) : bmName;
        console.log('[PluginJS] GoToBookmark:', target);
        if (window.Asc && window.Asc.plugin) {
            window.Asc.plugin.executeMethod("GoToBookmark", [target], function(result) {
                console.log('[PluginJS] GoToBookmark result:', result);
            });
        }
    }

    // Plugin initialization
    window.Asc.plugin.init = function() {
        console.log("[PluginJS] Navigator plugin initialized");
    };

    // Receive messages from React app via postMessage -> onExternalPluginMessage
    // Frontend sends: { type: 'onExternalPluginMessage', pluginGuid: '...', data: { type: 'NAV_PARA', ... } }
    window.Asc.plugin.onExternalPluginMessage = function(data) {
        console.log("[PluginJS] Received:", JSON.stringify(data));
        if (!data || !data.type) return;

        switch (data.type) {
            case "NAV_PARA":
                // data.paraIndex is a number, build bookmark name
                goToBookmark('audit_para_' + data.paraIndex);
                break;
            case "NAV_BOOKMARK":
                // data.bookmarkName is a string, could be raw index or full name
                goToBookmark(data.bookmarkName);
                break;
            case "LOCATE_TEXT":
                // Navigate to paragraph, then search to highlight
                goToBookmark(data.bookmarkName || ('audit_para_' + data.paraIndex));
                if (data.text) {
                    console.log('[PluginJS] LOCATE_TEXT: search highlight for:', data.text);
                    if (window.Asc && window.Asc.plugin) {
                        window.Asc.plugin.executeMethod("SearchAndReplace", [{
                            searchString: data.text,
                            replaceString: data.text,
                            matchCase: false
                        }], function(result) {
                            console.log('[PluginJS] LOCATE_TEXT result:', result);
                            // Report back to React
                            window.Asc.plugin.executeMethod("GetSelectionText", [], function(selectedText) {
                                window.parent.postMessage({
                                    type: "auditNavResponse",
                                    data: {
                                        type: "LOCATE_RESULT",
                                        found: (result !== false && result !== null),
                                        selectedText: selectedText || ""
                                    }
                                }, "*");
                            });
                        });
                    }
                }
                break;
            case "REPLACE_TEXT":
                // Navigate to bookmark, then replace text
                goToBookmark(data.bookmarkName || ('audit_para_' + data.paraIndex));
                if (data.originalText && data.newText) {
                    console.log('[PluginJS] REPLACE_TEXT:', data.originalText, '=>', data.newText);
                    if (window.Asc && window.Asc.plugin) {
                        window.Asc.plugin.executeMethod("SearchAndReplace", [{
                            searchString: data.originalText,
                            replaceString: data.newText,
                            matchCase: false
                        }], function(result) {
                            console.log('[PluginJS] REPLACE_TEXT result:', result);
                            window.Asc.plugin.executeMethod("GetSelectionText", [], function(selectedText) {
                                window.parent.postMessage({
                                    type: "auditNavResponse",
                                    data: {
                                        type: "REPLACE_RESULT",
                                        success: (result !== false && result !== null),
                                        selectedText: selectedText || ""
                                    }
                                }, "*");
                            });
                        });
                    }
                }
                break;
            default:
                console.warn("[PluginJS] Unknown message type:", data.type);
        }
    };

    window.Asc.plugin.button = function() {};
})(window, undefined);
"""
    return tmpl


@app.post("/doc/upload")
async def doc_upload(request: Request, file: UploadFile = File(...), vectorize: bool = False):
    """Upload a document to the editable docs directory.
    For .docx files, automatically inserts paragraph bookmarks for audit navigation.
    If vectorize=True, also parse and ingest into ChromaDB for RAG search."""
    _get_request_user(request, required=True)
    filename = file.filename or ""
    ext = _get_file_ext(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型：.{ext}")

    safe_name = f"{uuid.uuid4().hex}_{filename}"
    file_path = DOCS_DIR / safe_name

    content = await file.read()
    paragraph_map = []

    # Insert bookmarks into .docx files for paragraph navigation
    if ext == "docx":
        try:
            content, paragraph_map = _add_bookmarks_to_docx(content, safe_name)
            logger.info(f"doc_upload: inserted {len(paragraph_map)} bookmarks into {filename}")
        except Exception as e:
            logger.warning(f"Bookmark insertion failed for {filename}: {e} — saving without bookmarks")

    with open(file_path, "wb") as f:
        f.write(content)

    # Save paragraph index map alongside the document (JSON sidecar)
    if paragraph_map:
        meta_path = DOCS_DIR / f"{safe_name}.meta.json"
        with open(meta_path, "w", encoding="utf-8") as mf:
            json.dump({
                "filename": filename,
                "saved_as": safe_name,
                "paragraphs": paragraph_map,
                "size": len(content),
            }, mf, ensure_ascii=False, indent=2)

    # ── Optional: auto-vectorize into ChromaDB ──────────────────────────────
    vector_chunks = 0
    if vectorize:
        try:
            text = ""
            if ext == "docx":
                import docx
                doc = docx.Document(io.BytesIO(content))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text.strip())
                text = "\n".join(paragraphs)
            elif ext == "pdf":
                import pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    pages = [page.extract_text() or "" for page in pdf.pages]
                text = "\n\n".join(p for p in pages if p.strip())
            elif ext in ("txt", "md"):
                text = content.decode("utf-8", errors="replace")

            text = re.sub(r"\n{3,}", "\n\n", text).strip()
            if text:
                store = _get_vectorstore(create_if_missing=True)
                if store:
                    doc_id_prefix = safe_name.split("_")[0]
                    # 保留页码信息：PDF 按页切，DOCX 按字符切
                    doc_pages = pages if ext == "pdf" else None
                    chunks, metadatas, ids = _chunk_text_with_pages(
                        text, pages=doc_pages, source=filename, doc_id_prefix=doc_id_prefix,
                    )
                    store.add_texts(texts=chunks, metadatas=metadatas, ids=ids)
                    vector_chunks = len(chunks)
                    logger.info(f"doc_upload vectorize: {filename} → {vector_chunks} chunks into ChromaDB")
        except Exception as ve:
            logger.warning(f"doc_upload vectorize failed for {filename}: {ve}")

    message_parts = []
    if paragraph_map:
        message_parts.append(f"已插入 {len(paragraph_map)} 个导航书签")
    if vector_chunks > 0:
        message_parts.append(f"已向量化 {vector_chunks} 个片段入库")
    if not message_parts:
        message_parts.append("文件上传成功")

    return JSONResponse({
        "success": True,
        "filename": filename,
        "saved_as": safe_name,
        "size": len(content),
        "paragraph_count": len(paragraph_map),
        "vectorized": vector_chunks > 0,
        "chunks": vector_chunks,
        "message": "，".join(message_parts),
    })


@app.get("/doc/list")
async def doc_list(request: Request):
    """List all documents in the editable docs directory."""
    _get_request_user(request, required=True)
    files = []
    if DOCS_DIR.exists():
        for f in DOCS_DIR.iterdir():
            if f.is_file() and not (
                f.name.endswith(".meta.json") or
                f.name.endswith(".suggestions.json") or
                f.name.endswith(".review.json")
            ):
                stat = f.stat()
                files.append({
                    "name": f.name,
                    "size": stat.st_size,
                    "modified": stat.st_mtime,
                })
    return JSONResponse({"files": files, "count": len(files)})


@app.get("/doc/download/{saved_name}")
async def doc_download(request: Request, saved_name: str):
    """Download a document by its saved name."""
    _get_request_user(request, required=True)
    file_path = DOCS_DIR / saved_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(
        path=file_path,
        filename=file_path.name,
        media_type="application/octet-stream",
    )


@app.delete("/doc/delete/{saved_name}")
async def doc_delete(request: Request, saved_name: str):
    """Delete a document by its saved name."""
    _get_request_user(request, required=True)
    file_path = DOCS_DIR / saved_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    file_path.unlink()
    for sidecar_suffix in (".meta.json", ".suggestions.json", ".review.json"):
        sidecar_path = DOCS_DIR / f"{saved_name}{sidecar_suffix}"
        if sidecar_path.exists():
            sidecar_path.unlink()
    return JSONResponse({"success": True, "message": "文件已删除"})


@app.get("/doc/extract_bookmarks/{saved_name}")
async def doc_extract_bookmarks(request: Request, saved_name: str):
    """Extract paragraph index map from a previously uploaded document."""
    _get_request_user(request, required=True)
    meta_path = DOCS_DIR / f"{saved_name}.meta.json"
    if not meta_path.exists():
        raise HTTPException(status_code=404, detail="该文档未包含段落索引，请确认上传的是 .docx 文件")
    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)
    return JSONResponse(meta)


@app.post("/doc/edit_url")
async def doc_edit_url(request: Request):
    """
    Generate an OnlyOffice editor URL for a document.
    Optionally pass audit issues to include a navigation panel inside the editor.
    Body: { "saved_name": "...", "issues": [...] }"""
    _get_request_user(request, required=True)
    body = await request.json()
    saved_name = body.get("saved_name")
    issues = body.get("issues", [])

    if not saved_name:
        raise HTTPException(status_code=400, detail="missing saved_name")

    file_path = DOCS_DIR / saved_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")

    browser_backend_base = _get_browser_backend_base(request)
    # Build URLs for the OnlyOffice config.
    # OnlyOffice DocumentServer runs in Docker. The container needs to download
    # the file and call back to the backend. When backend runs on the host (local dev,
    # not in Docker), we must use host.docker.internal so the container can reach it.
    # When backend IS inside Docker (docker-compose), use the Docker service name.
    import socket
    _is_docker_service = BACKEND_INTERNAL_HOST != "host.docker.internal"
    if _is_docker_service:
        _oo_backend_url = INTERNAL_BACKEND_URL  # e.g. http://ia-audit-backend:8000
    else:
        _oo_backend_url = "http://host.docker.internal:8000"  # host machine from container POV

    file_url = f"{_oo_backend_url}/doc/download/{saved_name}"
    callback_url = f"{_oo_backend_url}/doc/callback"
    # Use LAN IP so other devices on the network can reach the file.
    public_url = f"{browser_backend_base}/doc/download/{saved_name}"

    ext = _get_file_ext(saved_name)
    document_type = _get_document_type(ext)
    # OnlyOffice doc key MUST be ASCII-only (max 128 chars).
    # Use the UUID prefix (which is already pure hex) as the stable part,
    # then append a millisecond timestamp so each new edit session gets a fresh key
    # (prevents the "document already loaded" / 400 WebSocket error).
    uuid_prefix = saved_name.split('_')[0] if '_' in saved_name else saved_name[:32]
    import time as _time
    doc_key = f"{uuid_prefix}_{int(_time.time() * 1000) % 1_000_000}"
    plugin_config_url = f"{browser_backend_base}/doc/plugin/audit_navigator/config.json?v={doc_key}"

    plugin_config = {
        "autostart": ["asc.{823A43AE-971A-4C2E-8041-356C197BA3C8}"],
        "pluginsData": [plugin_config_url],
    }

    # Build editor config with Track Changes enabled by default
    editor_config = {
        "callbackUrl": callback_url,
        "mode": "edit",
        "lang": "zh-CN",
        "user": {
            "id": "audit_user",
            "name": "合规审计员"
        },
        "customization": {
            "trackChanges": True,
            "forcesave": True,
            "commentAuthorOnly": False,
            "review": {
                "showReviewChanges": True,
                "reviewDisplay": "markup",
                "trackChanges": True
            }
        },
        "plugins": plugin_config,
    }

    oo_payload = {
        "document": {
            "fileType": ext,
            "key": doc_key,
            "title": saved_name,
            "url": file_url,
            "directUrl": public_url,
        },
        "documentType": document_type,
        "editorConfig": editor_config,
    }

    # Sign the payload exactly once
    doc_token = jwt.encode(oo_payload, OO_JWT_SECRET, algorithm="HS256")
    return JSONResponse({
        "success": True,
        "savedName": saved_name,
        "issuesCount": len(issues),
        "hasNavigator": True, # the UI handles navigator regardless of plugin
        "message": "编辑链接已生成",
        "document": oo_payload["document"],
        "documentType": oo_payload["documentType"],
        "editorConfig": oo_payload["editorConfig"],
        "token": doc_token,
    })



@app.get("/doc/editor_page/{saved_name}")
async def doc_editor_page(saved_name: str, issues: str = ""):
    """
    Serve a standalone HTML page that loads OnlyOffice editor directly.
    This bypasses the iframe data URL approach which has origin issues.
    """
    issues_list = json.loads(issues) if issues else []

    file_path = DOCS_DIR / saved_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")

    file_url = f"http://{PUBLIC_HOST}:8000/doc/download/{saved_name}"
    callback_url = f"http://{PUBLIC_HOST}:8000/doc/callback"
    public_url = f"http://{PUBLIC_HOST}:8000/doc/download/{saved_name}"
    plugin_url = None

    if issues_list:
        try:
            plugin_js = _generate_audit_plugin_js(saved_name, issues_list)
            plugin_file_name = f"{saved_name}.plugin.js"
            plugin_path = PLUGINS_DIR / plugin_file_name
            with open(plugin_path, "w", encoding="utf-8") as pf:
                pf.write(plugin_js)
            plugin_url = f"http://{PUBLIC_HOST}:8000/doc/plugin_js/{plugin_file_name}"
        except Exception as e:
            logger.warning(f"Plugin generation failed: {e}")

    doc_key = saved_name
    ext = _get_file_ext(saved_name)

    payload = {
        "document": {
            "fileType": ext,
            "key": doc_key,
            "title": saved_name,
            "url": file_url,
            "directUrl": public_url,
        },
        "editorConfig": {
            "callbackUrl": callback_url,
            "mode": "edit",
            "lang": "zh-CN",
            "user": {
                "id": "audit_user",
                "name": "合规审计员"
            },
            "customization": {
                "trackChanges": True,
                "review": {
                    "showReviewChanges": True,
                    "reviewDisplay": "markup",
                    "trackChanges": True
                }
            }
        },
        "type": "desktop",
    }

        # Use the proper audit_navigator plugin which handles NAV_PARA/NAV_BOOKMARK messages
    # Use host IP so BOTH Docker (via host.docker.internal alias) and browser can resolve
    _backend_host = os.getenv("BACKEND_INTERNAL_HOST", "localhost")
    # host.docker.internal -> Docker can resolve, browser can't.
    # Use the actual host IP that Docker's extra_hosts resolves.
    # Docker's host-gateway maps to the host's actual network IP.
    _plugin_base = f"http://{_backend_host}:8000"
    payload["editorConfig"]["plugins"] = {
        "autostart": [
            {"guid": "asc.{823A43AE-971A-4C2E-8041-356C197BA3C8}", "url": f"{_plugin_base}/doc/plugin/audit_navigator/index.html"}
        ]
    }

    config_json = json.dumps(payload, ensure_ascii=False)

    # Build audit issue panel JS for in-editor navigation
    issue_panel_js = ""
    if issues_list:
        severity_colors = {"high": "#cf222e", "medium": "#bf8700", "low": "#1a7f37"}
        issue_items_html = ""
        for i, issue in enumerate(issues_list):
            sev = issue.get("severity", "medium")
            color = severity_colors.get(sev, "#656d76")
            issue_items_html += f"""
    var item = document.createElement('div');
    item.style.cssText = 'padding:7px 8px;cursor:pointer;border-radius:0 4px 4px 0;border-left:3px solid {color};margin-bottom:2px';
    item.innerHTML = '<div style="font-size:12px">#{i+1} {issue.get("desc", "")}</div>';
    (function(pi) {{ item.onclick = function() {{ jumpToBookmark(pi); }}; }})({issue.get("para_index")});
    panel.appendChild(item);"""

        issue_panel_js = f"""
  function jumpToBookmark(pi) {{
    var bm = "audit_para_" + pi;
    try {{ if (window.asc && window.asc.api) {{ window.asc.api.asc_goToBookmark(bm); }} }} catch(e) {{}}
  }}
  function buildPanel() {{
    if (!ISSUES || ISSUES.length === 0) return;
    var panel = document.createElement("div");
    panel.id = "audit-issue-panel";
    panel.style.cssText = "position:fixed;right:10px;top:60px;width:260px;max-height:80vh;background:#fff;border:1px solid #d0d7de;border-radius:8px;padding:12px;z-index:9999;box-shadow:0 4px 12px rgba(0,0,0,0.15);font-family:-apple-system,BlinkMacSystemFont,sans-serif;font-size:13px;overflow-y:auto";
    panel.innerHTML = '<div style="font-size:14px;font-weight:700;margin-bottom:8px;padding-bottom:8px;border-bottom:1px solid #eaecef">审核问题导航 <span style="font-size:11px;color:#656d76;margin-left:auto">' + ISSUES.length + '项</span></div>';
    {issue_items_html}
    document.body.appendChild(panel);
    if (ISSUES.length > 0 && ISSUES[0].para_index !== undefined) {{
      setTimeout(function() {{ jumpToBookmark(ISSUES[0].para_index); }}, 3000);
    }}
  }}
  window.addEventListener("load", function() {{ buildPanel(); }});
"""

    html = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <title>OnlyOffice - {saved_name}</title>
  <style>
    body {{ margin: 0; padding: 0; }}
    #placeholder {{ width: 100vw; height: 100vh; }}
  </style>
</head>
<body>
  <div id="placeholder"></div>
  <script src="http://{PUBLIC_HOST}:8081/web-apps/apps/api/documents/api.js"></script>
  <script>
    var CONFIG = {config_json};
    var ISSUES = {json.dumps(issues_list)};
    {issue_panel_js}
    function init() {{
      if (typeof DocsAPI !== "undefined") {{
        new DocsAPI.DocEditor("placeholder", CONFIG);
      }} else {{
        setTimeout(init, 1000);
      }}
    }}
    if (document.readyState === "complete" || document.readyState === "interactive") {{
      init();
    }} else {{
      window.addEventListener("load", init);
    }}
  </script>
</body>
</html>"""

    return HTMLResponse(content=html)


# ── Audit Navigator Plugin (unified, correct API implementation) ────────────────

@app.get("/doc/plugin/audit_navigator")
async def audit_plugin_root():
    """Redirect /doc/plugin/audit_navigator to index.html for OnlyOffice plugin loading."""
    from fastapi.responses import RedirectResponse
    return RedirectResponse(url="/doc/plugin/audit_navigator/index.html", status_code=302)

@app.get("/doc/plugin/audit_navigator/config.json")
async def audit_plugin_config(request: Request):
    """Return the OnlyOffice Audit Navigator plugin config."""
    browser_backend_base = _get_browser_backend_base(request)
    plugin_base = f"{browser_backend_base}/doc/plugin/audit_navigator/"
    config = {
        "name": "AuditNavigator",
        "guid": "asc.{823A43AE-971A-4C2E-8041-356C197BA3C8}",
        "version": "1.0",
        "baseUrl": plugin_base,
        "variations": [{
            "description": "合同审查风险导航插件",
            "url": "index.html",
            "icons": ["icon.png"],
            "isViewer": False,
            "isDisplayedInViewer": False,
            "EditorsSupport": ["word"],
            "isVisual": False,
            "isInsideMode": True,
            "initDataType": "none",
            "initData": "",
            "buttons": []
        }]
    }
    return Response(
        content=json.dumps(config),
        media_type="application/json",
        headers={"Cache-Control": "no-store, no-cache, must-revalidate, max-age=0"},
    )


@app.get("/doc/plugin/audit_navigator/index.html")
async def audit_plugin_html():
    """Audit Navigator plugin — correct Asc.plugin.executeMethod usage."""
    html = """<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
</head>
<body>
<script>
(function(window, undefined) {
    var _lastBookmark = "";

    window.Asc.plugin.init = function() {
        console.log("[AuditNav] Plugin initialized");
        try {
            window.parent && window.parent.postMessage(JSON.stringify({
                type: "auditNavResponse",
                data: { type: "PLUGIN_READY" }
            }), "*");
        } catch (e) {
            console.warn("[AuditNav] Ready notify failed:", e);
        }
    };

    // Extract paraIndex from bookmark name or build bookmark name from paraIndex
    function bmNameToParaIndex(bmName) {
        if (!bmName) return null;
        if (/^\\d+$/.test(bmName)) return parseInt(bmName, 10);
        var m = bmName.match(/^audit_para_(\\d+)$/);
        if (m) return parseInt(m[1], 10);
        return null;
    }

    // Build bookmark name from paraIndex
    function paraIndexToBmName(idx) {
        return 'audit_para_' + idx;
    }

    // Receive messages from React app via postMessage -> onExternalPluginMessage
    // Frontend sends: { type: 'NAV_PARA', paraIndex: n } or { type: 'NAV_BOOKMARK', bookmarkName: 'xxx' }
    window.Asc.plugin.onExternalPluginMessage = function(data) {
        console.log("[AuditNav] Received:", JSON.stringify(data));
        if (!data || !data.type) return;

        switch (data.type) {
            // Frontend sends NAV_PARA (number) -> convert to bookmark name
            case "NAV_PARA":
                if (typeof data.paraIndex === 'number') {
                    jumpToBookmark(paraIndexToBmName(data.paraIndex));
                }
                break;
            // Frontend sends NAV_BOOKMARK (string) -> bookmark name
            case "NAV_BOOKMARK":
                if (data.bookmarkName) {
                    jumpToBookmark(data.bookmarkName);
                }
                break;
            // Legacy JUMP_TO_BOOKMARK (for backward compatibility)
            case "JUMP_TO_BOOKMARK":
                jumpToBookmark(data.bookmarkName);
                break;
            case "LOCATE_TEXT":
                // data: { paraIndex, text, bookmarkName }
                if (data.paraIndex != null) {
                    jumpToBookmark(paraIndexToBmName(data.paraIndex));
                } else if (data.bookmarkName) {
                    jumpToBookmark(data.bookmarkName);
                }
                if (data.text) {
                    locateText(data.text, data.bookmarkName || (data.paraIndex != null ? paraIndexToBmName(data.paraIndex) : null));
                }
                break;
            case "REPLACE_TEXT":
                // data: { paraIndex, originalText, newText, bookmarkName }
                if (data.paraIndex != null) {
                    jumpToBookmark(paraIndexToBmName(data.paraIndex));
                } else if (data.bookmarkName) {
                    jumpToBookmark(data.bookmarkName);
                }
                if (data.originalText && data.newText) {
                    replaceText(data.originalText, data.newText, data.bookmarkName);
                }
                break;
            case "GET_SELECTION":
                getSelectionAndReport(data.bookmarkName);
                break;
            default:
                console.warn("[AuditNav] Unknown type:", data.type);
        }
    };

    function jumpToBookmark(bmName) {
        if (!bmName) return;
        _lastBookmark = bmName;
        console.log("[AuditNav] GoToBookmark:", bmName);
        window.Asc.plugin.executeMethod("GoToBookmark", [bmName], function(result) {
            console.log("[AuditNav] GoToBookmark result:", result);
            reportResult("NAV_BOOKMARK_RESULT", result, bmName);
        });
    }

    // Legacy fallback text replacement (no highlight, just text replace)
    function replaceText(original, replacement, bmName) {
        if (!original || !replacement) return;
        console.log("[AuditNav] Replace via plugin fallback:", original, "=>", replacement);
        window.Asc.plugin.executeMethod("SearchAndReplace", [{
            searchString: original,
            replaceString: replacement,
            matchCase: false
        }], function(result) {
            reportResult("REPLACE_RESULT", result, bmName || _lastBookmark);
        });
    }

    // Legacy fallback locate text (no highlight, just text find+select)
    function locateText(searchText, bmName) {
        if (!searchText) return;
        _lastBookmark = bmName || _lastBookmark;
        console.log("[AuditNav] Locate via plugin fallback:", searchText);
        // We use SearchNext instead of SearchAndReplace for locating
        window.Asc.plugin.executeMethod("SearchNext", [{
            searchString: searchText,
            matchCase: false
        }, true], function(result) {
            reportResult("LOCATE_RESULT", result, _lastBookmark);
        });
    }

    function reportResult(type, success, bmName) {
        var resp = { type: type, success: !!success, bookmarkName: bmName || _lastBookmark };
        window.parent.postMessage({ type: "auditNavResponse", data: resp }, "*");
    }

    function getSelectionAndReport(bmName) {
        _lastBookmark = bmName || _lastBookmark;
        window.Asc.plugin.executeMethod("GetSelectionText", [], function(text) {
            var payload = {
                type: "SELECTION_REPORT",
                bookmarkName: _lastBookmark,
                selectedText: text || "(无选中文本)"
            };
            console.log("[AuditNav] Selection report:", payload);
            window.parent.postMessage({ type: "auditNavResponse", data: payload }, "*");
        });
    }

    window.Asc.plugin.button = function() {};
})(window, undefined);
</script>
</body>
</html>"""
    return Response(
        content=html,
        media_type="text/html",
        headers={"Cache-Control": "no-store, no-cache, must-revalidate, max-age=0"},
    )


@app.get("/doc/plugin/audit_navigator/icon.png")
async def audit_plugin_icon():
    import base64
    icon_data = base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAACklEQVR4nGMAAQAABQABDQottAAAAABJRU5ErkJggg==")
    return Response(content=icon_data, media_type="image/png")


@app.post("/doc/callback")
async def doc_callback(request: Request):
    """Handle OnlyOffice save callback — download edited file and replace original."""
    try:
        body = await request.json()
    except Exception:
        body = {}

    status = body.get("status", 0)
    if status == 2:
        download_url = body.get("url", "")
        if download_url:
            try:
                async with httpx.AsyncClient(timeout=30.0, trust_env=False) as client:
                    resp = await client.get(download_url)
                    resp.raise_for_status()
                    content = resp.content

                key = body.get("key", "")
                if key and (DOCS_DIR / key).exists():
                    with open(DOCS_DIR / key, "wb") as f:
                        f.write(content)
                    logger.info(f"OnlyOffice callback: file {key} updated successfully")
                    return JSONResponse({"error": 0, "message": "saved"})
            except Exception as e:
                logger.error(f"OnlyOffice callback download failed: {e}")
                return JSONResponse({"error": 1, "message": str(e)})

    return JSONResponse({"error": 0})



class SelectionRequest(BaseModel):
    saved_name: str
    bookmark_name: str

class SuggestionRequest(BaseModel):
    saved_name: str
    bookmark_name: str
    original_text: str
    suggested_text: str
    reason: str = ""


class ReviewedDocExportRequest(BaseModel):
    saved_name: str
    edits: dict[str, str]

@app.post("/doc/selection")
async def doc_selection(request: Request, body: SelectionRequest):
    """Return the paragraph text for a given bookmark (for confirming the selection)."""
    _get_request_user(request, required=True)
    saved_name = body.saved_name
    bookmark_name = body.bookmark_name

    # Load meta to get para_index from bookmark_name
    meta_path = DOCS_DIR / f"{saved_name}.meta.json"
    if not meta_path.exists():
        raise HTTPException(status_code=404, detail="文档 meta 信息不存在")

    try:
        with open(meta_path, "r", encoding="utf-8") as f:
            meta = json.load(f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    # Find the paragraph with matching bookmark_name
    para_entry = None
    for p in meta.get("paragraphs", []):
        if p.get("bookmark_name") == bookmark_name:
            para_entry = p
            break

    if not para_entry:
        raise HTTPException(status_code=404, detail=f"未找到书签 '{bookmark_name}'")

    return JSONResponse({
        "saved_name": saved_name,
        "bookmark_name": bookmark_name,
        "para_index": para_entry["para_index"],
        "text_preview": para_entry.get("text_preview", ""),
    })


@app.post("/doc/submit_suggestion")
async def doc_submit_suggestion(request: Request, body: SuggestionRequest):
    """Record a text-replacement suggestion for a given bookmark."""
    _get_request_user(request, required=True)
    saved_name = body.saved_name
    bookmark_name = body.bookmark_name

    # Load existing suggestions or create new
    suggestions_file = DOCS_DIR / f"{saved_name}.suggestions.json"
    suggestions = []
    if suggestions_file.exists():
        try:
            with open(suggestions_file, "r", encoding="utf-8") as f:
                suggestions = json.load(f)
        except Exception:
            suggestions = []

    # Add new suggestion
    suggestions.append({
        "bookmark_name": bookmark_name,
        "original_text": body.original_text,
        "suggested_text": body.suggested_text,
        "reason": body.reason,
        "status": "pending",
    })

    # Save
    try:
        with open(suggestions_file, "w", encoding="utf-8") as f:
            json.dump(suggestions, f, ensure_ascii=False, indent=2)
        return JSONResponse({"success": True, "message": "建议已提交", "total": len(suggestions)})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/doc/export_reviewed")
async def doc_export_reviewed(request: Request, body: ReviewedDocExportRequest):
    """Export a reviewed .docx copy with accepted/custom paragraph edits as tracked changes."""
    _get_request_user(request, required=True)
    export_saved_name, download_name, applied_count = _export_reviewed_docx(body.saved_name, body.edits)
    return JSONResponse({
        "success": True,
        "saved_name": export_saved_name,
        "filename": download_name,
        "applied_count": applied_count,
        "download_url": f"/doc/download/{export_saved_name}",
        "message": f"已生成留痕审查版文件，共写入 {applied_count} 处修订",
    })


CONTRACT_DATA_DIR = Path(os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "contracts"))
CONTRACT_DATA_DIR.mkdir(parents=True, exist_ok=True)


def _compute_para_id(text: str) -> str:
    """Stable paragraph ID: first 8 chars of MD5 of normalized text."""
    normalized = re.sub(r'\s+', '', text.strip())
    return hashlib.md5(normalized.encode('utf-8')).hexdigest()[:12]


def _compute_text_hash(text: str) -> str:
    """Content fingerprint: full SHA256 of normalized text."""
    normalized = re.sub(r'\s+', '', text.strip())
    return hashlib.sha256(normalized.encode('utf-8')).hexdigest()


def _load_docx_paragraphs(docx_path: Path) -> List[dict]:
    """Load all non-empty paragraphs from a .docx file with stable IDs."""
    from docx import Document
    paragraphs = []
    doc = Document(docx_path)
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append({
            "para_index": idx,
            "para_id": _compute_para_id(text),
            "text": text,
            "hash": _compute_text_hash(text),
        })
    return paragraphs


def _replace_paragraph_text(para, new_text: str):
    """Replace paragraph text while preserving the paragraph style container."""
    runs = list(para.runs)
    if runs:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
        return

    para.text = new_text


def _enable_docx_track_revisions(document):
    """Enable track revisions in the exported Word document."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = document.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.append(OxmlElement("w:trackRevisions"))


def _append_tracked_text(parent, change_tag: str, text: str, change_id: int, author: str, date_value: str):
    """Append a WordprocessingML insertion/deletion change block."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    change = OxmlElement(change_tag)
    change.set(qn("w:id"), str(change_id))
    change.set(qn("w:author"), author)
    change.set(qn("w:date"), date_value)

    text_tag = "w:delText" if change_tag == "w:del" else "w:t"
    lines = str(text or "").splitlines() or [""]
    for line_index, line in enumerate(lines):
        if line_index > 0:
            br_run = OxmlElement("w:r")
            br_run.append(OxmlElement("w:br"))
            change.append(br_run)

        run = OxmlElement("w:r")
        text_node = OxmlElement(text_tag)
        text_node.set(qn("xml:space"), "preserve")
        text_node.text = line
        run.append(text_node)
        change.append(run)

    parent.append(change)


def _mark_paragraph_review_change(para, new_text: str, change_id: int):
    """Replace a paragraph with Word track-changes deletion + insertion."""
    original_text = para.text
    parent = para._p

    for child in list(parent):
        if child.tag.endswith("}pPr"):
            continue
        parent.remove(child)

    date_value = datetime.utcnow().isoformat(timespec="seconds") + "Z"
    author = "AI 合同审查"

    if original_text:
        _append_tracked_text(parent, "w:del", original_text, change_id, author, date_value)
    _append_tracked_text(parent, "w:ins", new_text, change_id + 1, author, date_value)


def _export_reviewed_docx(saved_name: str, edits: dict[str, str]) -> tuple[str, str, int]:
    """Write paragraph-level edits into a new reviewed .docx file with tracked changes."""
    from docx import Document

    source_path = DOCS_DIR / saved_name
    if not source_path.exists():
        raise HTTPException(status_code=404, detail="原始文档不存在")

    if _get_file_ext(saved_name) != "docx":
        raise HTTPException(status_code=400, detail="当前仅支持导出 .docx 审查版文件")

    normalized_edits = {}
    for key, value in (edits or {}).items():
        try:
            para_index = int(key)
        except (TypeError, ValueError):
            continue
        text = str(value or "").strip()
        if not text:
            continue
        normalized_edits[para_index] = text

    if not normalized_edits:
        raise HTTPException(status_code=400, detail="没有可导出的修改内容")

    document = Document(source_path)
    _enable_docx_track_revisions(document)
    applied = 0
    for para_index, replacement in normalized_edits.items():
        if 0 <= para_index < len(document.paragraphs):
            _mark_paragraph_review_change(document.paragraphs[para_index], replacement, applied * 2 + 1)
            applied += 1

    if applied == 0:
        raise HTTPException(status_code=422, detail="没有匹配到可回写的段落")

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    source_stem = Path(saved_name).stem
    export_name = f"{source_stem}_留痕审查版_{timestamp}.docx"
    export_path = DOCS_DIR / export_name
    document.save(export_path)

    export_meta = {
        "source_saved_name": saved_name,
        "exported_saved_name": export_name,
        "applied_count": applied,
        "exported_at": datetime.now().isoformat(),
        "edits": normalized_edits,
    }
    with open(DOCS_DIR / f"{export_name}.review.json", "w", encoding="utf-8") as mf:
        json.dump(export_meta, mf, ensure_ascii=False, indent=2)

    return export_name, export_path.name, applied


def _search_legal_provisions(query: str, top_k: int = 5) -> List[dict]:
    """Vector search in ChromaDB for relevant legal provisions."""
    store = _get_vectorstore()
    if store is None:
        return []
    try:
        results = store.similarity_search_with_score(query, k=top_k)
        provisions = []
        for doc, score in results:
            provisions.append({
                "content": doc.page_content[:500],
                "source": doc.metadata.get("source", ""),
                "score": float(score),
            })
        return provisions
    except Exception as e:
        logger.warning(f"Legal provisions search failed: {e}")
        return []


def _get_contract_issues_file(saved_name: str) -> Path:
    """Path to the issues JSON sidecar for a contract."""
    return CONTRACT_DATA_DIR / f"{saved_name}.issues.json"


# ── Request/Response models ────────────────────────────────────────────────────

class ContractMapRequest(BaseModel):
    saved_name: str


class ContractAnalyzeRequest(BaseModel):
    saved_name: str
    doc_structure: Optional[List[dict]] = None  # from map_doc_structure
    extra_questions: Optional[List[str]] = None  # user追加的问题


class ContractReAnalyzeRequest(BaseModel):
    saved_name: str
    doc_structure: List[dict]
    extra_questions: List[str]
    previous_issues: Optional[List[dict]] = None

  # for diff tracking# ── Endpoints ──────────────────────────────────────────────────────────────────

@app.post("/contract/map_doc_structure")
@app.post("/api/contract/map_doc_structure")
async def contract_map_doc_structure(body: ContractMapRequest):
    """
    Parse a .docx file into a list of paragraph structures.
    Returns stable para_id + text_hash for each non-empty paragraph.
    """
    saved_name = body.saved_name
    file_path = DOCS_DIR / saved_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")

    ext = _get_file_ext(saved_name)
    if ext != "docx":
        raise HTTPException(status_code=400, detail="仅支持 .docx 文件")

    try:
        paragraphs = _load_docx_paragraphs(file_path)
        if not paragraphs:
            raise HTTPException(status_code=422, detail="文档中未找到有效段落")

        # Compute overall document hash
        all_text = "".join(p["text"] for p in paragraphs)
        doc_hash = _compute_text_hash(all_text[:5000])

        return JSONResponse({
            "success": True,
            "saved_name": saved_name,
            "doc_hash": doc_hash,
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs,
        })
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"contract_map_doc_structure failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/contract/analyze")
@app.post("/api/contract/analyze")
async def contract_analyze(body: ContractAnalyzeRequest):
    """
    Two-stage contract review:
    Stage 1: Classify contract clauses (rights/obligations, breach, payment, termination, etc.)
    Stage 2: For each clause, search ChromaDB for relevant legal provisions,
             then generate risk review points.
    """
    saved_name = body.saved_name
    doc_structure = body.doc_structure
    extra_questions = body.extra_questions or []

    # Load doc_structure if not provided
    if not doc_structure:
        map_resp = await contract_map_doc_structure(ContractMapRequest(saved_name=saved_name))
        raw = map_resp.body.decode()
        doc_structure = json.loads(raw)["paragraphs"]

    # Build clause text for the LLM
    clauses_text = "\n".join(
        f"[{i+1}] [{p['para_id']}] {p['text']}"
        for i, p in enumerate(doc_structure)
    )

    # Stage 1: Clause classification + risk identification
    clause_prompt = f"""你是资深合同法专家。请对以下合同进行结构化解构，识别所有关键条款并分类。

合同内容：
{clauses_text}

请按以下 JSON 数组格式输出，每条为一个审查维度（必须至少覆盖：合同主体、权利义务、违约责任、付款条款、解除条款、其他重要条款）：

[
  {{
    "category": "权利义务",
    "para_ids": ["a9f23feca546"],
    "clause_text": "相关条款的原文摘要（50字以内）",
    "risk_summary": "该条款的核心风险点一句话描述"
  }},
  ...
]

要求：
- category 只允许：合同主体 | 权利义务 | 违约责任 | 付款条款 | 解除条款 | 知识产权 | 争议解决 | 其他重要条款
- para_ids: 必须精确填入合同内容中对应段落的 12 位 ID（如 ["a9f23feca546"]），没有则填空数组
- 必须覆盖全部8个分类（没有的也要写 "无相关内容"
- 输出必须是合法的单个JSON数组，不要有任何额外文字"""
    try:
        resp = llm._generate(
            [clause_prompt],
            enable_thinking=False
        )
        clause_result_text = resp.generations[0].message.content.strip()
        # Extract JSON from response
        import re as _re
        json_match = _re.search(r'\[.*\]', clause_result_text, _re.DOTALL)
        if json_match:
            clauses_data = json.loads(json_match.group())
        else:
            clauses_data = []
    except Exception as e:
        logger.error(f"Stage1 clause classification failed: {e}")
        clauses_data = []

    # Stage 2: For each clause, RAG search + risk generation
    review_points = []
    issue_id_counter = 1

    for clause in clauses_data:
        category = clause.get("category", "其他重要条款")
        clause_text = clause.get("clause_text", "")
        risk_summary = clause.get("risk_summary", "")
        para_ids = clause.get("para_ids", [])

        if not clause_text or clause_text == "无相关内容":
            continue

        # RAG search for relevant legal provisions
        search_query = f"{category} {clause_text} {risk_summary}"
        provisions = _search_legal_provisions(search_query, top_k=3)

        provisions_text = ""
        if provisions:
            provisions_text = "相关法律条文：\n" + "\n".join(
                f"- {p['content'][:200]}（来源：{p['source']}）"
                for p in provisions
            )
        else:
            provisions_text = "未检索到直接匹配的法条"

        # Match para_ids back to paragraph data for stable IDs + hash
        matched_paras = [p for p in doc_structure if p["para_id"] in para_ids]
        
        # Fallback: if LLM didn't return valid para_ids, use text similarity to find the best match
        if not matched_paras and clause_text and clause_text != "无相关内容":
            best_p = None
            best_score = 0
            # Simple keyword overlap scoring
            keywords = [kw for kw in clause_text if '\u4e00' <= kw <= '\u9fff'] 
            for p in doc_structure:
                score = sum(1 for kw in keywords if kw in p["text"])
                # Boost score if it's a direct substring
                if clause_text in p["text"] or p["text"] in clause_text:
                    score += 100
                if score > best_score:
                    best_score = score
                    best_p = p
            if best_p and best_score > 0:
                matched_paras = [best_p]

        primary_para = matched_paras[0] if matched_paras else doc_structure[0]

        severity_prompt = f"""基于以下合同条款和检索到的法律依据，判断风险等级。

条款类别：{category}
条款摘要：{clause_text}
风险描述：{risk_summary}
{provisions_text}

只输出一个词：high / medium / low"""
        try:
            sev_resp = llm._generate([severity_prompt], enable_thinking=False)
            severity = sev_resp.generations[0].message.content.strip().lower()
            if severity not in ("high", "medium", "low"):
                severity = "medium"
        except Exception:
            severity = "medium"

        # Generate modification suggestion
        suggest_prompt = f"""你是资深合同法专家。基于以下信息，生成具体的修改建议。

条款类别：{category}
条款摘要：{clause_text}
风险描述：{risk_summary}
风险等级：{severity}
{provisions_text}

输出格式（直接输出JSON，不要任何额外文字）：
{{
  "issue_desc": "一句话描述该风险",
  "suggested_text": "修改后的建议文本（如果可以修改的话）",
  "reason": "为什么这样修改"
}}"""
        try:
            sug_resp = llm._generate([suggest_prompt], enable_thinking=False)
            sug_text = sug_resp.generations[0].message.content.strip()
            json_match2 = _re.search(r'\{.*\}', sug_text, _re.DOTALL)
            if json_match2:
                suggestion = json.loads(json_match2.group())
            else:
                suggestion = {"issue_desc": risk_summary, "suggested_text": "", "reason": "需人工审核"}
        except Exception as e:
            logger.warning(f"Suggestion generation failed: {e}")
            suggestion = {"issue_desc": risk_summary, "suggested_text": "", "reason": "生成失败"}

        # Build rule text from provisions
        rule_text = "\n".join(
            f"- {p['content'][:300]}"
            for p in provisions[:3]
        ) if provisions else ""

        review_points.append({
            "id": issue_id_counter,
            "para_id": primary_para["para_id"],
            "para_index": primary_para["para_index"],
            "text_hash": primary_para["hash"],
            "originalText": primary_para["text"],
            "category": category,
            "severity": severity,
            "issueDesc": suggestion.get("issue_desc", risk_summary),
            "suggestedText": suggestion.get("suggested_text", ""),
            "reason": suggestion.get("reason", ""),
            "rule": rule_text,
            "status": "open",
        })
        issue_id_counter += 1

    # Handle extra questions from user
    for q in extra_questions:
        # Find most relevant paragraph by keyword matching
        best_match = None
        best_score = 0
        q_keywords = _re.findall(r'[\u4e00-\u9fff]+', q)
        for p in doc_structure:
            score = sum(1 for kw in q_keywords if kw in p["text"])
            if score > best_score:
                best_score = score
                best_match = p

        if best_match:
            provisions = _search_legal_provisions(q, top_k=2)
            rule_text = "\n".join(p["content"][:300] for p in provisions) if provisions else ""

            review_points.append({
                "id": issue_id_counter,
                "para_id": best_match["para_id"],
                "para_index": best_match["para_index"],
                "text_hash": best_match["hash"],
                "originalText": best_match["text"],
                "category": "用户关注",
                "severity": "medium",
                "issueDesc": q,
                "suggestedText": "",
                "reason": "用户追加问题",
                "rule": rule_text,
                "status": "open",
            })
            issue_id_counter += 1

# Enrich review_points with bookmark_name from meta.json
    meta_path = DOCS_DIR / f"{saved_name}.meta.json"
    idx_to_bm = {}
    if meta_path.exists():
        try:
            with open(meta_path, "r", encoding="utf-8") as mf:
                meta = json.load(mf)
            idx_to_bm = {p["para_index"]: p["bookmark_name"] for p in meta.get("paragraphs", [])}
        except Exception:
            pass
        for rp in review_points:
            rp["bookmark_name"] = idx_to_bm.get(rp["para_index"], f"audit_para_{rp['para_index']}")
    issues_path = _get_contract_issues_file(saved_name)
    try:
        with open(issues_path, "w", encoding="utf-8") as f:
            json.dump({
                "saved_name": saved_name,
                "doc_structure": doc_structure,
                "issues": review_points,
            }, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.warning(f"Failed to save issues: {e}")

    return JSONResponse({
        "success": True,
        "saved_name": saved_name,
        "review_points_count": len(review_points),
        "review_points": review_points,
        "message": f"审查完成，发现 {len(review_points)} 个审查点",
    })


@app.post("/contract/re_analyze")
@app.post("/api/contract/re_analyze")
async def contract_re_analyze(body: ContractReAnalyzeRequest):
    """
    Re-analyze contract with user's extra questions.
    Diff against previous issues: resolved / new / modified.
    """
    saved_name = body.saved_name
    doc_structure = body.doc_structure
    extra_questions = body.extra_questions or []
    previous_issues = body.previous_issues or []

    if not extra_questions:
        return JSONResponse({"success": True, "review_points": previous_issues, "diff": {
            "new": [], "resolved": [], "modified": []
        }})

    # Run new analysis with extra questions (reuse the analyze logic)
    try:
        new_resp = await contract_analyze(ContractAnalyzeRequest(
            saved_name=saved_name,
            doc_structure=doc_structure,
            extra_questions=extra_questions,
        ))
        new_issues = json.loads(new_resp.body.decode())["review_points"]
    except Exception as e:
        logger.error(f"Re-analyze failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

    # Build previous issue lookup by para_id + text_hash
    prev_lookup = {(i.get("para_id"), i.get("text_hash")): i for i in previous_issues}

    diff = {"new": [], "resolved": [], "modified": []}
    all_seen = set()

    for ni in new_issues:
        key = (ni.get("para_id"), ni.get("text_hash"))
        all_seen.add(key)
        if key in prev_lookup:
            prev_i = prev_lookup[key]
            # Check if description changed significantly
            if ni.get("issueDesc") != prev_i.get("issueDesc"):
                diff["modified"].append({**ni, "previous_issue": prev_i})
            # else unchanged, skip
        else:
            diff["new"].append(ni)

    # Resolved = previous issues whose para_id+hash no longer appear
    for pi in previous_issues:
        key = (pi.get("para_id"), pi.get("text_hash"))
        if key not in all_seen and pi.get("status") != "resolved":
            diff["resolved"].append({**pi, "status": "resolved"})

    # Merge: keep old (unresolved) + add new + add modified
    merged = []
    seen_keys = set()
    for ni in new_issues:
        key = (ni.get("para_id"), ni.get("text_hash"))
        if key not in seen_keys:
            merged.append(ni)
            seen_keys.add(key)
    for pi in previous_issues:
        key = (pi.get("para_id"), pi.get("text_hash"))
        if key not in seen_keys and pi.get("status") != "resolved":
            merged.append({**pi, "status": "resolved"})
            seen_keys.add(key)

    # Save updated issues
    issues_path = _get_contract_issues_file(saved_name)
    try:
        with open(issues_path, "w", encoding="utf-8") as f:
            json.dump({
                "saved_name": saved_name,
                "doc_structure": doc_structure,
                "issues": merged,
            }, f, ensure_ascii=False, indent=2)
    except Exception:
        logger.exception("合同审查结果保存失败")

    return JSONResponse({
        "success": True,
        "review_points": merged,
        "diff": diff,
        "message": f"追加审查完成，新增 {len(diff['new'])} 个，修改 {len(diff['modified'])} 个，化解 {len(diff['resolved'])} 个",
    })


@app.get("/contract/issues/{saved_name}")
async def contract_get_issues(saved_name: str):
    """Get saved review issues for a contract."""
    issues_path = _get_contract_issues_file(saved_name)
    if not issues_path.exists():
        raise HTTPException(status_code=404, detail="未找到审查记录")
    with open(issues_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return JSONResponse(data)


@app.delete("/contract/issues/{saved_name}")
async def contract_delete_issues(saved_name: str):
    """Delete saved review issues for a contract."""
    issues_path = _get_contract_issues_file(saved_name)
    if issues_path.exists():
        issues_path.unlink()
    return JSONResponse({"success": True})


# ── Contract Draft (AI 起草) ────────────────────────────────────────────────

class ContractDraftRequest(BaseModel):
    contract_type: str = "采购合同"  # e.g. 采购合同, 服务合同, 劳动合同, 租赁合同
    requirements: str  # user's requirements in natural language


class ContractDraftExportRequest(BaseModel):
    markdown: str
    title: str = "合同草案"


CONTRACT_DRAFT_PROMPT = """你是资深合同法律师。请根据用户需求起草一份完整、专业的中文合同。

合同类型：{contract_type}
用户需求：{requirements}

请按以下结构输出 Markdown 格式的完整合同：

## 合同标题
（使用 # 一级标题，格式如：＃ 采购合同）

## 合同双方
（甲方、乙方基本信息，留空用 [请填写] 标出）

## 一、合同标的
（清晰描述合同标的物/服务内容）

## 二、合同价款与支付方式
（金额、币种、支付节点、发票要求）

## 三、双方权利与义务
（分甲方、乙方列出）

## 四、交付/履行期限与地点

## 五、质量与验收标准

## 六、违约责任
（分情形写明违约金、赔偿计算方式）

## 七、合同的变更与解除

## 八、保密条款

## 九、知识产权（如适用）

## 十、争议解决
（协商→仲裁/诉讼，写明管辖法院或仲裁机构）

## 十一、不可抗力

## 十二、其他约定
（通知送达、合同份数、生效条件等）

## 签署栏
（甲方/乙方签章、日期、联系方式）

要求：
1. 条款要具体、可执行，不要笼统模板语言
2. 用 [请填写] 标出需要用户填入的具体信息
3. 金额、日期、人名等具体信息用占位符
4. 违约条款要明确违约金比例或计算方式
5. 语言专业、严谨，符合中国合同法的规范
6. 输出纯 Markdown，不要有额外解释"""


@app.post("/contract/draft")
@app.post("/api/contract/draft")
async def contract_draft(body: ContractDraftRequest):
    """AI generates a contract draft in Markdown format based on user requirements."""
    prompt = CONTRACT_DRAFT_PROMPT.format(
        contract_type=body.contract_type,
        requirements=body.requirements,
    )
    try:
        resp = llm._generate([prompt], enable_thinking=True)
        markdown = resp.generations[0].message.content.strip()
        # Extract title from first # heading
        title = body.contract_type
        for line in markdown.split("\n"):
            line = line.strip()
            if line.startswith("# ") and not line.startswith("## "):
                title = line[2:].strip()
                break
        return JSONResponse({
            "success": True,
            "markdown": markdown,
            "title": title,
        })
    except Exception as e:
        logger.error(f"Contract draft generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"合同生成失败：{str(e)}")


@app.post("/contract/draft/export")
@app.post("/api/contract/draft/export")
async def contract_draft_export(body: ContractDraftExportRequest):
    """Convert AI-generated contract markdown to a professionally formatted Word document."""
    import re as _re
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from io import BytesIO

    markdown = body.markdown
    title = body.title or "合同草案"

    CJK_FONT = "仿宋"
    HEADING_FONT = "黑体"
    TITLE_FONT = "方正小标宋简体"

    def _cjk_run(run, font=CJK_FONT, size=Pt(12), bold=False, color=None):
        run.font.name = font
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        for k in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(k), font)
        rPr.insert(0, rFonts)
        run.font.size = size
        run.bold = bold
        if color:
            run.font.color.rgb = color

    def _add_styled_para(doc, text, font=CJK_FONT, size=Pt(12), bold=False,
                         alignment=None, indent=True, color=None, space_after=Pt(6)):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        if alignment is not None:
            p.alignment = alignment
        p.paragraph_format.space_after = space_after
        run = p.add_run(text)
        _cjk_run(run, font, size, bold=bold, color=color)
        return p

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # ── 红头标题 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(0)
    run = title_para.add_run(title)
    _cjk_run(run, TITLE_FONT, Pt(22), bold=True, color=RGBColor(180, 35, 24))

    # 红线
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_p.paragraph_format.space_before = Pt(2)
    line_p.paragraph_format.space_after = Pt(18)
    run = line_p.add_run("━" * 35)
    _cjk_run(run, CJK_FONT, Pt(10), color=RGBColor(180, 35, 24))

    # ── 解析 Markdown 并转换为 Word 格式 ──
    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Skip the main title (already rendered as 红头)
        if line.startswith("# ") and not line.startswith("## "):
            i += 1
            continue

        # Section headings (##)
        if line.startswith("## "):
            heading_text = line[3:].strip()
            h = doc.add_heading(level=2)
            run = h.add_run(heading_text)
            _cjk_run(run, HEADING_FONT, Pt(15), bold=True)
            i += 1
            continue

        # Sub-section headings (###)
        if line.startswith("### "):
            sub_text = line[4:].strip()
            h = doc.add_heading(level=3)
            run = h.add_run(sub_text)
            _cjk_run(run, HEADING_FONT, Pt(13), bold=True)
            i += 1
            continue

        # Numbered list items (1. / 2. / 一、etc.)
        if _re.match(r'^[\d一二三四五六七八九十]+[\.\、）)]', line):
            _add_styled_para(doc, line)
            i += 1
            continue

        # Bullet points
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            _add_styled_para(doc, text)
            i += 1
            continue

        # Table detection: lines with | (pipe)
        if "|" in line and line.count("|") >= 2:
            # Collect table rows
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            # Parse table
            rows_data = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split("|")[1:-1]]
                rows_data.append(cells)
            if len(rows_data) >= 2:
                # Skip separator row (---|---)
                header = rows_data[0]
                body_rows = [r for r in rows_data[1:] if not all(_re.match(r'^[-:]+$', c) for c in r)]
                if header and body_rows:
                    t = doc.add_table(rows=1 + len(body_rows), cols=len(header))
                    t.style = "Table Grid"
                    for j, h_cell in enumerate(header):
                        cell = t.rows[0].cells[j]
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(h_cell)
                        _cjk_run(run, HEADING_FONT, Pt(10.5), bold=True)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for ri, row in enumerate(body_rows):
                        for j, val in enumerate(row[:len(header)]):
                            cell = t.rows[ri + 1].cells[j]
                            cell.text = ""
                            run = cell.paragraphs[0].add_run(val)
                            _cjk_run(run, CJK_FONT, Pt(10.5))
                    doc.add_paragraph("")
            continue

        # Regular paragraph
        # Strip markdown bold/italic markers
        clean = _re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        clean = _re.sub(r'\*(.+?)\*', r'\1', clean)
        clean = _re.sub(r'__(.+?)__', r'\1', clean)
        clean = _re.sub(r'_(.+?)_', r'\1', clean)
        _add_styled_para(doc, clean)
        i += 1

    # ── 页脚/签署栏分隔 ──
    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(12)
    run = footer_p.add_run("— 合同草案由 AI 生成，请经法律审核后签署 —")
    _cjk_run(run, CJK_FONT, Pt(9), color=RGBColor(150, 150, 150))

    # Save to docs directory
    safe_title = _re.sub(r'[\\/*?:"<>|]', '', title)
    safe_name = f"{uuid.uuid4().hex}_{safe_title}.docx"
    file_path = DOCS_DIR / safe_name

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    with open(file_path, "wb") as f:
        f.write(buf.read())

    logger.info(f"Contract draft exported: {safe_name} ({title})")
    return JSONResponse({
        "success": True,
        "filename": f"{safe_title}.docx",
        "saved_as": safe_name,
        "download_url": f"/doc/download/{safe_name}",
    })


# ═══════════════════════════════════════════════════════════════════════
# 会议计时器 + 通知 + 跨会议搜索 + 仪表盘
# ═══════════════════════════════════════════════════════════════════════

@app.post("/api/meetings/{meeting_id}/timer/{action}")
async def meeting_timer(
    request: Request,
    meeting_id: str,
    action: str,
    duration_minutes: Optional[int] = None,
):
    """会议计时器：start / pause / reset / set-duration"""
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    now = _now_text()
    with _db_connect() as conn:
        meeting = conn.execute("SELECT * FROM meetings WHERE id = ?", (safe_id,)).fetchone()
    if not meeting:
        raise HTTPException(status_code=404, detail="会议不存在")
    _check_meeting_access(user, dict(meeting))
    if action == "start":
        with _db_connect() as conn:
            conn.execute("UPDATE meetings SET timer_started_at = ? WHERE id = ?", (now, safe_id))
    elif action == "reset":
        with _db_connect() as conn:
            conn.execute("UPDATE meetings SET timer_started_at = '' WHERE id = ?", (safe_id,))
    elif action == "set-duration" and duration_minutes is not None:
        with _db_connect() as conn:
            conn.execute("UPDATE meetings SET agenda_duration_minutes = ? WHERE id = ?", (duration_minutes, safe_id))
    _invalidate_meetings_cache()
    return JSONResponse({
        "success": True,
        "meetingId": safe_id,
        "action": action,
        "timerStartedAt": now if action == "start" else (meeting["timer_started_at"] or ""),
        "durationMinutes": duration_minutes if action == "set-duration" else (meeting["agenda_duration_minutes"] or 15),
    })


@app.get("/api/notifications")
async def get_notifications(request: Request):
    """获取当前用户的通知。顺便触发待办到期检查。"""
    user = _get_request_user(request, required=True)
    user_id = user.get("id") or user.get("username", "")
    _check_todo_deadlines()  # 每次拉通知时顺带检查
    with _db_connect() as conn:
        rows = conn.execute(
            "SELECT * FROM notifications WHERE user_id = ? OR user_id = '' ORDER BY created_at DESC LIMIT 50",
            (user_id,),
        ).fetchall()
    return JSONResponse({"notifications": [dict(row) for row in rows]})


@app.post("/api/notifications/{notification_id}/read")
async def mark_notification_read(request: Request, notification_id: str):
    """标记通知为已读。"""
    _get_request_user(request, required=True)
    with _db_connect() as conn:
        conn.execute("UPDATE notifications SET read = 1 WHERE id = ?", (notification_id,))
    return JSONResponse({"success": True})


def _create_notification(user_id: str, type_: str, title: str, body: str, meeting_id: str = ""):
    """创建一条通知（内部调用）。"""
    try:
        with _db_connect() as conn:
            conn.execute(
                "INSERT INTO notifications (id, user_id, type, title, body, meeting_id, created_at, read) VALUES (?, ?, ?, ?, ?, ?, ?, 0)",
                (f"notif_{uuid.uuid4().hex[:10]}", user_id, type_, title, body, meeting_id, _now_text()),
            )
    except Exception:
        pass  # 通知失败不影响主流程


def _check_todo_deadlines():
    """检查待办截止时间，对即将到期和已逾期的待办创建通知。
    每次调用检查一次，避免重复通知（同一待办每天最多通知一次）。
    """
    import hashlib
    today = datetime.now().strftime("%Y-%m-%d")
    try:
        with _db_connect() as conn:
            # 查找未完成、有截止时间的待办
            rows = conn.execute(
                """SELECT * FROM meeting_todos
                   WHERE status NOT IN ('已完成', '已取消')
                   AND deadline != '' AND deadline IS NOT NULL"""
            ).fetchall()
            for r in rows:
                deadline = r["deadline"] or ""
                if not deadline or len(deadline) < 10:
                    continue
                deadline_date = deadline[:10]
                todo_id = r["id"]
                owner = r["owner"] or ""
                task = r["task"] or ""
                meeting_id = r["meeting_id"] or ""
                # 生成去重 key：todo_id + 日期
                dedup_key = hashlib.md5(f"deadline_{todo_id}_{today}".encode()).hexdigest()[:16]
                # 检查今天是否已通知
                existing = conn.execute(
                    "SELECT id FROM notifications WHERE id = ?", (f"notif_{dedup_key}",)
                ).fetchone()
                if existing:
                    continue
                if deadline_date < today:
                    # 已逾期
                    days_overdue = (datetime.now() - datetime.strptime(deadline_date, "%Y-%m-%d")).days
                    conn.execute(
                        """INSERT OR IGNORE INTO notifications
                           (id, user_id, type, title, body, meeting_id, created_at, read)
                           VALUES (?, '', 'warning', ?, ?, ?, ?, 0)""",
                        (f"notif_{dedup_key}", f"待办已逾期 {days_overdue} 天",
                         f"{owner}：{task[:50]}", meeting_id, _now_text())
                    )
                elif deadline_date == today:
                    # 今天到期
                    conn.execute(
                        """INSERT OR IGNORE INTO notifications
                           (id, user_id, type, title, body, meeting_id, created_at, read)
                           VALUES (?, '', 'info', ?, ?, ?, ?, 0)""",
                        (f"notif_{dedup_key}", "待办今日到期",
                         f"{owner}：{task[:50]}", meeting_id, _now_text())
                    )
                else:
                    # 检查是否 3 天内到期
                    days_left = (datetime.strptime(deadline_date, "%Y-%m-%d") - datetime.now()).days
                    if 0 < days_left <= 3:
                        conn.execute(
                            """INSERT OR IGNORE INTO notifications
                               (id, user_id, type, title, body, meeting_id, created_at, read)
                               VALUES (?, '', 'info', ?, ?, ?, ?, 0)""",
                            (f"notif_{dedup_key}", f"待办还有 {days_left} 天到期",
                             f"{owner}：{task[:50]}", meeting_id, _now_text())
                        )
    except Exception as e:
        logger.warning("待办到期检查失败: %s", e)


@app.get("/api/meetings/search")
async def cross_meeting_search(request: Request, q: str = "", limit: int = 30):
    """全文检索：搜索会议标题、议题、转写原文、发言人、AI 纪要。"""
    _get_request_user(request, required=True)
    keyword = (q or "").strip()
    if not keyword or len(keyword) < 2:
        return JSONResponse({"results": [], "query": keyword})
    like = f"%{keyword}%"
    results_map: dict = {}  # meetingId -> result dict

    def _add_result(meeting_id, meeting_title, meeting_date, match_type, match_text):
        """添加或合并搜索结果。"""
        if meeting_id in results_map:
            r = results_map[meeting_id]
            r["matchCount"] += 1
            if match_type not in r["matchTypes"]:
                r["matchTypes"].append(match_type)
        else:
            # 截取匹配片段，高亮关键词
            snippet = match_text[:120] if match_text else ""
            results_map[meeting_id] = {
                "meetingId": meeting_id,
                "meetingTitle": meeting_title or "未命名会议",
                "meetingDate": meeting_date or "",
                "matchType": match_type,
                "matchTypes": [match_type],
                "matchText": snippet,
                "matchCount": 1,
            }

    try:
        with _db_connect() as conn:
            # 1. 搜索会议标题、项目、议题
            for row in conn.execute(
                "SELECT id, title, project, agenda, meeting_date FROM meetings WHERE title LIKE ? OR project LIKE ? OR agenda LIKE ? LIMIT ?",
                (like, like, like, limit)
            ).fetchall():
                match_field = "标题" if keyword in (row["title"] or "") else ("项目" if keyword in (row["project"] or "") else "议题")
                _add_result(row["id"], row["title"], row["meeting_date"], match_field, row["title"] or row["agenda"] or "")

            # 2. 搜索转写原文 + 修正转写 + 发言人
            for row in conn.execute(
                "SELECT mt.meeting_id, mt.transcript, mt.corrected_transcript, mt.speaker_name, "
                "m.title as meeting_title, m.meeting_date "
                "FROM meeting_transcripts mt JOIN meetings m ON mt.meeting_id = m.id "
                "WHERE mt.transcript LIKE ? OR mt.corrected_transcript LIKE ? OR mt.speaker_name LIKE ? LIMIT ?",
                (like, like, like, limit)
            ).fetchall():
                if keyword in (row["speaker_name"] or ""):
                    mtype = "发言人"
                    snippet = row["speaker_name"]
                elif keyword in (row["corrected_transcript"] or ""):
                    mtype = "转写(修正)"
                    snippet = row["corrected_transcript"]
                else:
                    mtype = "转写"
                    snippet = row["transcript"]
                _add_result(row["meeting_id"], row["meeting_title"], row["meeting_date"], mtype, snippet or "")

            # 3. 搜索 AI 纪要（generated_records_json）
            for row in conn.execute(
                "SELECT id, title, meeting_date, generated_records_json FROM meetings WHERE generated_records_json LIKE ? LIMIT ?",
                (like, limit)
            ).fetchall():
                gr_text = row["generated_records_json"] or ""
                # 判断命中的是 summary / decisions / minutes / todos
                import json as _json
                try:
                    gr = _json.loads(gr_text)
                except Exception:
                    gr = {}
                mtype = "纪要"
                snippet = ""
                for field_name, field_label in [("summary", "摘要"), ("decisions", "决议"), ("minutes", "纪要"), ("todos", "待办")]:
                    items = gr.get(field_name) or []
                    for item in items:
                        text = item if isinstance(item, str) else _json.dumps(item, ensure_ascii=False)
                        if keyword in text:
                            mtype = field_label
                            snippet = text[:120]
                            break
                    if snippet:
                        break
                if not snippet:
                    snippet = gr_text[:120]
                _add_result(row["id"], row["title"], row["meeting_date"], mtype, snippet)

    except Exception as e:
        logger.error("全文检索失败: %s", e)
        return JSONResponse({"results": [], "query": keyword, "total": 0, "error": str(e)})

    # 按匹配数排序
    sorted_results = sorted(results_map.values(), key=lambda x: x["matchCount"], reverse=True)[:limit]
    return JSONResponse({"results": sorted_results, "query": keyword, "total": len(sorted_results)})


# ══════════════════════════════════════════════════════════════════════════════════
# 数据导出 API
# ══════════════════════════════════════════════════════════════════════════════════

@app.get("/api/export/meetings")
async def export_all_meetings(request: Request):
    """导出所有会议数据为 JSON（含转写、纪要、待办）。"""
    user = _get_request_user(request, required=True)
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可导出数据")
    try:
        meetings_data = []
        with _db_connect() as conn:
            meetings_rows = conn.execute("SELECT * FROM meetings ORDER BY created_at DESC").fetchall()
            for m in meetings_rows:
                mid = m["id"]
                # 转写
                transcripts = [dict(r) for r in conn.execute(
                    "SELECT * FROM meeting_transcripts WHERE meeting_id = ? ORDER BY client_time, id", (mid,)
                ).fetchall()]
                # 待办
                todos = [dict(r) for r in conn.execute(
                    "SELECT * FROM meeting_todos WHERE meeting_id = ? ORDER BY created_at", (mid,)
                ).fetchall()]
                # 版本（如有）
                versions = [dict(r) for r in conn.execute(
                    "SELECT id, version, editor, edit_summary, created_at FROM meeting_record_versions WHERE meeting_id = ? ORDER BY version", (mid,)
                ).fetchall()]
                # generatedRecords
                gr = {}
                gr_json = m["generated_records_json"] or ""
                if gr_json:
                    try:
                        gr = json.loads(gr_json)
                    except Exception:
                        gr = {}
                meetings_data.append({
                    "id": mid,
                    "title": m["title"],
                    "project": m["project"],
                    "agenda": m["agenda"],
                    "meetingDate": m["meeting_date"],
                    "meetingType": m["meeting_type"],
                    "creator": m["creator"],
                    "phase": m["phase"],
                    "createdAt": m["created_at"],
                    "transcripts": transcripts,
                    "generatedRecords": gr,
                    "todos": todos,
                    "versions": versions,
                })
        export = {
            "exportedAt": _now_text(),
            "meetingCount": len(meetings_data),
            "transcriptCount": sum(len(m["transcripts"]) for m in meetings_data),
            "todoCount": sum(len(m["todos"]) for m in meetings_data),
            "meetings": meetings_data,
        }
        # 返回 JSON 文件下载
        import io
        buf = io.BytesIO()
        buf.write(json.dumps(export, ensure_ascii=False, indent=2).encode("utf-8"))
        buf.seek(0)
        filename = f"ai616_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        from fastapi.responses import StreamingResponse
        return StreamingResponse(
            buf,
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        logger.error("数据导出失败: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/export/meetings/{meeting_id}")
async def export_single_meeting(request: Request, meeting_id: str):
    """导出单个会议完整数据包（JSON + 录音，ZIP 格式）。"""
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    try:
        import io, zipfile
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            with _db_connect() as conn:
                # 会议信息
                m = conn.execute("SELECT * FROM meetings WHERE id = ?", (safe_id,)).fetchone()
                if not m:
                    raise HTTPException(status_code=404, detail="会议不存在")
                meeting_dict = dict(m)
                # 转写
                transcripts = [dict(r) for r in conn.execute(
                    "SELECT * FROM meeting_transcripts WHERE meeting_id = ? ORDER BY client_time, id", (safe_id,)
                ).fetchall()]
                # 待办
                todos = [dict(r) for r in conn.execute(
                    "SELECT * FROM meeting_todos WHERE meeting_id = ?", (safe_id,)
                ).fetchall()]
                # generatedRecords
                gr = {}
                gr_json = meeting_dict.get("generated_records_json") or ""
                if gr_json:
                    try:
                        gr = json.loads(gr_json)
                    except Exception:
                        gr = {}
                # 写入 data.json
                data = {
                    "meeting": {k: v for k, v in meeting_dict.items() if k != "generated_records_json"},
                    "generatedRecords": gr,
                    "transcripts": transcripts,
                    "todos": todos,
                    "exportedAt": _now_text(),
                }
                zf.writestr("data.json", json.dumps(data, ensure_ascii=False, indent=2))
                # 收集录音文件
                audio_dir = MEETING_FILES_DIR / "recordings" / safe_id
                if audio_dir.exists():
                    for f in audio_dir.iterdir():
                        if f.is_file():
                            zf.write(f, f"recordings/{f.name}")
                # 收集材料文件
                materials_dir = MEETING_FILES_DIR / safe_id
                if materials_dir.exists():
                    for f in materials_dir.iterdir():
                        if f.is_file() and f.suffix in ('.docx', '.pdf', '.xlsx'):
                            zf.write(f, f"materials/{f.name}")
        buf.seek(0)
        filename = f"{safe_id}_export_{datetime.now().strftime('%Y%m%d')}.zip"
        from fastapi.responses import StreamingResponse
        return StreamingResponse(
            buf,
            media_type="application/zip",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error("单会议导出失败: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/dashboard/stats")
async def dashboard_stats(request: Request):
    """仪表盘统计数据。"""
    _get_request_user(request, required=True)
    with _db_connect() as conn:
        total = conn.execute("SELECT COUNT(*) as cnt FROM meetings").fetchone()["cnt"]
        active = conn.execute(
            "SELECT COUNT(*) as cnt FROM meetings WHERE phase NOT IN ('已归档', '')"
        ).fetchone()["cnt"]
        archived = conn.execute(
            "SELECT COUNT(*) as cnt FROM meetings WHERE phase = '已归档'"
        ).fetchone()["cnt"]
        total_ts = conn.execute("SELECT COUNT(*) as cnt FROM meeting_transcripts").fetchone()["cnt"]
        recent = conn.execute(
            "SELECT id, title, project, phase, meeting_date, updated_at FROM meetings ORDER BY updated_at DESC LIMIT 5"
        ).fetchall()
    return JSONResponse({
        "totalMeetings": total,
        "activeMeetings": active,
        "archivedMeetings": archived,
        "totalTranscripts": total_ts,
        "recentMeetings": [dict(row) for row in recent],
    })


# ═══════════════════════════════════════════════════════════════════════
# 议题计时器 — 预设时长、会中倒计时、提前结束/延长
# ═══════════════════════════════════════════════════════════════════════

@app.post("/api/meetings/{meeting_id}/agenda-timer/{agenda_id}")
async def agenda_timer_action(
    request: Request,
    meeting_id: str,
    agenda_id: str,
    action: str = "start",  # start | extend | advance | reset
    extend_minutes: int = 5,
):
    """议题计时器操作。"""
    user = _get_request_user(request, required=True)
    safe_id = _safe_meeting_id(meeting_id)
    now = _now_text()
    with MEETINGS_LOCK:
        meetings = _load_meetings()
        meeting = meetings.get(safe_id)
        if not meeting:
            raise HTTPException(status_code=404, detail="会议不存在")
        _check_meeting_access(user, meeting)
        drafts = list(meeting.get("agendaDrafts") or [])

        if action == "start":
            # 标记该议题为活跃，记录开始时间
            for d in drafts:
                if d.get("id") == agenda_id:
                    d["timerStartedAt"] = now
                    d["timerExtended"] = (d.get("timerExtended") or 0)
                else:
                    d["timerStartedAt"] = ""
            meeting["activeAgendaId"] = agenda_id

        elif action == "extend":
            for d in drafts:
                if d.get("id") == agenda_id:
                    d["timerExtended"] = (d.get("timerExtended") or 0) + extend_minutes

        elif action == "advance":
            # 结束当前议题，激活下一个
            current_idx = next((i for i, d in enumerate(drafts) if d.get("id") == agenda_id), -1)
            for d in drafts:
                d["timerStartedAt"] = ""
            if current_idx >= 0 and current_idx + 1 < len(drafts):
                next_draft = drafts[current_idx + 1]
                next_draft["timerStartedAt"] = now
                next_draft["timerExtended"] = (next_draft.get("timerExtended") or 0)
                meeting["activeAgendaId"] = next_draft.get("id")
            else:
                meeting["activeAgendaId"] = ""

        elif action == "reset":
            for d in drafts:
                d["timerStartedAt"] = ""
                d["timerExtended"] = 0
            meeting["activeAgendaId"] = ""

        meeting["agendaDrafts"] = drafts
        meeting["updatedAt"] = now
        meetings[safe_id] = meeting
        _save_meetings(meetings)

    return JSONResponse({
        "success": True,
        "agendaId": agenda_id,
        "action": action,
        "activeAgendaId": meeting.get("activeAgendaId", ""),
        "agendaDrafts": meeting.get("agendaDrafts", []),
    })


def _check_port_available(port: int) -> bool:
    """检查端口是否可用（SO_REUSEADDR 允许重用 TIME_WAIT 端口）"""
    import socket
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
        try:
            s.bind(('0.0.0.0', port))
            return True
        except OSError:
            return False

def _get_port_owner(port: int) -> str:
    """获取占用端口的进程信息"""
    import subprocess
    try:
        result = subprocess.run(
            ['ss', '-tlnp', f'sport = :{port}'],
            capture_output=True, text=True, timeout=5
        )
        return result.stdout.strip() if result.stdout else "未知进程"
    except Exception:
        return "无法获取"

def _safe_create_task(coro, name: str = ""):
    """create_task with exception logging — prevents silent task failures."""
    async def _wrapper():
        try:
            return await coro
        except asyncio.CancelledError:
            raise
        except Exception:
            logger.exception(f"Background task [{name}] failed")
    return asyncio.create_task(_wrapper(), name=name)

if __name__ == "__main__":
    import uvicorn
    import sys
    import time as _time

    PORT = 8002

    # 等待端口释放（最多 15 秒），避免 systemd Restart=on-failure crash loop
    for _attempt in range(15):
        if _check_port_available(PORT):
            break
        if _attempt == 0:
            print(f"⏳ 端口 {PORT} 被占用，等待释放...", file=sys.stderr)
        _time.sleep(1)
    else:
        owner_info = _get_port_owner(PORT)
        print(f"\n❌ 端口 {PORT} 持续被占用（15秒超时）", file=sys.stderr)
        print(f"占用进程:\n{owner_info}", file=sys.stderr)
        sys.exit(1)

    # SQLite + WAL 模式下单 worker 即可高效处理并发读，
    # 多 worker 会导致缓存不一致和 SSE 跨进程失效。
    # 传 app 对象而非字符串，避免 uvicorn fork 子进程导致 systemd 端口冲突。
    uvicorn.run(app, host="0.0.0.0", port=PORT, workers=1,
                timeout_keep_alive=65, timeout_graceful_shutdown=10,
                ws_ping_interval=20, ws_ping_timeout=20)
