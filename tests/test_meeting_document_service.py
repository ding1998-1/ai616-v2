from pathlib import Path
from zipfile import ZipFile

import pytest

from backend.services import meeting_document_service
from backend.services.meeting_document_service import (
    ProofreadRequiredError,
    build_evidence_manifest,
    generate_document_bundle,
    render_evidence_markdown,
    render_formal_markdown,
)


def test_formal_word_marks_authorized_human_override():
    lines = meeting_document_service._formal_record_lines({
        "minutes": [{"agenda": "预算调整", "keyPoints": []}],
        "latestFormalOverride": {
            "operator": "会议秘书",
            "time": "2026-09-02 10:00:00",
            "reason": "已核对录音原文并确认内容",
        },
    })
    text = "\n".join(line for line, _ in lines)
    assert "人工核验说明" in text
    assert "会议秘书" in text
    assert "已核对录音原文并确认内容" in text


def test_empty_keypoints_fall_back_to_verified_map_evidence_without_placeholders():
    records = {
        "minutes": [
            {
                "agenda": "公共收益小程序演示与数据对接",
                "keyPoints": [],
                "basis": {"evidenceValid": False, "quotes": [{"text": "错误挂接内容"}]},
            },
            {"agenda": "没有可靠内容的空议题", "keyPoints": []},
        ],
        "mapResults": [{
            "output": {
                "topics": [{
                    "title": "公共收益管理小程序演示",
                    "evidence": "演示公共收益小程序...需要提供业主花名册完成数据联通",
                }],
            },
        }],
        "decisions": [],
        "risks": [],
        "todos": [],
    }

    enterprise = "\n".join(
        text for _, text in meeting_document_service._enterprise_minutes_blocks({}, records)
    )
    generic = "\n".join(
        text for text, _ in meeting_document_service._template_minutes_lines(records)
    )
    detailed = "\n".join(
        text for text, _ in meeting_document_service._formal_record_lines(records)
    )

    for output in (enterprise, generic, detailed):
        assert "需要提供业主花名册完成数据联通" in output
        assert "暂无已确认的讨论要点" not in output
        assert "没有可靠内容的空议题" not in output
        assert "错误挂接内容" not in output


def _records(proofread=True):
    return {
        "pipeline": "records-v2",
        "proofreadPassed": proofread,
        "generationSnapshot": {
            "provider": "local",
            "model": "qwen-test",
            "pipelineVersion": "records-v2",
            "promptVersion": "records-v2-map-reduce-v1",
            "schemaVersion": "meeting-records-v2",
            "glossaryVersion": "1",
            "chunkPolicy": "audio-file-boundary+time-4000-chars",
            "mapCallCount": 2,
            "reduceCallCount": 1,
        },
        "minutes": [{
            "agenda": "预算调整",
            "status": "已记录",
            "keyPoints": ["确认按程序补充材料"],
            "basis": {
                "timeRange": "00:01:00-00:02:00",
                "quotes": [{"text": "确认按程序补充材料", "segmentId": "s1"}],
            },
        }],
        "decisions": [{
            "content": "同意补充预算材料",
            "type": "决定",
            "status": "待确认",
            "basis": {
                "timeRange": "00:01:00-00:02:00",
                "quotes": [{"text": "确认按程序补充材料", "segmentId": "s1"}],
            },
        }],
        "risks": [{
            "content": "超过权限的支出需履行审批程序",
            "severity": "高",
            "basis": {
                "timeRange": "00:04:00-00:04:10",
                "quotes": [{"text": "需要重新履行审批程序", "segmentId": "s2"}],
            },
        }],
        "disclosures": [{
            "content": "向管理层披露预算变化",
            "audience": "管理层",
            "basis": {
                "timeRange": "00:04:00-00:04:10",
                "quotes": [{"text": "需要重新履行审批程序", "segmentId": "s2"}],
            },
        }],
        "todos": [{
            "task": "补充预算材料",
            "owner": "张三",
            "deadline": "待定",
            "basis": {
                "timeRange": "00:01:00-00:02:00",
                "quotes": [{"text": "确认按程序补充材料", "segmentId": "s1"}],
            },
        }],
        "mapResults": [{
            "chunkId": "chunk-0001",
            "fileId": "audio-1",
            "timeRange": "00:00:00-00:05:00",
            "chunkSegments": [{"segmentId": "s1"}, {"segmentId": "s2"}],
        }],
        "evidenceExceptions": [{
            "field": "decisions",
            "reason": "unsupported_ai_claim",
            "item": {"content": "没有原文支持的虚构决议"},
        }],
    }


def _chronicle():
    return [
        {"id": "s1", "fileId": "audio-1", "start": 60, "end": 120, "speaker": "张三", "text": "确认按程序补充材料"},
        {"id": "s2", "fileId": "audio-1", "start": 250, "end": 260, "speaker": "李四", "text": "需要重新履行审批程序"},
        {"id": "s3", "fileId": "audio-1", "start": 400, "end": 410, "speaker": "李四", "text": "这是一条应保留的完整原始证据长句"},
        {"id": "s4", "fileId": "audio-1", "start": 420, "end": 421, "speaker": "未知", "text": "噪音行", "isNoise": True, "noiseReason": "背景杂音"},
    ]


def test_manifest_keeps_valid_and_excluded_rows_and_marks_two_minute_gap():
    manifest = build_evidence_manifest(_records(), _chronicle())
    assert manifest["coverage"]["coverageRatio"] == 1.0
    assert len(manifest["rows"]) == 3
    assert len(manifest["excludedRows"]) == 1
    assert manifest["gapMarkers"][0]["durationSeconds"] > 120
    assert manifest["excludedRows"][0]["excludeReason"] == "背景杂音"
    assert manifest["reverseIndex"]["s1"] == ["会议纪要[1]", "决议[1]", "待办[1]"]


def test_formal_markdown_is_distilled_and_evidence_markdown_is_lossless():
    records = _records()
    manifest = build_evidence_manifest(records, _chronicle())
    formal = render_formal_markdown({"title": "预算例会", "type": "办公会"}, records)
    evidence = render_evidence_markdown({"title": "预算例会"}, records, manifest)
    assert "这是一条应保留的完整原始证据长句" not in formal
    assert "本正式件仅包含经校对的蒸馏内容" in formal
    assert "这是一条应保留的完整原始证据长句" in evidence
    assert "噪音行" in evidence
    assert "录音断档" in evidence
    assert "chunk-0001" in evidence
    assert "没有原文支持的虚构决议" in evidence
    assert "没有原文支持的虚构决议" not in formal


def test_formal_document_requires_proofread(tmp_path: Path):
    with pytest.raises(ProofreadRequiredError):
        generate_document_bundle(
            "m1", {"title": "例会"}, _records(proofread=False), _chronicle(), tmp_path,
            timestamp="20260821160000",
        )


def test_document_bundle_writes_two_independent_docx_files(tmp_path: Path):
    bundle = generate_document_bundle(
        "m1", {
            "title": "预算例会",
            "type": "办公会",
            "date": "2026-08-21 14:00-16:30",
            "location": "七楼会议室",
            "host": "张三",
            "recorder": "李四",
            "participants": [{"name": "王五"}, {"name": "赵六"}],
            "compiler": "李四",
        },
        _records(), _chronicle(), tmp_path, timestamp="20260821160000",
    )
    formal_path = Path(bundle["formal"]["path"])
    evidence_path = Path(bundle["evidence"]["path"])
    assert formal_path.exists() and evidence_path.exists()
    assert formal_path != evidence_path
    assert bundle["coverage"]["coverageRatio"] == 1.0
    from docx import Document

    formal_doc = Document(formal_path)
    formal_text = "\n".join(
        [paragraph.text for paragraph in formal_doc.paragraphs]
        + [cell.text for table in formal_doc.tables for row in table.rows for cell in row.cells]
    )
    evidence_text = "\n".join(paragraph.text for paragraph in Document(evidence_path).paragraphs)
    assert formal_path.name == "m1_会议纪要及记录_20260821160000.docx"
    assert "会 议 纪 要" in formal_text
    assert "会 议 记 录" in formal_text
    assert len(formal_doc.sections) == 2
    assert len(formal_doc.tables) == 2
    assert "会议性质" in formal_text and "办公会" in formal_text
    assert "会议主题" in formal_text and "预算例会" in formal_text
    assert "待办事项总结" in formal_text and "补充预算材料" in formal_text
    assert "会议名称" in formal_text and "预算例会" in formal_text
    assert "会议地点" in formal_text and "七楼会议室" in formal_text
    assert "主持人" in formal_text and "张三" in formal_text
    assert "记录人" in formal_text and "李四" in formal_text
    assert "参加单位及人员：王五、赵六" in formal_text
    assert "编制人" in formal_text
    assert "生成参数快照" not in formal_text
    assert "这是一条应保留的完整原始证据长句" not in formal_text
    assert "这是一条应保留的完整原始证据长句" in evidence_text
    assert "背景杂音" in evidence_text
    assert "没有原文支持的虚构决议" in evidence_text
    with ZipFile(formal_path) as package:
        assert "customXml/item1.xml" in package.namelist()
        assert "word/theme/theme1.xml" in package.namelist()


def test_document_template_changes_real_word_heading(tmp_path: Path):
    bundle = generate_document_bundle(
        "m2", {"title": "重大事项会"}, _records(), _chronicle(), tmp_path,
        timestamp="20260831120000", template_id="major",
    )
    from docx import Document

    text = "\n".join(paragraph.text for paragraph in Document(bundle["formal"]["path"]).paragraphs)
    assert "三重一大会议纪要" in text
    assert bundle["templateId"] == "major"


def test_enterprise_redhead_template_preserves_artwork_and_appends_record(tmp_path: Path):
    bundle = generate_document_bundle(
        "m-red",
        {
            "title": "重点项目推进会",
            "organization": "示范集团有限公司",
            "issuerDepartment": "集团办公室",
            "date": "2026-09-02 09:00-11:00",
            "location": "第一会议室",
            "host": "陈坚",
            "recorder": "李明",
            "participants": ["项目中心", "财务部", "审计部"],
            "reportTo": "集团领导、相关部门",
            "printNote": "（内部资料，按需印发）",
        },
        _records(),
        _chronicle(),
        tmp_path,
        timestamp="20260902110000",
        template_id="enterprise",
    )
    from docx import Document

    formal_path = Path(bundle["formal"]["path"])
    document = Document(formal_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    combined_text = f"{paragraph_text}\n{table_text}"
    assert bundle["templateId"] == "enterprise"
    assert bundle["templateTitle"] == "政企红头会议纪要"
    assert len(document.sections) == 2
    assert len(document.tables) == 1
    assert "示范集团有限公司" in combined_text
    assert "会议议题：重点项目推进会" in combined_text
    assert "一、预算调整" in combined_text
    assert "会议决议" in combined_text
    assert "风险与披露事项" in combined_text
    assert "待办事项" in combined_text
    assert "会 议 记 录" in combined_text
    assert "集团办公室" in combined_text
    assert "集团领导、相关部门" in combined_text
    assert "20XX" not in combined_text
    assert "XXXX" not in combined_text
    with ZipFile(formal_path) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        assert "会议纪要" in document_xml
        assert "FF0000" in document_xml
        assert "wp:anchor" in document_xml
        assert "v:shape" in document_xml
        assert "（内部资料，按需印发）" in document_xml
        assert "customXml/item1.xml" in package.namelist()


def test_combined_formal_word_expands_action_rows_without_truncation(tmp_path: Path):
    records = _records()
    records["todos"] = [
        {
            "task": f"落实第{index}项工作并提交完整说明材料",
            "owner": f"责任人{index}",
            "deadline": f"2026-09-{index + 10:02d}",
            "basis": {
                "timeRange": "00:01:00-00:02:00",
                "quotes": [{"text": "确认按程序补充材料", "segmentId": "s1"}],
            },
        }
        for index in range(1, 8)
    ]
    bundle = generate_document_bundle(
        "m3",
        {
            "title": "长待办测试会议",
            "type": "总经理办公会",
            "date": "2026-09-02 09:00-11:00",
            "location": "第一会议室",
            "recorder": "会议秘书",
            "participants": ["甲", "乙", "丙"],
        },
        records,
        _chronicle(),
        tmp_path,
        timestamp="20260902110000",
    )
    from docx import Document
    from docx.oxml.ns import qn

    document = Document(bundle["formal"]["path"])
    minutes_table = document.tables[0]
    assert len(minutes_table.rows) == 14
    assert [int(column.get(qn("w:w"))) for column in minutes_table._tbl.tblGrid] == [
        1271, 2835, 1418, 2772,
    ]
    assert "落实第7项工作并提交完整说明材料" in minutes_table.rows[-1].cells[1].text
    assert minutes_table.rows[-1].cells[2].text == "责任人7"
    assert minutes_table.rows[-1].cells[3].text == "2026-09-17"
    assert len(document.sections) == 2
    assert len(document.tables) == 2
