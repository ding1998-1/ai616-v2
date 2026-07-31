#!/usr/bin/env python3
"""
批量重新向量化脚本 — 使用新的按页切分逻辑

将指定目录下所有 PDF/DOCX/TXT 重新解析并写入 ChromaDB，
PDF 保留页码信息，DOCX 按字符切分。
"""

import os
import sys
import io
import json
import re
import uuid
import time
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(PROJECT_ROOT))

import pdfplumber
import docx
import fitz  # PyMuPDF — for rendering scanned PDF pages
import base64
import httpx

# 直接导入后端模块
from backend.config import PERSIST_DIR
from backend.db import _json_loads, _json_dumps

# ChromaDB 设置
os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
os.environ.setdefault("HF_HUB_OFFLINE", "1")

import chromadb
from sentence_transformers import SentenceTransformer

SOURCE_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else PROJECT_ROOT / "城投合规资料" / "宁波产权交易中心"
CHUNK_SIZE = 500
OVERLAP = 50

# 加载模型和 ChromaDB
print("加载 embedding 模型...")
model = SentenceTransformer("shibing624/text2vec-base-chinese")

print(f"连接 ChromaDB: {PERSIST_DIR}")
client = chromadb.PersistentClient(path=PERSIST_DIR)
collection = client.get_or_create_collection("langchain")

# 先获取已有文档的 source 列表，用于去重
existing = collection.get()
existing_sources = set()
if existing and existing.get("metadatas"):
    for meta in existing["metadatas"]:
        src = meta.get("source", "")
        if src:
            existing_sources.add(src)

# 删除已有同 source 的文档（避免重复）
sources_to_clear = set()
for f in SOURCE_DIR.rglob("*"):
    if f.suffix.lower() in (".pdf", ".docx", ".txt", ".md") and not f.name.startswith("."):
        sources_to_clear.add(f.name)

if sources_to_clear & existing_sources:
    print(f"清除 {len(sources_to_clear & existing_sources)} 个已有文档...")
    for src in (sources_to_clear & existing_sources):
        existing_ids = [
            existing["ids"][i] for i, meta in enumerate(existing["metadatas"])
            if meta.get("source") == src
        ]
        if existing_ids:
            collection.delete(ids=existing_ids)
    print("清除完成")


def chunk_text_with_pages(text: str, pages: list = None, source: str = "", doc_id_prefix: str = ""):
    """按页切分（PDF）或按字符切分（DOCX/TXT）。"""
    chunks, metadatas, ids = [], [], []

    if pages:
        total = len(pages)
        for page_num, page_text in enumerate(pages, start=1):
            page_text = page_text.strip()
            if not page_text:
                continue
            start = 0
            while start < len(page_text):
                end = min(start + CHUNK_SIZE, len(page_text))
                chunk_text = page_text[start:end]
                if chunk_text.strip():
                    cid = f"{doc_id_prefix}_p{page_num}_c{len(chunks)}"
                    chunks.append(chunk_text)
                    metadatas.append({
                        "source": source, "doc_id": doc_id_prefix,
                        "page": page_num, "chunk": len(chunks) - 1,
                        "total_pages": total,
                    })
                    ids.append(cid)
                start += CHUNK_SIZE - OVERLAP
    else:
        start = 0
        while start < len(text):
            end = min(start + CHUNK_SIZE, len(text))
            chunk_text = text[start:end]
            if chunk_text.strip():
                cid = f"{doc_id_prefix}_c{len(chunks)}"
                chunks.append(chunk_text)
                metadatas.append({
                    "source": source, "doc_id": doc_id_prefix,
                    "chunk": len(chunks) - 1,
                })
                ids.append(cid)
            start += CHUNK_SIZE - OVERLAP

    return chunks, metadatas, ids


def ocr_page_image(pix, api_key: str, page_num: int) -> str:
    """用 DashScope 百炼 OCR 识别单页扫描件图片。"""
    img_bytes = pix.tobytes("png")
    data_url = f"data:image/png;base64,{base64.b64encode(img_bytes).decode('ascii')}"
    payload = {
        "model": "qwen-vl-ocr",
        "messages": [{
            "role": "user",
            "content": [
                {"type": "text", "text": "请对图片做OCR，提取所有可见中文和数字。只返回纯文本，不要解释。"},
                {"type": "image_url", "image_url": {"url": data_url}},
            ],
        }],
        "temperature": 0,
    }
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    try:
        with httpx.Client(timeout=httpx.Timeout(60.0, connect=10.0), trust_env=False) as client:
            resp = client.post(
                "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions",
                headers=headers, json=payload,
            )
        if resp.status_code >= 400:
            return ""
        data = resp.json()
        choices = data.get("choices") if isinstance(data, dict) else None
        if choices:
            msg = choices[0].get("message") or {}
            content = msg.get("content")
            if isinstance(content, str):
                return content
            if isinstance(content, list):
                return "\n".join(str(p.get("text", "")) for p in content if isinstance(p, dict))
        return ""
    except Exception as e:
        print(f"  ⚠️ OCR第{page_num}页失败: {e}")
        return ""


def parse_file(filepath: Path, ocr_api_key: str = ""):
    """解析文件为文本 + 页面列表。扫描件PDF用OCR回退。"""
    ext = filepath.suffix.lower()
    text, pages = "", None

    if ext == ".pdf":
        with pdfplumber.open(filepath) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        text = "\n\n".join(p for p in pages if p.strip())

        # 扫描件回退：pdfplumber 提取为空时用 OCR
        if not text.strip() and ocr_api_key:
            print("  🔍 扫描件，OCR中...", end=" ", flush=True)
            doc = fitz.open(filepath)
            ocr_pages = []
            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=200)
                if pix.width * pix.height > 4000 * 4000:
                    pix = page.get_pixmap(dpi=120)  # 大图降分辨率
                result = ocr_page_image(pix, ocr_api_key, i + 1)
                ocr_pages.append(result)
            doc.close()
            pages = ocr_pages
            text = "\n\n".join(p for p in pages if p.strip())
        elif not text.strip():
            pages = None  # 无 OCR key，跳过
    elif ext == ".docx":
        doc = docx.Document(str(filepath))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
    elif ext in (".txt", ".md"):
        text = filepath.read_text(encoding="utf-8", errors="replace")

    return text, pages


# ═══ 主流程 ═══════════════════════════════════════════════════════════════════

files = sorted(
    [f for f in SOURCE_DIR.rglob("*") if f.suffix.lower() in (".pdf", ".docx", ".txt", ".md") and not f.name.startswith(".")]
)

# OCR API key
OCR_API_KEY = os.environ.get("DASHSCOPE_API_KEY", "")
if not OCR_API_KEY:
    print("⚠️ 未设置 DASHSCOPE_API_KEY，扫描件PDF将跳过")

print(f"\n共 {len(files)} 个文件待处理\n")

total_chunks = 0
failed = []

for i, filepath in enumerate(files, 1):
    rel = filepath.relative_to(SOURCE_DIR)
    print(f"[{i:2d}/{len(files)}] {rel}", end=" ... ", flush=True)

    try:
        text, pages = parse_file(filepath, OCR_API_KEY)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if not text:
            print("⚠️ 空内容，跳过")
            continue

        doc_id = f"reingest_{uuid.uuid4().hex[:10]}"
        chunks, metadatas, ids = chunk_text_with_pages(
            text, pages=pages, source=filepath.name, doc_id_prefix=doc_id,
        )

        if not chunks:
            print("⚠️ 无有效 chunk，跳过")
            continue

        # Embed + add
        embeddings = model.encode(chunks, normalize_embeddings=True).tolist()
        collection.add(
            embeddings=embeddings,
            documents=chunks,
            metadatas=metadatas,
            ids=ids,
        )

        page_info = ""
        if pages:
            page_info = f"，{len(pages)} 页"
        print(f"✅ {len(chunks)} chunks{page_info}")
        total_chunks += len(chunks)

    except Exception as e:
        print(f"❌ {e}")
        failed.append(str(rel))

print(f"\n══════════════════════════")
print(f"完成: {total_chunks} chunks 入库")
print(f"失败: {len(failed)}")
if failed:
    for f in failed:
        print(f"  ❌ {f}")

print(f"\nChromaDB 总数: {collection.count()} chunks")
