"""普通文档服务。

该服务只负责上传、解析、下载、段落索引和审查版导出。
OnlyOffice 在线编辑器相关的 edit_url/editor/plugin/callback/selection/suggestion
接口不属于本模块，故这里不会创建或依赖这些接口。
"""

from __future__ import annotations

import hashlib
import io
import json
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Iterable

from fastapi import HTTPException


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = PROJECT_ROOT / "data" / "docs"
DOCS_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {"docx", "xlsx", "pptx", "doc", "xls", "ppt", "pdf", "txt", "md"}
PARSE_EXTENSIONS = {"docx", "pdf", "txt", "md"}
_SAFE_NAME_RE = re.compile(r"[^a-zA-Z0-9_.\-\u4e00-\u9fff ]")


def file_ext(filename: str) -> str:
    return Path(str(filename or "")).suffix.lower().lstrip(".")


def safe_original_name(filename: str, fallback: str = "document") -> str:
    """保留用户可读文件名，同时去掉路径和危险字符。"""
    name = Path(str(filename or fallback)).name
    name = _SAFE_NAME_RE.sub("_", name).strip(" .")
    return name[:180] or fallback


def resolve_document(saved_name: str) -> Path:
    """解析存储文件名并阻断目录穿越。"""
    name = Path(str(saved_name or "")).name
    if not name or name != saved_name or name in {".", ".."}:
        raise HTTPException(status_code=400, detail="文件名无效")
    path = (DOCS_DIR / name).resolve()
    try:
        path.relative_to(DOCS_DIR.resolve())
    except ValueError as exc:
        raise HTTPException(status_code=403, detail="禁止访问目录外文件") from exc
    return path


def _extract_docx(raw: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(raw))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def _extract_pdf(raw: bytes) -> str:
    import pdfplumber

    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return "\n\n".join(page for page in pages if page.strip())


def parse_document_bytes(filename: str, raw: bytes) -> str:
    """解析普通文档为纯文本；不调用在线编辑器或外部回调。"""
    if not raw:
        raise HTTPException(status_code=400, detail="文件内容为空")
    ext = file_ext(filename)
    try:
        if ext == "docx":
            text = _extract_docx(raw)
        elif ext == "pdf":
            text = _extract_pdf(raw)
        elif ext in {"txt", "md"}:
            text = raw.decode("utf-8", errors="replace")
        else:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型：.{ext}，仅支持 .docx / .pdf / .txt / .md",
            )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"文件解析失败：{exc}") from exc
    return re.sub(r"\n{3,}", "\n\n", text or "").strip()


def compute_para_id(text: str) -> str:
    normalized = re.sub(r"\s+", "", str(text or "").strip())
    return hashlib.md5(normalized.encode("utf-8")).hexdigest()[:12]


def compute_text_hash(text: str) -> str:
    normalized = re.sub(r"\s+", "", str(text or "").strip())
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def _add_bookmarks_to_docx(raw: bytes, saved_name: str) -> tuple[bytes, list[dict]]:
    """给 docx 非空段落加稳定的导航书签，供普通文档审查定位使用。"""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document(io.BytesIO(raw))
    paragraph_map: list[dict] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        bookmark_name = f"audit_para_{index}"
        paragraph_map.append(
            {
                "para_index": index,
                "bookmark_name": bookmark_name,
                "text_preview": text[:80] + ("..." if len(text) > 80 else ""),
            }
        )
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(index))
        start.set(qn("w:name"), bookmark_name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(index))
        paragraph._p.insert(0, start)
        paragraph._p.insert(1, end)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue(), paragraph_map


def upload_document(filename: str, raw: bytes) -> dict:
    """保存普通文档并返回前端需要的元数据。"""
    original = safe_original_name(filename)
    ext = file_ext(original)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型：.{ext}")
    if not raw:
        raise HTTPException(status_code=400, detail="文件内容为空")

    saved_name = f"{uuid.uuid4().hex}_{original}"
    content = raw
    paragraphs: list[dict] = []
    if ext == "docx":
        try:
            content, paragraphs = _add_bookmarks_to_docx(raw, saved_name)
        except Exception:
            # 书签是增强能力，不能阻止原始文档保存；解析时仍可正常使用原文件。
            content = raw
            paragraphs = []

    path = resolve_document(saved_name)
    path.write_bytes(content)
    if paragraphs:
        meta = {
            "filename": original,
            "saved_as": saved_name,
            "paragraphs": paragraphs,
            "size": len(content),
        }
        (DOCS_DIR / f"{saved_name}.meta.json").write_text(
            json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    return {
        "success": True,
        "filename": original,
        "saved_as": saved_name,
        "size": len(content),
        "paragraph_count": len(paragraphs),
        "message": f"已插入 {len(paragraphs)} 个导航书签" if paragraphs else "文件上传成功",
    }


def list_documents() -> list[dict]:
    files = []
    for path in sorted(DOCS_DIR.iterdir(), key=lambda item: item.stat().st_mtime, reverse=True):
        if not path.is_file() or path.name.endswith((".meta.json", ".suggestions.json", ".review.json")):
            continue
        stat = path.stat()
        files.append({"name": path.name, "size": stat.st_size, "modified": stat.st_mtime})
    return files


def delete_document(saved_name: str) -> None:
    path = resolve_document(saved_name)
    if not path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    path.unlink()
    for suffix in (".meta.json", ".suggestions.json", ".review.json"):
        sidecar = DOCS_DIR / f"{path.name}{suffix}"
        if sidecar.exists():
            sidecar.unlink()


def extract_bookmarks(saved_name: str) -> dict:
    resolve_document(saved_name)
    meta_path = DOCS_DIR / f"{Path(saved_name).name}.meta.json"
    if not meta_path.exists():
        raise HTTPException(status_code=404, detail="该文档未包含段落索引，请确认上传的是 .docx 文件")
    try:
        return json.loads(meta_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise HTTPException(status_code=500, detail="段落索引损坏") from exc


def _enable_track_revisions(document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = document.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.append(OxmlElement("w:trackRevisions"))


def _append_tracked_text(parent, tag: str, text: str, change_id: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    change = OxmlElement(tag)
    change.set(qn("w:id"), str(change_id))
    change.set(qn("w:author"), "AI 合同审查")
    change.set(qn("w:date"), datetime.utcnow().isoformat(timespec="seconds") + "Z")
    text_tag = "w:delText" if tag == "w:del" else "w:t"
    for index, line in enumerate(str(text or "").splitlines() or [""]):
        if index:
            break_run = OxmlElement("w:r")
            break_run.append(OxmlElement("w:br"))
            change.append(break_run)
        run = OxmlElement("w:r")
        node = OxmlElement(text_tag)
        node.set(qn("xml:space"), "preserve")
        node.text = line
        run.append(node)
        change.append(run)
    parent.append(change)


def export_reviewed_docx(saved_name: str, edits: dict) -> tuple[str, str, int]:
    """生成带 Word 修订标记的审查版副本。"""
    from docx import Document

    source = resolve_document(saved_name)
    if not source.exists():
        raise HTTPException(status_code=404, detail="原始文档不存在")
    if file_ext(saved_name) != "docx":
        raise HTTPException(status_code=400, detail="当前仅支持导出 .docx 审查版文件")

    normalized: dict[int, str] = {}
    for key, value in (edits or {}).items():
        try:
            index = int(key)
        except (TypeError, ValueError):
            continue
        if str(value or "").strip():
            normalized[index] = str(value).strip()
    if not normalized:
        raise HTTPException(status_code=400, detail="没有可导出的修改内容")

    document = Document(source)
    _enable_track_revisions(document)
    applied = 0
    for index, replacement in normalized.items():
        if not (0 <= index < len(document.paragraphs)):
            continue
        paragraph = document.paragraphs[index]
        original = paragraph.text
        parent = paragraph._p
        for child in list(parent):
            if not child.tag.endswith("}pPr"):
                parent.remove(child)
        if original:
            _append_tracked_text(parent, "w:del", original, applied * 2 + 1)
        _append_tracked_text(parent, "w:ins", replacement, applied * 2 + 2)
        applied += 1
    if applied == 0:
        raise HTTPException(status_code=422, detail="没有匹配到可回写的段落")

    export_name = f"{source.stem}_留痕审查版_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    export_path = resolve_document(export_name)
    document.save(export_path)
    (DOCS_DIR / f"{export_name}.review.json").write_text(
        json.dumps(
            {
                "source_saved_name": saved_name,
                "exported_saved_name": export_name,
                "applied_count": applied,
                "exported_at": datetime.now().isoformat(),
                "edits": {str(key): value for key, value in normalized.items()},
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return export_name, export_path.name, applied
