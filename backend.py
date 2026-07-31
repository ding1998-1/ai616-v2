from fastapi import FastAPI, HTTPException, Request, UploadFile, File
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List, Any, AsyncIterator
import json
import logging
import os
import asyncio
import httpx
import io
import re
import asyncio
import contextvars
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

from audit_persistence import persistence

# Context variable to hold the current FastAPI Request for cancellation checks
try:
    from sentence_transformers import SentenceTransformer, util
    import torch
    # 全局初始化语义校验模型（paraphrase-multilingual-mpnet-base-v2）
    # 该模型在处理多语言（中英等）语义相似度上表现优异
    validator_model = SentenceTransformer('paraphrase-multilingual-mpnet-base-v2')
    if torch.cuda.is_available():
        validator_model = validator_model.to('cuda')
    elif torch.backends.mps.is_available():
        validator_model = validator_model.to('mps')
    logger.info("【Init】语义校验模型加载成功。")
except Exception as e:
    validator_model = None
    logger.error(f"【Init】语义校验模型加载失败: {e}")

# Context variable to hold the current FastAPI Request for cancellation checks
current_request = contextvars.ContextVar("current_request", default=None)

# ── LangChain imports ────────────────────────────────────────────────────────
from langchain.agents import create_react_agent, AgentExecutor
from langchain_core.prompts import PromptTemplate
from langchain_core.tools import tool
from langchain_core.language_models.chat_models import BaseChatModel
from langchain_core.messages import AIMessage, BaseMessage, HumanMessage, SystemMessage, AIMessageChunk
from langchain_core.outputs import ChatGeneration, ChatResult, ChatGenerationChunk
from langchain_chroma import Chroma
# 使用直接包装 sentence_transformers 绕开 langchain_huggingface 版本冲突
try:
    from sentence_transformers import SentenceTransformer as _SentTransformer
    class _STLangChainEmbeddings:
        def __init__(self, model_name: str):
            self._model = _SentTransformer(model_name)
        def embed_documents(self, texts: List[str]) -> List[List[float]]:
            return self._model.encode(texts, normalize_embeddings=True).tolist()
        def embed_query(self, text: str) -> List[float]:
            return self._model.encode([text], normalize_embeddings=True).tolist()[0]
    _embedding_fn_cls = _STLangChainEmbeddings
except Exception as _emb_import_err:
    _embedding_fn_cls = None
    _emb_import_err_msg = str(_emb_import_err)


# Custom Exception to break the ReAct loop
class ReportGeneratedException(Exception):
    def __init__(self, report_md: str):
        self.report_md = report_md

# Initialize FastAPI
app = FastAPI(title="三重一大合规审核 API")

# Global semaphore to limit concurrent LLM requests
llm_semaphore = asyncio.Semaphore(1)

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Request models ────────────────────────────────────────────────────────────
class ChatRequest(BaseModel):
    matter_type: str
    material_text: str

class ChatResponse(BaseModel):
    success: bool
    message: str
    report: Optional[str] = None

class KBQueryRequest(BaseModel):
    query: str

# ── Custom Qwen LLM (supports reasoning_content) ─────────────────────────────
class QwenThinkingLLM(BaseChatModel):
    """兼容 reasoning_content 字段的 Qwen 客户端"""
    api_base: str = "http://192.168.66.44:8088/v1"
    model_name: str = "default"                      # ← 与技术文档一致，使用 curl 验证过的 model 名
    temperature: float = 0.1
    max_tokens: int = 2048                           # ← 与技术文档一致
    timeout: float = 180.0

    @property
    def _llm_type(self) -> str:
        return "qwen-thinking"

    def _convert_messages(self, messages: List[BaseMessage]) -> List[dict]:
        result = []
        for m in messages:
            if isinstance(m, SystemMessage):
                result.append({"role": "system", "content": m.content})
            elif isinstance(m, HumanMessage):
                result.append({"role": "user", "content": m.content})
            elif isinstance(m, AIMessage):
                result.append({"role": "assistant", "content": m.content or ""})
            else:
                result.append({"role": "user", "content": str(m.content)})
        return result

    def _generate(self, messages: List[BaseMessage], stop=None, run_manager=None, enable_thinking: bool = True, **kwargs) -> ChatResult:
        payload = {
            "model": self.model_name,
            "messages": self._convert_messages(messages),
            "temperature": self.temperature,
            "max_tokens": self.max_tokens,
            "enable_thinking": enable_thinking,
        }
        with httpx.Client(timeout=self.timeout) as client:
            resp = client.post(f"{self.api_base}/chat/completions", json=payload)
            resp.raise_for_status()
            data = resp.json()

        msg = data["choices"][0]["message"]
        # 重要：Agent 调用必须使用 content 字段。
        # reasoning_content 是模型内部思考链，不能传给 ReAct 框架解析，
        # 否则框架会把思考内容当成 Action 解析，导致循环重试。
        text = msg.get("content") or ""
        if not text:
            # 如果 content 为空，说明模型只输出了思考链而没有实质回答，记录警告便于调试
            rc = msg.get("reasoning_content", "")
            logger.warning(
                f"【LLM】content 为空，模型可能仅输出了思考链（reasoning_content 长度={len(rc)}）。"
                f"建议检查模型是否支持当前请求格式，当前返回空字符串。"
            )
        return ChatResult(generations=[ChatGeneration(message=AIMessage(content=text))])

    async def _agenerate(self, messages, stop=None, run_manager=None, **kwargs):
        return await asyncio.get_event_loop().run_in_executor(
            None, lambda: self._generate(messages, stop, run_manager, **kwargs)
        )

    async def _astream(
        self, messages: List[BaseMessage], stop=None, run_manager=None, enable_thinking: bool = True, **kwargs
    ) -> AsyncIterator[ChatGenerationChunk]:
        """Stream chunks from Qwen, tagging reasoning vs content separately."""
        payload = {
            "model": self.model_name,
            "messages": self._convert_messages(messages),
            "temperature": self.temperature,
            "max_tokens": self.max_tokens,
            "stream": True,
            "enable_thinking": enable_thinking,
        }
        async with httpx.AsyncClient(timeout=self.timeout) as client:
            async with client.stream("POST", f"{self.api_base}/chat/completions", json=payload) as response:
                response.raise_for_status()
                async for line in response.aiter_lines():
                    # Check for client disconnection context
                    req = current_request.get()
                    if req and getattr(req, "is_disconnected", None):
                        if await req.is_disconnected():
                            logger.info("【LLM】Client disconnected, aborting LLM stream.")
                            raise asyncio.CancelledError()

                    line = line.strip()
                    if not line or not line.startswith("data: "):
                        continue
                    data_str = line[6:]
                    if data_str == "[DONE]":
                        break
                    try:
                        data = json.loads(data_str)
                        delta = data["choices"][0].get("delta", {})
                        r_text = delta.get("reasoning_content") or ""
                        c_text = delta.get("content") or ""

                        # Yield thinking chunk (reasoning_content) tagged as 'thinking'
                        if r_text:
                            chunk = ChatGenerationChunk(
                                message=AIMessageChunk(
                                    content=r_text,
                                    additional_kwargs={"chunk_type": "thinking"},
                                )
                            )
                            if run_manager:
                                await run_manager.on_llm_new_token(r_text, chunk=chunk)
                            yield chunk

                        # Yield answer chunk (content) tagged as 'content'
                        if c_text:
                            chunk = ChatGenerationChunk(
                                message=AIMessageChunk(
                                    content=c_text,
                                    additional_kwargs={"chunk_type": "content"},
                                )
                            )
                            if run_manager:
                                await run_manager.on_llm_new_token(c_text, chunk=chunk)
                            yield chunk
                    except Exception:
                        pass

llm = QwenThinkingLLM()

# ── Compliance rules DB ───────────────────────────────────────────────────────
# 严格基于《中国核能电力股份有限公司 “三重一大”决策制度实施办法》（CG-AB-210 Rev.H）
RULES_DB = {
    "重大决策": {
        "强制要求": "必须经党委前置研究讨论；集体决策；法律审查；会议纪要存档。",
        "禁止事项": "禁止个人或少数人决定；无会议纪要；未经法律审查。",
        "决策程序": ["提出书面建议书", "党支部审查列入", "承办部门拟方案", "征求意见", "院办公室报告", "院务会议集体讨论表决", "实施与监督"],
        "责任主体": "党委书记/董事长主持；主管领导论证；法律合规部审查；纪检监督部监督。"
    },
    "重大项目安排": {
        "强制要求": "必须可行性报告、风险评估、法律审查；党委前置。",
        "禁止事项": "禁止超预算、无审批。",
        "决策程序": ["项目审查", "专家论证", "征求意见", "会议决策", "公示", "正式审批"],
        "责任主体": "战略规划部论证；法律合规部审查。"
    },
    "大额度资金运作": {
        "强制要求": "必须资金使用计划；双人签字或集体审批。",
        "禁止事项": "禁止私下转账、无审计记录。",
        "决策程序": ["安排预算", "党组集体研究", "公开公示", "资金拨付"],
        "责任主体": "财务部门执行；审计部监督。"
    },
    "重要人事任免": {
        "强制要求": "坚持党管干部；事先征求纪检意见；集体决定；任前公示；试用期考核。",
        "禁止事项": "禁止个人决定。",
        "决策程序": ["民主推荐", "组织考察", "会议决定", "任前公示", "试用1年", "正式任免"],
        "责任主体": "人力资源部考察；纪检监督部意见。"
    }
}

# ── Agent Tools ───────────────────────────────────────────────────────────────
@tool
def extract_rules(matter_type: str) -> str:
    """从制度提取当前事项的强制要求、禁止事项、决策程序、责任主体"""
    data = RULES_DB.get(matter_type.strip(), {"error": "未匹配事项类型"})
    return json.dumps(data, ensure_ascii=False, indent=2)

@tool
def validate_material(material_text: str, rules_text: str) -> str:
    """与规则交叉校验，返回每条规则状态 + 证据（基于 SentenceTransformer 语义相似度）"""
    try:
        rules = json.loads(rules_text)
        if "error" in rules:
            return json.dumps({"error": rules["error"]}, ensure_ascii=False)
        
        if validator_model is None:
            return json.dumps({"error": "语义校验模型未加载，请检查后台日志。"}, ensure_ascii=False)

        report = []
        # 将材料按句分割
        sentences = [s.strip() for s in material_text.replace('\n', ' ').split('。') if s.strip()]
        if not sentences:
            return json.dumps({"error": "材料内容为空或无法解析句子"}, ensure_ascii=False)

        # 预计算句子向量
        sentence_embeddings = validator_model.encode(sentences, convert_to_tensor=True)

        for key in ["强制要求", "禁止事项"]:
            if key not in rules:
                continue
            
            rule_content = rules[key]
            # 计算规则与所有句子的相似度
            rule_embedding = validator_model.encode(rule_content, convert_to_tensor=True)
            cos_sims = util.cos_sim(rule_embedding, sentence_embeddings)[0]
            
            max_sim = cos_sims.max().item()
            best_idx = cos_sims.argmax().item()
            
            status = "合规" if max_sim > 0.72 else "⚠️ 不合规"
            evidence = sentences[best_idx] if max_sim > 0.5 else "无明显证据"
            
            report.append({
                "规则类型": key,
                "规则要求": rule_content,
                "状态": status,
                "置信度": round(max_sim, 2),
                "证据原文": evidence
            })
            
        return json.dumps(report, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Error in validate_material: {e}")
        return json.dumps({"error": str(e)}, ensure_ascii=False)

@tool
def check_procedure_completeness(material_text: str, rules_text: str) -> str:
    """检查决策程序是否完整覆盖"""
    try:
        rules = json.loads(rules_text)
        if "error" in rules:
            return json.dumps({"error": rules["error"]}, ensure_ascii=False)
        steps = rules.get("决策程序", [])
        report = [{"环节": step, "状态": "已覆盖" if step in material_text else "缺失"} for step in steps]
        return json.dumps(report, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)

@tool
def identify_responsibility(material_text: str, rules_text: str) -> str:
    """检查责任主体和监督是否落实"""
    try:
        rules = json.loads(rules_text)
        if "error" in rules:
            return json.dumps({"error": rules["error"]}, ensure_ascii=False)
        subject = rules.get("责任主体", "")
        if subject and subject in material_text:
            return "责任主体明确，监督机制提及。"
        return "责任主体或监督缺失，请补充。"
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)

@tool
def generate_compliance_report(results: str) -> str:
    """综合所有工具结果，生成国企公文规范的 Markdown 合规审核报告"""
    try:
        prompt = f"""你是资深三重一大合规审核专家，请严格按照国企公文规范，基于以下工具执行结果，输出一份排版精美的 Markdown 合规审核报告。

【工具执行结果原始数据（仅供参考，不得在报告中出现）】
{results}

=========
【输出格式要求】
1. 报告全程使用中文。
2. 输出格式为纯 Markdown 文本，章节用 ## 标题区隔，数据区域用表格呈现。
3. 在报告最开头，直接输出以下风险雷达 XML 数据块（根据实际审查结果，将 status 改为 green/yellow/red）：
   <risk_radar>
   <item status="green">党委前置审查(通过)</item>
   <item status="yellow">程序完整性(缺少某环节)</item>
   <item status="red">大额资金审批(违规)</item>
   </risk_radar>
4. 紧接 XML 块之后，输出以下六个章节：
   - ## 一、审核基本信息（表格形式：审核类型、审核日期、审核结论）
   - ## 二、风险等级评定（使用 ⚠️ 高风险 / 🔶 中风险 / 🟢 低风险，说明评定理由）
   - ## 三、违规事项与证据清单（务必在叙述违规/合规项时，用特殊的 Markdown 链接语法标出证据原文。例如：[此项目缺乏财务审计报告](evidence:"未见财务部门签字的审计文本")）
   - ## 四、程序完整性核查（Markdown 表格：程序环节 | 状态 | 备注）
   - ## 五、责任主体认定（说明责任人/监督部门落实情况）
   - ## 六、整改建议（分条陈述，每条格式：**建议N**：具体措施。关键要求：在每条具体建议的末尾，必须加上一个特殊的生成按钮链接。例如：[✨ 一键生成补正模板](remediate:"为这个项目起草缺少前置审查环节的情况说明及补发文的模版")）
5. 报告末尾附：> 📋 本报告由 AI 合规审核系统自动生成，仅供参考，最终结论以人工复核为准。

现在从 <risk_radar> 开始输出："""
        response = llm.invoke(prompt)
        # Force a stop of the ReAct agent loop by raising our custom exception
        raise ReportGeneratedException(response.content)
    except Exception as e:
        logger.error(f"Error in generate_compliance_report: {e}")
        # If it's our own stop exception, re-raise it so the upstream try-except catches it
        if isinstance(e, ReportGeneratedException):
            raise e
        return f"报告生成失败: {e}"

tools = [extract_rules, validate_material, check_procedure_completeness, identify_responsibility, generate_compliance_report]

# ── Removed ReAct Agent due to infinite loop bugs ──
# We will use a deterministic 5-step procedural pipeline instead in audit_stream

# ── ChromaDB setup ────────────────────────────────────────────────────────────
PERSIST_DIR = "/Users/macos/Documents/ai 合规 demo/chroma_db"
vectorstore = None

try:
    if os.path.exists(PERSIST_DIR):
        logger.info(f"Loading ChromaDB from {PERSIST_DIR}")
        if _embedding_fn_cls is not None:
            _embed_fn = _embedding_fn_cls("shibing624/text2vec-base-chinese")
            import chromadb
            _client = chromadb.PersistentClient(path=PERSIST_DIR)
            vectorstore = Chroma(
                client=_client,
                collection_name="langchain",
                embedding_function=_embed_fn,
            )
        else:
            logger.warning(f"Embedding 模块加载失败：{_emb_import_err_msg}，知识库功能不可用。")
        logger.info("ChromaDB loaded successfully.")
    else:
        logger.warning(f"ChromaDB not found at {PERSIST_DIR}. Run ingest_knowledge.py first.")
except Exception as e:
    logger.error(f"Failed to load ChromaDB: {e}")

# ── API Routes ────────────────────────────────────────────────────────────────
@app.get("/")
async def root():
    return {"message": "三重一大合规审核 API 服务已启动"}

@app.get("/matter-types")
async def get_matter_types():
    return {"matter_types": list(RULES_DB.keys())}

def sanitize_report(text: str) -> str:
    """Strip any ReAct/JSON artifacts that may leak into the final report."""
    # Remove fenced code blocks (```...```)
    text = re.sub(r'```[\s\S]*?```', '', text)
    # Remove Action / Observation / Thought lines (ReAct scaffolding)
    text = re.sub(r'^(Action|Observation|Thought|Action Input)\s*:.*$', '', text, flags=re.MULTILINE)
    # Remove lines that are pure JSON objects/arrays
    text = re.sub(r'^\s*[\{\[][\s\S]*?[\}\]]\s*$', '', text, flags=re.MULTILINE)
    # Remove "Final Answer:" prefix if the model echoed it
    text = re.sub(r'^Final Answer\s*:\s*', '', text, flags=re.MULTILINE)
    # Collapse excess blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


@app.get("/api/audit_history")
async def get_audit_history():
    """获取审核历史记录"""
    return {"success": True, "history": persistence.get_history()}


@app.post("/audit_stream")
async def audit_stream(http_request: Request, body: ChatRequest):
    """流式执行合规审核 (SSE)"""
    logger.info(f"收到流式审核请求 - 事项类型: {body.matter_type}")
    material = body.material_text  # Allow full context length
    query = f"审核事项：{body.matter_type}\n材料内容：{material}"
    
    # Set the current request in context so the LLM can check for disconnection
    current_request.set(http_request)

    async def event_generator():
        try:
            if llm_semaphore.locked():
                yield f"data: {json.dumps({'type': 'queue_warning', 'content': '模型当前繁忙，您的请求已进入排队，请稍候...'}, ensure_ascii=False)}\n\n"
            
            async with llm_semaphore:
                # ── Deterministic Procedural Pipeline ──
                # Step 1: extract_rules
                yield f"data: {json.dumps({'type': 'tool_start', 'tool': 'extract_rules'}, ensure_ascii=False)}\n\n"
                rules_res = extract_rules.invoke({"matter_type": body.matter_type})
                yield f"data: {json.dumps({'type': 'tool_end', 'tool': 'extract_rules'}, ensure_ascii=False)}\n\n"

                # Step 2: validate_material
                yield f"data: {json.dumps({'type': 'tool_start', 'tool': 'validate_material'}, ensure_ascii=False)}\n\n"
                val_res = validate_material.invoke({"material_text": material, "rules_text": rules_res})
                yield f"data: {json.dumps({'type': 'tool_end', 'tool': 'validate_material'}, ensure_ascii=False)}\n\n"

                # Step 3: check_procedure_completeness
                yield f"data: {json.dumps({'type': 'tool_start', 'tool': 'check_procedure_completeness'}, ensure_ascii=False)}\n\n"
                proc_res = check_procedure_completeness.invoke({"material_text": material, "rules_text": rules_res})
                yield f"data: {json.dumps({'type': 'tool_end', 'tool': 'check_procedure_completeness'}, ensure_ascii=False)}\n\n"

                # Step 4: identify_responsibility
                yield f"data: {json.dumps({'type': 'tool_start', 'tool': 'identify_responsibility'}, ensure_ascii=False)}\n\n"
                resp_res = identify_responsibility.invoke({"material_text": material, "rules_text": rules_res})
                yield f"data: {json.dumps({'type': 'tool_end', 'tool': 'identify_responsibility'}, ensure_ascii=False)}\n\n"

                # Step 5: generate_compliance_report (streamed directly)
                yield f"data: {json.dumps({'type': 'tool_start', 'tool': 'generate_compliance_report'}, ensure_ascii=False)}\n\n"
                
                combined_results = f"""
1. 规则提取结果：\n{rules_res}
2. 材料合规校验结果：\n{val_res}
3. 程序完整性核查：\n{proc_res}
4. 责任主体落实：\n{resp_res}
                """
                
                prompt_text = f"""你是资深三重一大合规审核专家，请严格按照国企公文规范，基于以下前置工具执行结果，输出一份排版精美的 Markdown 合规审核报告。材料内容详见下文。

【材料内容】
{material}

【工具执行结果原始数据（仅供参考，不得在报告中出现）】
{combined_results}

=========
【输出格式要求】
1. 报告全程使用中文。
2. 输出格式为纯 Markdown 文本，章节用 ## 标题区隔，数据区域用表格呈现。
3. 在报告最开头，直接输出以下风险雷达 XML 数据块（根据实际审查结果，将 status 改为 green/yellow/red）：
   <risk_radar>
   <item status="green">党委前置审查(通过)</item>
   <item status="yellow">程序完整性(缺少某环节)</item>
   <item status="red">大额资金审批(违规)</item>
   </risk_radar>
4. 紧接 XML 块之后，输出以下六个章节：
   - ## 一、审核基本信息（表格形式：审核类型、审核日期、审核结论）
   - ## 二、风险等级评定（使用 ⚠️ 高风险 / 🔶 中风险 / 🟢 低风险，说明评定理由）
   - ## 三、违规事项与证据清单（务必在叙述违规/合规项时，用特殊的 Markdown 链接语法标出证据原文。例如：[此项目缺乏财务审计报告](evidence:"未见财务部门签字的审计文本")）
   - ## 四、程序完整性核查（Markdown 表格：程序环节 | 状态 | 备注）
   - ## 五、责任主体认定（说明责任人/监督部门落实情况）
   - ## 六、整改建议（分条陈述，每条格式：**建议N**：具体措施。关键要求：在每条具体建议的末尾，必须加上一个特殊的生成按钮链接。例如：[✨ 一键生成补正模板](remediate:"为这个项目起草缺少前置审查环节的情况说明及补发文的模版")）
5. 报告末尾附：> 📋 本报告由 AI 合规审核系统自动生成，仅供参考，最终结论以人工复核为准。

现在从 <risk_radar> 开始输出："""
                
                messages = [
                    SystemMessage(content="你是资深三重一大合规审核专家。直接输出报告内容，不要做任何前置检查或规划。"),
                    HumanMessage(content=prompt_text),
                    AIMessage(content="<risk_radar>\n")
                ]
                
                full_report = "<risk_radar>\n"
                async for chunk in llm._astream(messages):
                    if await http_request.is_disconnected():
                        logger.info("【Audit】前台断开，中断生成。")
                        raise asyncio.CancelledError()
                    
                    chunk_text = chunk.message.content
                    chunk_type = chunk.message.additional_kwargs.get("chunk_type", "content")
                    
                    if chunk_text:
                        if chunk_type == "thinking":
                            yield f"data: {json.dumps({'type': 'thinking_chunk', 'content': chunk_text}, ensure_ascii=False)}\n\n"
                        else:
                            full_report += chunk_text
                            yield f"data: {json.dumps({'type': 'llm_chunk', 'content': chunk_text}, ensure_ascii=False)}\n\n"
                
                # Report finished
                yield f"data: {json.dumps({'type': 'tool_end', 'tool': 'generate_compliance_report'}, ensure_ascii=False)}\n\n"
                
                clean = sanitize_report(full_report)
                
                # Save to persistent history
                try:
                    persistence.save_audit(
                        matter_type=body.matter_type,
                        material=material,
                        report=clean,
                        results={
                            "rules": rules_res,
                            "validation": val_res,
                            "procedure": proc_res,
                            "responsibility": resp_res
                        }
                    )
                    logger.info("【Audit】记录已成功存档。")
                except Exception as save_err:
                    logger.error(f"【Audit】记录存档失败: {save_err}")

                yield f"data: {json.dumps({'type': 'report', 'content': clean}, ensure_ascii=False)}\n\n"
                yield 'data: {"type": "done"}\n\n'

        except Exception as e:
            logger.error(f"Audit stream error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'detail': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")

@app.post("/kb_stream")
async def kb_stream(http_request: Request, body: KBQueryRequest):
    """流式查询企业知识库 (RAG)"""
    logger.info(f"收到知识库查询: {body.query}")
    
    # Set the current request in context so the LLM can check for disconnection
    current_request.set(http_request)

    async def kb_generator():
        try:
            if vectorstore is None:
                raise ValueError("知识库未初始化，请联系管理员建立索引。")

            # 1. Retrieve relevant docs
            yield f"data: {json.dumps({'type': 'tool_start', 'tool': '检索本地文档库(MMR)'}, ensure_ascii=False)}\n\n"
            docs = vectorstore.max_marginal_relevance_search(body.query, k=4, fetch_k=12)
            context = "\n\n".join([f"【参考资料 {i+1}】\n{d.page_content}" for i, d in enumerate(docs)])
            yield f"data: {json.dumps({'type': 'tool_end', 'tool': f'检索完毕，找到 {len(docs)} 条相关资料'}, ensure_ascii=False)}\n\n"

            prompt_text = f"""请根据以下内部资料回答用户的提问。

【内部资料开始】
{context}
【内部资料结束】

要求：
1. 请详细、专业地回答。
2. 答案必须严格基于上述资料，不得捏造不存在的政策；如果资料中没有相关直接信息，请说明当前本地知识库并未涉及此部分细节。
3. 在回答中标出参考了哪些资料（如参考资料1）。
4. 必须全部使用中文（简体）进行思考和回答，禁止出现英文。

用户提问：{body.query}"""
            messages = [
                SystemMessage(content="你是一个专业的城投企业合规法务专家。请始终严格使用中文（简体）回答问题，包括思考过程和最终结论，绝对禁止使用英文。"),
                HumanMessage(content=prompt_text)
            ]

            # 3. Stream LLM response
            full_response = ""
            
            if llm_semaphore.locked():
                yield f"data: {json.dumps({'type': 'queue_warning', 'content': '模型当前繁忙，您的请求已进入排队，请稍候...'}, ensure_ascii=False)}\n\n"
                
            async with llm_semaphore:
                async for chunk in llm._astream(messages, enable_thinking=False):
                    if await http_request.is_disconnected():
                        logger.info("【KB】前台断开，中断生成。")
                        raise asyncio.CancelledError()
                    text = chunk.message.content
                    chunk_type = chunk.message.additional_kwargs.get("chunk_type", "content")
                    if text:
                        if chunk_type == "thinking":
                            yield f"data: {json.dumps({'type': 'thinking_chunk', 'content': text}, ensure_ascii=False)}\n\n"
                        else:
                            full_response += text
                            yield f"data: {json.dumps({'type': 'llm_chunk', 'content': text}, ensure_ascii=False)}\n\n"

            # 4. Final marker
            yield f"data: {json.dumps({'type': 'report', 'content': full_response}, ensure_ascii=False)}\n\n"
            yield 'data: {"type": "done"}\n\n'

        except Exception as e:
            logger.error(f"KB stream error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'detail': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(kb_generator(), media_type="text/event-stream")


# ── File parsing endpoint ──────────────────────────────────────────────────────
@app.post("/api/parse_file")
async def parse_file(file: UploadFile = File(...)):
    """Parse a Word (.docx) or PDF file and return extracted plain text."""
    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    raw = await file.read()
    
    # Save file to persistent storage
    try:
        unique_name = persistence.save_file(filename, raw)
        logger.info(f"文件已保存至: {unique_name}")
    except Exception as e:
        logger.warning(f"文件保存失败: {e}")

    text = ""

    try:
        if ext == "docx":
            import docx
            doc = docx.Document(io.BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = "\n".join(paragraphs)

        elif ext == "pdf":
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                pages = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
            text = "\n\n".join(pages)

        elif ext in ("txt", "md"):
            text = raw.decode("utf-8", errors="replace")

        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件类型：.{ext}，仅支持 .docx / .pdf / .txt")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文件解析失败: {e}")
        raise HTTPException(status_code=500, detail=f"文件解析失败：{str(e)}")

    # Clean up whitespace
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    char_count = len(text)
    logger.info(f"文件解析完成: {filename}, {char_count} 字符")
    return JSONResponse({"text": text, "filename": filename, "char_count": char_count})


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
